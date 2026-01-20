"""
AI-Powered Risk Document Ingestion (Complete)
- Reads PDF/DOCX/XLSX/TXT
- Extracts risk data using AI (OpenAI or Ollama - configurable)
- Fills missing fields with focused prompts
- Returns normalized, DB-ready JSON for `risk` table

Configuration:
- Set RISK_AI_PROVIDER='openai' or 'ollama' in environment or settings
- Default: 'ollama' if both configured, else 'openai'
- OpenAI requires: OPENAI_API_KEY
- Ollama requires: OLLAMA_BASE_URL (uses optimized quantized models)
"""

import os
import re
import json
import math
import time
import tempfile
from datetime import date, datetime
from typing import Any, Optional

import requests
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser

# RBAC imports
from ...rbac.decorators import rbac_required

# Phase 2 Optimizations
from ...utils.ai_cache import cached_llm_call
from ...utils.document_preprocessor import preprocess_document, calculate_document_hash
from ...utils.few_shot_prompts import get_field_extraction_prompt
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)
# Phase 3 Optimizations
from ...utils.rag_system import (
    add_document_to_rag,
    retrieve_relevant_context,
    build_rag_prompt,
    is_rag_available,
    get_rag_stats
)
from ...utils.model_router import (
    route_model,
    track_system_load,
    get_current_system_load
)
from ...utils.file_compression import decompress_if_needed
from ...routes.Global.s3_fucntions import create_direct_mysql_client
from ...utils.request_queue import (
    rate_limit_decorator,
    process_with_queue,
    get_queue_status
)

# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import Risk  # , Users  # (Users not needed here but you can import if required)


# =========================
# AI PROVIDER CONFIG (OpenAI or Ollama)
# =========================
# AI Provider Configuration - Use Django settings
from django.conf import settings

# Provider selection: 'openai' or 'ollama' (default: 'ollama' if both configured, else 'openai')
AI_PROVIDER = getattr(settings, 'RISK_AI_PROVIDER', os.environ.get('RISK_AI_PROVIDER', 'ollama')).lower()

# OpenAI Configuration
OPENAI_API_KEY = getattr(settings, 'OPENAI_API_KEY', None)
OPENAI_API_URL = "https://api.openai.com/v1/chat/completions"
OPENAI_MODEL = getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini')

# Ollama Configuration (Optimized)
OLLAMA_BASE_URL = getattr(settings, 'OLLAMA_BASE_URL', 'http://13.205.15.232:11434').rstrip('/')
OLLAMA_TIMEOUT = getattr(settings, 'OLLAMA_TIMEOUT', 600)
OLLAMA_TEMPERATURE = getattr(settings, 'OLLAMA_TEMPERATURE', 0.1)
OLLAMA_SEED = getattr(settings, 'OLLAMA_SEED', 42)
OLLAMA_MODEL_DEFAULT = getattr(settings, 'OLLAMA_MODEL', 'llama3.2:3b-instruct-q4_K_M')
OLLAMA_MODEL_FAST = 'llama3.2:1b-instruct-q4_K_M'  # For simple tasks
OLLAMA_MODEL_COMPLEX = 'llama3:8b-instruct-q4_K_M'  # For complex reasoning

# Auto-select provider if not explicitly set
if AI_PROVIDER not in ['openai', 'ollama']:
    # Auto-detect: prefer Ollama if configured, else OpenAI
    if OLLAMA_BASE_URL and OLLAMA_MODEL_DEFAULT:
        AI_PROVIDER = 'ollama'
        print("🔍 Auto-selected Ollama as AI provider (Ollama configured)")
    elif OPENAI_API_KEY:
        AI_PROVIDER = 'openai'
        print("🔍 Auto-selected OpenAI as AI provider (OpenAI configured)")
    else:
        AI_PROVIDER = 'openai'  # Default fallback
        print("⚠️  WARNING: No AI provider fully configured, defaulting to OpenAI")

# Print configuration
print(f"\n🤖 AI Provider Configuration:")
print(f"   Selected Provider: {AI_PROVIDER.upper()}")

if AI_PROVIDER == 'openai':
    if not OPENAI_API_KEY:
        print("⚠️  WARNING: OPENAI_API_KEY not found in Django settings!")
        print("   Please set OPENAI_API_KEY in your .env file.")
    else:
        print(f"🌐 OpenAI Configuration:")
        print(f"   API URL: {OPENAI_API_URL}")
        print(f"   Model: {OPENAI_MODEL}")
        print(f"   API Key: {'*' * (len(OPENAI_API_KEY) - 4) + OPENAI_API_KEY[-4:]}")
elif AI_PROVIDER == 'ollama':
    print(f"🚀 Ollama Configuration (OPTIMIZED):")
    print(f"   Base URL: {OLLAMA_BASE_URL}")
    print(f"   Default Model: {OLLAMA_MODEL_DEFAULT}")
    print(f"   Fast Model: {OLLAMA_MODEL_FAST}")
    print(f"   Complex Model: {OLLAMA_MODEL_COMPLEX}")
    print(f"   Temperature: {OLLAMA_TEMPERATURE}")
    print(f"   Timeout: {OLLAMA_TIMEOUT}s")

# Only the columns you want to fill (you said: do NOT fetch RiskId, ComplianceId, FrameworkId)
RISK_DB_FIELDS = [
    "RiskTitle",
    "Criticality",
    "PossibleDamage",
    "Category",
    "RiskType",
    "BusinessImpact",
    "RiskDescription",
    "RiskLikelihood",
    "RiskImpact",
    "RiskExposureRating",
    "RiskPriority",
    "RiskMitigation",
    "CreatedAt",
    "RiskMultiplierX",
    "RiskMultiplierY",
]

# Canonical choices / constraints to stabilize LLM outputs
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES    = ["Low", "Medium", "High", "Critical"]
CATEGORY_HINTS      = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]
RISKTYPE_HINTS      = ["Current", "Residual", "Inherent", "Emerging", "Accepted"]
DATE_FORMAT_HINT    = "YYYY-MM-DD (ISO)"

# Field-specific micro-prompts (used when a single field is missing/invalid)
# NOTE: RiskTitle is NEVER inferred by AI - it must always come from the document
FIELD_PROMPTS = {
    "Criticality": f"Return one of: {CRITICALITY_CHOICES}.",
    "PossibleDamage": "Describe concrete damages (data loss, downtime, penalties, reputation). Be concise (1–2 sentences).",
    "Category": f"Return one category from this list (best fit): {CATEGORY_HINTS}. If none fits, pick the closest.",
    "RiskType": f"Return one of: {RISKTYPE_HINTS}.",
    "BusinessImpact": "Explain business impact in business terms (SLA breach, revenue, compliance). 1–2 sentences.",
    "RiskDescription": "Write a precise description (1–3 sentences) of how/why the risk arises in this context.",
    "RiskLikelihood": "Return an integer 1–10 (1=rare, 10=almost certain).",
    "RiskImpact": "Return an integer 1–10 (1=negligible, 10=catastrophic).",
    "RiskExposureRating": "Return a float (0–100). If missing, use Likelihood*Impact as proxy.",
    "RiskPriority": f"Return one of: {PRIORITY_CHOICES}. Base it on exposure + criticality.",
    "RiskMitigation": "Return 2–4 actionable mitigation steps as one paragraph or a bullet-style JSON list.",
    "CreatedAt": f"Return a plausible assessment date in {DATE_FORMAT_HINT}; if unknown, use today's date.",
    "RiskMultiplierX": "Return a float in 0.1–1.5 reflecting org weighting factor X (defaults ~0.5 if unknown).",
    "RiskMultiplierY": "Return a float in 0.1–1.5 reflecting org weighting factor Y (defaults ~0.5 if unknown).",
}

# Strict JSON schema block the LLM must follow
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON array. No markdown, no code blocks, no explanations.
Start with [ and end with ]. Use proper JSON syntax with double quotes.

Example structure (return array of risks like this):
[
  {{
    "RiskTitle": "Brief risk title here",
    "Criticality": "Medium",
    "PossibleDamage": "Description of damage",
    "Category": "Operational",
    "RiskType": "Current",
    "BusinessImpact": "Business impact description",
    "RiskDescription": "Risk description",
    "RiskLikelihood": 5,
    "RiskImpact": 6,
    "RiskExposureRating": 30.0,
    "RiskPriority": "Medium",
    "RiskMitigation": "Mitigation steps",
    "CreatedAt": "{date.today().isoformat()}",
    "RiskMultiplierX": 0.5,
    "RiskMultiplierY": 0.5,
    "_meta": {{
      "per_field": {{
        "RiskTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found in document"}},
        "Criticality": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from context"}}
      }}
    }}
  }}
]

Rules:
- Criticality must be: Low, Medium, High, or Critical
- RiskPriority must be: Low, Medium, High, or Critical
- RiskLikelihood and RiskImpact must be integers 1-10
- RiskExposureRating must be float 0-100
- Dates must be YYYY-MM-DD format
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON array, nothing else
"""

# =========================
# UTILITIES / VALIDATORS
# =========================
def _json_from_llm_text(text: str) -> Any:
    """Extract the first valid JSON array/object from the LLM response."""
    # Remove markdown code blocks if present
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    
    # Try to find JSON array or object
    m = re.search(r"(\[.*\]|\{.*\})", text, flags=re.S)
    block = m.group(1) if m else text.strip()
    
    # Clean up common JSON issues
    # Remove trailing commas before closing braces/brackets
    block = re.sub(r',(\s*[}\]])', r'\1', block)
    
    return json.loads(block)

def _calculate_optimal_context_size(text_length: int, task_complexity: str = "medium") -> int:
    """
    Calculate optimal context size based on text length and task complexity.
    Optimized for Ollama models.
    """
    base_sizes = {
        "simple": 1000,
        "medium": 2000,
        "complex": 4000
    }
    
    base = base_sizes.get(task_complexity, 2000)
    
    if text_length < 1000:
        return min(1000, base)
    elif text_length < 5000:
        return min(2000, base)
    elif text_length < 10000:
        return min(3000, base)
    else:
        return min(4000, base)

def _select_ollama_model_by_complexity(text_length: int, num_risks: int = 1) -> str:
    """
    Select the best Ollama model based on task complexity.
    """
    if text_length < 2000 and num_risks == 1:
        return OLLAMA_MODEL_FAST
    elif text_length > 10000 or num_risks > 5:
        return OLLAMA_MODEL_COMPLEX
    else:
        return OLLAMA_MODEL_DEFAULT

def _call_ollama_json_internal(prompt: str, model: str = None, retries: int = 3, timeout: int = None) -> Any:
    """
    Internal Ollama API call (without caching) - used by cached wrapper.
    
    Args:
        prompt: The prompt to send
        model: Model name (auto-selected if None)
        retries: Number of retry attempts
        timeout: Request timeout (uses default if None)
    """
    if timeout is None:
        timeout = OLLAMA_TIMEOUT
    
    if model is None:
        model = _select_ollama_model_by_complexity(len(prompt))
    
    # Optimize context size
    optimal_context = _calculate_optimal_context_size(len(prompt))
    if len(prompt) > optimal_context:
        mid = optimal_context // 2
        prompt = prompt[:mid] + "\n\n[... content truncated for performance ...]\n\n" + prompt[-mid:]
        print(f"📏 Optimized context: {len(prompt)} chars (reduced from larger size)")
    
    url = f"{OLLAMA_BASE_URL}/api/generate"
    
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": OLLAMA_TEMPERATURE,
            "top_p": 0.9,
            "top_k": 40,
            "num_predict": 2000,
            "seed": OLLAMA_SEED,
            "repeat_penalty": 1.1,
        },
        "format": "json"
    }
    
    print(f"🤖 Calling Ollama API (OPTIMIZED)")
    print(f"   Model: {model}")
    print(f"   Prompt length: {len(prompt)} chars")
    print(f"   URL: {url}")
    
    for attempt in range(retries):
        print(f"🤖 Attempt {attempt + 1}/{retries}...")
        resp = None
        try:
            start_time = time.time()
            resp = requests.post(url, json=payload, timeout=timeout)
            resp.raise_for_status()
            elapsed = time.time() - start_time
            
            print(f"✅ Ollama API responded with status {resp.status_code} in {elapsed:.2f}s")
            
            response_data = resp.json()
            raw = response_data.get("response", "")
            
            print(f"📝 Response length: {len(raw)} chars")
            print(f"📝 First 200 chars: {raw[:200]}...")
            
            result = _json_from_llm_text(raw)
            print(f"✅ Successfully parsed JSON from Ollama response")
            return result
            
        except json.JSONDecodeError as je:
            print(f"❌ JSON parsing error on attempt {attempt + 1}: {je}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            print(f"❌ All retries exhausted. Raw response: {raw[:500] if 'raw' in locals() else 'N/A'}...")
            raise RuntimeError(f"Failed to parse JSON from Ollama response after {retries} attempts")
            
        except requests.exceptions.HTTPError as he:
            print(f"❌ HTTP error on attempt {attempt + 1}: {he}")
            if resp is None:
                raise RuntimeError(f"Ollama API HTTP error: {he}")
            
            print(f"🔍 Status Code: {resp.status_code}")
            try:
                error_response = resp.json()
                error_message = error_response.get('error', 'Unknown error')
                print(f"🔍 Ollama Error: {error_message}")
            except:
                print(f"🔍 Raw response: {resp.text[:500]}")
            
            if resp.status_code == 404:
                raise RuntimeError(f"Model '{model}' not found on Ollama server. Please check available models.")
            elif resp.status_code >= 500:
                if attempt < retries - 1:
                    time.sleep(2)
                    continue
            
            raise RuntimeError(f"Ollama API HTTP error: {he}")
            
        except requests.exceptions.ConnectionError as ce:
            print(f"❌ Connection error on attempt {attempt + 1}: {ce}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"Failed to connect to Ollama API: {ce}")
            
        except requests.exceptions.Timeout as te:
            print(f"❌ Timeout error on attempt {attempt + 1}: {te}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"Ollama API request timed out: {te}")
            
        except Exception as e:
            print(f"❌ Unexpected error on attempt {attempt + 1}: {type(e).__name__}: {e}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            raise RuntimeError(f"Unexpected error calling Ollama API: {e}")
    
    raise RuntimeError(f"Failed to get response from Ollama API after {retries} attempts")

def call_ollama_json(prompt: str, model: str = None, retries: int = 3, timeout: int = None, 
                     document_hash: str = None, use_cache: bool = True) -> Any:
    """
    Call Ollama API expecting JSON response (OPTIMIZED with Phase 2 caching).
    
    Args:
        prompt: The prompt to send
        model: Model name (auto-selected if None)
        retries: Number of retry attempts
        timeout: Request timeout (uses default if None)
        document_hash: Optional document hash for cache key
        use_cache: Whether to use Redis caching (default: True)
    """
    if model is None:
        model = _select_ollama_model_by_complexity(len(prompt))
    
    # Use cached wrapper if caching enabled
    if use_cache:
        # Determine TTL based on prompt length (documents = 24h, queries = 1h)
        ttl = 86400 if len(prompt) > 2000 else 3600
        return cached_llm_call(
            llm_function=_call_ollama_json_internal,
            model_name=model,
            prompt=prompt,
            document_hash=document_hash,
            ttl=ttl,
            use_cache=use_cache,
            model=model,
            retries=retries,
            timeout=timeout
        )
    else:
        return _call_ollama_json_internal(prompt, model, retries, timeout)

def _call_openai_json_internal(prompt: str, retries: int = 3, timeout: int = 120) -> Any:
    """Internal OpenAI API call (without caching) - used by cached wrapper."""
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY environment variable is not set")
    
    # OpenAI API format
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    }
    
    # Models that support json_object response_format:
    # gpt-4-turbo-preview, gpt-4-0125-preview, gpt-3.5-turbo-1106, gpt-4-turbo, gpt-4o
    # Note: gpt-4o-mini does NOT support json_object format
    # Clean model name - strip quotes and whitespace
    model_clean = str(OPENAI_MODEL).strip().strip('"').strip("'")
    model_lower = model_clean.lower()
    
    print(f"🔍 Model check - Original: '{OPENAI_MODEL}', Cleaned: '{model_clean}', Lower: '{model_lower}'")
    
    # Explicitly exclude gpt-4o-mini first (it contains "gpt-4o" but doesn't support json_object)
    if "gpt-4o-mini" in model_lower:
        supports_json_format = False
        print(f"🔍 Model '{model_clean}' is gpt-4o-mini - NOT adding response_format")
    else:
        # Check if model exactly matches or starts with a supported model
        models_with_json_support = [
            "gpt-4-turbo-preview", "gpt-4-0125-preview", "gpt-3.5-turbo-1106", 
            "gpt-4-turbo", "gpt-4o"
        ]
        supports_json_format = any(
            model_lower == model.lower() or model_lower.startswith(model.lower() + "-")
            for model in models_with_json_support
        )
        print(f"🔍 Model '{model_clean}' supports_json_format check result: {supports_json_format}")
    
    payload = {
        "model": model_clean,  # Use cleaned model name
        "messages": [
            {"role": "system", "content": "You are a GRC (Governance, Risk, and Compliance) analyst expert. Always return valid JSON responses as requested."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.1
    }
    
    # Only add response_format for models that support it
    if supports_json_format:
        payload["response_format"] = {"type": "json_object"}
        print(f"✅ Adding response_format: json_object (model '{model_clean}' supports it)")
    else:
        print(f"⚠️  NOT adding response_format (model '{model_clean}' does not support json_object)")
    
    print(f"🤖 Calling OpenAI API at {OPENAI_API_URL}")
    print(f"🤖 Model: {model_clean}")
    print(f"🤖 Prompt length: {len(prompt)} chars")
    print(f"🔍 Payload keys: {list(payload.keys())}")
    if 'response_format' in payload:
        print(f"🔍 response_format value: {payload['response_format']}")
    else:
        print(f"🔍 response_format: NOT in payload")
    
    for attempt in range(retries):
        print(f"🤖 Attempt {attempt + 1}/{retries}...")
        resp = None
        try:
            resp = requests.post(OPENAI_API_URL, json=payload, headers=headers, timeout=timeout)
            resp.raise_for_status()
            print(f"✅ OpenAI API responded with status {resp.status_code}")
            
            # OpenAI response format
            response_data = resp.json()
            raw = response_data.get("choices", [{}])[0].get("message", {}).get("content", "")
            print(f"📝 Response length: {len(raw)} chars")
            print(f"📝 First 200 chars: {raw[:200]}...")
            
            result = _json_from_llm_text(raw)
            print(f"✅ Successfully parsed JSON from OpenAI response")
            return result
            
        except json.JSONDecodeError as je:
            print(f"❌ JSON parsing error on attempt {attempt + 1}: {je}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            print(f"❌ All retries exhausted. Raw response: {raw[:500] if 'raw' in locals() else 'N/A'}...")
            raise RuntimeError(f"Failed to parse JSON from OpenAI response after {retries} attempts")
            
        except requests.exceptions.HTTPError as he:
            print(f"❌ HTTP error on attempt {attempt + 1}: {he}")
            
            # Ensure resp is available
            if resp is None:
                print(f"⚠️  Response object is None - error occurred before response was received")
                raise RuntimeError(f"OpenAI API HTTP error: {he}")
            
            print(f"🔍 Status Code: {resp.status_code}")
            
            # Log the actual error response from OpenAI - be more robust
            try:
                # Try to get response text first
                response_text = resp.text if hasattr(resp, 'text') else 'N/A'
                print(f"🔍 Raw response text (first 1000 chars): {response_text[:1000]}")
                
                # Try to parse as JSON
                try:
                    error_response = resp.json()
                    error_message = error_response.get('error', {})
                    if isinstance(error_message, dict):
                        error_detail = error_message.get('message', 'Unknown error')
                        error_type = error_message.get('type', 'Unknown type')
                        error_code = error_message.get('code', 'Unknown code')
                        print(f"🔍 OpenAI Error Details:")
                        print(f"   Type: {error_type}")
                        print(f"   Code: {error_code}")
                        print(f"   Message: {error_detail}")
                        print(f"   Full error object: {error_message}")
                    else:
                        print(f"🔍 OpenAI Error Response (non-dict): {error_response}")
                except (ValueError, AttributeError, json.JSONDecodeError) as json_err:
                    print(f"⚠️  Response is not valid JSON: {json_err}")
                    print(f"   Response text: {response_text[:500]}")
            except Exception as e:
                print(f"⚠️  Could not parse error response: {e}")
                print(f"   Exception type: {type(e).__name__}")
                import traceback
                print(f"   Traceback: {traceback.format_exc()}")
            
            if resp.status_code == 401:
                raise RuntimeError("OpenAI API authentication failed. Please check your OPENAI_API_KEY.")
            elif resp.status_code == 429:
                print(f"⚠️  Rate limit exceeded. Waiting 5 seconds...")
                if attempt < retries - 1:
                    time.sleep(5)
                    continue
            elif resp.status_code >= 500:
                print(f"⚠️  OpenAI server error. Retrying...")
                if attempt < retries - 1:
                    time.sleep(2)
                    continue
            elif resp.status_code == 400:
                # For 400 errors, log the payload to help debug
                print(f"🔍 Debugging 400 error - Payload being sent:")
                print(f"   Model: {payload.get('model')}")
                print(f"   Has response_format: {'response_format' in payload}")
                if 'response_format' in payload:
                    print(f"   response_format value: {payload.get('response_format')}")
                print(f"   Messages count: {len(payload.get('messages', []))}")
                print(f"   Temperature: {payload.get('temperature')}")
            
            raise RuntimeError(f"OpenAI API HTTP error: {he}")
            
        except requests.exceptions.ConnectionError as ce:
            print(f"❌ Connection error on attempt {attempt + 1}: {ce}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"Failed to connect to OpenAI API: {ce}")
            
        except requests.exceptions.Timeout as te:
            print(f"❌ Timeout error on attempt {attempt + 1}: {te}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"OpenAI API request timed out: {te}")
            
        except Exception as e:
            print(f"❌ Unexpected error on attempt {attempt + 1}: {type(e).__name__}: {e}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            raise RuntimeError(f"Unexpected error calling OpenAI API: {e}")
    
    raise RuntimeError(f"Failed to get response from OpenAI API after {retries} attempts")

def call_openai_json(prompt: str, retries: int = 3, timeout: int = 120, 
                     document_hash: str = None, use_cache: bool = True) -> Any:
    """
    Call OpenAI API expecting JSON response (with Phase 2 caching).
    
    Args:
        prompt: The prompt to send
        retries: Number of retry attempts
        timeout: Request timeout
        document_hash: Optional document hash for cache key
        use_cache: Whether to use Redis caching (default: True)
    """
    # Use cached wrapper if caching enabled
    if use_cache:
        # Determine TTL based on prompt length (documents = 24h, queries = 1h)
        ttl = 86400 if len(prompt) > 2000 else 3600
        return cached_llm_call(
            llm_function=_call_openai_json_internal,
            model_name=OPENAI_MODEL,
            prompt=prompt,
            document_hash=document_hash,
            ttl=ttl,
            use_cache=use_cache,
            retries=retries,
            timeout=timeout
        )
    else:
        return _call_openai_json_internal(prompt, retries, timeout)

def clamp_int(v, lo, hi) -> Optional[int]:
    if v is None: return None
    try:
        return max(lo, min(hi, int(v)))
    except Exception:
        # try to find a number in string
        m = re.search(r"\b(\d{1,2})\b", str(v))
        if m:
            return max(lo, min(hi, int(m.group(1))))
        return None

def as_float_or_none(v) -> Optional[float]:
    if v is None: return None
    try:
        return float(v)
    except Exception:
        try:
            # extract first float from string
            m = re.search(r"[-+]?\d*\.?\d+", str(v))
            return float(m.group(0)) if m else None
        except Exception:
            return None

def as_date_or_none(s: str) -> Optional[str]:
    if not s: return None
    s = str(s).strip()
    return s if re.match(r"^\d{4}-\d{2}-\d{2}$", s) else None

def compute_exposure(lk, im) -> Optional[float]:
    if lk is None or im is None: return None
    try:
        return round(float(lk) * float(im), 2)
    except Exception:
        return None

def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None


# =========================
# FILE EXTRACTORS
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF (pdfplumber preferred; PyPDF2 fallback)."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        print(f"[PDF] Extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX] Extraction error: {e}")
        return ""

def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        print(f"[XLSX] Extraction error: {e}")
        return ""

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""


# =========================
# DOCUMENT PARSING & DETECTION
# =========================
def detect_and_parse_risk_blocks(text: str) -> list[dict]:
    """
    Step 1: Detect how many risks are in the document by finding 'Risk X:' patterns.
    Step 2: For each risk, extract whatever fields are explicitly present.
    Returns a list of partially-filled risk dictionaries with only extracted fields.
    """
    print(f"🔍 Detecting risk blocks in document...")
    
    # Find all risk blocks using pattern "Risk X: Title"
    risk_pattern = r'Risk\s+(\d+):\s*([^\n]+)'
    risk_matches = list(re.finditer(risk_pattern, text, re.IGNORECASE))
    
    if not risk_matches:
        print(f"⚠️  No risks found with 'Risk X:' pattern")
        return []
    
    print(f"✅ Found {len(risk_matches)} risk(s) in document")
    
    risks = []
    for i, match in enumerate(risk_matches):
        risk_num = match.group(1)
        risk_title = match.group(2).strip()
        
        # Extract the text block for this risk (from this match to next risk or end of document)
        start_pos = match.end()
        if i + 1 < len(risk_matches):
            end_pos = risk_matches[i + 1].start()
        else:
            end_pos = len(text)
        
        risk_block = text[start_pos:end_pos]
        
        print(f"  📋 Risk {risk_num}: {risk_title[:60]}...")
        
        # Extract fields that are explicitly present in the block
        risk_data = {
            "RiskTitle": risk_title,  # Always extracted from document
            "_extracted_fields": ["RiskTitle"],  # Track what was extracted
            "_risk_block": risk_block  # Keep for AI context
        }
        
        # Field mapping: document field name -> DB field name
        field_mappings = {
            r'Description:\s*([^\n]+(?:\n(?!(?:Risk|Possible Damage|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'RiskDescription',
            r'Possible Damage:\s*([^\n]+(?:\n(?!(?:Risk|Description|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'PossibleDamage',
            r'Risk Priority:\s*([^\n]+)': 'RiskPriority',
            r'Status:\s*([^\n]+)': 'RiskExposureRating',
            r'Mitigation:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Status|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'RiskMitigation',
            r'Created At:\s*([^\n]+)': 'CreatedAt',
            r'Risk Type:\s*([^\n]+)': 'RiskType',
            r'Risk Likelihood:\s*([^\n]+)': 'RiskLikelihood',
            r'Risk Impact:\s*([^\n]+)': 'RiskImpact',
            r'Category:\s*([^\n]+)': 'Category',
            r'Criticality:\s*([^\n]+)': 'Criticality',
            r'Business Impact:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'BusinessImpact',
        }
        
        # Extract each field if present
        for pattern, field_name in field_mappings.items():
            match = re.search(pattern, risk_block, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if value:  # Only add if not empty
                    risk_data[field_name] = value
                    risk_data["_extracted_fields"].append(field_name)
                    print(f"    ✓ {field_name}: {value[:50]}...")
        
        risks.append(risk_data)
    
    return risks


# =========================
# AI EXTRACTION CORE
# =========================
def infer_single_field(field_name: str, current_record: dict, document_context: str, 
                       document_hash: str = None) -> tuple[Any, dict]:
    """
    Focused prompt for ONE field using AI (supports both OpenAI and Ollama).
    Uses Phase 2 few-shot prompts and caching.
    Returns: (value, metadata_dict)
    """
    provider_name = AI_PROVIDER.upper()
    print(f"🤖 AI PREDICTING FIELD: {field_name} (using {provider_name} with Phase 2+3 optimizations)")
    
    # Optimize context for Ollama
    if AI_PROVIDER == 'ollama':
        context_size = _calculate_optimal_context_size(len(document_context), "simple")
        optimized_context = document_context[:context_size] if len(document_context) > context_size else document_context
    else:
        optimized_context = document_context[:3000]  # OpenAI default
    
    # Phase 3: Try to retrieve relevant context from RAG
    rag_context = None
    if is_rag_available():
        try:
            # Search for relevant context about this field
            query = f"What is the {field_name} for this risk?"
            retrieved = retrieve_relevant_context(query, n_results=3)
            if retrieved:
                rag_context = retrieved
                print(f"   📚 Phase 3 RAG: Retrieved {len(retrieved)} relevant document chunks")
        except Exception as e:
            print(f"   ⚠️  RAG retrieval failed: {e}")
    
    # Use few-shot prompt template (Phase 2 optimization)
    try:
        mini = get_field_extraction_prompt(
            field_name=field_name,
            document_text=optimized_context,
            field_prompts=FIELD_PROMPTS
        )
        # Add current record context
        mini += f"\n\nCurrent risk (partial):\n{json.dumps({k: current_record.get(k) for k in RISK_DB_FIELDS if current_record.get(k)}, indent=2)}"
        print(f"   📚 Using few-shot prompt template for {field_name}")
    except Exception as e:
        print(f"   ⚠️  Few-shot prompt failed, using basic prompt: {e}")
        # Fallback to basic prompt
        guidance = FIELD_PROMPTS.get(field_name, "Return a concise, professional value.")
        mini = f"""
You are a GRC analyst. Infer ONLY the field "{field_name}" for this risk.
Return JSON: {{"value": <scalar or string>, "confidence": 0.0-1.0, "rationale": "brief explanation"}}

Context (document):
\"\"\"{optimized_context}\"\"\"

Current risk (partial):
{json.dumps({k: current_record.get(k) for k in RISK_DB_FIELDS if current_record.get(k)}, indent=2)}

Rules:
- {guidance}
- If you cannot infer, return {{"value": null, "confidence": 0.0, "rationale": "Not enough information"}}.
- Always include a brief rationale explaining your decision.
- Return ONLY valid JSON, no markdown, no code blocks.
"""
    
    # Phase 3: Enhance prompt with RAG context if available
    if rag_context:
        mini = build_rag_prompt(
            user_query=mini,
            retrieved_context=rag_context,
            base_prompt=None
        )
    
    try:
        if AI_PROVIDER == 'ollama':
            print(f"   📤 Sending prompt to Ollama for {field_name}...")
            # Select appropriate model for this field
            model = _select_ollama_model_by_complexity(len(optimized_context), 1)
            out = call_ollama_json(mini, model=model, document_hash=document_hash)
            model_used = model
        else:
            print(f"   📤 Sending prompt to OpenAI for {field_name}...")
            out = call_openai_json(mini, document_hash=document_hash)
            model_used = OPENAI_MODEL
        
        v = out.get("value") if isinstance(out, dict) else None
        confidence = out.get("confidence", 0.7) if isinstance(out, dict) else 0.7
        rationale = out.get("rationale", "AI predicted based on document context") if isinstance(out, dict) else "AI predicted based on document context"
        print(f"   ✅ AI PREDICTED {field_name}: '{v}' (confidence: {confidence:.2f})")
    except Exception as e:
        print(f"   ❌ AI FAILED to predict {field_name}: {str(e)}")
        v = None
        confidence = 0.0
        rationale = f"AI prediction failed: {str(e)}"
        model_used = None

    # Create metadata for this field
    metadata = {
        "source": "AI_GENERATED",
        "confidence": confidence,
        "rationale": rationale,
        "provider": AI_PROVIDER,
        "model_used": model_used
    }

    # normalize after inference
    if field_name in ("RiskLikelihood", "RiskImpact"):
        v = clamp_int(v, 1, 10) or 5
    elif field_name == "RiskExposureRating":
        lk = clamp_int(current_record.get("RiskLikelihood"), 1, 10)
        im = clamp_int(current_record.get("RiskImpact"), 1, 10)
        computed = compute_exposure(lk, im)
        v = computed if computed is not None else 0.0
    elif field_name in ("RiskMultiplierX", "RiskMultiplierY"):
        vf = as_float_or_none(v)
        v = vf if vf is not None else 0.5
    elif field_name == "CreatedAt":
        v = as_date_or_none(v) or date.today().isoformat()
    elif field_name == "Criticality":
        v = normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
    elif field_name == "RiskPriority":
        v = normalize_choice(v, PRIORITY_CHOICES) or "Medium"
    elif field_name == "RiskType":
        v = normalize_choice(v, RISKTYPE_HINTS) or "Current"
    elif isinstance(v, str):
        v = v.strip() or None
    
    return v, metadata


def fallback_risk_extraction(text: str) -> list[dict]:
    """
    Minimal pattern-based fallback when AI fails completely.
    Produces at least one record using generic defaults.
    """
    risks = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    risk_keywords = ["risk", "threat", "vulnerability", "hazard", "danger", "exposure"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in risk_keywords):
            if current:
                risks.append(current)
            count += 1
            current = {
                "RiskTitle": f"Risk {count}: {ln[:100]}",
                "RiskDescription": ln,
                "Category": "Operational",
                "Criticality": "Medium",
                "RiskPriority": "Medium",
                "RiskType": "Current",
                "BusinessImpact": "Potential business impact",
                "PossibleDamage": "Potential damages vary (reputation, compliance, downtime).",
                "RiskMitigation": "Mitigation steps to be defined after assessment.",
                "RiskLikelihood": 5,
                "RiskImpact": 5,
                "RiskExposureRating": 25.0,
                "CreatedAt": date.today().isoformat(),
                "RiskMultiplierX": 0.5,
                "RiskMultiplierY": 0.5,
            }

    if current:
        risks.append(current)

    if not risks:
        risks.append({
            "RiskTitle": "Document Risk Analysis",
            "RiskDescription": f"Automated risk derived from document (length={len(text)} chars).",
            "Category": "Operational",
            "Criticality": "Medium",
            "RiskPriority": "Medium",
            "RiskType": "Current",
            "BusinessImpact": "Potential business impact in operations/compliance.",
            "PossibleDamage": "Data loss, downtime, penalties, reputation.",
            "RiskMitigation": "Review & implement standard mitigations for identified weaknesses.",
            "RiskLikelihood": 5,
            "RiskImpact": 5,
            "RiskExposureRating": 25.0,
            "CreatedAt": date.today().isoformat(),
            "RiskMultiplierX": 0.5,
            "RiskMultiplierY": 0.5,
        })

    return risks


def parse_risks_from_text(text: str, document_hash: str = None) -> list[dict]:
    """
    NEW APPROACH:
    1. First detect how many risks are in the document by finding 'Risk X:' patterns
    2. Extract whatever fields are explicitly present in each risk block
    3. Use AI to fill ONLY the missing fields (not to create risk titles)
    """
    print(f"📊 parse_risks_from_text() called with {len(text)} chars of text")
    
    # Step 1: Detect and parse risk blocks from document
    detected_risks = detect_and_parse_risk_blocks(text)
    
    if not detected_risks:
        print(f"⚠️  No structured risks found. Falling back to old AI extraction...")
        # Fallback to old behavior if no "Risk X:" pattern found
        return fallback_risk_extraction(text)
    
    print(f"✅ Detected {len(detected_risks)} risk(s), now processing each...")
    
    # Step 2: For each detected risk, normalize extracted fields and fill missing ones
    completed_risks = []
    for idx, risk_data in enumerate(detected_risks, 1):
        print(f"\n🔧 Processing Risk {idx}: {risk_data.get('RiskTitle', 'Unknown')[:50]}...")
        
        # Start with empty item for all DB fields
        item = {k: None for k in RISK_DB_FIELDS}
        
        # Copy extracted fields
        for field in RISK_DB_FIELDS:
            if field in risk_data and risk_data[field]:
                item[field] = risk_data[field]
        
        # Get the risk block for AI context
        risk_block = risk_data.get("_risk_block", "")
        extracted_fields = risk_data.get("_extracted_fields", [])
        
        print(f"  📝 Extracted fields: {', '.join(extracted_fields)}")
        
        # Initialize metadata structure
        if "_meta" not in item:
            item["_meta"] = {}
        if "per_field" not in item["_meta"]:
            item["_meta"]["per_field"] = {}
        
        # Mark extracted fields in metadata
        for field in extracted_fields:
            if field in item and item[field]:
                item["_meta"]["per_field"][field] = {
                    "source": "EXTRACTED",
                    "confidence": 0.95,
                    "rationale": "Found explicitly in document"
                }
        
        # Normalize extracted fields
        if item.get("RiskLikelihood"):
            item["RiskLikelihood"] = clamp_int(item["RiskLikelihood"], 1, 10)
        if item.get("RiskImpact"):
            item["RiskImpact"] = clamp_int(item["RiskImpact"], 1, 10)
        if item.get("RiskExposureRating"):
            item["RiskExposureRating"] = as_float_or_none(item["RiskExposureRating"])
        if item.get("Criticality"):
            item["Criticality"] = normalize_choice(item["Criticality"], CRITICALITY_CHOICES)
        if item.get("RiskPriority"):
            item["RiskPriority"] = normalize_choice(item["RiskPriority"], PRIORITY_CHOICES)
        if item.get("RiskType"):
            item["RiskType"] = normalize_choice(item["RiskType"], RISKTYPE_HINTS)
        if item.get("CreatedAt"):
            item["CreatedAt"] = as_date_or_none(item["CreatedAt"])
        if item.get("RiskMultiplierX"):
            item["RiskMultiplierX"] = as_float_or_none(item["RiskMultiplierX"])
        if item.get("RiskMultiplierY"):
            item["RiskMultiplierY"] = as_float_or_none(item["RiskMultiplierY"])
        
        # Step 3: Use AI to fill ONLY missing fields (NEVER RiskTitle - must be in document)
        missing_fields = [f for f in RISK_DB_FIELDS if item.get(f) in (None, "", []) and f != "RiskTitle"]
        if missing_fields:
            print(f"  🤖 Missing fields: {', '.join(missing_fields)}")
            print(f"  🤖 Using AI to infer missing fields...")
            
            for field in missing_fields:
                print(f"    🔍 Inferring {field}...")
                value, metadata = infer_single_field(field, item, risk_block or text[:3000], document_hash=document_hash)
                item[field] = value
                # Store AI generation metadata
                item["_meta"]["per_field"][field] = metadata
                print(f"    🏷️  Marked {field} as AI_GENERATED in metadata")
        else:
            print(f"  ✅ All fields extracted from document!")
        
        # Final normalization and defaults
        item["RiskLikelihood"] = item["RiskLikelihood"] or 5
        item["RiskImpact"] = item["RiskImpact"] or 5
        
        # Compute exposure if not present
        if not item.get("RiskExposureRating"):
            item["RiskExposureRating"] = compute_exposure(item["RiskLikelihood"], item["RiskImpact"]) or 25.0
        
        item["RiskExposureRating"] = float(max(0.0, min(100.0, item["RiskExposureRating"])))
        item["RiskMultiplierX"] = item["RiskMultiplierX"] or 0.5
        item["RiskMultiplierY"] = item["RiskMultiplierY"] or 0.5
        item["CreatedAt"] = item["CreatedAt"] or date.today().isoformat()
        item["Criticality"] = item["Criticality"] or "Medium"
        item["RiskPriority"] = item["RiskPriority"] or "Medium"
        item["RiskType"] = item["RiskType"] or "Current"
        
        # RiskTitle must ALWAYS come from document - never generate it
        if not item.get("RiskTitle"):
            raise ValueError(f"RiskTitle is missing for risk {idx}. All risk titles must be present in the document as 'Risk X: Title'.")
        
        # Debug: Print metadata summary
        ai_fields = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "AI_GENERATED"]
        extracted = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "EXTRACTED"]
        
        print(f"  📊 Metadata Summary:")
        print(f"     🤖 AI Generated: {len(ai_fields)} fields - {ai_fields}")
        print(f"     📄 Extracted: {len(extracted)} fields - {extracted}")
        
        completed_risks.append(item)
        print(f"  ✅ Risk {idx} completed!")
    
    return completed_risks


# =========================
# DJANGO API ENDPOINTS
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@rate_limit_decorator(requests_per_minute=10, requests_per_hour=100)  # Phase 3: Rate limiting
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def upload_and_process_risk_document(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    print(f"📤 Upload request for risk document")
    print(f"📤 Request data: {request.POST}")
    print(f"📤 Request files: {request.FILES}")
    print(f"📤 User ID: {request.POST.get('user_id', 'unknown')}")

    # CORS preflight support
    if request.method == 'OPTIONS':
        response = HttpResponse()
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        response['Access-Control-Max-Age'] = '86400'
        return response

    try:
        if 'file' not in request.FILES:
            resp = JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        
        # Handle compressed files (.gz extension)
        # Remove .gz extension to get original file extension
        original_filename = file_name
        if file_name.endswith('.gz'):
            original_filename = file_name[:-3]  # Remove .gz
            print(f"📦 Detected compressed file: {file_name} → {original_filename}")
        
        ext = os.path.splitext(original_filename)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            resp = JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        # Create the ai_uploads/risk directory if it doesn't exist
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'risk')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save the file with timestamp to avoid conflicts
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        # Write the uploaded file to disk
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        print(f"✅ File saved to: {file_path}")

        # Decompress if needed (client-side compression)
        compression_metadata = None
        file_path, was_compressed, compression_stats = decompress_if_needed(file_path)
        if was_compressed:
            compression_metadata = compression_stats
            # Update extension after decompression (remove .gz)
            ext = os.path.splitext(file_path)[1].lower()
            print(f"📦 Decompressed file: {compression_stats['ratio']}% reduction, saved {compression_stats['bandwidth_saved_kb']} KB")

        # Upload to S3 for backup and cloud storage
        s3_url = None
        s3_key = None
        user_id = request.POST.get('user_id', '1')
        try:
            print(f"☁️ Uploading file to S3...")
            s3_client = create_direct_mysql_client()
            connection_test = s3_client.test_connection()
            if connection_test.get('overall_success', False):
                # Generate unique filename for S3
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                s3_filename = f"risk_ai_{timestamp}_{os.path.basename(file_path)}"
                upload_result = s3_client.upload(
                    file_path=file_path,
                    user_id=user_id,
                    custom_file_name=s3_filename,
                    module='Risk'
                )
                if upload_result.get('success'):
                    s3_url = upload_result['file_info']['url']
                    s3_key = upload_result['file_info'].get('s3Key', '')
                    print(f"✅ File uploaded to S3: {s3_url}")
                else:
                    print(f"⚠️ S3 upload failed: {upload_result.get('error', 'Unknown error')}")
            else:
                print(f"⚠️ S3 service unavailable, continuing with local file")
        except Exception as s3_error:
            print(f"⚠️ S3 upload error (continuing with local file): {str(s3_error)}")

        try:
            # Step 1: Extract text from the saved file
            print(f"🔍 STEP 1: Starting text extraction from {ext} file...")
            raw_text = extract_text_from_file(file_path, ext)
            
            if not raw_text or len(raw_text.strip()) < 50:
                print(f"❌ ERROR: Could not extract meaningful text. Length: {len(raw_text) if raw_text else 0}")
                resp = JsonResponse({'status': 'error', 'message': 'Could not extract meaningful text from document'}, status=400)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp

            print(f"✅ STEP 1A COMPLETE: Extracted {len(raw_text)} characters from document")
            
            # Step 1B: Preprocess document (Phase 2 optimization)
            print(f"🔍 STEP 1B: Preprocessing document (Phase 2 optimization)...")
            text, preprocess_metadata = preprocess_document(raw_text, max_length=8000)
            print(f"✅ STEP 1B COMPLETE: Preprocessed document")
            print(f"   Original length: {preprocess_metadata['original_length']} chars")
            print(f"   Processed length: {preprocess_metadata['processed_length']} chars")
            if preprocess_metadata['was_truncated']:
                print(f"   ⚠️  Document was truncated ({preprocess_metadata['reduction_percent']:.1f}% reduction)")
            
            # Calculate document hash for caching (Phase 2)
            document_hash = calculate_document_hash(text)
            print(f"📝 Document hash: {document_hash[:16]}... (for caching)")
            
            print(f"📄 First 200 chars: {text[:200]}...")
            
            # Step 2: Check AI provider configuration
            print(f"🔍 STEP 2: Checking AI provider configuration...")
            if AI_PROVIDER == 'openai' and not OPENAI_API_KEY:
                print(f"❌ ERROR: OPENAI_API_KEY is not set")
                resp = JsonResponse({
                    'status': 'error', 
                    'message': 'OPENAI_API_KEY environment variable is not set. Please configure your OpenAI API key or switch to Ollama.'
                }, status=503)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp
            elif AI_PROVIDER == 'ollama' and not OLLAMA_BASE_URL:
                print(f"❌ ERROR: OLLAMA_BASE_URL is not set")
                resp = JsonResponse({
                    'status': 'error', 
                    'message': 'OLLAMA_BASE_URL environment variable is not set. Please configure your Ollama server URL.'
                }, status=503)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp
            
            print(f"✅ STEP 2 COMPLETE: {AI_PROVIDER.upper()} provider is configured")
            
            # Step 3: Process with AI (Phase 2+3 optimizations)
            # Phase 3: Use intelligent model routing
            start_time = time.time()
            
            provider_info = f"{AI_PROVIDER.upper()} ({OPENAI_MODEL if AI_PROVIDER == 'openai' else OLLAMA_MODEL_DEFAULT})"
            print(f"🤖 STEP 3: Calling {provider_info} to extract risks (Phase 2+3: cached + few-shot + RAG + routing)...")
            
            # Phase 3: Process with queuing (if needed)
            request_id = f"risk_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(file_name)}"
            
            def process_document():
                # MULTI-TENANCY: Extract tenant_id from request
                tenant_id = get_tenant_id_from_request(request)
                
                return parse_risks_from_text(text, document_hash=document_hash)
            
            # Use queuing for heavy processing
            if len(text) > 10000:  # Large documents use queue
                print(f"📋 Large document detected, using Phase 3 queuing...")
                risks = process_with_queue(request_id, process_document)
            else:
                risks = process_document()
            
            # Track processing time for system load monitoring (Phase 3)
            processing_time = time.time() - start_time
            track_system_load(processing_time, len(text))
            
            # Phase 3: Add document to RAG for future context retrieval
            if is_rag_available():
                try:
                    add_document_to_rag(
                        document_text=text,
                        document_id=f"risk_doc_{document_hash[:16]}",
                        metadata={
                            "type": "risk_assessment",
                            "filename": file_name,
                            "uploaded_at": datetime.now().isoformat(),
                            "num_risks": len(risks) if 'risks' in locals() else 0
                        }
                    )
                    print(f"✅ Phase 3 RAG: Document added to knowledge base")
                except Exception as e:
                    print(f"⚠️  Phase 3 RAG: Failed to add document: {e}")
            
            print(f"✅ STEP 3 COMPLETE: AI extracted {len(risks)} risk(s) from document")
            for idx, risk in enumerate(risks, 1):
                print(f"  Risk {idx}: {risk.get('RiskTitle', 'Untitled')[:50]}...")

            # Phase 3: Include RAG and routing stats in response
            phase3_metadata = {
                "rag_available": is_rag_available(),
                "rag_stats": get_rag_stats() if is_rag_available() else None,
                "system_load": get_current_system_load(),
                "processing_time": processing_time,
                "model_routing": "enabled"
            }
            
            response_data = {
                'status': 'success',
                'message': f'Successfully extracted {len(risks)} risk(s)',
                'document_name': file_name,
                'saved_path': safe_filename,
                'extracted_text_length': len(text),
                'preprocessing_metadata': preprocess_metadata,
                'phase3_metadata': phase3_metadata,  # Phase 3 stats
                'risks': risks
            }
            
            # Include compression metadata if file was compressed
            if compression_metadata:
                response_data['compression_metadata'] = compression_metadata
            
            # Include S3 info if uploaded successfully
            if s3_url:
                response_data['s3_url'] = s3_url
                response_data['s3_key'] = s3_key
            
            resp = JsonResponse(response_data)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp
        except Exception as process_error:
            # Clean up the file if processing fails
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        resp = JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)
        resp['Access-Control-Allow-Origin'] = '*'
        return resp


@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def save_extracted_risks(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        data = json.loads(request.body or "{}")
        risks_data = data.get('risks', [])
        if not risks_data:
            return JsonResponse({'status': 'error', 'message': 'No risks provided'}, status=400)

        saved = []
        errors = []

        for idx, r in enumerate(risks_data):
            try:
                kwargs = {
                    # not setting: RiskId, ComplianceId, FrameworkId
                    'RiskTitle': r.get('RiskTitle', f'Risk {idx+1}'),
                    'Criticality': r.get('Criticality', 'Medium'),
                    'PossibleDamage': r.get('PossibleDamage', ''),
                    'Category': r.get('Category', ''),
                    'RiskType': r.get('RiskType', 'Current'),
                    'BusinessImpact': r.get('BusinessImpact', ''),
                    'RiskDescription': r.get('RiskDescription', ''),
                    'RiskLikelihood': int(r.get('RiskLikelihood') or 5),
                    'RiskImpact': int(r.get('RiskImpact') or 5),
                    'RiskExposureRating': float(r.get('RiskExposureRating') or 25.0),
                    'RiskPriority': r.get('RiskPriority', 'Medium'),
                    'RiskMitigation': r.get('RiskMitigation', ''),
                    'CreatedAt': r.get('CreatedAt', date.today().isoformat()),
                    'RiskMultiplierX': float(r.get('RiskMultiplierX') or 0.5),
                    'RiskMultiplierY': float(r.get('RiskMultiplierY') or 0.5),
                }
                risk = Risk.objects.create(**kwargs)
                saved.append({'risk_id': getattr(risk, "RiskId", None), 'risk_title': risk.RiskTitle})
            except Exception as ex:
                errors.append({'risk_index': idx, 'title': r.get('RiskTitle'), 'error': str(ex)})

        resp = {
            'status': 'success',
            'message': f'Saved {len(saved)} risk(s)' + (f' with {len(errors)} error(s)' if errors else ''),
            'saved': saved
        }
        if errors:
            resp['errors'] = errors
        return JsonResponse(resp)

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON payload'}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error saving risks: {str(e)}'}, status=500)


@api_view(['GET'])
@permission_classes([AllowAny])
@rbac_required(required_permission='view_all_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def test_openai_connection(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        if not OPENAI_API_KEY:
            return JsonResponse({
                'status': 'error',
                'message': 'OPENAI_API_KEY is not set',
                'model': OPENAI_MODEL,
                'api_url': OPENAI_API_URL
            }, status=500)
        
        out = call_openai_json('Return JSON: {"ok": true}')
        return JsonResponse({
            'status': 'success', 
            'openai_reply': out, 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error', 
            'message': f'OpenAI error: {e}', 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        }, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def test_file_upload(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        if 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file found in request'}, status=400)
        f = request.FILES['file']
        return JsonResponse({
            'status': 'success',
            'message': 'File upload test successful',
            'file_name': f.name,
            'file_size': f.size,
            'content_type': f.content_type
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Test upload error: {str(e)}'}, status=500)
