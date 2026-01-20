"""
AI-Powered Risk Document Ingestion (OPTIMIZED VERSION)
- Uses Ollama with quantized models
- Dynamic context sizing
- Optimized prompts
- Better performance
"""

import os
import re
import json
import math
import time
import tempfile
from datetime import date, datetime
from typing import Any, Optional

import requests
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser

# RBAC imports
from ...rbac.decorators import rbac_required
# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)


# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import Risk

# =========================
# OLLAMA CONFIG (OPTIMIZED)
# =========================
from django.conf import settings

# Ollama Configuration - Optimized
OLLAMA_BASE_URL = getattr(settings, 'OLLAMA_BASE_URL', 'http://13.205.15.232:11434')
OLLAMA_TIMEOUT = getattr(settings, 'OLLAMA_TIMEOUT', 600)
OLLAMA_TEMPERATURE = getattr(settings, 'OLLAMA_TEMPERATURE', 0.1)
OLLAMA_SEED = getattr(settings, 'OLLAMA_SEED', 42)

# Optimized model selection based on task complexity
# Use quantized models for better performance
OLLAMA_MODEL_DEFAULT = getattr(settings, 'OLLAMA_MODEL', 'llama3.2:3b-instruct-q4_K_M')
OLLAMA_MODEL_FAST = 'llama3.2:1b-instruct-q4_K_M'  # For simple tasks
OLLAMA_MODEL_COMPLEX = 'llama3:8b-instruct-q4_K_M'  # For complex reasoning

# Remove trailing slash from URL
OLLAMA_BASE_URL = OLLAMA_BASE_URL.rstrip('/')

print(f"🌐 Ollama Configuration (OPTIMIZED):")
print(f"   Base URL: {OLLAMA_BASE_URL}")
print(f"   Default Model: {OLLAMA_MODEL_DEFAULT}")
print(f"   Fast Model: {OLLAMA_MODEL_FAST}")
print(f"   Complex Model: {OLLAMA_MODEL_COMPLEX}")
print(f"   Temperature: {OLLAMA_TEMPERATURE}")
print(f"   Timeout: {OLLAMA_TIMEOUT}s")

# Only the columns you want to fill
RISK_DB_FIELDS = [
    "RiskTitle",
    "Criticality",
    "PossibleDamage",
    "Category",
    "RiskType",
    "BusinessImpact",
    "RiskDescription",
    "RiskLikelihood",
    "RiskImpact",
    "RiskExposureRating",
    "RiskPriority",
    "RiskMitigation",
    "CreatedAt",
    "RiskMultiplierX",
    "RiskMultiplierY",
]

# Canonical choices / constraints
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES    = ["Low", "Medium", "High", "Critical"]
CATEGORY_HINTS      = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]
RISKTYPE_HINTS      = ["Current", "Residual", "Inherent", "Emerging", "Accepted"]
DATE_FORMAT_HINT    = "YYYY-MM-DD (ISO)"

# Field-specific micro-prompts (optimized for Ollama)
FIELD_PROMPTS = {
    "Criticality": f"Return one of: {CRITICALITY_CHOICES}.",
    "PossibleDamage": "Describe concrete damages (data loss, downtime, penalties, reputation). Be concise (1–2 sentences).",
    "Category": f"Return one category from this list (best fit): {CATEGORY_HINTS}. If none fits, pick the closest.",
    "RiskType": f"Return one of: {RISKTYPE_HINTS}.",
    "BusinessImpact": "Explain business impact in business terms (SLA breach, revenue, compliance). 1–2 sentences.",
    "RiskDescription": "Write a precise description (1–3 sentences) of how/why the risk arises in this context.",
    "RiskLikelihood": "Return an integer 1–10 (1=rare, 10=almost certain).",
    "RiskImpact": "Return an integer 1–10 (1=negligible, 10=catastrophic).",
    "RiskExposureRating": "Return a float (0–100). If missing, use Likelihood*Impact as proxy.",
    "RiskPriority": f"Return one of: {PRIORITY_CHOICES}. Base it on exposure + criticality.",
    "RiskMitigation": "Return 2–4 actionable mitigation steps as one paragraph or a bullet-style JSON list.",
    "CreatedAt": f"Return a plausible assessment date in {DATE_FORMAT_HINT}; if unknown, use today's date.",
    "RiskMultiplierX": "Return a float in 0.1–1.5 reflecting org weighting factor X (defaults ~0.5 if unknown).",
    "RiskMultiplierY": "Return a float in 0.1–1.5 reflecting org weighting factor Y (defaults ~0.5 if unknown).",
}

# Optimized JSON schema block for Ollama
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON object or array. No markdown, no code blocks, no explanations.
Use proper JSON syntax with double quotes.

Example structure (return object like this):
{{
  "value": <your_value_here>,
  "confidence": 0.8,
  "rationale": "brief explanation"
}}

Or for array of risks:
[
  {{
    "RiskTitle": "Brief risk title here",
    "Criticality": "Medium",
    "PossibleDamage": "Description of damage",
    "Category": "Operational",
    "RiskType": "Current",
    "BusinessImpact": "Business impact description",
    "RiskDescription": "Risk description",
    "RiskLikelihood": 5,
    "RiskImpact": 6,
    "RiskExposureRating": 30.0,
    "RiskPriority": "Medium",
    "RiskMitigation": "Mitigation steps",
    "CreatedAt": "{date.today().isoformat()}",
    "RiskMultiplierX": 0.5,
    "RiskMultiplierY": 0.5
  }}
]

Rules:
- Criticality must be: Low, Medium, High, or Critical
- RiskPriority must be: Low, Medium, High, or Critical
- RiskLikelihood and RiskImpact must be integers 1-10
- RiskExposureRating must be float 0-100
- Dates must be YYYY-MM-DD format
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON, nothing else
"""

# =========================
# UTILITIES / VALIDATORS
# =========================
def _json_from_llm_text(text: str) -> Any:
    """Extract the first valid JSON array/object from the LLM response."""
    # Remove markdown code blocks if present
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    
    # Try to find JSON array or object
    m = re.search(r"(\[.*\]|\{.*\})", text, flags=re.S)
    block = m.group(1) if m else text.strip()
    
    # Clean up common JSON issues
    block = re.sub(r',(\s*[}\]])', r'\1', block)
    
    return json.loads(block)

def _calculate_optimal_context_size(text_length: int, task_complexity: str = "medium") -> int:
    """
    Calculate optimal context size based on text length and task complexity.
    Optimized for Ollama models.
    """
    # Base context sizes for different models
    base_sizes = {
        "simple": 1000,      # Fast model
        "medium": 2000,      # Default model
        "complex": 4000      # Complex model
    }
    
    base = base_sizes.get(task_complexity, 2000)
    
    # Adjust based on text length (don't send more than needed)
    if text_length < 1000:
        return min(1000, base)
    elif text_length < 5000:
        return min(2000, base)
    elif text_length < 10000:
        return min(3000, base)
    else:
        return min(4000, base)  # Cap at 4k for performance

def _select_model_by_complexity(text_length: int, num_risks: int = 1) -> str:
    """
    Select the best model based on task complexity.
    """
    # Simple: short text, single risk
    if text_length < 2000 and num_risks == 1:
        return OLLAMA_MODEL_FAST
    # Complex: long text or multiple risks
    elif text_length > 10000 or num_risks > 5:
        return OLLAMA_MODEL_COMPLEX
    # Default: medium complexity
    else:
        return OLLAMA_MODEL_DEFAULT

def call_ollama_json(prompt: str, model: str = None, retries: int = 3, timeout: int = None, 
                     use_streaming: bool = False) -> Any:
    """
    Call Ollama API expecting JSON response (OPTIMIZED).
    
    Args:
        prompt: The prompt to send
        model: Model name (auto-selected if None)
        retries: Number of retry attempts
        timeout: Request timeout (uses default if None)
        use_streaming: Whether to use streaming (for better UX)
    """
    if timeout is None:
        timeout = OLLAMA_TIMEOUT
    
    if model is None:
        # Auto-select model based on prompt length
        model = _select_model_by_complexity(len(prompt))
    
    # Optimize context size
    optimal_context = _calculate_optimal_context_size(len(prompt))
    if len(prompt) > optimal_context:
        # Truncate prompt intelligently (keep beginning and end)
        mid = optimal_context // 2
        prompt = prompt[:mid] + "\n\n[... content truncated for performance ...]\n\n" + prompt[-mid:]
        print(f"📏 Optimized context: {len(prompt)} chars (reduced from larger size)")
    
    # Ollama API format
    url = f"{OLLAMA_BASE_URL}/api/generate"
    
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": use_streaming,
        "options": {
            "temperature": OLLAMA_TEMPERATURE,
            "top_p": 0.9,
            "top_k": 40,
            "num_predict": 2000,  # Optimized for faster responses
            "seed": OLLAMA_SEED,
            "repeat_penalty": 1.1,
        },
        "format": "json"  # Request JSON format (Ollama supports this)
    }
    
    print(f"🤖 Calling Ollama API (OPTIMIZED)")
    print(f"   Model: {model}")
    print(f"   Prompt length: {len(prompt)} chars")
    print(f"   Streaming: {use_streaming}")
    print(f"   URL: {url}")
    
    for attempt in range(retries):
        print(f"🤖 Attempt {attempt + 1}/{retries}...")
        resp = None
        try:
            start_time = time.time()
            resp = requests.post(url, json=payload, timeout=timeout)
            resp.raise_for_status()
            elapsed = time.time() - start_time
            
            print(f"✅ Ollama API responded with status {resp.status_code} in {elapsed:.2f}s")
            
            # Handle streaming response
            if use_streaming:
                full_response = ""
                for line in resp.iter_lines():
                    if line:
                        try:
                            chunk = json.loads(line)
                            if 'response' in chunk:
                                full_response += chunk['response']
                            if chunk.get('done', False):
                                break
                        except json.JSONDecodeError:
                            continue
                raw = full_response
            else:
                # Non-streaming response
                response_data = resp.json()
                raw = response_data.get("response", "")
            
            print(f"📝 Response length: {len(raw)} chars")
            print(f"📝 First 200 chars: {raw[:200]}...")
            
            result = _json_from_llm_text(raw)
            print(f"✅ Successfully parsed JSON from Ollama response")
            return result
            
        except json.JSONDecodeError as je:
            print(f"❌ JSON parsing error on attempt {attempt + 1}: {je}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            print(f"❌ All retries exhausted. Raw response: {raw[:500] if 'raw' in locals() else 'N/A'}...")
            raise RuntimeError(f"Failed to parse JSON from Ollama response after {retries} attempts")
            
        except requests.exceptions.HTTPError as he:
            print(f"❌ HTTP error on attempt {attempt + 1}: {he}")
            if resp is None:
                raise RuntimeError(f"Ollama API HTTP error: {he}")
            
            print(f"🔍 Status Code: {resp.status_code}")
            try:
                error_response = resp.json()
                error_message = error_response.get('error', 'Unknown error')
                print(f"🔍 Ollama Error: {error_message}")
            except:
                print(f"🔍 Raw response: {resp.text[:500]}")
            
            if resp.status_code == 404:
                raise RuntimeError(f"Model '{model}' not found on Ollama server. Please check available models.")
            elif resp.status_code >= 500:
                if attempt < retries - 1:
                    time.sleep(2)
                    continue
            
            raise RuntimeError(f"Ollama API HTTP error: {he}")
            
        except requests.exceptions.ConnectionError as ce:
            print(f"❌ Connection error on attempt {attempt + 1}: {ce}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"Failed to connect to Ollama API: {ce}")
            
        except requests.exceptions.Timeout as te:
            print(f"❌ Timeout error on attempt {attempt + 1}: {te}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"Ollama API request timed out: {te}")
            
        except Exception as e:
            print(f"❌ Unexpected error on attempt {attempt + 1}: {type(e).__name__}: {e}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            raise RuntimeError(f"Unexpected error calling Ollama API: {e}")
    
    raise RuntimeError(f"Failed to get response from Ollama API after {retries} attempts")

# Reuse the same utility functions from original
def clamp_int(v, lo, hi) -> Optional[int]:
    if v is None: return None
    try:
        return max(lo, min(hi, int(v)))
    except Exception:
        m = re.search(r"\b(\d{1,2})\b", str(v))
        if m:
            return max(lo, min(hi, int(m.group(1))))
        return None

def as_float_or_none(v) -> Optional[float]:
    if v is None: return None
    try:
        return float(v)
    except Exception:
        try:
            m = re.search(r"[-+]?\d*\.?\d+", str(v))
            return float(m.group(0)) if m else None
        except Exception:
            return None

def as_date_or_none(s: str) -> Optional[str]:
    if not s: return None
    s = str(s).strip()
    return s if re.match(r"^\d{4}-\d{2}-\d{2}$", s) else None

def compute_exposure(lk, im) -> Optional[float]:
    if lk is None or im is None: return None
    try:
        return round(float(lk) * float(im), 2)
    except Exception:
        return None

def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None

# =========================
# FILE EXTRACTORS (Same as original)
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF (pdfplumber preferred; PyPDF2 fallback)."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        print(f"[PDF] Extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX] Extraction error: {e}")
        return ""

def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        print(f"[XLSX] Extraction error: {e}")
        return ""

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""

# =========================
# DOCUMENT PARSING & DETECTION (Same as original)
# =========================
def detect_and_parse_risk_blocks(text: str) -> list[dict]:
    """Detect how many risks are in the document by finding 'Risk X:' patterns."""
    print(f"🔍 Detecting risk blocks in document...")
    
    risk_pattern = r'Risk\s+(\d+):\s*([^\n]+)'
    risk_matches = list(re.finditer(risk_pattern, text, re.IGNORECASE))
    
    if not risk_matches:
        print(f"⚠️  No risks found with 'Risk X:' pattern")
        return []
    
    print(f"✅ Found {len(risk_matches)} risk(s) in document")
    
    risks = []
    for i, match in enumerate(risk_matches):
        risk_num = match.group(1)
        risk_title = match.group(2).strip()
        
        start_pos = match.end()
        if i + 1 < len(risk_matches):
            end_pos = risk_matches[i + 1].start()
        else:
            end_pos = len(text)
        
        risk_block = text[start_pos:end_pos]
        
        print(f"  📋 Risk {risk_num}: {risk_title[:60]}...")
        
        risk_data = {
            "RiskTitle": risk_title,
            "_extracted_fields": ["RiskTitle"],
            "_risk_block": risk_block
        }
        
        field_mappings = {
            r'Description:\s*([^\n]+(?:\n(?!(?:Risk|Possible Damage|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'RiskDescription',
            r'Possible Damage:\s*([^\n]+(?:\n(?!(?:Risk|Description|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'PossibleDamage',
            r'Risk Priority:\s*([^\n]+)': 'RiskPriority',
            r'Status:\s*([^\n]+)': 'RiskExposureRating',
            r'Mitigation:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Status|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'RiskMitigation',
            r'Created At:\s*([^\n]+)': 'CreatedAt',
            r'Risk Type:\s*([^\n]+)': 'RiskType',
            r'Risk Likelihood:\s*([^\n]+)': 'RiskLikelihood',
            r'Risk Impact:\s*([^\n]+)': 'RiskImpact',
            r'Category:\s*([^\n]+)': 'Category',
            r'Criticality:\s*([^\n]+)': 'Criticality',
            r'Business Impact:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'BusinessImpact',
        }
        
        for pattern, field_name in field_mappings.items():
            match = re.search(pattern, risk_block, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if value:
                    risk_data[field_name] = value
                    risk_data["_extracted_fields"].append(field_name)
                    print(f"    ✓ {field_name}: {value[:50]}...")
        
        risks.append(risk_data)
    
    return risks

# =========================
# AI EXTRACTION CORE (OPTIMIZED)
# =========================
def infer_single_field(field_name: str, current_record: dict, document_context: str) -> tuple[Any, dict]:
    """Focused prompt for ONE field using AI (OPTIMIZED for Ollama)."""
    print(f"🤖 AI PREDICTING FIELD: {field_name} (OPTIMIZED)")
    
    guidance = FIELD_PROMPTS.get(field_name, "Return a concise, professional value.")
    
    # Optimize context size
    context_size = _calculate_optimal_context_size(len(document_context), "simple")
    optimized_context = document_context[:context_size] if len(document_context) > context_size else document_context
    
    # Select appropriate model
    model = _select_model_by_complexity(len(optimized_context), 1)
    
    mini = f"""
You are a GRC analyst. Infer ONLY the field "{field_name}" for this risk.
Return JSON: {{"value": <scalar or string>, "confidence": 0.0-1.0, "rationale": "brief explanation"}}

Context (document):
\"\"\"{optimized_context}\"\"\"

Current risk (partial):
{json.dumps({k: current_record.get(k) for k in RISK_DB_FIELDS if current_record.get(k)}, indent=2)}

Rules:
- {guidance}
- If you cannot infer, return {{"value": null, "confidence": 0.0, "rationale": "Not enough information"}}.
- Always include a brief rationale explaining your decision.
- Return ONLY valid JSON, no markdown, no code blocks.
"""
    try:
        print(f"   📤 Sending prompt to Ollama ({model}) for {field_name}...")
        out = call_ollama_json(mini, model=model)
        v = out.get("value") if isinstance(out, dict) else None
        confidence = out.get("confidence", 0.7) if isinstance(out, dict) else 0.7
        rationale = out.get("rationale", "AI predicted based on document context") if isinstance(out, dict) else "AI predicted based on document context"
        print(f"   ✅ AI PREDICTED {field_name}: '{v}' (confidence: {confidence:.2f})")
    except Exception as e:
        print(f"   ❌ AI FAILED to predict {field_name}: {str(e)}")
        v = None
        confidence = 0.0
        rationale = f"AI prediction failed: {str(e)}"

    metadata = {
        "source": "AI_GENERATED",
        "confidence": confidence,
        "rationale": rationale,
        "model_used": model
    }

    # Normalize after inference (same as original)
    if field_name in ("RiskLikelihood", "RiskImpact"):
        v = clamp_int(v, 1, 10) or 5
    elif field_name == "RiskExposureRating":
        lk = clamp_int(current_record.get("RiskLikelihood"), 1, 10)
        im = clamp_int(current_record.get("RiskImpact"), 1, 10)
        computed = compute_exposure(lk, im)
        v = computed if computed is not None else 0.0
    elif field_name in ("RiskMultiplierX", "RiskMultiplierY"):
        vf = as_float_or_none(v)
        v = vf if vf is not None else 0.5
    elif field_name == "CreatedAt":
        v = as_date_or_none(v) or date.today().isoformat()
    elif field_name == "Criticality":
        v = normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
    elif field_name == "RiskPriority":
        v = normalize_choice(v, PRIORITY_CHOICES) or "Medium"
    elif field_name == "RiskType":
        v = normalize_choice(v, RISKTYPE_HINTS) or "Current"
    elif isinstance(v, str):
        v = v.strip() or None
    
    return v, metadata

def fallback_risk_extraction(text: str) -> list[dict]:
    """Minimal pattern-based fallback when AI fails completely."""
    risks = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    risk_keywords = ["risk", "threat", "vulnerability", "hazard", "danger", "exposure"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in risk_keywords):
            if current:
                risks.append(current)
            count += 1
            current = {
                "RiskTitle": f"Risk {count}: {ln[:100]}",
                "RiskDescription": ln,
                "Category": "Operational",
                "Criticality": "Medium",
                "RiskPriority": "Medium",
                "RiskType": "Current",
                "BusinessImpact": "Potential business impact",
                "PossibleDamage": "Potential damages vary (reputation, compliance, downtime).",
                "RiskMitigation": "Mitigation steps to be defined after assessment.",
                "RiskLikelihood": 5,
                "RiskImpact": 5,
                "RiskExposureRating": 25.0,
                "CreatedAt": date.today().isoformat(),
                "RiskMultiplierX": 0.5,
                "RiskMultiplierY": 0.5,
            }

    if current:
        risks.append(current)

    if not risks:
        risks.append({
            "RiskTitle": "Document Risk Analysis",
            "RiskDescription": f"Automated risk derived from document (length={len(text)} chars).",
            "Category": "Operational",
            "Criticality": "Medium",
            "RiskPriority": "Medium",
            "RiskType": "Current",
            "BusinessImpact": "Potential business impact in operations/compliance.",
            "PossibleDamage": "Data loss, downtime, penalties, reputation.",
            "RiskMitigation": "Review & implement standard mitigations for identified weaknesses.",
            "RiskLikelihood": 5,
            "RiskImpact": 5,
            "RiskExposureRating": 25.0,
            "CreatedAt": date.today().isoformat(),
            "RiskMultiplierX": 0.5,
            "RiskMultiplierY": 0.5,
        })

    return risks

def parse_risks_from_text(text: str) -> list[dict]:
    """
    OPTIMIZED VERSION:
    1. Detect risks in document
    2. Extract fields explicitly present
    3. Use optimized Ollama to fill ONLY missing fields
    """
    print(f"📊 parse_risks_from_text() called with {len(text)} chars of text (OPTIMIZED)")
    
    # Step 1: Detect and parse risk blocks
    detected_risks = detect_and_parse_risk_blocks(text)
    
    if not detected_risks:
        print(f"⚠️  No structured risks found. Falling back to old AI extraction...")
        return fallback_risk_extraction(text)
    
    print(f"✅ Detected {len(detected_risks)} risk(s), now processing each (OPTIMIZED)...")
    
    # Step 2: Process each risk with optimized AI
    completed_risks = []
    for idx, risk_data in enumerate(detected_risks, 1):
        print(f"\n🔧 Processing Risk {idx}: {risk_data.get('RiskTitle', 'Unknown')[:50]}... (OPTIMIZED)")
        
        item = {k: None for k in RISK_DB_FIELDS}
        
        for field in RISK_DB_FIELDS:
            if field in risk_data and risk_data[field]:
                item[field] = risk_data[field]
        
        risk_block = risk_data.get("_risk_block", "")
        extracted_fields = risk_data.get("_extracted_fields", [])
        
        print(f"  📝 Extracted fields: {', '.join(extracted_fields)}")
        
        if "_meta" not in item:
            item["_meta"] = {}
        if "per_field" not in item["_meta"]:
            item["_meta"]["per_field"] = {}
        
        for field in extracted_fields:
            if field in item and item[field]:
                item["_meta"]["per_field"][field] = {
                    "source": "EXTRACTED",
                    "confidence": 0.95,
                    "rationale": "Found explicitly in document"
                }
        
        # Normalize extracted fields
        if item.get("RiskLikelihood"):
            item["RiskLikelihood"] = clamp_int(item["RiskLikelihood"], 1, 10)
        if item.get("RiskImpact"):
            item["RiskImpact"] = clamp_int(item["RiskImpact"], 1, 10)
        if item.get("RiskExposureRating"):
            item["RiskExposureRating"] = as_float_or_none(item["RiskExposureRating"])
        if item.get("Criticality"):
            item["Criticality"] = normalize_choice(item["Criticality"], CRITICALITY_CHOICES)
        if item.get("RiskPriority"):
            item["RiskPriority"] = normalize_choice(item["RiskPriority"], PRIORITY_CHOICES)
        if item.get("RiskType"):
            item["RiskType"] = normalize_choice(item["RiskType"], RISKTYPE_HINTS)
        if item.get("CreatedAt"):
            item["CreatedAt"] = as_date_or_none(item["CreatedAt"])
        if item.get("RiskMultiplierX"):
            item["RiskMultiplierX"] = as_float_or_none(item["RiskMultiplierX"])
        if item.get("RiskMultiplierY"):
            item["RiskMultiplierY"] = as_float_or_none(item["RiskMultiplierY"])
        
        # Step 3: Use optimized AI to fill missing fields
        missing_fields = [f for f in RISK_DB_FIELDS if item.get(f) in (None, "", []) and f != "RiskTitle"]
        if missing_fields:
            print(f"  🤖 Missing fields: {', '.join(missing_fields)}")
            print(f"  🤖 Using OPTIMIZED Ollama to infer missing fields...")
            
            for field in missing_fields:
                print(f"    🔍 Inferring {field}...")
                value, metadata = infer_single_field(field, item, risk_block or text[:3000])
                item[field] = value
                item["_meta"]["per_field"][field] = metadata
                print(f"    🏷️  Marked {field} as AI_GENERATED in metadata")
        else:
            print(f"  ✅ All fields extracted from document!")
        
        # Final normalization
        item["RiskLikelihood"] = item["RiskLikelihood"] or 5
        item["RiskImpact"] = item["RiskImpact"] or 5
        
        if not item.get("RiskExposureRating"):
            item["RiskExposureRating"] = compute_exposure(item["RiskLikelihood"], item["RiskImpact"]) or 25.0
        
        item["RiskExposureRating"] = float(max(0.0, min(100.0, item["RiskExposureRating"])))
        item["RiskMultiplierX"] = item["RiskMultiplierX"] or 0.5
        item["RiskMultiplierY"] = item["RiskMultiplierY"] or 0.5
        item["CreatedAt"] = item["CreatedAt"] or date.today().isoformat()
        item["Criticality"] = item["Criticality"] or "Medium"
        item["RiskPriority"] = item["RiskPriority"] or "Medium"
        item["RiskType"] = item["RiskType"] or "Current"
        
        if not item.get("RiskTitle"):
            raise ValueError(f"RiskTitle is missing for risk {idx}. All risk titles must be present in the document as 'Risk X: Title'.")
        
        ai_fields = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "AI_GENERATED"]
        extracted = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "EXTRACTED"]
        
        print(f"  📊 Metadata Summary:")
        print(f"     🤖 AI Generated: {len(ai_fields)} fields - {ai_fields}")
        print(f"     📄 Extracted: {len(extracted)} fields - {extracted}")
        
        completed_risks.append(item)
        print(f"  ✅ Risk {idx} completed! (OPTIMIZED)")
    
    return completed_risks

# =========================
# DJANGO API ENDPOINTS (Same interface as original)
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def upload_and_process_risk_document_optimized(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
print(f"📤 Upload request for risk document (OPTIMIZED VERSION)")

    if request.method == 'OPTIONS':
        response = HttpResponse()
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        response['Access-Control-Max-Age'] = '86400'
        return response

    try:
        if 'file' not in request.FILES:
            resp = JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            resp = JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'risk')
        os.makedirs(upload_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        print(f"✅ File saved to: {file_path}")

        try:
            print(f"🔍 STEP 1: Starting text extraction from {ext} file... (OPTIMIZED)")
            text = extract_text_from_file(file_path, ext)
            
            if not text or len(text.strip()) < 50:
                print(f"❌ ERROR: Could not extract meaningful text. Length: {len(text) if text else 0}")
                resp = JsonResponse({'status': 'error', 'message': 'Could not extract meaningful text from document'}, status=400)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp

            print(f"✅ STEP 1 COMPLETE: Extracted {len(text)} characters from document")
            
            print(f"🤖 STEP 2: Calling OPTIMIZED Ollama model to extract risks...")
            risks = parse_risks_from_text(text)
            
            print(f"✅ STEP 2 COMPLETE: OPTIMIZED AI extracted {len(risks)} risk(s) from document")
            for idx, risk in enumerate(risks, 1):
                print(f"  Risk {idx}: {risk.get('RiskTitle', 'Untitled')[:50]}...")

            resp = JsonResponse({
                'status': 'success',
                'message': f'Successfully extracted {len(risks)} risk(s) (OPTIMIZED)',
                'document_name': file_name,
                'saved_path': safe_filename,
                'extracted_text_length': len(text),
                'risks': risks,
                'version': 'optimized'
            })
            resp['Access-Control-Allow-Origin'] = '*'
            return resp
        except Exception as process_error:
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        resp = JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)
        resp['Access-Control-Allow-Origin'] = '*'
        return resp













