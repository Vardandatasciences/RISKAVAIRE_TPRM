"""
AI-Powered Risk Instance Document Ingestion (Complete)
- Reads PDF/DOCX/XLSX/TXT
- Extracts risk instance data using AI (OpenAI or Ollama - configurable)
- Fills missing fields with focused prompts
- Returns normalized, DB-ready JSON for `risk_instance` table

Configuration:
- Set RISK_AI_PROVIDER='openai' or 'ollama' in environment or Django settings
- Default: 'ollama' if Ollama is configured, else 'openai'
- OpenAI requires: OPENAI_API_KEY
- Ollama requires: OLLAMA_BASE_URL (uses optimized quantized models)
"""

import os
import re
import json
import time
import tempfile
from datetime import date, datetime
from typing import Any, Optional

import requests
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser

# RBAC imports
from ...rbac.decorators import rbac_required
# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)


# Phase 2 Optimizations (reuse same utilities as risk_ai_doc)
from ...utils.ai_cache import cached_llm_call
from ...utils.document_preprocessor import preprocess_document, calculate_document_hash
from ...utils.few_shot_prompts import get_field_extraction_prompt

# Phase 3 Optimizations (reuse same utilities as risk_ai_doc)
from ...utils.rag_system import (
    add_document_to_rag,
    retrieve_relevant_context,
    build_rag_prompt,
    is_rag_available,
    get_rag_stats,
)
from ...utils.model_router import (
    route_model,
    track_system_load,
    get_current_system_load,
)
from ...utils.request_queue import (
    rate_limit_decorator,
    process_with_queue,
    get_queue_status
)
from ...utils.file_compression import decompress_if_needed
from ...routes.Global.s3_fucntions import create_direct_mysql_client

# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import RiskInstance  # Import RiskInstance model


# =========================
# AI PROVIDER CONFIG (OpenAI or Ollama)
# =========================
# Use the same provider configuration and helpers as risk_ai_doc
from django.conf import settings
from .risk_ai_doc import (
    AI_PROVIDER,
    call_ollama_json,
    call_openai_json,  # Phase 2: cached OpenAI wrapper
    _calculate_optimal_context_size,
    _select_ollama_model_by_complexity,
    OLLAMA_BASE_URL,
    OLLAMA_MODEL_DEFAULT,
    OLLAMA_MODEL_FAST,
    OLLAMA_MODEL_COMPLEX,
    OPENAI_API_KEY,
    OPENAI_API_URL,
    OPENAI_MODEL,
)

if AI_PROVIDER == 'openai':
    if not OPENAI_API_KEY:
        print("[WARNING]  WARNING: OPENAI_API_KEY not found in Django settings!")
        print("   Please set OPENAI_API_KEY in your .env file.")
    else:
        print(f"[API] OpenAI Configuration for Risk Instance AI:")
        print(f"   API URL: {OPENAI_API_URL}")
        print(f"   Model: {OPENAI_MODEL}")
        print(f"   API Key: {'*' * (len(OPENAI_API_KEY) - 4) + OPENAI_API_KEY[-4:]}")
elif AI_PROVIDER == 'ollama':
    print("[INFO] Ollama Configuration for Risk Instance AI (OPTIMIZED):")
    print(f"   Base URL: {OLLAMA_BASE_URL}")
    print(f"   Default Model: {OLLAMA_MODEL_DEFAULT}")
    print(f"   Fast Model: {OLLAMA_MODEL_FAST}")
    print(f"   Complex Model: {OLLAMA_MODEL_COMPLEX}")

# RiskInstance DB fields (excluding auto-generated IDs: RiskInstanceId, UserId, ReportedBy, ReviewerId, IncidentId, ComplianceId, RiskId)
RISK_INSTANCE_DB_FIELDS = [
    "RiskTitle",
    "RiskDescription",
    "PossibleDamage",
    "RiskPriority",
    "Criticality",
    "Category",
    "Origin",
    "RiskLikelihood",
    "RiskImpact",
    "RiskExposureRating",
    "RiskMultiplierX",
    "RiskMultiplierY",
    "Appetite",
    "RiskResponseType",
    "RiskResponseDescription",
    "RiskMitigation",
    "RiskType",
    "RiskOwner",
    "BusinessImpact",
    "RiskStatus",
    "MitigationStatus",
    "ModifiedMitigations",
    "RiskFormDetails",
    "Reviewer",
]

# Canonical choices / constraints to stabilize LLM outputs
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES = ["Low", "Medium", "High", "Critical"]
CATEGORY_HINTS = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]
RISKTYPE_HINTS = ["Current", "Residual", "Inherent", "Emerging", "Accepted"]
ORIGIN_HINTS = ["Internal", "External", "Third-Party", "Regulatory", "Market", "Operational"]
APPETITE_HINTS = ["Low", "Medium", "High"]
RESPONSE_TYPE_HINTS = ["Avoid", "Mitigate", "Transfer", "Accept"]
RISK_STATUS_CHOICES = ["Not Assigned", "Assigned", "Approved", "Rejected"]
MITIGATION_STATUS_CHOICES = ["Pending", "Yet to Start", "Work In Progress", "Revision Required by Reviewer", "Revision Required by User", "Completed"]
DATE_FORMAT_HINT = "YYYY-MM-DD (ISO)"

# Field-specific micro-prompts (used when a single field is missing/invalid)
# NOTE: RiskTitle is NEVER inferred by AI - it must always come from the document
FIELD_PROMPTS = {
    "RiskDescription": "Write a precise, detailed description (2–4 sentences) of how/why the risk instance occurred, what happened, when it happened, and who was affected. Include specific details about the incident.",
    "PossibleDamage": "Describe concrete, specific damages that occurred or could occur (data loss volume, downtime duration, financial penalties, reputation impact). Be detailed and quantitative where possible (2–3 sentences).",
    "RiskPriority": f"Return one of: {PRIORITY_CHOICES}. Analyze the severity based on exposure, criticality, business impact, and urgency. Consider both likelihood and potential damage.",
    "Criticality": f"Return one of: {CRITICALITY_CHOICES}. Consider the severity of impact on business operations, compliance requirements, financial impact, and reputational damage.",
    "Category": f"Return one category from this list (best fit): {CATEGORY_HINTS}. Analyze the nature of the risk and choose the most appropriate category based on the root cause and impact area.",
    "Origin": f"Return one of: {ORIGIN_HINTS}. Identify where the risk originated from - was it internal processes, external threats, third-party vendors, regulatory changes, or market conditions?",
    "RiskLikelihood": "Return an integer 1–10 (1=rare/almost never, 5=possible/moderate, 10=almost certain/very frequent). Base this on historical data, environmental factors, and current controls.",
    "RiskImpact": "Return an integer 1–10 (1=negligible/minimal, 5=moderate/noticeable, 10=catastrophic/severe). Consider financial, operational, reputational, and compliance impacts.",
    "RiskExposureRating": "Return a float (0–100) representing overall risk exposure. Calculate as: (Likelihood × Impact × average of multipliers). Higher values indicate greater exposure.",
    "RiskMultiplierX": "Return a float in 0.1–2.0 reflecting organizational weighting factor X for likelihood adjustment based on industry, size, or environmental factors (default ~1.0 if unknown).",
    "RiskMultiplierY": "Return a float in 0.1–2.0 reflecting organizational weighting factor Y for impact adjustment based on risk tolerance and organizational resilience (default ~1.0 if unknown).",
    "Appetite": f"Return one of: {APPETITE_HINTS}. Determine the organization's tolerance level for this type of risk based on strategic objectives, regulatory requirements, and industry standards.",
    "RiskResponseType": f"Return one of: {RESPONSE_TYPE_HINTS}. Choose the most appropriate response: Avoid (eliminate risk), Mitigate (reduce likelihood/impact), Transfer (insurance/outsource), Accept (acknowledge and monitor).",
    "RiskResponseDescription": "Provide a detailed 2-3 sentence description of the chosen risk response strategy, explaining specific actions, rationale, timeline, and expected outcomes.",
    "RiskMitigation": "Return 3–5 specific, actionable mitigation steps as a JSON array of objects with 'step' (action title) and 'description' (detailed implementation) fields. Each step should be practical, measurable, and time-bound. Example: [{\"step\": \"Implement Multi-Factor Authentication\", \"description\": \"Deploy MFA across all user accounts within 30 days to prevent unauthorized access\"}]",
    "RiskType": f"Return one of: {RISKTYPE_HINTS}. Current=existing risk, Residual=risk after controls, Inherent=risk before controls, Emerging=newly identified, Accepted=acknowledged without action.",
    "RiskOwner": "Return the specific name, role, or department responsible for managing this risk instance (e.g., 'John Smith - IT Security Manager', 'Compliance Department', 'Chief Risk Officer'). Be specific.",
    "BusinessImpact": "Provide a detailed 2-4 sentence explanation of business impact using business terms: revenue loss, SLA breaches, customer churn, regulatory penalties, operational disruption, market share impact, brand reputation damage. Include quantitative estimates where possible.",
    "RiskStatus": f"Return one of: {RISK_STATUS_CHOICES}. Assess the current state: Not Assigned (new, unassigned), Assigned (owner designated), Approved (accepted by management), Rejected (dismissed as non-risk).",
    "MitigationStatus": f"Return one of: {MITIGATION_STATUS_CHOICES}. Determine progress: Pending (awaiting action), Yet to Start (planned but not begun), Work In Progress (actively being addressed), Revision Required by Reviewer (needs reviewer changes), Revision Required by User (needs user changes), Completed (fully implemented).",
    "ModifiedMitigations": "Return a JSON array documenting any changes made to original mitigation plans. Format: [{\"date\": \"YYYY-MM-DD\", \"changed_by\": \"Name/Role\", \"changes\": \"Description of modifications\", \"reason\": \"Why changes were needed\"}]. Return empty array [] if no modifications yet.",
    "RiskFormDetails": "Return a JSON object containing additional structured risk assessment details. Format: {\"assessment_method\": \"Qualitative/Quantitative/Mixed\", \"data_sources\": [\"source1\", \"source2\"], \"stakeholders_consulted\": [\"person1\", \"person2\"], \"assessment_date\": \"YYYY-MM-DD\", \"next_review_date\": \"YYYY-MM-DD\", \"additional_notes\": \"Any relevant notes\"}. Infer reasonable values based on context.",
    "Reviewer": "Return the name or role of the person who reviewed/approved this risk instance (e.g., 'Sarah Johnson', 'Senior Risk Analyst', 'Compliance Manager'). Default to 'Pending Review' if not yet reviewed.",
}

# Strict JSON schema block the LLM must follow
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON array. No markdown, no code blocks, no explanations.
Start with [ and end with ]. Use proper JSON syntax with double quotes.

Example structure (return array of risk instances like this):
[
  {{
    "RiskTitle": "Specific risk instance title here",
    "RiskDescription": "Detailed description of the risk instance occurrence",
    "PossibleDamage": "Concrete description of damages",
    "RiskPriority": "Medium",
    "Criticality": "Medium",
    "Category": "Operational",
    "Origin": "Internal",
    "RiskLikelihood": 5,
    "RiskImpact": 6,
    "RiskExposureRating": 30.0,
    "RiskMultiplierX": 1.0,
    "RiskMultiplierY": 1.0,
    "Appetite": "Medium",
    "RiskResponseType": "Mitigate",
    "RiskResponseDescription": "Detailed response strategy description",
    "RiskMitigation": [{{"step": "Step 1", "description": "Detailed action description"}}],
    "RiskType": "Current",
    "RiskOwner": "Risk Manager Name/Role",
    "BusinessImpact": "Detailed business impact explanation",
    "RiskStatus": "Not Assigned",
    "MitigationStatus": "Pending",
    "ModifiedMitigations": [],
    "RiskFormDetails": {{"assessment_method": "Qualitative", "data_sources": ["Document"], "assessment_date": "{date.today().isoformat()}"}},
    "Reviewer": "Pending Review",
    "_meta": {{
      "per_field": {{
        "RiskTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found explicitly in document"}},
        "Criticality": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from severity indicators"}},
        "RiskLikelihood": {{"source": "AI_GENERATED", "confidence": 0.6, "rationale": "Estimated based on frequency mentions"}}
      }}
    }}
  }}
]

Rules:
- Criticality must be: {CRITICALITY_CHOICES}
- RiskPriority must be: {PRIORITY_CHOICES}
- RiskLikelihood and RiskImpact must be integers 1-10
- RiskExposureRating must be float 0-100
- RiskMultiplierX and RiskMultiplierY must be floats 0.1-2.0
- RiskMitigation must be a JSON array of objects with "step" and "description"
- ModifiedMitigations must be a JSON array (empty [] if no modifications)
- RiskFormDetails must be a JSON object
- RiskStatus must be one of: {RISK_STATUS_CHOICES}
- MitigationStatus must be one of: {MITIGATION_STATUS_CHOICES}
- Appetite must be one of: {APPETITE_HINTS}
- RiskResponseType must be one of: {RESPONSE_TYPE_HINTS}
- Origin must be one of: {ORIGIN_HINTS}
- RiskType must be one of: {RISKTYPE_HINTS}
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON array, nothing else
- Include _meta.per_field for EVERY field showing source (EXTRACTED or AI_GENERATED), confidence (0.0-1.0), and rationale
"""

def clamp_int(v, lo, hi) -> Optional[int]:
    if v is None: return None
    try:
        return max(lo, min(hi, int(v)))
    except Exception:
        # try to find a number in string
        m = re.search(r"\b(\d{1,2})\b", str(v))
        if m:
            return max(lo, min(hi, int(m.group(1))))
        return None

def as_float_or_none(v) -> Optional[float]:
    if v is None: return None
    try:
        return float(v)
    except Exception:
        try:
            # extract first float from string
            m = re.search(r"[-+]?\d*\.?\d+", str(v))
            return float(m.group(0)) if m else None
        except Exception:
            return None

def as_date_or_none(s: str) -> Optional[str]:
    if not s: return None
    s = str(s).strip()
    return s if re.match(r"^\d{4}-\d{2}-\d{2}$", s) else None

def compute_exposure(lk, im) -> Optional[float]:
    if lk is None or im is None: return None
    try:
        return round(float(lk) * float(im), 2)
    except Exception:
        return None

def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None


# =========================
# FILE EXTRACTORS
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF (pdfplumber preferred; PyPDF2 fallback)."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        print(f"[PDF] Extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX] Extraction error: {e}")
        return ""

def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        print(f"[XLSX] Extraction error: {e}")
        return ""

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""


# =========================
# DOCUMENT PARSING & DETECTION
# =========================
def detect_and_parse_risk_instance_blocks(text: str) -> list[dict]:
    """
    Step 1: Detect how many risk instances are in the document by finding 'Risk X:' patterns.
    Step 2: For each risk instance, extract whatever fields are explicitly present.
    Returns a list of partially-filled risk instance dictionaries with only extracted fields.
    """
    print(f"[DEBUG] Detecting risk instance blocks in document...")
    
    # Find all risk blocks using pattern "Risk X: Title" or "Risk Instance X: Title"
    risk_pattern = r'Risk(?:\s+Instance)?\s+(\d+):\s*([^\n]+)'
    risk_matches = list(re.finditer(risk_pattern, text, re.IGNORECASE))
    
    if not risk_matches:
        print(f"[WARNING]  No risk instances found with 'Risk X:' or 'Risk Instance X:' pattern")
        return []
    
    print(f"[OK] Found {len(risk_matches)} risk instance(s) in document")
    
    risk_instances = []
    for i, match in enumerate(risk_matches):
        risk_num = match.group(1)
        risk_title = match.group(2).strip()
        
        # Extract the text block for this risk instance (from this match to next risk or end of document)
        start_pos = match.end()
        if i + 1 < len(risk_matches):
            end_pos = risk_matches[i + 1].start()
        else:
            end_pos = len(text)
        
        risk_block = text[start_pos:end_pos]
        
        print(f"  [LIST] Risk Instance {risk_num}: {risk_title[:60]}...")
        
        # Extract fields that are explicitly present in the block
        risk_data = {
            "RiskTitle": risk_title,  # Always extracted from document
            "_extracted_fields": ["RiskTitle"],  # Track what was extracted
            "_risk_block": risk_block  # Keep for AI context
        }
        
        # Field mapping: document field name -> DB field name
        field_mappings = {
            r'(?:Risk\s+)?Description:\s*([^\n]+(?:\n(?!(?:Risk|Possible Damage|Risk Priority|Criticality|Category|Origin|Status|Mitigation|Risk Type|Risk Owner|Business Impact|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'RiskDescription',
            r'Possible Damage:\s*([^\n]+(?:\n(?!(?:Risk|Description|Risk Priority|Criticality|Category|Origin|Status|Mitigation|Risk Type|Risk Owner|Business Impact|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'PossibleDamage',
            r'Risk Priority:\s*([^\n]+)': 'RiskPriority',
            r'Criticality:\s*([^\n]+)': 'Criticality',
            r'Category:\s*([^\n]+)': 'Category',
            r'Origin:\s*([^\n]+)': 'Origin',
            r'Risk Likelihood:\s*([^\n]+)': 'RiskLikelihood',
            r'Risk Impact:\s*([^\n]+)': 'RiskImpact',
            r'(?:Risk\s+)?Exposure(?:\s+Rating)?:\s*([^\n]+)': 'RiskExposureRating',
            r'Multiplier\s*X:\s*([^\n]+)': 'RiskMultiplierX',
            r'Multiplier\s*Y:\s*([^\n]+)': 'RiskMultiplierY',
            r'Appetite:\s*([^\n]+)': 'Appetite',
            r'Risk Response Type:\s*([^\n]+)': 'RiskResponseType',
            r'Risk Response(?:\s+Description)?:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Criticality|Category|—))[^\n]+)*)': 'RiskResponseDescription',
            r'Mitigation:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Criticality|Category|Origin|Risk Type|Risk Owner|Business Impact|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'RiskMitigation',
            r'Risk Type:\s*([^\n]+)': 'RiskType',
            r'Risk Owner:\s*([^\n]+)': 'RiskOwner',
            r'Business Impact:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Criticality|Category|Origin|Mitigation|Risk Type|Risk Owner|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'BusinessImpact',
            r'Risk Status:\s*([^\n]+)': 'RiskStatus',
            r'Mitigation Status:\s*([^\n]+)': 'MitigationStatus',
            r'Modified Mitigations:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|—))[^\n]+)*)': 'ModifiedMitigations',
            r'Risk Form Details:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|—))[^\n]+)*)': 'RiskFormDetails',
            r'Reviewer:\s*([^\n]+)': 'Reviewer',
        }
        
        # Extract each field if present
        for pattern, field_name in field_mappings.items():
            match = re.search(pattern, risk_block, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if value:  # Only add if not empty
                    risk_data[field_name] = value
                    risk_data["_extracted_fields"].append(field_name)
                    print(f"    [EMOJI] {field_name}: {value[:50]}...")
        
        risk_instances.append(risk_data)
    
    return risk_instances


# =========================
# AI EXTRACTION CORE
# =========================
def infer_single_field(
    field_name: str,
    current_record: dict,
    document_context: str,
    document_hash: str = None,
) -> tuple[Any, dict]:
    """
    Focused prompt for ONE field using AI (supports both OpenAI and Ollama).
    Uses Phase 2 few-shot prompts, Phase 3 RAG, and caching.
    Returns: (value, metadata_dict)
    """
    provider_name = AI_PROVIDER.upper()
    print(f"[AI] AI PREDICTING FIELD: {field_name} (using {provider_name} with Phase 2+3 optimizations)")

    # Optimize context size for Ollama; keep previous behavior for OpenAI
    if AI_PROVIDER == 'ollama':
        context_size = _calculate_optimal_context_size(len(document_context), "simple")
        optimized_context = document_context[:context_size] if len(document_context) > context_size else document_context
    else:
        optimized_context = document_context[:3000]

    # Phase 3: Try to retrieve relevant context from RAG
    rag_context = None
    if is_rag_available():
        try:
            query = f"What is the {field_name} for this risk instance?"
            retrieved = retrieve_relevant_context(query, n_results=3)
            if retrieved:
                rag_context = retrieved
                print(f"   [DATA] Phase 3 RAG: Retrieved {len(retrieved)} relevant document chunks")
        except Exception as e:
            print(f"   [WARNING]  RAG retrieval failed: {e}")

    # Phase 2: Use few-shot prompt template (reuse shared helper)
    try:
        mini = get_field_extraction_prompt(
            field_name=field_name,
            document_text=optimized_context,
            field_prompts=FIELD_PROMPTS,
        )
        # Add current record context
        mini += f"\n\nCurrent risk instance (partial):\n{json.dumps({k: current_record.get(k) for k in RISK_INSTANCE_DB_FIELDS if current_record.get(k)}, indent=2)}"
        print(f"   [DATA] Using few-shot prompt template for {field_name}")
    except Exception as e:
        print(f"   [WARNING]  Few-shot prompt failed, using basic prompt: {e}")
        guidance = FIELD_PROMPTS.get(field_name, "Return a concise, professional value.")
        mini = f"""
You are a GRC analyst. Infer ONLY the field "{field_name}" for this risk instance.
Return JSON: {{"value": <scalar or string or array>, "confidence": 0.0-1.0, "rationale": "brief explanation"}}.

Context (document):
\"\"\"{optimized_context}\"\"\"

Current risk instance (partial):
{json.dumps({k: current_record.get(k) for k in RISK_INSTANCE_DB_FIELDS if current_record.get(k)}, indent=2)}

Rules:
- {guidance}
- If you cannot infer, return {{"value": null, "confidence": 0.0, "rationale": "Not enough information"}}.
- Always include a brief rationale explaining your decision.
- Return ONLY valid JSON, no markdown, no code blocks.
"""

    # Phase 3: Enhance prompt with RAG context if available
    if rag_context:
        mini = build_rag_prompt(
            user_query=mini,
            retrieved_context=rag_context,
            base_prompt=None,
        )

    try:
        if AI_PROVIDER == 'ollama':
            print(f"   [UPLOAD] Sending prompt to Ollama for {field_name}...")
            model = _select_ollama_model_by_complexity(len(optimized_context), 1)
            out = call_ollama_json(mini, model=model, document_hash=document_hash)
            model_used = model
        else:
            print(f"   [UPLOAD] Sending prompt to OpenAI for {field_name}...")
            out = call_openai_json(mini, document_hash=document_hash)
            model_used = OPENAI_MODEL

        v = out.get("value") if isinstance(out, dict) else None
        confidence = out.get("confidence", 0.7) if isinstance(out, dict) else 0.7
        rationale = out.get("rationale", "AI predicted based on document context") if isinstance(out, dict) else "AI predicted based on document context"
        print(f"   [OK] AI PREDICTED {field_name}: '{v}' (confidence: {confidence:.2f})")
    except Exception as e:
        print(f"   [ERROR] AI FAILED to predict {field_name}: {str(e)}")
        v = None
        confidence = 0.0
        rationale = f"AI prediction failed: {str(e)}"
        model_used = None

    # Create metadata for this field
    metadata = {
        "source": "AI_GENERATED",
        "confidence": confidence,
        "rationale": rationale,
        "provider": AI_PROVIDER,
        "model_used": model_used,
    }

    # normalize after inference
    if field_name in ("RiskLikelihood", "RiskImpact"):
        v = clamp_int(v, 1, 10) or 5
    elif field_name == "RiskExposureRating":
        lk = clamp_int(current_record.get("RiskLikelihood"), 1, 10)
        im = clamp_int(current_record.get("RiskImpact"), 1, 10)
        computed = compute_exposure(lk, im)
        v = computed if computed is not None else 0.0
    elif field_name in ("RiskMultiplierX", "RiskMultiplierY"):
        vf = as_float_or_none(v)
        v = vf if vf is not None else 1.0
    elif field_name in ("CreatedAt", "MitigationDueDate"):
        v = as_date_or_none(v) or date.today().isoformat()
    elif field_name == "Criticality":
        v = normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
    elif field_name == "RiskPriority":
        v = normalize_choice(v, PRIORITY_CHOICES) or "Medium"
    elif field_name == "RiskType":
        v = normalize_choice(v, RISKTYPE_HINTS) or "Current"
    elif field_name == "Origin":
        v = normalize_choice(v, ORIGIN_HINTS) or "Internal"
    elif field_name == "Appetite":
        v = normalize_choice(v, APPETITE_HINTS) or "Medium"
    elif field_name == "RiskResponseType":
        v = normalize_choice(v, RESPONSE_TYPE_HINTS) or "Mitigate"
    elif field_name == "RiskStatus":
        v = normalize_choice(v, RISK_STATUS_CHOICES) or "Not Assigned"
    elif field_name == "MitigationStatus":
        v = normalize_choice(v, MITIGATION_STATUS_CHOICES) or "Pending"
    elif field_name == "RiskMitigation":
        # Should be a JSON array of mitigation steps
        if isinstance(v, list):
            pass
        elif isinstance(v, str):
            try:
                v = json.loads(v)
            except:
                v = [{"step": "Mitigation Step 1", "description": v}]
        else:
            v = [{"step": "Mitigation Step 1", "description": "To be defined"}]
    elif field_name == "ModifiedMitigations":
        # Should be a JSON array documenting modifications
        if isinstance(v, list):
            pass
        elif isinstance(v, str):
            try:
                v = json.loads(v)
            except:
                v = []
        else:
            v = []
    elif field_name == "RiskFormDetails":
        # Should be a JSON object with assessment details
        if isinstance(v, dict):
            pass
        elif isinstance(v, str):
            try:
                v = json.loads(v)
            except:
                v = {
                    "assessment_method": "Qualitative",
                    "data_sources": ["Document"],
                    "assessment_date": date.today().isoformat()
                }
        else:
            v = {
                "assessment_method": "Qualitative",
                "data_sources": ["Document"],
                "assessment_date": date.today().isoformat()
            }
    elif field_name == "Reviewer":
        v = str(v).strip() if v else "Pending Review"
    elif isinstance(v, str):
        v = v.strip() or None
    
    return v, metadata


def fallback_risk_instance_extraction(text: str) -> list[dict]:
    """
    Minimal pattern-based fallback when AI fails completely.
    Produces at least one record using generic defaults.
    """
    risk_instances = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    risk_keywords = ["risk", "threat", "vulnerability", "hazard", "danger", "exposure", "incident"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in risk_keywords):
            if current:
                risk_instances.append(current)
            count += 1
            current = {
                "RiskTitle": f"Risk Instance {count}: {ln[:100]}",
                "RiskDescription": ln,
                "PossibleDamage": "Potential damages vary (reputation, compliance, downtime).",
                "RiskPriority": "Medium",
                "Criticality": "Medium",
                "Category": "Operational",
                "Origin": "Internal",
                "RiskLikelihood": 5,
                "RiskImpact": 5,
                "RiskExposureRating": 25.0,
                "RiskMultiplierX": 1.0,
                "RiskMultiplierY": 1.0,
                "Appetite": "Medium",
                "RiskResponseType": "Mitigate",
                "RiskResponseDescription": "Response strategy to be defined after assessment.",
                "RiskMitigation": [{"step": "Step 1", "description": "Mitigation steps to be defined"}],
                "RiskType": "Current",
                "RiskOwner": "Risk Manager",
                "BusinessImpact": "Potential business impact in operations/compliance.",
                "RiskStatus": "Not Assigned",
                "MitigationStatus": "Pending",
                "ModifiedMitigations": [],
                "RiskFormDetails": {
                    "assessment_method": "Qualitative",
                    "data_sources": ["Document"],
                    "assessment_date": date.today().isoformat()
                },
                "Reviewer": "Pending Review",
            }

    if current:
        risk_instances.append(current)

    if not risk_instances:
        risk_instances.append({
            "RiskTitle": "Document Risk Instance Analysis",
            "RiskDescription": f"Automated risk instance derived from document (length={len(text)} chars).",
            "PossibleDamage": "Data loss, downtime, penalties, reputation.",
            "RiskPriority": "Medium",
            "Criticality": "Medium",
            "Category": "Operational",
            "Origin": "Internal",
            "RiskLikelihood": 5,
            "RiskImpact": 5,
            "RiskExposureRating": 25.0,
            "RiskMultiplierX": 1.0,
            "RiskMultiplierY": 1.0,
            "Appetite": "Medium",
            "RiskResponseType": "Mitigate",
            "RiskResponseDescription": "Response strategy to be defined after assessment.",
            "RiskMitigation": [{"step": "Step 1", "description": "Review & implement standard mitigations for identified weaknesses."}],
            "RiskType": "Current",
            "RiskOwner": "Risk Manager",
            "BusinessImpact": "Potential business impact in operations/compliance.",
            "RiskStatus": "Not Assigned",
            "MitigationStatus": "Pending",
            "ModifiedMitigations": [],
            "RiskFormDetails": {
                "assessment_method": "Qualitative",
                "data_sources": ["Document"],
                "assessment_date": date.today().isoformat()
            },
            "Reviewer": "Pending Review",
        })

    return risk_instances


def parse_risk_instances_from_text(text: str, document_hash: str = None) -> list[dict]:
    """
    NEW APPROACH:
    1. First detect how many risk instances are in the document by finding 'Risk X:' patterns
    2. Extract whatever fields are explicitly present in each risk instance block
    3. Use AI to fill ONLY the missing fields (not to create risk titles)
    """
    print(f"[STATS] parse_risk_instances_from_text() called with {len(text)} chars of text")
    
    # Step 1: Detect and parse risk instance blocks from document
    detected_risk_instances = detect_and_parse_risk_instance_blocks(text)
    
    if not detected_risk_instances:
        print(f"[WARNING]  No structured risk instances found. Falling back to old AI extraction...")
        # Fallback to old behavior if no "Risk X:" pattern found
        return fallback_risk_instance_extraction(text)
    
    print(f"[OK] Detected {len(detected_risk_instances)} risk instance(s), now processing each...")
    
    # Step 2: For each detected risk instance, normalize extracted fields and fill missing ones
    completed_risk_instances = []
    for idx, risk_data in enumerate(detected_risk_instances, 1):
        print(f"\n[PROC] Processing Risk Instance {idx}: {risk_data.get('RiskTitle', 'Unknown')[:50]}...")
        
        # Start with empty item for all DB fields
        item = {k: None for k in RISK_INSTANCE_DB_FIELDS}
        
        # Copy extracted fields
        for field in RISK_INSTANCE_DB_FIELDS:
            if field in risk_data and risk_data[field]:
                item[field] = risk_data[field]
        
        # Get the risk block for AI context
        risk_block = risk_data.get("_risk_block", "")
        extracted_fields = risk_data.get("_extracted_fields", [])
        
        print(f"  [INFO] Extracted fields: {', '.join(extracted_fields)}")
        
        # Initialize metadata structure
        if "_meta" not in item:
            item["_meta"] = {}
        if "per_field" not in item["_meta"]:
            item["_meta"]["per_field"] = {}
        
        # Mark extracted fields in metadata
        for field in extracted_fields:
            if field in item and item[field]:
                item["_meta"]["per_field"][field] = {
                    "source": "EXTRACTED",
                    "confidence": 0.95,
                    "rationale": "Found explicitly in document"
                }
        
        # Normalize extracted fields
        if item.get("RiskLikelihood"):
            item["RiskLikelihood"] = clamp_int(item["RiskLikelihood"], 1, 10)
        if item.get("RiskImpact"):
            item["RiskImpact"] = clamp_int(item["RiskImpact"], 1, 10)
        if item.get("RiskExposureRating"):
            item["RiskExposureRating"] = as_float_or_none(item["RiskExposureRating"])
        if item.get("Criticality"):
            item["Criticality"] = normalize_choice(item["Criticality"], CRITICALITY_CHOICES)
        if item.get("RiskPriority"):
            item["RiskPriority"] = normalize_choice(item["RiskPriority"], PRIORITY_CHOICES)
        if item.get("RiskType"):
            item["RiskType"] = normalize_choice(item["RiskType"], RISKTYPE_HINTS)
        if item.get("Origin"):
            item["Origin"] = normalize_choice(item["Origin"], ORIGIN_HINTS)
        if item.get("Appetite"):
            item["Appetite"] = normalize_choice(item["Appetite"], APPETITE_HINTS)
        if item.get("RiskResponseType"):
            item["RiskResponseType"] = normalize_choice(item["RiskResponseType"], RESPONSE_TYPE_HINTS)
        if item.get("RiskStatus"):
            item["RiskStatus"] = normalize_choice(item["RiskStatus"], RISK_STATUS_CHOICES)
        if item.get("MitigationStatus"):
            item["MitigationStatus"] = normalize_choice(item["MitigationStatus"], MITIGATION_STATUS_CHOICES)
        if item.get("RiskMultiplierX"):
            item["RiskMultiplierX"] = as_float_or_none(item["RiskMultiplierX"])
        if item.get("RiskMultiplierY"):
            item["RiskMultiplierY"] = as_float_or_none(item["RiskMultiplierY"])
        
        # Handle JSON fields
        if item.get("RiskMitigation"):
            mitigation = item["RiskMitigation"]
            if isinstance(mitigation, str):
                try:
                    item["RiskMitigation"] = json.loads(mitigation)
                except:
                    item["RiskMitigation"] = [{"step": "Step 1", "description": mitigation}]
            elif not isinstance(mitigation, list):
                item["RiskMitigation"] = None
        
        if item.get("ModifiedMitigations"):
            mod_mits = item["ModifiedMitigations"]
            if isinstance(mod_mits, str):
                try:
                    item["ModifiedMitigations"] = json.loads(mod_mits)
                except:
                    item["ModifiedMitigations"] = []
            elif not isinstance(mod_mits, list):
                item["ModifiedMitigations"] = None
        
        if item.get("RiskFormDetails"):
            form_details = item["RiskFormDetails"]
            if isinstance(form_details, str):
                try:
                    item["RiskFormDetails"] = json.loads(form_details)
                except:
                    item["RiskFormDetails"] = None
            elif not isinstance(form_details, dict):
                item["RiskFormDetails"] = None
        
        # Step 3: Use AI to fill ONLY missing fields (NEVER RiskTitle - must be in document)
        missing_fields = [f for f in RISK_INSTANCE_DB_FIELDS if item.get(f) in (None, "", []) and f != "RiskTitle"]
        if missing_fields:
            print(f"  [AI] Missing fields: {', '.join(missing_fields)}")
            print(f"  [AI] Using AI to infer missing fields...")

            for field in missing_fields:
                print(f"    [DEBUG] Inferring {field}...")
                value, metadata = infer_single_field(field, item, risk_block or text[:3000], document_hash=document_hash)
                item[field] = value
                # Store AI generation metadata
                item["_meta"]["per_field"][field] = metadata
                print(f"    [TAG]  Marked {field} as AI_GENERATED in metadata")
        else:
            print(f"  [OK] All fields extracted from document!")
        
        # Final normalization and defaults
        item["RiskLikelihood"] = item["RiskLikelihood"] or 5
        item["RiskImpact"] = item["RiskImpact"] or 5
        
        # Compute exposure if not present
        if not item.get("RiskExposureRating"):
            item["RiskExposureRating"] = compute_exposure(item["RiskLikelihood"], item["RiskImpact"]) or 25.0
        
        item["RiskExposureRating"] = float(max(0.0, min(100.0, item["RiskExposureRating"])))
        item["RiskMultiplierX"] = item["RiskMultiplierX"] or 1.0
        item["RiskMultiplierY"] = item["RiskMultiplierY"] or 1.0
        item["Criticality"] = item["Criticality"] or "Medium"
        item["RiskPriority"] = item["RiskPriority"] or "Medium"
        item["RiskType"] = item["RiskType"] or "Current"
        item["Origin"] = item["Origin"] or "Internal"
        item["Appetite"] = item["Appetite"] or "Medium"
        item["RiskResponseType"] = item["RiskResponseType"] or "Mitigate"
        item["RiskStatus"] = item["RiskStatus"] or "Not Assigned"
        item["MitigationStatus"] = item["MitigationStatus"] or "Pending"
        item["Reviewer"] = item["Reviewer"] or "Pending Review"
        
        # RiskTitle must ALWAYS come from document - never generate it
        if not item.get("RiskTitle"):
            raise ValueError(f"RiskTitle is missing for risk instance {idx}. All risk titles must be present in the document as 'Risk X: Title'.")
        
        # Ensure JSON fields have proper defaults
        if not isinstance(item.get("RiskMitigation"), list):
            item["RiskMitigation"] = [{"step": "Step 1", "description": "To be defined"}]
        if not isinstance(item.get("ModifiedMitigations"), list):
            item["ModifiedMitigations"] = []
        if not isinstance(item.get("RiskFormDetails"), dict):
            item["RiskFormDetails"] = {
                "assessment_method": "Qualitative",
                "data_sources": ["Document"],
                "assessment_date": date.today().isoformat()
            }
        
        # Debug: Print metadata summary
        ai_fields = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "AI_GENERATED"]
        extracted = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "EXTRACTED"]
        
        print(f"  [STATS] Metadata Summary:")
        print(f"     [AI] AI Generated: {len(ai_fields)} fields - {ai_fields}")
        print(f"     [DOC] Extracted: {len(extracted)} fields - {extracted}")
        
        completed_risk_instances.append(item)
        print(f"  [OK] Risk Instance {idx} completed!")
    
    return completed_risk_instances


# =========================
# DJANGO API ENDPOINTS
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@rate_limit_decorator(requests_per_minute=10, requests_per_hour=100)  # Phase 3: Rate limiting
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def upload_and_process_risk_instance_document(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    print(f"[UPLOAD] Upload request for risk instance document")
    print(f"[UPLOAD] Request data: {request.POST}")
    print(f"[UPLOAD] Request files: {request.FILES}")
    print(f"[UPLOAD] User ID: {request.POST.get('user_id', 'unknown')}")

    # CORS preflight support
    if request.method == 'OPTIONS':
        response = HttpResponse()
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        response['Access-Control-Max-Age'] = '86400'
        return response

    try:
        if 'file' not in request.FILES:
            resp = JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            resp = JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        # Create the ai_uploads/risk_instance directory if it doesn't exist
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'risk_instance')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save the file with timestamp to avoid conflicts
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        # Write the uploaded file to disk
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        # Decompress if needed (client-side compression)
        compression_metadata = None
        file_path, was_compressed, compression_stats = decompress_if_needed(file_path)
        if was_compressed:
            compression_metadata = compression_stats
            # Update extension after decompression (remove .gz)
            ext = os.path.splitext(file_path)[1].lower()
            print(f"[FILE] Decompressed file: {compression_stats['ratio']}% reduction, saved {compression_stats['bandwidth_saved_kb']} KB")
        
        # Upload to S3 for backup and cloud storage
        s3_url = None
        s3_key = None
        user_id = request.POST.get('user_id', '1')
        try:
            print(f"[EMOJI] Uploading file to S3...")
            s3_client = create_direct_mysql_client()
            connection_test = s3_client.test_connection()
            if connection_test.get('overall_success', False):
                # Generate unique filename for S3
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                s3_filename = f"risk_instance_{timestamp}_{os.path.basename(file_path)}"
                upload_result = s3_client.upload(
                    file_path=file_path,
                    user_id=user_id,
                    custom_file_name=s3_filename,
                    module='Risk'
                )
                if upload_result.get('success'):
                    s3_url = upload_result['file_info']['url']
                    s3_key = upload_result['file_info'].get('s3Key', '')
                    print(f"[OK] File uploaded to S3: {s3_url}")
                else:
                    print(f"[WARNING] S3 upload failed: {upload_result.get('error', 'Unknown error')}")
            else:
                print(f"[WARNING] S3 service unavailable, continuing with local file")
        except Exception as s3_error:
            print(f"[WARNING] S3 upload error (continuing with local file): {str(s3_error)}")
        
        print(f"[OK] File saved to: {file_path}")

        try:
            # Step 1: Extract text from the saved file
            print(f"[DEBUG] STEP 1: Starting text extraction from {ext} file...")
            raw_text = extract_text_from_file(file_path, ext)

            if not raw_text or len(raw_text.strip()) < 50:
                print(f"[ERROR] ERROR: Could not extract meaningful text. Length: {len(raw_text) if raw_text else 0}")
                resp = JsonResponse({'status': 'error', 'message': 'Could not extract meaningful text from document'}, status=400)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp

            print(f"[OK] STEP 1 COMPLETE: Extracted {len(raw_text)} characters from document")
            print(f"[DOC] First 200 chars: {raw_text[:200]}...")

            # Step 1B: Preprocess document (Phase 2 optimization)
            print(f"[DEBUG] STEP 1B: Preprocessing document (Phase 2 optimization for risk instance)...")
            text, preprocess_metadata = preprocess_document(raw_text, max_length=8000)
            print(f"[OK] STEP 1B COMPLETE: Preprocessed document")
            print(f"   Original length: {preprocess_metadata['original_length']} chars")
            print(f"   Processed length: {preprocess_metadata['processed_length']} chars")
            if preprocess_metadata['was_truncated']:
                print(f"   [WARNING]  Document was truncated ({preprocess_metadata['reduction_percent']:.1f}% reduction)")

            # Calculate document hash for caching / RAG (Phase 2/3)
            document_hash = calculate_document_hash(text)
            print(f"[INFO] Document hash (risk instance): {document_hash[:16]}...")

            # Step 2: Check AI provider configuration
            print(f"[DEBUG] STEP 2: Checking AI provider configuration for risk instance module...")
            if AI_PROVIDER == 'openai' and not OPENAI_API_KEY:
                print(f"[ERROR] ERROR: OPENAI_API_KEY is not set")
                resp = JsonResponse({
                    'status': 'error', 
                    'message': 'OPENAI_API_KEY environment variable is not set. Please configure your OpenAI API key or switch to Ollama.'
                }, status=503)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp
            elif AI_PROVIDER == 'ollama' and not OLLAMA_BASE_URL:
                print(f"[ERROR] ERROR: OLLAMA_BASE_URL is not set")
                resp = JsonResponse({
                    'status': 'error', 
                    'message': 'OLLAMA_BASE_URL environment variable is not set. Please configure your Ollama server URL.'
                }, status=503)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp
            
            print(f"[OK] STEP 2 COMPLETE: {AI_PROVIDER.upper()} provider is configured for risk instance module")
            
            # Step 3: Process with AI (Phase 2+3 optimizations)
            # Phase 3: Use intelligent model routing and queuing
            start_time = time.time()

            provider_info = f"{AI_PROVIDER.upper()} ({OPENAI_MODEL if AI_PROVIDER == 'openai' else OLLAMA_MODEL_DEFAULT})"
            print(f"[AI] STEP 3: Calling {provider_info} to extract risk instances (Phase 2+3: cached + few-shot + RAG + routing)...")

            def process_document():
                # MULTI-TENANCY: Extract tenant_id from request
                tenant_id = get_tenant_id_from_request(request)
                
                return parse_risk_instances_from_text(text, document_hash=document_hash)

            # Use queuing for heavy processing
            if len(text) > 10000:
                request_id = f"risk_instance_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(file_name)}"
                print(f"[LIST] Large risk instance document detected, using Phase 3 queuing...")
                risk_instances = process_with_queue(request_id, process_document)
            else:
                risk_instances = process_document()

            # Track processing time for system load monitoring (Phase 3)
            processing_time = time.time() - start_time
            track_system_load(processing_time, len(text))

            print(f"[OK] STEP 3 COMPLETE: AI extracted {len(risk_instances)} risk instance(s) from document")
            for idx, ri in enumerate(risk_instances, 1):
                print(f"  Risk Instance {idx}: {ri.get('RiskTitle', 'Untitled')[:50]}...")

            # Phase 3: Add document to RAG for future context retrieval
            if is_rag_available():
                try:
                    add_document_to_rag(
                        document_text=text,
                        document_id=f"risk_instance_doc_{document_hash[:16]}",
                        metadata={
                            "type": "risk_instance_assessment",
                            "filename": file_name,
                            "uploaded_at": datetime.now().isoformat(),
                            "num_risk_instances": len(risk_instances) if 'risk_instances' in locals() else 0,
                        },
                    )
                    print(f"[OK] Phase 3 RAG: Risk instance document added to knowledge base")
                except Exception as e:
                    print(f"[WARNING]  Phase 3 RAG (risk instance): Failed to add document: {e}")

            # Phase 3: Include RAG and routing stats in response
            phase3_metadata = {
                "rag_available": is_rag_available(),
                "rag_stats": get_rag_stats() if is_rag_available() else None,
                "system_load": get_current_system_load(),
                "processing_time": processing_time,
                "model_routing": "enabled",
            }

            response_data = {
                'status': 'success',
                'message': f'Successfully extracted {len(risk_instances)} risk instance(s)',
                'document_name': file_name,
                'saved_path': safe_filename,
                'extracted_text_length': len(text),
                'preprocessing_metadata': preprocess_metadata,
                'phase3_metadata': phase3_metadata,
                'risk_instances': risk_instances
            }
            
            # Include compression metadata if file was compressed
            if compression_metadata:
                response_data['compression_metadata'] = compression_metadata
            
            # Include S3 info if uploaded successfully
            if s3_url:
                response_data['s3_url'] = s3_url
                response_data['s3_key'] = s3_key
            
            resp = JsonResponse(response_data)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp
        except Exception as process_error:
            # Clean up the file if processing fails
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        resp = JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)
        resp['Access-Control-Allow-Origin'] = '*'
        return resp


@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def save_extracted_risk_instances(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        data = json.loads(request.body or "{}")
        risk_instances_data = data.get('risk_instances', [])
        if not risk_instances_data:
            return JsonResponse({'status': 'error', 'message': 'No risk instances provided'}, status=400)

        saved = []
        errors = []

        for idx, r in enumerate(risk_instances_data):
            try:
                # Parse RiskMitigation if it's a string
                risk_mitigation = r.get('RiskMitigation')
                if isinstance(risk_mitigation, str):
                    try:
                        risk_mitigation = json.loads(risk_mitigation)
                    except:
                        risk_mitigation = [{"step": "Step 1", "description": risk_mitigation}]
                elif not isinstance(risk_mitigation, list):
                    risk_mitigation = [{"step": "Step 1", "description": "To be defined"}]
                
                # Parse ModifiedMitigations if it's a string
                modified_mitigations = r.get('ModifiedMitigations', [])
                if isinstance(modified_mitigations, str):
                    try:
                        modified_mitigations = json.loads(modified_mitigations)
                    except:
                        modified_mitigations = []
                elif not isinstance(modified_mitigations, list):
                    modified_mitigations = []
                
                # Parse RiskFormDetails if it's a string
                risk_form_details = r.get('RiskFormDetails')
                if isinstance(risk_form_details, str):
                    try:
                        risk_form_details = json.loads(risk_form_details)
                    except:
                        risk_form_details = {
                            "assessment_method": "Qualitative",
                            "data_sources": ["Document"],
                            "assessment_date": date.today().isoformat()
                        }
                elif not isinstance(risk_form_details, dict):
                    risk_form_details = {
                        "assessment_method": "Qualitative",
                        "data_sources": ["Document"],
                        "assessment_date": date.today().isoformat()
                    }
                
                kwargs = {
                    # not setting: RiskInstanceId (auto), RiskId, IncidentId, ComplianceId, UserId, ReviewerId, ReportedBy
                    'RiskTitle': r.get('RiskTitle', f'Risk Instance {idx+1}'),
                    'RiskDescription': r.get('RiskDescription', ''),
                    'PossibleDamage': r.get('PossibleDamage', ''),
                    'RiskPriority': r.get('RiskPriority', 'Medium'),
                    'Criticality': r.get('Criticality', 'Medium'),
                    'Category': r.get('Category', ''),
                    'Origin': r.get('Origin', 'Internal'),
                    'RiskLikelihood': int(r.get('RiskLikelihood') or 5),
                    'RiskImpact': int(r.get('RiskImpact') or 5),
                    'RiskExposureRating': float(r.get('RiskExposureRating') or 25.0),
                    'RiskMultiplierX': float(r.get('RiskMultiplierX') or 1.0),
                    'RiskMultiplierY': float(r.get('RiskMultiplierY') or 1.0),
                    'Appetite': r.get('Appetite', 'Medium'),
                    'RiskResponseType': r.get('RiskResponseType', 'Mitigate'),
                    'RiskResponseDescription': r.get('RiskResponseDescription', ''),
                    'RiskMitigation': risk_mitigation,
                    'RiskType': r.get('RiskType', 'Current'),
                    'RiskOwner': r.get('RiskOwner', ''),
                    'BusinessImpact': r.get('BusinessImpact', ''),
                    'RiskStatus': r.get('RiskStatus', 'Not Assigned'),
                    'MitigationStatus': r.get('MitigationStatus', 'Pending'),
                    'ModifiedMitigations': modified_mitigations,
                    'RiskFormDetails': risk_form_details,
                    'Reviewer': r.get('Reviewer', 'Pending Review'),
                }
                risk_instance = RiskInstance.objects.create(**kwargs)
                saved.append({'risk_instance_id': getattr(risk_instance, "RiskInstanceId", None), 'risk_title': risk_instance.RiskTitle})
            except Exception as ex:
                errors.append({'risk_instance_index': idx, 'title': r.get('RiskTitle'), 'error': str(ex)})

        resp = {
            'status': 'success',
            'message': f'Saved {len(saved)} risk instance(s)' + (f' with {len(errors)} error(s)' if errors else ''),
            'saved': saved
        }
        if errors:
            resp['errors'] = errors
        return JsonResponse(resp)

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON payload'}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error saving risk instances: {str(e)}'}, status=500)


@api_view(['GET'])
@permission_classes([AllowAny])
@rbac_required(required_permission='view_all_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def test_openai_connection_risk_instance(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        if not OPENAI_API_KEY:
            return JsonResponse({
                'status': 'error',
                'message': 'OPENAI_API_KEY is not set',
                'model': OPENAI_MODEL,
                'api_url': OPENAI_API_URL
            }, status=500)
        
        out = call_openai_json('Return JSON: {"ok": true}')
        return JsonResponse({
            'status': 'success', 
            'openai_reply': out, 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error', 
            'message': f'OpenAI error: {e}', 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        }, status=500)

