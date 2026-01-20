"""
AI-Powered Incident Document Ingestion
- Reads PDF/DOCX/XLSX/TXT
- Extracts incident data using OpenAI GPT models
- Fills missing fields with focused prompts
- Returns normalized, DB-ready JSON for `incidents` table
"""

import os
import re
import json
import time
from datetime import date, datetime
from typing import Any, Optional

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser
import requests

# Phase 2 Optimizations
from ...utils.ai_cache import cached_llm_call
from ...utils.document_preprocessor import preprocess_document, calculate_document_hash
from ...utils.few_shot_prompts import get_field_extraction_prompt

# Phase 3 Optimizations
from ...utils.rag_system import (
    add_document_to_rag,
    retrieve_relevant_context,
    build_rag_prompt,
    is_rag_available,
    get_rag_stats
)
from ...utils.model_router import (
    route_model,
    track_system_load,
    get_current_system_load
)
from ...utils.request_queue import (
    rate_limit_decorator,
    process_with_queue,
    get_queue_status
)
from ...utils.file_compression import decompress_if_needed
from ...routes.Global.s3_fucntions import create_direct_mysql_client

# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import Incident

# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)


# =========================
# AI PROVIDER CONFIG (OpenAI or Ollama)
# =========================
# Reuse shared AI provider config from risk_ai_doc
from django.conf import settings
from ..Risk.risk_ai_doc import (
    AI_PROVIDER,
    call_ollama_json,
    call_openai_json,
    _calculate_optimal_context_size,
    _select_ollama_model_by_complexity,
    OLLAMA_BASE_URL,
    OLLAMA_MODEL_DEFAULT,
    OLLAMA_MODEL_FAST,
    OLLAMA_MODEL_COMPLEX,
    OPENAI_API_KEY,
    OPENAI_API_URL,
    OPENAI_MODEL,
)

print("\n🤖 Incident Import AI Provider Configuration:")
print(f"   Selected Provider: {AI_PROVIDER.upper()}")
if AI_PROVIDER == 'openai':
    if not OPENAI_API_KEY:
        print("[WARNING] OPENAI_API_KEY not found in Django settings!")
        print("   Please set OPENAI_API_KEY in your .env file")
    else:
        print(f"[INFO] Incident AI OpenAI Configuration:")
        print(f"   Model: {OPENAI_MODEL}")
        print(f"   API Key: {'*' * (len(OPENAI_API_KEY) - 4)}{OPENAI_API_KEY[-4:]}")
elif AI_PROVIDER == 'ollama':
    print(f"[INFO] Incident AI Ollama Configuration:")
    print(f"   URL: {OLLAMA_BASE_URL}")
    print(f"   Default Model: {OLLAMA_MODEL_DEFAULT}")
    print(f"   Fast Model: {OLLAMA_MODEL_FAST}")
    print(f"   Complex Model: {OLLAMA_MODEL_COMPLEX}")

# Fields we want to extract from incidents table (excluding ID and date fields)
INCIDENT_DB_FIELDS = [
    "IncidentTitle",
    "Description",
    "Mitigation",
    "Origin",
    "Comments",
    "RiskCategory",
    "IncidentCategory",
    "RiskPriority",
    "Attachments",
    "Status",
    "RepeatedNot",
    "CostOfIncident",
    "ReopenedNot",
    "RejectionSource",
    "AffectedBusinessUnit",
    "SystemsAssetsInvolved",
    "GeographicLocation",
    "Criticality",
    "InitialImpactAssessment",
    "InternalContacts",
    "ExternalPartiesInvolved",
    "RegulatoryBodies",
    "RelevantPoliciesProceduresViolated",
    "ControlFailures",
    "LessonsLearned",
    "IncidentClassification",
    "PossibleDamage",
    "IncidentFormDetails",
]

# Canonical choices / constraints to stabilize LLM outputs
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES = ["Low", "Medium", "High", "Critical"]
STATUS_CHOICES = ["New", "In Progress", "Under Investigation", "Resolved", "Closed", "Escalated", "Risk Mitigated"]
ORIGIN_CHOICES = ["AUDIT_FINDING", "MANUAL", "AUTOMATED", "EXTERNAL_REPORT", "INTERNAL_DETECTION"]
REJECTION_SOURCE_CHOICES = ["INCIDENT", "RISK"]
INCIDENT_CATEGORY_HINTS = [
    "Security Breach", "Data Loss", "System Outage", "Compliance Violation", 
    "Operational Failure", "Third-Party Issue", "Human Error", "Natural Disaster",
    "Cyber Attack", "Privacy Incident", "Safety Incident", "Financial Loss"
]
RISK_CATEGORY_HINTS = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]

# Field-specific prompts (optimized for OpenAI GPT models)
FIELD_PROMPTS = {
    "IncidentTitle": "Extract or generate a clear, concise incident title (max 255 characters). Format: '[Incident Type] - [Key Impact] - [Timeframe if available]'. Example: 'Data Breach - 10,000 Customer Records Exposed - Q3 2024'. Be specific and professional.",
    
    "Description": "Extract or generate a comprehensive incident description covering: (1) What happened, (2) When it was detected, (3) How it was discovered, (4) Affected systems/processes, (5) Immediate consequences. Write 3-5 factual sentences with specific details like timestamps and system names.",
    
    "Mitigation": "Extract or generate a JSON array of mitigation steps. Each step must have: 'step' (action description), 'status' (Completed/Planned/In Progress), 'responsible' (team/person), 'deadline' (YYYY-MM-DD format). Return array with at least 2-3 concrete actions. Example: [{\"step\": \"Isolated affected servers\", \"status\": \"Completed\", \"responsible\": \"IT Security Team\", \"deadline\": \"2024-01-15\"}]",
    
    "Origin": f"Classify the incident origin. Return EXACTLY ONE of these values: {', '.join(ORIGIN_CHOICES)}. Guidelines: AUDIT_FINDING (discovered during audit), MANUAL (person reported), AUTOMATED (system detected), EXTERNAL_REPORT (outside party), INTERNAL_DETECTION (internal monitoring). Choose the best match.",
    
    "Comments": "Extract or provide 2-3 sentences of additional context: unusual circumstances, related incidents, influencing factors, or observations that add value beyond the main description. Be concise and informative.",
    
    "RiskCategory": f"Select the PRIMARY risk category from: {', '.join(RISK_CATEGORY_HINTS)}. Choose the most significant risk domain based on the core nature of the incident. Return the exact category name as listed.",
    
    "IncidentCategory": f"Choose the incident category from: {', '.join(INCIDENT_CATEGORY_HINTS)}. This reflects the TYPE of incident event (e.g., 'Cyber Attack' for phishing, even if it creates operational risk). Return the exact category name.",
    
    "RiskPriority": f"Assess priority level. Return EXACTLY ONE of: {', '.join(PRIORITY_CHOICES)}. Criteria: Critical (immediate threat to operations/safety), High (significant impact, prompt action needed), Medium (notable impact, manageable timeline), Low (minor impact, minimal urgency).",
    
    "Attachments": "Extract file names or document references mentioned in the incident. Return as semicolon-separated list (e.g., 'report.pdf;evidence.xlsx;screenshot.png'). If none mentioned, return empty string ''.",
    
    "Status": f"Determine current incident status. Return EXACTLY ONE of: {', '.join(STATUS_CHOICES)}. Guidelines: New (just reported), In Progress (being worked), Under Investigation (analyzing cause), Resolved (fixed not closed), Closed (fully completed), Escalated (elevated to higher authority), Risk Mitigated (risk addressed).",
    
    "RepeatedNot": "Determine if this is a recurring incident. Return boolean true if document mentions: 'recurring', 'happened before', 'previous occurrence', 'similar to incident X'. Return false if first-time or no indication of recurrence.",
    
    "CostOfIncident": "Extract or estimate financial cost/impact. Format: '$50,000', '€25K', or 'Estimated $100K-$150K'. If specific cost mentioned, use it. If impact described but no cost, provide reasonable estimate. If no financial info, return 'Not assessed'.",
    
    "ReopenedNot": "Determine if incident was reopened. Return boolean true ONLY if document explicitly states it was previously closed then reopened. Keywords: 'reopened', 'recurred after closure'. Return false for new incidents or ongoing ones.",
    
    "RejectionSource": f"Identify rejection/escalation source. Return EXACTLY ONE of: {', '.join(REJECTION_SOURCE_CHOICES)}, or null. Use 'INCIDENT' if rejected from incident workflow, 'RISK' if escalated from risk assessment. Return null if not applicable.",
    
    "AffectedBusinessUnit": "Extract specific business units/departments impacted. Be precise with actual names: 'Customer Service - EMEA Region', 'IT Infrastructure', 'Finance - Accounts Payable'. Multiple units: comma-separated. Organization-wide: 'Enterprise-Wide'. Unknown: 'To be determined'.",
    
    "SystemsAssetsInvolved": "List specific systems, applications, or infrastructure affected. Include technical details: version numbers, hostnames, identifiers. Example: 'SAP ERP Production (sap-prod-01), Customer DB v3.2, Payment Gateway API'. Comma-separated. Unknown: 'To be determined'.",
    
    "GeographicLocation": "Specify physical/logical location of incident. Include: country, region, city, data center, or office. Examples: 'London Office - UK', 'AWS US-East-1', 'Global - Multiple Regions'. Be as specific as possible.",
    
    "Criticality": f"Assess criticality level. Return EXACTLY ONE of: {', '.join(CRITICALITY_CHOICES)}. Criteria: Critical (threatens core business/safety), High (significant operational/financial/reputational impact), Medium (moderate impact with workarounds), Low (minimal impact).",
    
    "InitialImpactAssessment": "Provide structured initial assessment (3-4 sentences) covering: (1) immediate operational impact, (2) affected stakeholders/customers with numbers, (3) data/system integrity concerns, (4) preliminary scope. Be factual and quantitative where possible.",
    
    "InternalContacts": "List key internal personnel involved/notified. Format: 'John Smith (IT Manager), Jane Doe (CISO), Security Operations Team'. If names unavailable, list roles/departments. Comma-separated.",
    
    "ExternalPartiesInvolved": "Identify external organizations/vendors/partners involved in incident or response. Examples: 'Microsoft Support', 'Acme Cloud Services', 'External Auditors', 'Law Enforcement'. Include their role if mentioned. Empty string if purely internal.",
    
    "RegulatoryBodies": "List regulatory authorities/compliance bodies/government agencies requiring notification or involvement. Examples: 'SEC', 'GDPR DPA', 'FDA', 'PCI DSS Council'. Include notification requirements if mentioned. Empty string if none required.",
    
    "RelevantPoliciesProceduresViolated": "Identify specific policies/procedures/standards violated. Include policy names/numbers: 'Password Policy v2.3 (Section 4.2)', 'Change Management Procedure violation'. Be specific to identify systemic issues.",
    
    "ControlFailures": "Describe failed security controls/safeguards (2-3 sentences). Be technical and specific: 'Multi-factor authentication bypass', 'Firewall rule misconfiguration', 'Failed backup verification'. Explain why controls didn't prevent the incident.",
    
    "LessonsLearned": "Summarize key insights and learnings (2-4 sentences). What could be done differently? What worked well? What process improvements needed? Focus on actionable takeaways for future prevention.",
    
    "IncidentClassification": "Extract classification code or generate severity-based classification. Examples: 'CAT-1: Critical Security Incident', 'P1: Production Outage', 'Type A: Data Breach'. Format: 'Category: Description'. If not mentioned, infer appropriate classification.",
    
    "PossibleDamage": "Describe all potential damages (2-3 sentences): operational, financial, reputational, legal, compliance impacts. Include realized and avoided consequences with quantitative estimates. Example: 'Service outage 50K users 4hrs, revenue loss $200K, potential fines $500K, brand damage'.",
    
    "IncidentFormDetails": "Generate JSON object with incident details. Required keys: 'reported_by' (name/role), 'detection_method' (how discovered), 'response_time_minutes' (number), 'escalation_level' (L1/L2/L3), 'containment_status' (Contained/Not Contained), 'root_cause_category', 'affected_records_count' (number), 'recovery_time_objective' (duration). Fill based on context.",
}

# Strict JSON schema block
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON array. No markdown, no code blocks, no explanations.
Start with [ and end with ]. Use proper JSON syntax with double quotes.

Example structure (return array of incidents like this):
[
  {{
    "IncidentTitle": "Brief incident title here",
    "Description": "Detailed description",
    "Mitigation": [{{"step": "Action taken", "status": "Completed", "responsible": "Team", "deadline": "2024-01-01"}}],
    "Origin": "MANUAL",
    "Comments": "Additional context",
    "RiskCategory": "Operational",
    "IncidentCategory": "Security Breach",
    "RiskPriority": "High",
    "Attachments": "",
    "Status": "In Progress",
    "RepeatedNot": false,
    "CostOfIncident": "$50,000",
    "ReopenedNot": false,
    "RejectionSource": null,
    "AffectedBusinessUnit": "IT Department",
    "SystemsAssetsInvolved": "Production Server, Database",
    "GeographicLocation": "New York Office",
    "Criticality": "High",
    "InitialImpactAssessment": "Initial assessment details",
    "InternalContacts": "John Smith (IT Manager)",
    "ExternalPartiesInvolved": "",
    "RegulatoryBodies": "",
    "RelevantPoliciesProceduresViolated": "Security Policy Section 4.2",
    "ControlFailures": "Firewall misconfiguration",
    "LessonsLearned": "Key insights from incident",
    "IncidentClassification": "P1: Critical",
    "PossibleDamage": "Service outage, revenue loss",
    "IncidentFormDetails": {{"reported_by": "Security Team", "detection_method": "Automated", "response_time_minutes": 30}},
    "_meta": {{
      "per_field": {{
        "IncidentTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found in document"}},
        "Description": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from context"}}
      }}
    }}
  }}
]

Rules:
- Criticality must be: Low, Medium, High, or Critical
- RiskPriority must be: Low, Medium, High, or Critical
- Status must be one of: {STATUS_CHOICES}
- Origin must be one of: {ORIGIN_CHOICES}
- RepeatedNot and ReopenedNot must be boolean (true/false)
- Mitigation must be JSON array
- IncidentFormDetails must be JSON object
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON array, nothing else
"""


# =========================
# UTILITIES / VALIDATORS
# =========================
def _json_from_llm_text(text: str) -> Any:
    """Extract the first valid JSON array/object from the LLM response."""
    # Remove markdown code blocks
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    
    # Try to find JSON structure
    m = re.search(r"(\[.*\]|\{.*\})", text, flags=re.S)
    block = m.group(1) if m else text.strip()
    
    # Remove trailing commas (common JSON error)
    block = re.sub(r',(\s*[}\]])', r'\1', block)
    
    return json.loads(block)


# Note: call_ollama_json and call_openai_json are now imported from risk_ai_doc.py
# They include Phase 2 caching and Phase 3 optimizations


def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None


def as_boolean(v) -> bool:
    """Convert various inputs to boolean."""
    if isinstance(v, bool):
        return v
    if isinstance(v, str):
        return v.lower() in ('true', 'yes', '1', 't', 'y')
    return bool(v)


# =========================
# FILE EXTRACTORS
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        print(f"[PDF] Extraction error: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX] Extraction error: {e}")
        return ""


def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        print(f"[XLSX] Extraction error: {e}")
        return ""


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""


# =========================
# AI EXTRACTION CORE
# =========================
def infer_single_field(field_name: str, current_record: dict, document_context: str, 
                       document_hash: str = None) -> Any:
    """
    Focused prompt for ONE field using AI (supports both OpenAI and Ollama).
    Uses Phase 2 few-shot prompts and caching, Phase 3 RAG context.
    """
    provider_name = AI_PROVIDER.upper()
    print(f"🤖 AI PREDICTING FIELD: {field_name} (using {provider_name} with Phase 2+3 optimizations)")
    
    # Optimize context for Ollama
    if AI_PROVIDER == 'ollama':
        context_size = _calculate_optimal_context_size(len(document_context), "simple")
        optimized_context = document_context[:context_size] if len(document_context) > context_size else document_context
    else:
        optimized_context = document_context[:3000]  # OpenAI default
    
    # Phase 3: Try to retrieve relevant context from RAG
    rag_context = None
    if is_rag_available():
        try:
            query = f"What is the {field_name} for this incident?"
            retrieved = retrieve_relevant_context(query, n_results=3)
            if retrieved:
                rag_context = retrieved
                print(f"   📚 Phase 3 RAG: Retrieved {len(retrieved)} relevant document chunks")
        except Exception as e:
            print(f"   ⚠️  RAG retrieval failed: {e}")
    
    # Use few-shot prompt template (Phase 2 optimization)
    try:
        mini = get_field_extraction_prompt(
            field_name=field_name,
            document_text=optimized_context,
            field_prompts=FIELD_PROMPTS
        )
        # Add current record context
        filled_fields = {k: current_record.get(k) for k in INCIDENT_DB_FIELDS if current_record.get(k)}
        mini += f"\n\nALREADY EXTRACTED FIELDS:\n{json.dumps(filled_fields, indent=2)}"
        print(f"   📚 Using few-shot prompt template for {field_name}")
    except Exception as e:
        print(f"   ⚠️  Few-shot prompt failed, using basic prompt: {e}")
        # Fallback to basic prompt
        guidance = FIELD_PROMPTS.get(field_name, "Return a concise, professional value.")
        filled_fields = {k: current_record.get(k) for k in INCIDENT_DB_FIELDS if current_record.get(k)}
        mini = f"""Analyze the incident document and extract ONLY the "{field_name}" field.

DOCUMENT CONTEXT (first 3000 chars):
\"\"\"{optimized_context}\"\"\"

ALREADY EXTRACTED FIELDS:
{json.dumps(filled_fields, indent=2)}

INSTRUCTIONS FOR "{field_name}":
{guidance}

REQUIRED OUTPUT FORMAT:
Return ONLY a JSON object in this exact format:
{{
  "value": <extracted or inferred value>,
  "confidence": <number between 0.0 and 1.0>
}}

Rules:
1. If the field is explicitly mentioned in the document, extract it (confidence 0.8-1.0)
2. If you must infer based on context, do so (confidence 0.5-0.7)
3. If you cannot determine, return {{"value": null, "confidence": 0.0}}
4. Return ONLY the JSON object, no other text
"""
    
    # Phase 3: Enhance prompt with RAG context if available
    if rag_context:
        mini = build_rag_prompt(
            user_query=mini,
            retrieved_context=rag_context,
            base_prompt=None
        )
    
    try:
        if AI_PROVIDER == 'ollama':
            print(f"   📤 Sending prompt to Ollama for {field_name}...")
            model = _select_ollama_model_by_complexity(len(optimized_context), 1)
            out = call_ollama_json(mini, model=model, document_hash=document_hash)
        else:
            print(f"   📤 Sending prompt to OpenAI for {field_name}...")
            out = call_openai_json(mini, document_hash=document_hash)
        
        v = out.get("value") if isinstance(out, dict) else None
        confidence = out.get("confidence", 0.0) if isinstance(out, dict) else 0.0
        print(f"   ✅ AI PREDICTED {field_name}: '{v}' (confidence: {confidence:.2f})")
    except Exception as e:
        print(f"   ❌ AI FAILED to predict {field_name}: {str(e)}")
        v = None

    # Normalize after inference
    if field_name in ("RepeatedNot", "ReopenedNot"):
        return as_boolean(v)
    if field_name == "Criticality":
        return normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
    if field_name == "RiskPriority":
        return normalize_choice(v, PRIORITY_CHOICES) or "Medium"
    if field_name == "Status":
        return normalize_choice(v, STATUS_CHOICES) or "New"
    if field_name == "Origin":
        return normalize_choice(v, ORIGIN_CHOICES) or "MANUAL"
    if field_name == "RejectionSource":
        return normalize_choice(v, REJECTION_SOURCE_CHOICES) if v else None
    if field_name in ("Mitigation", "IncidentFormDetails"):
        # Try to parse as JSON
        if isinstance(v, (dict, list)):
            return v
        if isinstance(v, str):
            try:
                return json.loads(v)
            except:
                return None
        return None
    if isinstance(v, str):
        return v.strip() or None
    return v


def fallback_incident_extraction(text: str) -> list[dict]:
    """Minimal pattern-based fallback when AI fails completely."""
    incidents = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    incident_keywords = ["incident", "breach", "outage", "failure", "violation", "attack"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in incident_keywords):
            if current:
                incidents.append(current)
            count += 1
            current = {
                "IncidentTitle": f"Incident {count}: {ln[:100]}",
                "Description": ln,
                "Mitigation": [],
                "Origin": "MANUAL",
                "Comments": "",
                "RiskCategory": "Operational",
                "IncidentCategory": "Operational Failure",
                "RiskPriority": "Medium",
                "Attachments": "",
                "Status": "New",
                "RepeatedNot": False,
                "CostOfIncident": "Not assessed",
                "ReopenedNot": False,
                "RejectionSource": None,
                "AffectedBusinessUnit": "Unknown",
                "SystemsAssetsInvolved": "Unknown",
                "GeographicLocation": "Unknown",
                "Criticality": "Medium",
                "InitialImpactAssessment": "Requires investigation",
                "InternalContacts": "",
                "ExternalPartiesInvolved": "",
                "RegulatoryBodies": "",
                "RelevantPoliciesProceduresViolated": "",
                "ControlFailures": "",
                "LessonsLearned": "",
                "IncidentClassification": "Standard",
                "PossibleDamage": "To be assessed",
                "IncidentFormDetails": {},
            }

    if current:
        incidents.append(current)

    if not incidents:
        incidents.append({
            "IncidentTitle": "Document Incident Analysis",
            "Description": f"Automated incident derived from document (length={len(text)} chars).",
            "Mitigation": [],
            "Origin": "MANUAL",
            "Comments": "Extracted from uploaded document",
            "RiskCategory": "Operational",
            "IncidentCategory": "Operational Failure",
            "RiskPriority": "Medium",
            "Attachments": "",
            "Status": "New",
            "RepeatedNot": False,
            "CostOfIncident": "Not assessed",
            "ReopenedNot": False,
            "RejectionSource": None,
            "AffectedBusinessUnit": "To be determined",
            "SystemsAssetsInvolved": "To be determined",
            "GeographicLocation": "To be determined",
            "Criticality": "Medium",
            "InitialImpactAssessment": "Requires detailed assessment",
            "InternalContacts": "",
            "ExternalPartiesInvolved": "",
            "RegulatoryBodies": "",
            "RelevantPoliciesProceduresViolated": "",
            "ControlFailures": "",
            "LessonsLearned": "",
            "IncidentClassification": "Standard",
            "PossibleDamage": "To be assessed through investigation",
            "IncidentFormDetails": {},
        })

    return incidents


def parse_incidents_from_text(text: str, document_hash: str = None) -> list[dict]:
    """
    Extract ALL incident fields using AI with strict JSON schema.
    Phase 2: Uses document preprocessing and caching.
    Phase 3: Uses RAG context retrieval and model routing.
    """
    print(f"📊 parse_incidents_from_text() called with {len(text)} chars of text")
    
    # Phase 3: Try to retrieve relevant context from RAG
    rag_context = None
    if is_rag_available():
        try:
            retrieved = retrieve_relevant_context(text[:1000], n_results=3)  # Use first 1000 chars as query
            if retrieved:
                rag_context = retrieved
                print(f"   📚 Phase 3 RAG: Retrieved {len(retrieved)} relevant document chunks for incident extraction")
        except Exception as e:
            print(f"   ⚠️  RAG retrieval failed: {e}")
    
    # Optimize context size for Ollama
    if AI_PROVIDER == 'ollama':
        context_size = _calculate_optimal_context_size(len(text), "complex")
        optimized_text = text[:context_size] if len(text) > context_size else text
    else:
        optimized_text = text[:8000]  # OpenAI default
    
    # Improved prompt for AI
    base_prompt = f"""You are a GRC (Governance, Risk, Compliance) incident analyst. Analyze the following document and extract ALL incidents mentioned.

DOCUMENT TO ANALYZE:
\"\"\"{optimized_text}\"\"\"

EXTRACTION REQUIREMENTS:

1. IDENTIFY ALL INCIDENTS in the document (usually 1, sometimes multiple)

2. FOR EACH INCIDENT, extract these fields:
   - IncidentTitle: Clear title (max 255 chars)
   - Description: Comprehensive description (3-5 sentences)
   - Mitigation: JSON array of mitigation steps [{{"step": "...", "status": "...", "responsible": "...", "deadline": "YYYY-MM-DD"}}]
   - Origin: EXACTLY ONE OF: {', '.join(ORIGIN_CHOICES)}
   - Comments: Additional context (2-3 sentences)
   - RiskCategory: ONE OF: {', '.join(RISK_CATEGORY_HINTS)}
   - IncidentCategory: ONE OF: {', '.join(INCIDENT_CATEGORY_HINTS)}
   - RiskPriority: EXACTLY ONE OF: {', '.join(PRIORITY_CHOICES)}
   - Attachments: Semicolon-separated file names (or empty string)
   - Status: EXACTLY ONE OF: {', '.join(STATUS_CHOICES)}
   - RepeatedNot: boolean (true/false)
   - CostOfIncident: String like "$50,000" or "Not assessed"
   - ReopenedNot: boolean (true/false)
   - RejectionSource: {', '.join(REJECTION_SOURCE_CHOICES)} or null
   - AffectedBusinessUnit: Specific business units affected
   - SystemsAssetsInvolved: Systems/applications involved
   - GeographicLocation: Physical/logical location
   - Criticality: EXACTLY ONE OF: {', '.join(CRITICALITY_CHOICES)}
   - InitialImpactAssessment: Initial assessment (3-4 sentences)
   - InternalContacts: Internal personnel involved
   - ExternalPartiesInvolved: External organizations (or empty string)
   - RegulatoryBodies: Regulatory authorities (or empty string)
   - RelevantPoliciesProceduresViolated: Policies violated
   - ControlFailures: Security control failures (2-3 sentences)
   - LessonsLearned: Key learnings (2-4 sentences)
   - IncidentClassification: Classification code
   - PossibleDamage: Potential damages (2-3 sentences)
   - IncidentFormDetails: JSON object with keys: reported_by, detection_method, response_time_minutes, escalation_level, containment_status, root_cause_category, affected_records_count, recovery_time_objective

3. ADD METADATA for each field:
   - source: "EXTRACTED" (if explicitly in document) or "AI_GENERATED" (if inferred)
   - confidence: 0.0-1.0
   - rationale: Brief explanation

4. OUTPUT FORMAT:
Return a JSON object with an "incidents" key containing an array:
{{
  "incidents": [
    {{
      "IncidentTitle": "...",
      "Description": "...",
      ... (all fields above) ...,
      "_meta": {{
        "per_field": {{
          "IncidentTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found in document header"}},
          "Description": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from context"}}
        }}
      }}
    }}
  ]
}}

CRITICAL RULES:
- Return ONLY valid JSON, no markdown, no code blocks, no explanations
- Use double quotes for all strings
- Boolean values must be true/false (lowercase)
- No trailing commas
- Ensure all choice fields match EXACTLY the allowed values
- If a field cannot be determined, use reasonable defaults or empty strings
- For JSON fields (Mitigation, IncidentFormDetails), ensure proper nested structure

Begin analysis now and return the JSON object:"""
    
    # Phase 3: Enhance prompt with RAG context if available
    if rag_context:
        prompt = build_rag_prompt(
            user_query=base_prompt,
            retrieved_context=rag_context,
            base_prompt=None
        )
    else:
        prompt = base_prompt

    print(f"📝 Generated prompt for AI (length: {len(prompt)} chars)")
    
    try:
        provider_info = f"{AI_PROVIDER.upper()} ({OPENAI_MODEL if AI_PROVIDER == 'openai' else OLLAMA_MODEL_DEFAULT})"
        print(f"🚀 Calling {provider_info} to extract incidents (Phase 2+3: cached + few-shot + RAG + routing)...")
        print(f"📊 Processing document with {len(text)} characters")
        
        # Use shared AI functions with caching and routing
        if AI_PROVIDER == 'ollama':
            model = route_model(
                task_type="incident_extraction",
                text_length=len(text),
                num_risks=1,  # Estimate 1 incident per document
                accuracy_required="high",
                system_load=get_current_system_load(),
                provider="ollama"
            )
            response = call_ollama_json(prompt, model=model, document_hash=document_hash)
        else:
            response = call_openai_json(prompt, document_hash=document_hash)
        
        # Handle different response formats
        if isinstance(response, dict) and "incidents" in response:
            incidents = response["incidents"]
        elif isinstance(response, list):
            incidents = response
        else:
            raise ValueError("Unexpected response format from OpenAI")
        
        if not isinstance(incidents, list):
            raise ValueError("Incidents must be a JSON array")
        
        print(f"✅ OpenAI returned {len(incidents)} incident(s)")

        cleaned = []
        for idx, inc in enumerate(incidents, 1):
            print(f"📋 Processing incident {idx}/{len(incidents)}")
            item = {k: inc.get(k) for k in INCIDENT_DB_FIELDS}

            # Normalize all fields
            item["IncidentTitle"] = (item.get("IncidentTitle") or "").strip() or "Untitled Incident"
            item["Description"] = (item.get("Description") or "").strip() or None
            item["Comments"] = (item.get("Comments") or "").strip() or None
            item["Attachments"] = (item.get("Attachments") or "").strip() or ""
            
            item["Criticality"] = normalize_choice(item.get("Criticality"), CRITICALITY_CHOICES) or "Medium"
            item["RiskPriority"] = normalize_choice(item.get("RiskPriority"), PRIORITY_CHOICES) or "Medium"
            item["Status"] = normalize_choice(item.get("Status"), STATUS_CHOICES) or "New"
            item["Origin"] = normalize_choice(item.get("Origin"), ORIGIN_CHOICES) or "MANUAL"
            item["RejectionSource"] = normalize_choice(item.get("RejectionSource"), REJECTION_SOURCE_CHOICES) if item.get("RejectionSource") else None
            
            item["RepeatedNot"] = as_boolean(item.get("RepeatedNot"))
            item["ReopenedNot"] = as_boolean(item.get("ReopenedNot"))
            
            # Handle JSON fields
            mitigation = item.get("Mitigation")
            if isinstance(mitigation, str):
                try:
                    item["Mitigation"] = json.loads(mitigation)
                except:
                    item["Mitigation"] = []
            elif not isinstance(mitigation, list):
                item["Mitigation"] = []
            
            form_details = item.get("IncidentFormDetails")
            if isinstance(form_details, str):
                try:
                    item["IncidentFormDetails"] = json.loads(form_details)
                except:
                    item["IncidentFormDetails"] = {}
            elif not isinstance(form_details, dict):
                item["IncidentFormDetails"] = {}

            # Keep meta if present
            meta = inc.get("_meta") or {}
            item["_meta"] = meta

            # Fill any remaining missing fields
            print(f"🔍 Checking missing fields for incident: {item.get('IncidentTitle', 'Untitled')}")
            missing_fields = []
            for field in INCIDENT_DB_FIELDS:
                if item.get(field) in (None, "", []):
                    missing_fields.append(field)
            
            if missing_fields:
                print(f"   📝 Missing fields to predict: {missing_fields}")
                for field in missing_fields:
                    predicted_value = infer_single_field(field, item, text, document_hash=document_hash)
                    item[field] = predicted_value
                    # Mark as AI generated in metadata
                    if predicted_value is not None and predicted_value != "":
                        if "_meta" not in item:
                            item["_meta"] = {}
                        if "per_field" not in item["_meta"]:
                            item["_meta"]["per_field"] = {}
                        item["_meta"]["per_field"][field] = {
                            "source": "AI_GENERATED",
                            "confidence": 0.7,  # Default confidence for AI predictions
                            "rationale": f"AI predicted this value based on document context"
                        }
                        print(f"   🏷️  Marked {field} as AI_GENERATED in metadata")
            else:
                print(f"   ✅ All fields already populated")

            # Debug: Print metadata structure
            if "_meta" in item and "per_field" in item["_meta"]:
                ai_fields = [field for field, info in item["_meta"]["per_field"].items() 
                            if info.get("source") == "AI_GENERATED"]
                if ai_fields:
                    print(f"   🤖 AI Generated fields: {ai_fields}")
                else:
                    print(f"   📄 No AI generated fields for this incident")
            else:
                print(f"   📄 No metadata available for this incident")
            
            cleaned.append(item)

        return cleaned

    except Exception as e:
        print(f"AI extraction failed, using fallback extractor: {e}")
        base = fallback_incident_extraction(text)
        completed = []
        for inc in base:
            item = {k: inc.get(k) for k in INCIDENT_DB_FIELDS}
            # Ensure everything present
            for field in INCIDENT_DB_FIELDS:
                if item.get(field) in (None, "", []):
                    item[field] = infer_single_field(field, item, text, document_hash=document_hash)
            # Normalize again
            item["Criticality"] = normalize_choice(item.get("Criticality"), CRITICALITY_CHOICES) or "Medium"
            item["RiskPriority"] = normalize_choice(item.get("RiskPriority"), PRIORITY_CHOICES) or "Medium"
            item["Status"] = normalize_choice(item.get("Status"), STATUS_CHOICES) or "New"
            item["Origin"] = normalize_choice(item.get("Origin"), ORIGIN_CHOICES) or "MANUAL"
            item["RepeatedNot"] = as_boolean(item.get("RepeatedNot"))
            item["ReopenedNot"] = as_boolean(item.get("ReopenedNot"))
            completed.append(item)
        return completed


# =========================
# DJANGO API ENDPOINTS
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rate_limit_decorator(requests_per_minute=10, requests_per_hour=100)  # Phase 3: Rate limiting
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def upload_and_process_incident_document(request):
    """
    Upload and process incident document
    MULTI-TENANCY: Extracted incidents will be automatically assigned to user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    """
    Upload a document and process it to extract incident data.
    Phase 2: Uses document preprocessing, few-shot prompts, and caching.
    Phase 3: Uses RAG context retrieval, model routing, request queuing, and system load tracking.
    """
    print(f"📤 Upload request for incident document")

    try:
        if 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            return JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)

        # Create upload directory
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'incident')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        # Decompress if needed (client-side compression)
        compression_metadata = None
        file_path, was_compressed, compression_stats = decompress_if_needed(file_path)
        if was_compressed:
            compression_metadata = compression_stats
            # Update extension after decompression (remove .gz)
            ext = os.path.splitext(file_path)[1].lower()
            print(f"📦 Decompressed file: {compression_stats['ratio']}% reduction, saved {compression_stats['bandwidth_saved_kb']} KB")
        
        # Upload to S3 for backup and cloud storage
        s3_url = None
        s3_key = None
        user_id = request.POST.get('user_id', '1')
        try:
            print(f"☁️ Uploading file to S3...")
            s3_client = create_direct_mysql_client()
            connection_test = s3_client.test_connection()
            if connection_test.get('overall_success', False):
                # Generate unique filename for S3
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                s3_filename = f"incident_ai_{timestamp}_{os.path.basename(file_path)}"
                upload_result = s3_client.upload(
                    file_path=file_path,
                    user_id=user_id,
                    custom_file_name=s3_filename,
                    module='Incident'
                )
                if upload_result.get('success'):
                    s3_url = upload_result['file_info']['url']
                    s3_key = upload_result['file_info'].get('s3Key', '')
                    print(f"✅ File uploaded to S3: {s3_url}")
                else:
                    print(f"⚠️ S3 upload failed: {upload_result.get('error', 'Unknown error')}")
            else:
                print(f"⚠️ S3 service unavailable, continuing with local file")
        except Exception as s3_error:
            print(f"⚠️ S3 upload error (continuing with local file): {str(s3_error)}")
        
        print(f"✅ File saved to: {file_path}")

        try:
            # Step 1: Extract text
            print(f"🔍 STEP 1: Starting text extraction from {ext} file...")
            raw_text = extract_text_from_file(file_path, ext)
            
            if not raw_text or len(raw_text.strip()) < 50:
                return JsonResponse({'status': 'error', 'message': 'Could not extract meaningful text from document'}, status=400)

            print(f"✅ STEP 1A COMPLETE: Extracted {len(raw_text)} characters from document")
            
            # Step 1B: Preprocess document (Phase 2 optimization)
            print(f"🔍 STEP 1B: Preprocessing document (Phase 2 optimization)...")
            text, preprocess_metadata = preprocess_document(raw_text, max_length=8000)
            print(f"✅ STEP 1B COMPLETE: Preprocessed document")
            print(f"   Original length: {preprocess_metadata['original_length']} chars")
            print(f"   Processed length: {preprocess_metadata['processed_length']} chars")
            if preprocess_metadata['was_truncated']:
                print(f"   ⚠️  Document was truncated ({preprocess_metadata['reduction_percent']:.1f}% reduction)")
            
            # Calculate document hash for caching (Phase 2)
            document_hash = calculate_document_hash(text)
            print(f"📝 Document hash: {document_hash[:16]}... (for caching)")
            
            # Step 2: Check AI provider configuration
            print(f"🔍 STEP 2: Checking AI provider configuration...")
            if AI_PROVIDER == 'openai' and not OPENAI_API_KEY:
                print(f"❌ ERROR: OPENAI_API_KEY is not set")
                return JsonResponse({
                    'status': 'error', 
                    'message': 'OPENAI_API_KEY environment variable is not set. Please configure your OpenAI API key or switch to Ollama.'
                }, status=503)
            elif AI_PROVIDER == 'ollama' and not OLLAMA_BASE_URL:
                print(f"❌ ERROR: OLLAMA_BASE_URL is not set")
                return JsonResponse({
                    'status': 'error', 
                    'message': 'OLLAMA_BASE_URL environment variable is not set. Please configure your Ollama server URL.'
                }, status=503)
            
            print(f"✅ STEP 2 COMPLETE: {AI_PROVIDER.upper()} provider is configured")
            
            # Step 3: Process with AI (Phase 2+3 optimizations)
            start_time = time.time()
            
            provider_info = f"{AI_PROVIDER.upper()} ({OPENAI_MODEL if AI_PROVIDER == 'openai' else OLLAMA_MODEL_DEFAULT})"
            print(f"🤖 STEP 3: Calling {provider_info} to extract incidents (Phase 2+3: cached + few-shot + RAG + routing)...")
            
            # Phase 3: Process with queuing (if needed)
            request_id = f"incident_doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hash(file_name)}"
            
            def process_document():
                return parse_incidents_from_text(text, document_hash=document_hash)
            
            # Use queuing for heavy processing
            if len(text) > 10000:  # Large documents use queue
                print(f"📋 Large document detected, using Phase 3 queuing...")
                incidents = process_with_queue(request_id, process_document)
            else:
                incidents = process_document()
            
            # Track processing time for system load monitoring (Phase 3)
            processing_time = time.time() - start_time
            track_system_load(processing_time, len(text))
            
            # Phase 3: Add document to RAG for future context retrieval
            if is_rag_available():
                try:
                    add_document_to_rag(
                        document_text=text,
                        document_id=f"incident_doc_{document_hash[:16]}",
                        metadata={
                            "type": "incident_assessment",
                            "filename": file_name,
                            "uploaded_at": datetime.now().isoformat(),
                            "num_incidents": len(incidents) if 'incidents' in locals() else 0
                        }
                    )
                    print(f"✅ Phase 3 RAG: Document added to knowledge base")
                except Exception as e:
                    print(f"⚠️  Phase 3 RAG: Failed to add document: {e}")
            
            print(f"✅ STEP 3 COMPLETE: AI extracted {len(incidents)} incident(s) from document")
            for idx, incident in enumerate(incidents, 1):
                print(f"  Incident {idx}: {incident.get('IncidentTitle', 'Untitled')[:50]}...")

            # Phase 3: Include RAG and routing stats in response
            phase3_metadata = {
                "rag_available": is_rag_available(),
                "rag_stats": get_rag_stats() if is_rag_available() else None,
                "system_load": get_current_system_load(),
                "processing_time": processing_time,
                "model_routing": "enabled"
            }
            
            response_data = {
                'status': 'success',
                'message': f'Successfully extracted {len(incidents)} incident(s)',
                'document_name': file_name,
                'saved_path': safe_filename,
                'extracted_text_length': len(text),
                'preprocessing_metadata': preprocess_metadata,
                'phase3_metadata': phase3_metadata,  # Phase 3 stats
                'incidents': incidents
            }
            
            # Include compression metadata if file was compressed
            if compression_metadata:
                response_data['compression_metadata'] = compression_metadata
            
            # Include S3 info if uploaded successfully
            if s3_url:
                response_data['s3_url'] = s3_url
                response_data['s3_key'] = s3_key
            
            return JsonResponse(response_data)
        except Exception as process_error:
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
@csrf_exempt
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def save_extracted_incidents(request):
    """
    Save extracted/reviewed incidents to the database.
    MULTI-TENANCY: Incidents will be automatically assigned to user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    print(f"💾 Save incidents request received")
    
    try:
        # Try to get data from request.body first
        try:
            data = json.loads(request.body or "{}")
        except Exception as body_error:
            print(f"⚠️  Error reading request.body: {body_error}")
            # Fallback to request.data if available
            if hasattr(request, 'data') and request.data:
                data = request.data
                print(f"✅ Using request.data as fallback")
            else:
                raise body_error
        
        incidents_data = data.get('incidents', [])
        user_id = data.get('user_id', '1')
        
        print(f"💾 Processing {len(incidents_data)} incident(s) for user {user_id}")
        
        if not incidents_data:
            return JsonResponse({'status': 'error', 'message': 'No incidents provided'}, status=400)

        saved = []
        errors = []

        for idx, inc in enumerate(incidents_data):
            try:
                # Parse JSON fields
                mitigation = inc.get('Mitigation')
                if isinstance(mitigation, str):
                    try:
                        mitigation = json.loads(mitigation)
                    except:
                        mitigation = []
                
                form_details = inc.get('IncidentFormDetails')
                if isinstance(form_details, str):
                    try:
                        form_details = json.loads(form_details)
                    except:
                        form_details = {}
                
                # Set required fields with defaults, skip IDs and foreign keys
                kwargs = {
                    'IncidentTitle': inc.get('IncidentTitle', f'Incident {idx+1}'),
                    'Description': inc.get('Description', ''),
                    'Mitigation': mitigation,
                    # Date and Time are required in the model
                    'Date': date.today(),
                    'Time': datetime.now().time(),
                    # Skip UserId and other IDs - they can be null
                    'Origin': inc.get('Origin', 'MANUAL'),
                    'Comments': inc.get('Comments', ''),
                    'RiskCategory': inc.get('RiskCategory', ''),
                    'IncidentCategory': inc.get('IncidentCategory', ''),
                    'RiskPriority': inc.get('RiskPriority', 'Medium'),
                    'Attachments': inc.get('Attachments', ''),
                    'Status': 'Open',  # Always set status to 'Open' for new incidents
                    'RepeatedNot': inc.get('RepeatedNot', False),
                    'CostOfIncident': inc.get('CostOfIncident', ''),
                    'ReopenedNot': inc.get('ReopenedNot', False),
                    'RejectionSource': inc.get('RejectionSource'),
                    'AffectedBusinessUnit': inc.get('AffectedBusinessUnit', ''),
                    'SystemsAssetsInvolved': inc.get('SystemsAssetsInvolved', ''),
                    'GeographicLocation': inc.get('GeographicLocation', ''),
                    'Criticality': inc.get('Criticality', 'Medium'),
                    'InitialImpactAssessment': inc.get('InitialImpactAssessment', ''),
                    'InternalContacts': inc.get('InternalContacts', ''),
                    'ExternalPartiesInvolved': inc.get('ExternalPartiesInvolved', ''),
                    'RegulatoryBodies': inc.get('RegulatoryBodies', ''),
                    'RelevantPoliciesProceduresViolated': inc.get('RelevantPoliciesProceduresViolated', ''),
                    'ControlFailures': inc.get('ControlFailures', ''),
                    'LessonsLearned': inc.get('LessonsLearned', ''),
                    'IncidentClassification': inc.get('IncidentClassification', ''),
                    'PossibleDamage': inc.get('PossibleDamage', ''),
                    'IncidentFormDetails': form_details,
                }
                incident = Incident.objects.create(**kwargs)
                print(f"✅ Saved incident {idx+1}: {incident.IncidentTitle} (ID: {incident.IncidentId})")
                saved.append({'incident_id': incident.IncidentId, 'incident_title': incident.IncidentTitle})
            except Exception as ex:
                print(f"❌ Error saving incident {idx+1}: {str(ex)}")
                errors.append({'incident_index': idx, 'title': inc.get('IncidentTitle'), 'error': str(ex)})

        print(f"💾 Save complete: {len(saved)} saved, {len(errors)} errors")
        
        resp = {
            'status': 'success',
            'message': f'Saved {len(saved)} incident(s)' + (f' with {len(errors)} error(s)' if errors else ''),
            'saved': saved
        }
        if errors:
            resp['errors'] = errors
        return JsonResponse(resp)

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON payload'}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error saving incidents: {str(e)}'}, status=500)


@api_view(['GET'])
@permission_classes([AllowAny])
def test_openai_connection_incident(request):
    """Quick check that OpenAI API responds for incident module."""
    try:
        test_prompt = 'Return a JSON object with a single key "status" set to "ok" and a key "message" with value "OpenAI connection successful".'
        out = call_openai_json(test_prompt)
        return JsonResponse({
            'status': 'success', 
            'openai_reply': out, 
            'model': OPENAI_MODEL,
            'module': 'incident',
            'api_provider': 'OpenAI',
            'message': 'OpenAI API is working correctly for incident module'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error', 
            'message': f'OpenAI API error: {str(e)}', 
            'model': OPENAI_MODEL,
            'module': 'incident',
            'api_provider': 'OpenAI',
            'suggestion': 'Check OPENAI_API_KEY environment variable'
        }, status=500)