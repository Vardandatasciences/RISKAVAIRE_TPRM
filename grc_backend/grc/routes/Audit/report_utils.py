import os
import tempfile
import shutil
from django.db import connection
from .s3_functions import S3Client
from ...models import Audit
import logging
from ...routes.Global.notification_service import NotificationService
from django.utils import timezone
from datetime import datetime, timedelta
from .checklist_utils import update_lastchecklistitem_verified
import html
from ...routes.Global.logging_service import send_log
from django.db import connection
from datetime import datetime
from django.utils import timezone
from ...routes.Global.s3_fucntions import create_direct_mysql_client
from...rbac.decorators import (
    audit_view_reports_required,
    audit_conduct_required
)

# Set up logging
logger = logging.getLogger(__name__)

# Helper function for HTML escaping
def escape_html(content):
    """Safely escape HTML content"""
    if content is None:
        return ""
    return html.escape(str(content))

# Create Reports directory if it doesn't exist
REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'Reports')
os.makedirs(REPORTS_DIR, exist_ok=True)

@audit_view_reports_required
def generate_and_upload_report(audit_id, user_id='system'):
    """
    Generate an audit report and upload it to S3, then delete the local file
    
    Args:
        audit_id: ID of the audit to generate a report for
        user_id: ID of the user requesting the report
        
    Returns:
        Dict containing result with success status and file info if successful
    """
    try:
        # Log the start of report generation
        send_log(
            module="Report",
            actionType="GENERATE_REPORT_START",
            description=f"Starting report generation for audit ID {audit_id}",
            userId=user_id,
            entityType="Audit",
            entityId=str(audit_id)
        )
        
        # Get audit details
        try:
            audit = Audit.objects.get(pk=audit_id)
            logger.info(f"Generating report for audit {audit_id} with status {audit.Status}")
            
            # Log audit details retrieved
            send_log(
                module="Report",
                actionType="AUDIT_DETAILS_RETRIEVED",
                description=f"Retrieved audit details for report generation",
                userId=user_id,
                entityType="Audit",
                entityId=str(audit_id),
                additionalInfo={
                    "audit_status": audit.Status,
                    "audit_type": audit.AuditType
                }
            )
        except Audit.DoesNotExist:
            logger.error(f"Audit with ID {audit_id} not found")
            
            # Log audit not found error
            send_log(
                module="Report",
                actionType="AUDIT_NOT_FOUND",
                description=f"Audit with ID {audit_id} not found for report generation",
                userId=user_id,
                entityType="Audit",
                entityId=str(audit_id),
                logLevel="ERROR"
            )
            
            return {'success': False, 'error': f"Audit with ID {audit_id} not found"}
        
        # Create a temporary directory for processing
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Generate a file name with audit details
            file_name = f"Audit_Report_{audit_id}_{audit.AuditType or 'Unknown'}.docx"
            temp_file_path = os.path.join(temp_dir, file_name)
            
            # Call the internal report generation function
            # This mimics what the API endpoint does but returns the file path rather than an HTTP response
            from .report_views import generate_report_file
            
            # Log report file generation attempt
            send_log(
                module="Report",
                actionType="REPORT_FILE_GENERATION",
                description=f"Generating report file for audit ID {audit_id}",
                userId=user_id,
                entityType="Report",
                entityId=str(audit_id),
                additionalInfo={"file_path": temp_file_path}
            )
            
            report_file = generate_report_file(audit_id, temp_file_path)
            
            if not report_file or not os.path.exists(temp_file_path):
                logger.error(f"Failed to generate report for audit {audit_id}")
                
                # Log report generation failure
                send_log(
                    module="Report",
                    actionType="REPORT_GENERATION_FAILED",
                    description=f"Failed to generate report file for audit ID {audit_id}",
                    userId=user_id,
                    entityType="Report",
                    entityId=str(audit_id),
                    logLevel="ERROR"
                )
                
                return {'success': False, 'error': "Failed to generate report"}
            
            # Copy the file to the Reports directory
            local_file_path = os.path.join(REPORTS_DIR, file_name)
            shutil.copy2(temp_file_path, local_file_path)
            logger.info(f"Report saved locally at {local_file_path}")
            
            # Log report file saved locally
            send_log(
                module="Report",
                actionType="REPORT_SAVED_LOCALLY",
                description=f"Report saved locally for audit ID {audit_id}",
                userId=user_id,
                entityType="Report",
                entityId=str(audit_id),
                additionalInfo={"local_path": local_file_path}
            )
            
            # Initialize S3 client
            s3_client = S3Client()
            
            # Prepare metadata
            metadata = {
                'auditId': str(audit_id),
                'auditType': audit.AuditType or 'Unknown',
                'documentType': 'audit_report',
                'status': audit.Status or 'Unknown'
            }
            
            # Log S3 upload attempt
            send_log(
                module="Report",
                actionType="S3_UPLOAD_ATTEMPT",
                description=f"Attempting to upload report to S3 for audit ID {audit_id}",
                userId=user_id,
                entityType="Report",
                entityId=str(audit_id),
                additionalInfo=metadata
            )
            
            # Upload the file to S3
            upload_result = s3_client.upload_file(
                local_file_path, 
                user_id=user_id,
                module='Audit',
                **metadata  # Pass metadata as keyword arguments
            )
            
            if not isinstance(upload_result, dict) or 'error' in upload_result:
                error_msg = upload_result.get('error') if isinstance(upload_result, dict) else str(upload_result)
                logger.error(f"Failed to upload report to S3: {error_msg}")
                
                # Log S3 upload failure
                send_log(
                    module="Report",
                    actionType="S3_UPLOAD_FAILED",
                    description=f"Failed to upload report to S3: {error_msg}",
                    userId=user_id,
                    entityType="Report",
                    entityId=str(audit_id),
                    logLevel="ERROR",
                    additionalInfo={"error": error_msg}
                )
                
                return {'success': False, 'error': f"Failed to upload report: {error_msg}"}
            
            logger.info(f"Report uploaded to S3 successfully: {upload_result}")
            
            # Log S3 upload success
            send_log(
                module="Report",
                actionType="S3_UPLOAD_SUCCESS",
                description=f"Report uploaded to S3 successfully for audit ID {audit_id}",
                userId=user_id,
                entityType="Report",
                entityId=str(audit_id),
                additionalInfo={"s3_url": upload_result.get('file', {}).get('url')}
            )
            
            # Delete the local file after successful upload
            try:
                os.remove(local_file_path)
                logger.info(f"Local report file deleted: {local_file_path}")
                
                # Log local file deletion
                send_log(
                    module="Report",
                    actionType="LOCAL_FILE_DELETED",
                    description=f"Local report file deleted for audit ID {audit_id}",
                    userId=user_id,
                    entityType="Report",
                    entityId=str(audit_id)
                )
            except Exception as e:
                logger.warning(f"Failed to delete local report file: {str(e)}")
                
                # Log local file deletion failure
                send_log(
                    module="Report",
                    actionType="LOCAL_FILE_DELETE_FAILED",
                    description=f"Failed to delete local report file: {str(e)}",
                    userId=user_id,
                    entityType="Report",
                    entityId=str(audit_id),
                    logLevel="WARNING",
                    additionalInfo={"error": str(e)}
                )
            
            # Send notification about report upload
            try:
                notification_service = NotificationService()
                
                # Get user emails
                with connection.cursor() as cursor:
                    cursor.execute("""
                        SELECT 
                            auditor.Email as auditor_email,
                            reviewer.Email as reviewer_email,
                            auditor.UserName as auditor_name
                        FROM 
                            audit a
                        JOIN users auditor ON a.auditor = auditor.UserId
                        LEFT JOIN users reviewer ON a.reviewer = reviewer.UserId
                        WHERE a.AuditId = %s
                    """, [audit_id])
                    
                    user_row = cursor.fetchone()
                    if user_row:
                        auditor_email, reviewer_email, auditor_name = user_row
                        
                        # Notify reviewer about new report
                        if reviewer_email:
                            notification_data = {
                                'notification_type': 'reportUploaded',
                                'email': reviewer_email,
                                'email_type': 'gmail',
                                'template_data': [
                                    "Reviewer", 
                                    f"Audit #{audit_id}", 
                                    auditor_name or "Auditor", 
                                    timezone.now().strftime('%Y-%m-%d %H:%M')
                                ]
                            }
                            notification_service.send_multi_channel_notification(notification_data)
                            
                            # Log notification sent
                            send_log(
                                module="Report",
                                actionType="NOTIFICATION_SENT",
                                description=f"Notification sent to reviewer about report upload for audit ID {audit_id}",
                                userId=user_id,
                                entityType="Notification",
                                entityId=str(audit_id),
                                additionalInfo={"recipient": reviewer_email}
                            )
            except Exception as e:
                logger.error(f"Failed to send notification: {str(e)}")
                
                # Log notification failure
                send_log(
                    module="Report",
                    actionType="NOTIFICATION_FAILED",
                    description=f"Failed to send notification about report upload: {str(e)}",
                    userId=user_id,
                    entityType="Notification",
                    entityId=str(audit_id),
                    logLevel="ERROR",
                    additionalInfo={"error": str(e)}
                )
            
            # Log overall success
            send_log(
                module="Report",
                actionType="GENERATE_REPORT_SUCCESS",
                description=f"Report generated and uploaded successfully for audit ID {audit_id}",
                userId=user_id,
                entityType="Report",
                entityId=str(audit_id),
                additionalInfo={
                    "s3_url": upload_result.get('file', {}).get('url'),
                    "file_name": file_name
                }
            )
            
            return {
                'success': True,
                'message': 'Report generated and uploaded successfully',
                'file_info': upload_result.get('file'),
                'audit_id': audit_id
            }
            
        finally:
            # Clean up temporary directory
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        logger.exception(f"Error generating and uploading report: {str(e)}")
        
        # Log overall failure
        send_log(
            module="Report",
            actionType="GENERATE_REPORT_ERROR",
            description=f"Error generating and uploading report: {str(e)}",
            userId=user_id,
            entityType="Report",
            entityId=str(audit_id) if 'audit_id' in locals() else None,
            logLevel="ERROR",
            additionalInfo={"error": str(e)}
        )
        
        return {'success': False, 'error': str(e)}


@audit_view_reports_required
def generate_report_file(audit_id, output_file_path, version=None):
    """
    Internal function to generate a report file for an audit
    
    Args:
        audit_id: ID of the audit to generate a report for
        output_file_path: Path to save the report to
        version: Optional specific version to generate
        
    Returns:
        Path to the generated file if successful, None otherwise
    """
    try:
        # Log the start of report file generation
        send_log(
            module="Report",
            actionType="GENERATE_REPORT_FILE_START",
            description=f"Starting report file generation for audit ID {audit_id}",
            entityType="Audit",
            entityId=str(audit_id),
            additionalInfo={"output_path": output_file_path, "version": version}
        )
        
        # Import here to avoid circular imports
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
        import datetime
        
        # Get the audit details
        try:
            audit = Audit.objects.get(pk=audit_id)
            
            # Log audit details retrieved
            send_log(
                module="Report",
                actionType="REPORT_AUDIT_DETAILS",
                description=f"Retrieved audit details for report file generation",
                entityType="Audit",
                entityId=str(audit_id)
            )
        except Audit.DoesNotExist:
            logger.error(f"Audit with ID {audit_id} not found")
            
            # Log audit not found
            send_log(
                module="Report",
                actionType="REPORT_AUDIT_NOT_FOUND",
                description=f"Audit with ID {audit_id} not found for report file generation",
                entityType="Audit",
                entityId=str(audit_id),
                logLevel="ERROR"
            )
            
            return None
        
        # Get the findings (similar to report_views.py)
        findings = []
        if version and version.startswith('R'):
            # Log version-specific report generation
            send_log(
                module="Report",
                actionType="REPORT_VERSION_SPECIFIC",
                description=f"Generating report for specific version {version} of audit ID {audit_id}",
                entityType="AuditVersion",
                entityId=str(audit_id),
                additionalInfo={"version": version}
            )
            
            # Get specific version data
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT 
                        av.ExtractedInfo,
                        av.Date
                    FROM 
                        audit_version av
                    WHERE 
                        av.AuditId = %s AND av.Version = %s
                """, [audit_id, version])
                
                version_row = cursor.fetchone()
                if not version_row:
                    logger.error(f"Version {version} not found for audit {audit_id}")
                    
                    # Log version not found
                    send_log(
                        module="Report",
                        actionType="REPORT_VERSION_NOT_FOUND",
                        description=f"Version {version} not found for audit {audit_id}",
                        entityType="AuditVersion",
                        entityId=str(audit_id),
                        logLevel="ERROR"
                    )
                    
                    return None
                
                extracted_info = version_row[0]
                version_date = version_row[1]
                
                if extracted_info:
                    import json
                    # Parse JSON data
                    try:
                        version_data = json.loads(extracted_info)
                        
                        # Extract finding data from the version
                        for compliance_id, finding_data in version_data.items():
                            if compliance_id == '__metadata__' or compliance_id == 'overall_comments':
                                continue
                            
                            # Get compliance details
                            cursor.execute("""
                                SELECT 
                                    c.ComplianceItemDescription
                                FROM 
                                    compliance c
                                WHERE 
                                    c.ComplianceId = %s
                            """, [compliance_id])
                            
                            compliance_row = cursor.fetchone()
                            compliance_description = compliance_row[0] if compliance_row else "Unknown Compliance Item"
                            
                            # Convert version finding data
                            finding = {
                                'ComplianceId': compliance_id,
                                'MajorMinor': finding_data.get('major_minor', 'N/A'),
                                'Check': finding_data.get('check', 'N/A'),
                                'HowToVerify': finding_data.get('how_to_verify', 'N/A'),
                                'Evidence': finding_data.get('evidence', 'N/A'),
                                'DetailsOfFinding': finding_data.get('details_of_finding', 'N/A'),
                                'Impact': finding_data.get('impact', 'N/A'),
                                'Recommendation': finding_data.get('recommendation', 'N/A'),
                                'Comments': finding_data.get('comments', 'N/A'),
                                'CheckedDate': version_date,
                                'ComplianceItemDescription': compliance_description
                            }
                            findings.append(finding)
                        
                        # Log findings extracted from version
                        send_log(
                            module="Report",
                            actionType="REPORT_VERSION_FINDINGS",
                            description=f"Extracted {len(findings)} findings from version {version} for audit ID {audit_id}",
                            entityType="AuditVersion",
                            entityId=str(audit_id),
                            additionalInfo={"finding_count": len(findings)}
                        )
                            
                    except json.JSONDecodeError as e:
                        logger.error(f"Invalid JSON in ExtractedInfo: {str(e)}")
                        
                        # Log JSON parsing error
                        send_log(
                            module="Report",
                            actionType="REPORT_JSON_ERROR",
                            description=f"Invalid JSON in ExtractedInfo: {str(e)}",
                            entityType="AuditVersion",
                            entityId=str(audit_id),
                            logLevel="ERROR",
                            additionalInfo={"error": str(e)}
                        )
                        
                        return None
        else:
            # Log standard report generation
            send_log(
                module="Report",
                actionType="REPORT_STANDARD",
                description=f"Generating standard report for audit ID {audit_id}",
                entityType="Audit",
                entityId=str(audit_id)
            )
            
            # Fetch findings directly from audit_findings table
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT 
                        af.ComplianceId,
                        c.ComplianceItemDescription,
                        af.`Check`,
                        af.Evidence,
                        af.Comments
                    FROM 
                        audit_findings af
                    JOIN
                        compliance c ON af.ComplianceId = c.ComplianceId
                    WHERE 
                        af.AuditId = %s
                """, [audit_id])
                
                columns = [col[0] for col in cursor.description]
                findings = [dict(zip(columns, row)) for row in cursor.fetchall()]
                
                # Log findings retrieved
                send_log(
                    module="Report",
                    actionType="REPORT_FINDINGS",
                    description=f"Retrieved {len(findings)} findings for standard report of audit ID {audit_id}",
                    entityType="Audit",
                    entityId=str(audit_id),
                    additionalInfo={"finding_count": len(findings)}
                )
        
        # Check if we have findings
        if len(findings) == 0:
            logger.warning(f"No findings available for audit {audit_id}")
            
            # Log no findings available
            send_log(
                module="Report",
                actionType="REPORT_NO_FINDINGS",
                description=f"No findings available for audit {audit_id}",
                entityType="Audit",
                entityId=str(audit_id),
                logLevel="WARNING"
            )
            
            return None
        
        # Create the document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add professional title page
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run('AUDIT COMPLIANCE REPORT')
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.name = 'Arial'
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle_paragraph = doc.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(f'Audit ID: {audit_id}')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.name = 'Arial'
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if version:
            version_paragraph = doc.add_paragraph()
            version_run = version_paragraph.add_run(f'Version: {version}')
            version_run.font.size = Pt(12)
            version_run.font.name = 'Arial'
            version_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(12)
        date_run.font.name = 'Arial'
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Add Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run('This audit compliance report provides a comprehensive assessment of the organization\'s adherence to established policies, procedures, and regulatory requirements. The audit was conducted to evaluate compliance status, identify gaps, and provide recommendations for improvement.')
        
        # Add audit overview
        doc.add_heading('AUDIT OVERVIEW', level=2)
        
        # Create overview table
        overview_table = doc.add_table(rows=0, cols=2)
        overview_table.style = 'Table Grid'
        overview_table.autofit = False
        overview_table.columns[0].width = Inches(2)
        overview_table.columns[1].width = Inches(4)
        
        overview_data = [
            ('Audit ID', str(audit_id)),
            ('Status', getattr(audit, 'Status', 'Unknown')),
            ('Type', getattr(audit, 'AuditType', 'Unknown')),
            ('Due Date', getattr(audit, 'DueDate', 'N/A')),
            ('Review Status', getattr(audit, 'ReviewStatus', 'N/A')),
            ('Total Findings', str(len(findings))),
        ]
        
        if version:
            overview_data.append(('Report Version', version))
        
        for item, value in overview_data:
            row = overview_table.add_row()
            row.cells[0].text = item
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Add spacing
        
        doc.add_page_break()
        
        # Add detailed findings section
        doc.add_heading('DETAILED FINDINGS', level=1)
        
        # Group findings by severity
        major_findings = []
        minor_findings = []
        critical_findings = []
        
        for finding in findings:
            severity = finding.get('MajorMinor', '0')
            if severity == '1':
                major_findings.append(finding)
            elif severity == '2':
                critical_findings.append(finding)
            else:
                minor_findings.append(finding)
        
        # Add Critical Findings
        if critical_findings:
            doc.add_heading('CRITICAL FINDINGS', level=2)
            doc.add_paragraph('Critical findings represent significant compliance gaps that require immediate attention and remediation.')
            
            for i, finding in enumerate(critical_findings, 1):
                safe_description = escape_html(finding.get("ComplianceItemDescription", "N/A"))
                doc.add_heading(f'Finding {i}: {safe_description}', level=3)
                
                # Add finding details with proper formatting
                details_para = doc.add_paragraph()
                details_para.add_run('Finding Details: ').bold = True
                details_para.add_run(escape_html(finding.get('DetailsOfFinding', 'No details provided')))
                
                impact_para = doc.add_paragraph()
                impact_para.add_run('Impact: ').bold = True
                impact_para.add_run(escape_html(finding.get('Impact', 'Impact not specified')))
                
                recommendation_para = doc.add_paragraph()
                recommendation_para.add_run('Recommendation: ').bold = True
                recommendation_para.add_run(escape_html(finding.get('Recommendation', 'No recommendation provided')))
                
                doc.add_paragraph()  # Add spacing
        
        # Add Major Findings
        if major_findings:
            doc.add_heading('MAJOR FINDINGS', level=2)
            doc.add_paragraph('Major findings represent compliance gaps that should be addressed in the short term.')
            
            for i, finding in enumerate(major_findings, 1):
                safe_description = escape_html(finding.get("ComplianceItemDescription", "N/A"))
                doc.add_heading(f'Finding {i}: {safe_description}', level=3)
                
                details_para = doc.add_paragraph()
                details_para.add_run('Finding Details: ').bold = True
                details_para.add_run(escape_html(finding.get('DetailsOfFinding', 'No details provided')))
                
                impact_para = doc.add_paragraph()
                impact_para.add_run('Impact: ').bold = True
                impact_para.add_run(escape_html(finding.get('Impact', 'Impact not specified')))
                
                recommendation_para = doc.add_paragraph()
                recommendation_para.add_run('Recommendation: ').bold = True
                recommendation_para.add_run(escape_html(finding.get('Recommendation', 'No recommendation provided')))
                
                doc.add_paragraph()  # Add spacing
        
        # Add Minor Findings
        if minor_findings:
            doc.add_heading('MINOR FINDINGS', level=2)
            doc.add_paragraph('Minor findings represent areas for improvement that should be addressed as resources permit.')
            
            for i, finding in enumerate(minor_findings, 1):
                safe_description = escape_html(finding.get("ComplianceItemDescription", "N/A"))
                doc.add_heading(f'Finding {i}: {safe_description}', level=3)
                
                details_para = doc.add_paragraph()
                details_para.add_run('Finding Details: ').bold = True
                details_para.add_run(escape_html(finding.get('DetailsOfFinding', 'No details provided')))
                
                recommendation_para = doc.add_paragraph()
                recommendation_para.add_run('Recommendation: ').bold = True
                recommendation_para.add_run(escape_html(finding.get('Recommendation', 'No recommendation provided')))
                
                doc.add_paragraph()  # Add spacing
        
        doc.add_page_break()
        
        # Add compliance summary
        doc.add_heading('COMPLIANCE SUMMARY', level=1)
        
        # Calculate compliance statistics
        total_findings = len(findings)
        compliant_count = sum(1 for f in findings if f.get('Check') == '2')
        non_compliant_count = sum(1 for f in findings if f.get('Check') == '0')
        partially_compliant_count = sum(1 for f in findings if f.get('Check') == '1')
        
        compliance_rate = (compliant_count / total_findings * 100) if total_findings > 0 else 0
        
        # Add compliance statistics
        doc.add_heading('Compliance Statistics', level=2)
        
        stats_table = doc.add_table(rows=0, cols=2)
        stats_table.style = 'Table Grid'
        stats_table.autofit = False
        stats_table.columns[0].width = Inches(2.5)
        stats_table.columns[1].width = Inches(3.5)
        
        stats_data = [
            ('Total Items Audited', str(total_findings)),
            ('Compliant Items', str(compliant_count)),
            ('Partially Compliant Items', str(partially_compliant_count)),
            ('Non-Compliant Items', str(non_compliant_count)),
            ('Overall Compliance Rate', f'{compliance_rate:.1f}%'),
        ]
        
        for item, value in stats_data:
            row = stats_table.add_row()
            row.cells[0].text = item
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Add spacing
        
        # Add recommendations section
        doc.add_heading('RECOMMENDATIONS', level=1)
        
        recommendations_para = doc.add_paragraph()
        recommendations_para.add_run('Based on the audit findings, the following recommendations are provided to improve compliance and strengthen the organization\'s control environment:')
        
        doc.add_paragraph()
        
        # Add numbered recommendations
        if critical_findings:
            doc.add_heading('Priority 1 - Critical Findings', level=2)
            doc.add_paragraph('Address critical findings immediately to mitigate significant risks.')
        
        if major_findings:
            doc.add_heading('Priority 2 - Major Findings', level=2)
            doc.add_paragraph('Address major findings within the next quarter to improve compliance posture.')
        
        if minor_findings:
            doc.add_heading('Priority 3 - Minor Findings', level=2)
            doc.add_paragraph('Address minor findings as resources permit to enhance overall compliance.')
        
        doc.add_page_break()
        
        # Add conclusion
        doc.add_heading('CONCLUSION', level=1)
        
        conclusion_para = doc.add_paragraph()
        conclusion_para.add_run('This audit has identified areas of strength and opportunities for improvement in the organization\'s compliance program. ')
        
        if compliance_rate >= 80:
            conclusion_para.add_run('Overall, the organization demonstrates a strong commitment to compliance with most requirements being met.')
        elif compliance_rate >= 60:
            conclusion_para.add_run('The organization shows progress in compliance but has areas that require attention to achieve full compliance.')
        else:
            conclusion_para.add_run('The organization has significant compliance gaps that require immediate attention and dedicated resources for remediation.')
        
        conclusion_para.add_run(' Implementation of the recommendations provided in this report will strengthen the control environment and improve overall compliance posture.')
        
        # Save the document
        doc.save(output_file_path)
        
        # Log successful report file generation
        send_log(
            module="Report",
            actionType="GENERATE_REPORT_FILE_SUCCESS",
            description=f"Successfully generated report file for audit ID {audit_id}",
            entityType="Report",
            entityId=str(audit_id),
            additionalInfo={
                "output_path": output_file_path,
                "version": version,
                "finding_count": len(findings)
            }
        )
        
        return output_file_path
        
    except Exception as e:
        logger.exception(f"Error generating report file: {str(e)}")
        
        # Log error in report file generation
        send_log(
            module="Report",
            actionType="GENERATE_REPORT_FILE_ERROR",
            description=f"Error generating report file: {str(e)}",
            entityType="Report",
            entityId=str(audit_id) if 'audit_id' in locals() else None,
            logLevel="ERROR",
            additionalInfo={"error": str(e)}
        )
        
        return None 