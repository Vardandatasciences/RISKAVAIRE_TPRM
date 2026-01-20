from django.http import HttpResponse, FileResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework import status
from django.db import connection
from ...models import Audit
import os
import io
import tempfile
import shutil
from django.conf import settings
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ...routes.Global.notification_service import NotificationService
import json
from typing import Optional, Dict, Any
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.db import connection
from datetime import datetime, timedelta
from typing import Dict, Any, Optional
from django.utils import timezone
from ...rbac.permissions import (
    AuditViewPermission, AuditReviewPermission, AuditConductPermission
)
from ...rbac.decorators import (
    audit_view_reports_required,
    audit_review_required,
    audit_conduct_required
)
from .framework_filter_helper import get_active_framework_filter, apply_framework_filter_to_audits, get_framework_sql_filter

# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)

@api_view(['GET'])
@permission_classes([AuditViewPermission])
@audit_view_reports_required
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def generate_audit_report(request, audit_id):
    """
    Generate and download an audit report in DOCX format with tables for each finding
    MULTI-TENANCY: Only generates reports for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        print(f"DEBUG: generate_audit_report called for audit_id: {audit_id}")
        
        # Get user ID from request (JWT or session)
        from .audit_views import get_user_id_from_jwt
        user_id = get_user_id_from_jwt(request)
        print(f"DEBUG: User authenticated with user_id: {user_id}")
        print(f"DEBUG: Request user: {request.user}")
        
        # Check if a specific version is requested
        version = request.query_params.get('version')
        print(f"DEBUG: Version parameter: {version}")
        
        # Verify audit exists for tenant
        try:
            audit = Audit.objects.get(AuditId=audit_id, tenant_id=tenant_id)
        except Audit.DoesNotExist:
            return Response({"error": f"Audit {audit_id} not found"}, status=status.HTTP_404_NOT_FOUND)
        
        # If version is provided, check if it has an ApprovedRejected status
        if version:
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT av.ApprovedRejected
                    FROM audit_version av
                    JOIN audit a ON av.AuditId = a.AuditId
                    WHERE av.AuditId = %s AND av.Version = %s AND a.TenantId = %s
                """, [audit_id, version, tenant_id])
                
                version_row = cursor.fetchone()
                if not version_row:
                    return Response({"error": f"Version {version} not found for audit {audit_id}"}, 
                                   status=status.HTTP_404_NOT_FOUND)
                
                if version_row[0] is None:
                    return Response({"error": f"Version {version} does not have an approved or rejected status"}, 
                                   status=status.HTTP_400_BAD_REQUEST)
        
        # Create a temporary directory for the process
        temp_dir = tempfile.mkdtemp()
        output_file = os.path.join(temp_dir, f"audit_report_{audit_id}_{version if version else 'latest'}.docx")
        
        try:
            # Generate the report file, passing tenant_id
            report_file = generate_report_file(audit_id, output_file, version, tenant_id)
            
            if not report_file or not os.path.exists(output_file):
                print(f"ERROR: Failed to generate report for audit {audit_id}")
                return Response({"error": "Failed to generate report"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            # Determine file name for download
            if version:
                download_filename = f"audit_report_{audit_id}_v{version}.docx"
            else:
                download_filename = f"audit_report_{audit_id}.docx"
            
            # Return the file for download
            response = FileResponse(
                open(output_file, 'rb'),
                as_attachment=True,
                filename=download_filename
            )
            
            # Notify relevant users about report generation
            try:
                notification_service = NotificationService()
                
                # Get user details
                with connection.cursor() as cursor:
                    cursor.execute("""
                        SELECT 
                            auditor.Email as auditor_email,
                            reviewer.Email as reviewer_email,
                            auditor.UserName as auditor_name
                        FROM 
                            audit a
                        JOIN users auditor ON a.auditor = auditor.UserId AND auditor.TenantId = %s
                        LEFT JOIN users reviewer ON a.reviewer = reviewer.UserId AND reviewer.TenantId = %s
                        WHERE a.AuditId = %s AND a.TenantId = %s
                    """, [tenant_id, tenant_id, audit_id, tenant_id])
                    
                    user_row = cursor.fetchone()
                    if user_row:
                        auditor_email, reviewer_email, auditor_name = user_row
                        
                        # Notify the user who generated the report (could be auditor or reviewer)
                        user_email = request.user.Email if hasattr(request, 'user') and hasattr(request.user, 'Email') else auditor_email
                        
                        notification_data = {
                            'notification_type': 'reportGenerated',
                            'email': user_email,
                            'email_type': 'gmail',
                            'template_data': [
                                "User",  # Generic name since we don't know exactly who generated
                                f"Audit #{audit_id}",
                                version if version else "latest",
                                datetime.now().strftime('%Y-%m-%d %H:%M')
                            ]
                        }
                        notification_service.send_multi_channel_notification(notification_data)
            except Exception as e:
                print(f"Failed to send notification: {str(e)}")
            
            return response
            
        finally:
            # Clean up temporary directory after sending response
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        print(f"ERROR in generate_audit_report: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def get_audit_data(audit_id: int, tenant_id: int) -> Optional[Dict[str, Any]]:
    """
    Get all necessary audit data for report generation
    MULTI-TENANCY: Requires tenant_id parameter for data isolation
    """
    try:
        with connection.cursor() as cursor:
            # Get audit details, filtered by tenant
            cursor.execute("""
                SELECT 
                    a.Title, a.Scope, a.Objective, a.BusinessUnit,
                    a.Evidence as AuditEvidence, a.Comments as OverallAuditComments,
                    a.ReviewerComments as OverallReviewComments,
                    p.PolicyName as Policy, sp.SubPolicyName as SubPolicy, 
                    f.FrameworkName as Framework,
                    a.CompletionDate, a.ReviewDate
                FROM audit a
                LEFT JOIN policies p ON a.PolicyId = p.PolicyId AND p.TenantId = %s
                LEFT JOIN subpolicies sp ON a.SubPolicyId = sp.SubPolicyId AND sp.TenantId = %s
                LEFT JOIN frameworks f ON a.FrameworkId = f.FrameworkId AND f.TenantId = %s
                WHERE a.AuditId = %s AND a.TenantId = %s
            """, [tenant_id, tenant_id, tenant_id, audit_id, tenant_id])
            audit_data = cursor.fetchone()
            
            if not audit_data:
                return None
                
            # Get audit findings with compliance details, filtered by tenant
            cursor.execute("""
                SELECT 
                    af.ComplianceId,
                    af.MajorMinor as TypeOfFinding,
                    af.CheckedDate as ReviewDate,
                    af.Evidence as ComplianceEvidence,
                    af.HowToVerify,
                    af.Impact,
                    af.DetailsOfFinding,
                    af.Comments,
                    af.Check as ReviewStatus,
                    af.ReviewComments,
                    COALESCE(af.SeverityRating, 0) as SeverityRating,
                    af.PredictiveRisks,
                    af.CorrectiveActions,
                    af.UnderlyingCause,
                    af.WhyToVerify,
                    af.WhatToVerify,
                    af.SuggestedActionPlan,
                    af.AssignedDate as MitigationDate,
                    af.ResponsibleForPlan,
                    COALESCE(af.ReAudit, 0) as ReAudit,
                    af.ReAuditDate,
                    af.Check as ComplianceStatus,
                    c.ComplianceTitle,
                    c.ComplianceItemDescription
                FROM audit_findings af
                LEFT JOIN compliance c ON af.ComplianceId = c.ComplianceId AND c.TenantId = %s
                WHERE af.AuditId = %s AND af.TenantId = %s
            """, [tenant_id, audit_id, tenant_id])
            findings = cursor.fetchall()
            
            return {
                'audit': audit_data,
                'findings': findings
            }
    except Exception as e:
        print(f"Error getting audit data: {str(e)}")
        return None

def generate_report_file(audit_id: int, output_path: str, version=None, tenant_id=None) -> Optional[str]:
    """
    Generate a professional audit compliance report in Word format with proper formatting
    MULTI-TENANCY: Requires tenant_id parameter for data isolation
    """
    try:
        # Get audit data, filtered by tenant
        data = get_audit_data(audit_id, tenant_id)
        if not data:
            return None
            
        audit_data = data['audit']
        findings = data['findings']
        
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add professional title page
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run('AUDIT COMPLIANCE REPORT')
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.name = 'Arial'
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle_paragraph = doc.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(f'Audit ID: {audit_id}')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.name = 'Arial'
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
        date_run.font.size = Pt(12)
        date_run.font.name = 'Arial'
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Add Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        summary_para = doc.add_paragraph()
        summary_para.add_run('This audit compliance report provides a comprehensive assessment of the organization\'s adherence to established policies, procedures, and regulatory requirements. The audit was conducted to evaluate compliance status, identify gaps, and provide recommendations for improvement.')
        
        # Add audit overview
        doc.add_heading('AUDIT OVERVIEW', level=2)
        
        # Create overview table
        overview_table = doc.add_table(rows=0, cols=2)
        overview_table.style = 'Table Grid'
        overview_table.autofit = False
        overview_table.columns[0].width = Inches(2)
        overview_table.columns[1].width = Inches(4)
        
        overview_data = [
            ('Audit ID', str(audit_id)),
            ('Framework', audit_data[9] or 'N/A'),
            ('Policy', audit_data[7] or 'N/A'),
            ('Sub-Policy', audit_data[8] or 'N/A'),
            ('Completion Date', audit_data[10].strftime('%B %d, %Y') if audit_data[10] else 'N/A'),
            ('Review Date', audit_data[11].strftime('%B %d, %Y') if audit_data[11] else 'N/A'),
            ('Total Findings', str(len(findings))),
        ]
        
        for item, value in overview_data:
            row = overview_table.add_row()
            row.cells[0].text = item
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Add spacing
        
        # Add scope and objectives
        if audit_data[1]:  # Scope
            doc.add_heading('AUDIT SCOPE', level=2)
            scope_para = doc.add_paragraph()
            scope_para.add_run(audit_data[1])
        
        if audit_data[2]:  # Objective
            doc.add_heading('AUDIT OBJECTIVES', level=2)
            objective_para = doc.add_paragraph()
            objective_para.add_run(audit_data[2])
        
        doc.add_page_break()
        
        # Add detailed findings section
        doc.add_heading('DETAILED FINDINGS', level=1)
        
        # Map check status to readable format
        check_status_map = {
            '0': 'Not Compliant',
            '1': 'Partially Compliant',
            '2': 'Compliant',
            '3': 'Not Applicable'
        }
        
        # Map major/minor to readable format
        major_minor_map = {
            '0': 'Minor',
            '1': 'Major',
            '2': 'Critical'
        }
        
        # Group findings by severity
        major_findings = []
        minor_findings = []
        critical_findings = []
        
        for finding in findings:
            severity = finding[1] if finding[1] else '0'
            if severity == '1':
                major_findings.append(finding)
            elif severity == '2':
                critical_findings.append(finding)
            else:
                minor_findings.append(finding)
        
        # Add Critical Findings
        if critical_findings:
            doc.add_heading('CRITICAL FINDINGS', level=2)
            doc.add_paragraph('Critical findings represent significant compliance gaps that require immediate attention and remediation.')
            
            for i, finding in enumerate(critical_findings, 1):
                doc.add_heading(f'Finding {i}: {finding[22]}', level=3)
                
                # Add finding details with proper formatting
                details_para = doc.add_paragraph()
                details_para.add_run('Finding Details: ').bold = True
                details_para.add_run(finding[6] or 'No details provided')
                
                impact_para = doc.add_paragraph()
                impact_para.add_run('Impact: ').bold = True
                impact_para.add_run(finding[5] or 'Impact not specified')
                
                recommendation_para = doc.add_paragraph()
                recommendation_para.add_run('Recommendation: ').bold = True
                recommendation_para.add_run(finding[16] or 'No recommendation provided')
                
                doc.add_paragraph()  # Add spacing
        
        # Add Major Findings
        if major_findings:
            doc.add_heading('MAJOR FINDINGS', level=2)
            doc.add_paragraph('Major findings represent compliance gaps that should be addressed in the short term.')
            
            for i, finding in enumerate(major_findings, 1):
                doc.add_heading(f'Finding {i}: {finding[22]}', level=3)
                
                details_para = doc.add_paragraph()
                details_para.add_run('Finding Details: ').bold = True
                details_para.add_run(finding[6] or 'No details provided')
                
                impact_para = doc.add_paragraph()
                impact_para.add_run('Impact: ').bold = True
                impact_para.add_run(finding[5] or 'Impact not specified')
                
                recommendation_para = doc.add_paragraph()
                recommendation_para.add_run('Recommendation: ').bold = True
                recommendation_para.add_run(finding[16] or 'No recommendation provided')
                
                doc.add_paragraph()  # Add spacing
        
        # Add Minor Findings
        if minor_findings:
            doc.add_heading('MINOR FINDINGS', level=2)
            doc.add_paragraph('Minor findings represent areas for improvement that should be addressed as resources permit.')
            
            for i, finding in enumerate(minor_findings, 1):
                doc.add_heading(f'Finding {i}: {finding[22]}', level=3)
                
                details_para = doc.add_paragraph()
                details_para.add_run('Finding Details: ').bold = True
                details_para.add_run(finding[6] or 'No details provided')
                
                recommendation_para = doc.add_paragraph()
                recommendation_para.add_run('Recommendation: ').bold = True
                recommendation_para.add_run(finding[16] or 'No recommendation provided')
                
                doc.add_paragraph()  # Add spacing
        
        doc.add_page_break()
        
        # Add compliance summary
        doc.add_heading('COMPLIANCE SUMMARY', level=1)
        
        # Calculate compliance statistics
        total_findings = len(findings)
        compliant_count = sum(1 for f in findings if f[21] == '2')
        non_compliant_count = sum(1 for f in findings if f[21] == '0')
        partially_compliant_count = sum(1 for f in findings if f[21] == '1')
        
        compliance_rate = (compliant_count / total_findings * 100) if total_findings > 0 else 0
        
        # Add compliance statistics
        doc.add_heading('Compliance Statistics', level=2)
        
        stats_table = doc.add_table(rows=0, cols=2)
        stats_table.style = 'Table Grid'
        stats_table.autofit = False
        stats_table.columns[0].width = Inches(2.5)
        stats_table.columns[1].width = Inches(3.5)
        
        stats_data = [
            ('Total Items Audited', str(total_findings)),
            ('Compliant Items', str(compliant_count)),
            ('Partially Compliant Items', str(partially_compliant_count)),
            ('Non-Compliant Items', str(non_compliant_count)),
            ('Overall Compliance Rate', f'{compliance_rate:.1f}%'),
        ]
        
        for item, value in stats_data:
            row = stats_table.add_row()
            row.cells[0].text = item
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Add spacing
        
        # Add recommendations section
        doc.add_heading('RECOMMENDATIONS', level=1)
        
        recommendations_para = doc.add_paragraph()
        recommendations_para.add_run('Based on the audit findings, the following recommendations are provided to improve compliance and strengthen the organization\'s control environment:')
        
        doc.add_paragraph()
        
        # Add numbered recommendations
        if critical_findings:
            doc.add_heading('Priority 1 - Critical Findings', level=2)
            doc.add_paragraph('Address critical findings immediately to mitigate significant risks.')
        
        if major_findings:
            doc.add_heading('Priority 2 - Major Findings', level=2)
            doc.add_paragraph('Address major findings within the next quarter to improve compliance posture.')
        
        if minor_findings:
            doc.add_heading('Priority 3 - Minor Findings', level=2)
            doc.add_paragraph('Address minor findings as resources permit to enhance overall compliance.')
        
        doc.add_page_break()
        
        # Add conclusion
        doc.add_heading('CONCLUSION', level=1)
        
        conclusion_para = doc.add_paragraph()
        conclusion_para.add_run('This audit has identified areas of strength and opportunities for improvement in the organization\'s compliance program. ')
        
        if compliance_rate >= 80:
            conclusion_para.add_run('Overall, the organization demonstrates a strong commitment to compliance with most requirements being met.')
        elif compliance_rate >= 60:
            conclusion_para.add_run('The organization shows progress in compliance but has areas that require attention to achieve full compliance.')
        else:
            conclusion_para.add_run('The organization has significant compliance gaps that require immediate attention and dedicated resources for remediation.')
        
        conclusion_para.add_run(' Implementation of the recommendations provided in this report will strengthen the control environment and improve overall compliance posture.')
        
        # Add appendices if there's additional information
        if audit_data[4] or audit_data[5] or audit_data[6]:  # Evidence, Comments, Review Comments
            doc.add_page_break()
            doc.add_heading('APPENDICES', level=1)
            
            if audit_data[4]:  # Audit Evidence
                doc.add_heading('Audit Evidence', level=2)
                evidence_para = doc.add_paragraph()
                evidence_para.add_run(audit_data[4])
            
            if audit_data[5]:  # Overall Audit Comments
                doc.add_heading('Overall Audit Comments', level=2)
                comments_para = doc.add_paragraph()
                comments_para.add_run(audit_data[5])
            
            if audit_data[6]:  # Overall Review Comments
                doc.add_heading('Overall Review Comments', level=2)
                review_para = doc.add_paragraph()
                review_para.add_run(audit_data[6])
        
        # Save document
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def create_incidents_for_findings(audit_id: int, tenant_id: int) -> None:
    """
    Create incidents for non-compliant and partially compliant audit findings
    MULTI-TENANCY: Requires tenant_id parameter for data isolation
    """
    try:
        with connection.cursor() as cursor:
            # Get relevant audit findings and compliance details, filtered by tenant
            cursor.execute("""
                SELECT 
                    af.ComplianceId,
                    af.Check,
                    af.Comments,
                    af.user_id,
                    c.ComplianceTitle,
                    c.ComplianceItemDescription,
                    c.PossibleDamage,
                    c.Mitigation
                FROM audit_findings af
                JOIN compliance c ON af.ComplianceId = c.ComplianceId AND c.TenantId = %s
                WHERE af.AuditId = %s 
                AND af.TenantId = %s
                AND (
                    (af.Check = '0') OR  -- Not Compliant
                    (af.Check = '1')     -- Partially Compliant
                )
            """, [tenant_id, audit_id, tenant_id])
            
            findings = cursor.fetchall()
            
            current_datetime = datetime.now()
            current_date = current_datetime.date()
            current_time = current_datetime.time()
            
            # Create incidents for each finding
            for finding in findings:
                compliance_id = finding[0]
                check_status = finding[1]
                comments = finding[2]
                user_id = finding[3]
                compliance_title = finding[4]
                compliance_desc = finding[5]
                possible_damage = finding[6]
                mitigation = finding[7]
                
                # Check if incident already exists for this finding, filtered by tenant
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM incidents i
                    JOIN audit a ON i.AuditId = a.AuditId
                    WHERE i.AuditId = %s AND i.ComplianceId = %s AND a.TenantId = %s
                """, [audit_id, compliance_id, tenant_id])
                
                if cursor.fetchone()[0] > 0:
                    print(f"Incident already exists for AuditId {audit_id} and ComplianceId {compliance_id}")
                    continue
                
                # Get FrameworkId from the audit, filtered by tenant
                cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s AND TenantId = %s", [audit_id, tenant_id])
                framework_row = cursor.fetchone()
                framework_id = framework_row[0] if framework_row else None
                
                # Create new incident with TenantId
                cursor.execute("""
                    INSERT INTO incidents (
                        IncidentTitle,
                        Description,
                        PossibleDamage,
                        Mitigation,
                        AuditId,
                        ComplianceId,
                        Date,
                        Time,
                        UserId,
                        Origin,
                        Comments,
                        Status,
                        FrameworkId,
                        TenantId
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                    )
                """, [
                    f"Non-Compliance of {compliance_title}",
                    compliance_desc,
                    possible_damage,
                    mitigation,
                    audit_id,
                    compliance_id,
                    current_date,
                    current_time,
                    user_id,
                    "Audit Finding",
                    comments,
                    "Open",
                    framework_id,
                    tenant_id  # MULTI-TENANCY: Add TenantId to incident (note: variable is tenant_id but column is TenantId)
                ])
                
                print(f"Created incident for ComplianceId {compliance_id} in AuditId {audit_id}")
                
    except Exception as e:
        print(f"Error creating incidents: {str(e)}")
        import traceback
        traceback.print_exc()

@csrf_exempt
@api_view(['POST'])
@permission_classes([AuditReviewPermission])
@audit_review_required
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def approve_audit_and_create_incidents(request, audit_id):
    """
    API endpoint to approve audit and create incidents for non-compliant findings
    MULTI-TENANCY: Only approves audits and creates incidents for user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        # Verify audit exists for tenant
        try:
            audit = Audit.objects.get(AuditId=audit_id, tenant_id=tenant_id)
        except Audit.DoesNotExist:
            return Response({"error": f"Audit {audit_id} not found"}, status=status.HTTP_404_NOT_FOUND)
        
        with connection.cursor() as cursor:
            # Update audit status, filtered by tenant
            cursor.execute("""
                UPDATE audit 
                SET Status = 'Approved',
                    CompletionDate = %s
                WHERE AuditId = %s AND TenantId = %s
            """, [datetime.now(), audit_id, tenant_id])
            
            # Update LastChecklistItemVerified in audit_findings, filtered by tenant
            cursor.execute("""
                UPDATE audit_findings
                SET LastChecklistItemVerified = %s
                WHERE AuditId = %s AND TenantId = %s
            """, [datetime.now(), audit_id, tenant_id])
            
            # Create incidents for non-compliant and partially compliant findings, filtered by tenant
            create_incidents_for_findings(audit_id, tenant_id)
            
            return Response({
                "message": "Audit approved and incidents created successfully",
                "audit_id": audit_id
            }, status=status.HTTP_200_OK)
            
    except Exception as e:
        print(f"Error in approve_audit_and_create_incidents: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({
            "error": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR) 