import os
import json
import logging
import requests
from typing import Dict, List, Optional, Any
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import pandas as pd
from io import BytesIO
import tempfile

# Import our custom modules
from tprm_backend.s3 import create_direct_mysql_client
from tprm_backend.lamma import simple_ollama

logger = logging.getLogger(__name__)


class DocumentProcessingService:
    """Service for processing documents with OCR and AI extraction"""
    
    def __init__(self):
        self.s3_client = None
        self._init_s3_client()
    
    def _init_s3_client(self):
        """Initialize S3 client"""
        try:
            # Configure MySQL connection for S3 client using udm database
            db_config = settings.DATABASES.get('udm', settings.DATABASES['default'])
            mysql_config = {
                'host': db_config['HOST'],
                'user': db_config['USER'],
                'password': db_config['PASSWORD'],
                'database': db_config['NAME'],
                'port': int(db_config['PORT']),
            }
            
            self.s3_client = create_direct_mysql_client(mysql_config)
            logger.info(f"[OK] S3 client initialized successfully with database: {mysql_config['database']}")
            
        except Exception as e:
            logger.error(f"[ERROR] Failed to initialize S3 client: {e}")
            self.s3_client = None
    
    def _clean_json_response(self, raw_response: str) -> str:
        """
        Clean and extract JSON from AI response that may contain markdown or preamble text.
        
        Args:
            raw_response: Raw response from AI that may contain markdown blocks and preamble
            
        Returns:
            Cleaned JSON string ready for parsing
        """
        logger.info("[INFO] Starting JSON cleaning from AI response")
        logger.debug(f"[DEBUG] Raw response length: {len(raw_response)} chars")
        
        cleaned = raw_response.strip()
        
        # Step 1: Remove any markdown code blocks
        if '```json' in cleaned:
            logger.info("[INFO] Detected ```json markdown block, extracting...")
            # Extract JSON from ```json ... ``` blocks
            start_idx = cleaned.find('```json') + 7
            end_idx = cleaned.find('```', start_idx)
            if end_idx > start_idx:
                cleaned = cleaned[start_idx:end_idx].strip()
                logger.info(f"[INFO] Extracted from markdown block: {len(cleaned)} chars")
        elif '```' in cleaned:
            logger.info("[INFO] Detected ``` markdown block, extracting...")
            # Extract JSON from ``` ... ``` blocks
            start_idx = cleaned.find('```') + 3
            end_idx = cleaned.find('```', start_idx)
            if end_idx > start_idx:
                cleaned = cleaned[start_idx:end_idx].strip()
                logger.info(f"[INFO] Extracted from markdown block: {len(cleaned)} chars")
        
        # Step 2: Remove any preamble text before JSON
        # Find the first { or [ character (start of JSON)
        json_start = -1
        for i, char in enumerate(cleaned):
            if char in ['{', '[']:
                json_start = i
                break
        
        if json_start > 0:
            logger.info(f"[INFO] Removing {json_start} chars of preamble text")
            cleaned = cleaned[json_start:].strip()
        
        # Step 3: Remove any text after JSON ends
        # Find the last } or ] character (end of JSON)
        json_end = -1
        for i in range(len(cleaned) - 1, -1, -1):
            if cleaned[i] in ['}', ']']:
                json_end = i + 1
                break
        
        if json_end > 0 and json_end < len(cleaned):
            logger.info(f"[INFO] Removing {len(cleaned) - json_end} chars of text after JSON")
            cleaned = cleaned[:json_end].strip()
        
        # Final cleanup
        cleaned = cleaned.strip()
        logger.info(f"[INFO] Final cleaned JSON length: {len(cleaned)} chars")
        logger.debug(f"[DEBUG] First 200 chars: {cleaned[:200]}")
        
        return cleaned
    
    def upload_to_s3(self, file_path: str, user_id: str = "ocr-user") -> Dict:
        """Upload document to S3"""
        if not self.s3_client:
            logger.error("[ERROR] S3 client is not initialized")
            return {
                'success': False,
                'error': 'S3 client not available. Please check S3 configuration and database connection.'
            }
        
        try:
            file_name = os.path.basename(file_path)
            logger.info(f"[INFO] Attempting to upload file to S3: {file_name}")
            
            result = self.s3_client.upload(file_path, user_id, file_name)
            
            if not isinstance(result, dict):
                logger.error(f"[ERROR] S3 upload returned invalid response type: {type(result)}")
                return {
                    'success': False,
                    'error': f'S3 client returned invalid response type: {type(result).__name__}'
                }
            
            if result.get('success'):
                logger.info(f"[OK] File uploaded to S3: {file_name}")
                return result
            else:
                error_msg = result.get('error', 'Unknown S3 upload error')
                logger.error(f"[ERROR] S3 upload failed: {error_msg}")
                
                # Provide more specific error messages based on the error type
                if 'timeout' in error_msg.lower() or 'connection' in error_msg.lower():
                    error_msg = 'S3 service is currently unavailable. Please try again later.'
                elif 'not found' in error_msg.lower() or '404' in error_msg.lower():
                    error_msg = 'S3 service endpoint not found. Please check configuration.'
                elif 'permission' in error_msg.lower() or 'unauthorized' in error_msg.lower():
                    error_msg = 'S3 service access denied. Please check credentials.'
                elif 'network' in error_msg.lower() or 'dns' in error_msg.lower():
                    error_msg = 'Network error connecting to S3 service. Please check your internet connection.'
                
                return {
                    'success': False,
                    'error': error_msg,
                    'original_error': result.get('error', 'Unknown error')
                }
                
        except Exception as e:
            error_msg = str(e)
            logger.error(f"[ERROR] S3 upload exception: {type(e).__name__}: {error_msg}")
            import traceback
            logger.error(f"[ERROR] Traceback: {traceback.format_exc()}")
            
            # Provide more specific error messages based on exception type
            if 'timeout' in error_msg.lower():
                error_msg = 'S3 upload timed out. The service may be slow or unavailable.'
            elif 'connection' in error_msg.lower():
                error_msg = 'Cannot connect to S3 service. Please check your network connection.'
            elif 'name resolution' in error_msg.lower() or 'dns' in error_msg.lower():
                error_msg = 'S3 service hostname cannot be resolved. Please check DNS settings.'
            elif 'permission' in error_msg.lower():
                error_msg = 'Permission denied accessing S3 service. Please check credentials.'
            else:
                error_msg = f'S3 upload failed: {type(e).__name__}: {error_msg}'
            
            return {
                'success': False,
                'error': error_msg,
                'original_error': str(e)
            }
    
    def extract_text_from_pdf(self, file_path: str) -> Dict:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text_content = ""
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            
            # Calculate confidence based on text length and page count
            confidence = min(95.0, max(60.0, (len(text_content) / (page_count * 100)) * 100))
            
            return {
                'success': True,
                'text': text_content,
                'confidence': confidence,
                'page_count': page_count,
                'engine': 'PyMuPDF'
            }
            
        except Exception as e:
            logger.error(f"[ERROR] PDF text extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_text_from_image(self, file_path: str) -> Dict:
        """Extract text from image using Tesseract OCR"""
        try:
            # Open image and perform OCR
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            # Get confidence data if available
            try:
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidences) / len(confidences) if confidences else 60.0
            except:
                avg_confidence = 60.0
            
            return {
                'success': True,
                'text': text,
                'confidence': avg_confidence,
                'engine': 'Tesseract'
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Image OCR error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_text_from_excel(self, file_path: str) -> Dict:
        """Extract text from Excel file"""
        try:
            # Read Excel file
            excel_file = pd.ExcelFile(file_path)
            text_content = ""
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_content += f"Sheet: {sheet_name}\n"
                text_content += df.to_string() + "\n\n"
            
            return {
                'success': True,
                'text': text_content,
                'confidence': 90.0,  # Excel files have high confidence
                'engine': 'Pandas'
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Excel text extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def extract_text_from_docx(self, file_path: str) -> Dict:
        """Extract text from DOCX file"""
        try:
            try:
                from docx import Document
            except ImportError:
                logger.error("[ERROR] python-docx not installed. Install with: pip install python-docx")
                return {
                    'success': False,
                    'error': 'python-docx library not installed'
                }
            
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
                text_content += "\n"
            
            # Calculate confidence based on text length
            confidence = min(95.0, max(70.0, (len(text_content) / 100) * 100)) if text_content else 70.0
            
            return {
                'success': True,
                'text': text_content,
                'confidence': confidence,
                'engine': 'python-docx'
            }
            
        except Exception as e:
            logger.error(f"[ERROR] DOCX text extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def perform_ocr(self, file_path: str) -> Dict:
        """Perform OCR on document based on file type"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return self.extract_text_from_image(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self.extract_text_from_excel(file_path)
        elif file_ext in ['.docx']:
            return self.extract_text_from_docx(file_path)
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_ext}'
            }
    
    def extract_sla_data_with_ai(self, ocr_text: str) -> Dict:
        """Extract SLA data using AI (LLaMA)"""
        try:
            prompt = f"""
            Extract SLA (Service Level Agreement) information from the following document text.
            Return the data in JSON format with these exact fields:
            
            {{
                "sla_name": "string",
                "vendor_id": "string",
                "contract_id": "string", 
                "sla_type": "string (e.g., AVAILABILITY, PERFORMANCE, SECURITY)",
                "effective_date": "YYYY-MM-DD",
                "expiry_date": "YYYY-MM-DD",
                "status": "string (PENDING, ACTIVE, EXPIRED)",
                "business_service_impacted": "string",
                "reporting_frequency": "string (e.g., monthly, quarterly)",
                "baseline_period": "string",
                "improvement_targets": "JSON object",
                "penalty_threshold": "decimal number",
                "credit_threshold": "decimal number", 
                "measurement_methodology": "string",
                "exclusions": "string",
                "force_majeure_clauses": "string",
                "compliance_framework": "string",
                "audit_requirements": "string",
                "document_versioning": "string",
                "compliance_score": "decimal number",
                "priority": "string (LOW, MEDIUM, HIGH)",
                "approval_status": "string (PENDING, APPROVED, REJECTED)",
                "metrics": [
                    {{
                        "metric_name": "string",
                        "target_value": "string",
                        "measurement_unit": "string",
                        "measurement_frequency": "string",
                        "penalty_clause": "string",
                        "measurement_methodology": "string"
                    }}
                ]
            }}
            
            Document text:
            {ocr_text}
            
            Extract the information and return only valid JSON:
            """
            
            # Call LLaMA for extraction
            extracted_json = simple_ollama(prompt, max_tokens=1000)
            
            if not extracted_json:
                return {
                    'success': False,
                    'error': 'AI extraction failed - no response from LLaMA'
                }
            
            # Clean and parse the response
            try:
                # Use the helper function to clean JSON response
                cleaned = self._clean_json_response(extracted_json)
                
                # Parse JSON
                extracted_data = json.loads(cleaned)
                
                return {
                    'success': True,
                    'data': extracted_data,
                    'confidence': 85.0
                }
                
            except json.JSONDecodeError as e:
                logger.error(f"[ERROR] JSON parsing error: {e}")
                logger.error(f"Raw AI response: {extracted_json}")
                
                # Try to extract basic information manually
                return self._extract_basic_sla_info(ocr_text)
                
        except Exception as e:
            logger.error(f"[ERROR] AI extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def extract_contract_data_with_ai(self, ocr_text: str) -> Dict:
        """Extract Contract details (matching CreateContract.vue OCR fields) using AI (LLaMA).

        This does not modify SLA extraction. It returns only the fields that
        the frontend expects for contract OCR auto-fill.
        """
        try:
            # Define the exact contract fields expected by the frontend
            expected_schema = {
                "contract_title": "string",
                "contract_number": "string",
                "contract_type": "string (MASTER_AGREEMENT | SOW | PURCHASE_ORDER | SERVICE_AGREEMENT | LICENSE | NDA | AMENDMENT)",
                "contract_kind": "string (MAIN | SUBCONTRACT | AMENDMENT)",
                "contract_category": "string (goods | services | technology | consulting | others)",

                "vendor_name": "string",

                "contract_value": "number",
                "currency": "string (USD | EUR | GBP | CAD | AUD)",
                "liability_cap": "number",

                "start_date": "YYYY-MM-DD",
                "end_date": "YYYY-MM-DD",
                "notice_period_days": "integer",
                "auto_renewal": "boolean",
                "renewal_terms": "string",

                "status": "string (PENDING_ASSIGNMENT | DRAFT | ACTIVE)",
                "workflow_stage": "string (under_review | draft | submitted)",
                "priority": "string (low | medium | high | urgent)",
                "compliance_status": "string (under_review | compliant | non_compliant | exempt)",

                "dispute_resolution_method": "string (negotiation | mediation | arbitration | litigation)",
                "governing_law": "string",
                "termination_clause_type": "string (convenience | cause | both | none)",
                "contract_risk_score": "number",
                "compliance_framework": "string (SOC2 | GDPR | CCPA | ISO27001 | PCI DSS | HIPAA | Other)",
                
                "amendment_number": "string (for amendments only)",
                "amendment_date": "YYYY-MM-DD (for amendments only)",
                "effective_date": "YYYY-MM-DD (for amendments only)",
                "amendment_type": "string (financial | scope | timeline | terms | other - for amendments only)",
                "financial_impact": "number (for amendments only)",
                "impact_type": "string (increase | decrease | no_change - for amendments only)",
                "amendment_reason": "string (for amendments only)",
                "changes_summary": "string (for amendments only)",

                "terms": [
                    {
                        "category": "string (Payment | Delivery | Performance | Liability | Termination | Intellectual Property | Confidentiality)",
                        "title": "string",
                        "text": "string",
                        "is_standard": "boolean"
                    }
                ],
                "clauses": [
                    {
                        "name": "string",
                        "type": "string (standard | renewal | termination)",
                        "text": "string",
                        "notice_period_days": "integer or null",
                        "auto_renew": "boolean or null",
                        "renewal_terms": "string",
                        "termination_notice_period": "integer or null",
                        "early_termination_fee": "number or null",
                        "termination_conditions": "string"
                    }
                ]
            }

            example_schema = json.dumps(expected_schema, indent=2)

            prompt = f"""
            You are an expert contracts analyst. Extract ALL contract/amendment information from the following text.
            Be comprehensive and extract every field you can identify. Fill ALL fields in the schema.
            
            Schema to fill:
            {example_schema}

            EXTRACTION INSTRUCTIONS:
            1. Extract ALL basic contract fields (title, number, type, vendor, dates, value, etc.)
            2. For AMENDMENTS: Extract amendment_number, amendment_date, effective_date, amendment_type, financial_impact, impact_type, amendment_reason, changes_summary
            3. Extract ALL terms and clauses - be EXTREMELY thorough, do not skip any
            4. For renewal clauses: set type='renewal' and include notice_period_days, auto_renew, renewal_terms
            5. For termination clauses: set type='termination' and include termination_notice_period, early_termination_fee, termination_conditions
            6. For other clauses: set type='standard'
            7. Extract ALL terms with proper categories (Payment, Delivery, Performance, Liability, Termination, Intellectual Property, Confidentiality)
            
            DATA FORMATTING:
            - Dates: YYYY-MM-DD format
            - Numbers: numeric values only (no currency symbols)
            - Booleans: true/false only
            - Strings: clean text without extra formatting
            
            FIELD MAPPING:
            - contract_title: Extract from document title/header
            - contract_number: Look for contract ID, agreement number, etc.
            - contract_type: Determine from document type (MASTER_AGREEMENT, SOW, AMENDMENT, etc.)
            - vendor_name: Extract vendor/company name
            - contract_value: Extract monetary amounts
            - start_date/end_date: Extract contract period dates
            - renewal_terms: Extract from renewal sections
            - termination_clause_type: Determine type (convenience, cause, both)
            - governing_law: Extract jurisdiction/law
            - dispute_resolution_method: Extract resolution method
            - compliance_framework: Extract compliance standards (SOC2, GDPR, etc.)
            
            AMENDMENT FIELDS (if this is an amendment):
            - amendment_number: Extract amendment ID (e.g., "Amendment No. 1", "AMEND-001")
            - amendment_date: Extract date amendment was created/signed
            - effective_date: Extract when amendment becomes effective
            - amendment_type: Classify as financial, scope, timeline, terms, or other
            - financial_impact: Extract monetary change amount
            - impact_type: Determine if increase, decrease, or no_change
            - amendment_reason: Extract reason for amendment
            - changes_summary: Extract summary of changes made
            
            TERMS EXTRACTION - CRITICAL:
            - Extract EVERY SINGLE term and condition from the document
            - Look in sections like "Terms and Conditions", "Obligations", "Responsibilities"
            - Categorize each term properly: Payment, Delivery, Performance, Liability, Termination, Intellectual Property, Confidentiality
            - Each term MUST have: category, title, text, is_standard
            - Example: If you see "Payment shall be made within 30 days", create a Payment term
            - Do NOT skip terms - extract ALL of them
            
            CLAUSES EXTRACTION - CRITICAL:
            - Extract EVERY SINGLE clause from the document
            - Classify renewal clauses as type='renewal'
            - Classify termination clauses as type='termination'
            - Include all clause details (notice periods, fees, conditions)
            - Each clause MUST have: name, type, text
            - Do NOT skip clauses - extract ALL of them
            
            RENEWAL CLAUSES - Look for sections like "Renewal Terms", "4. Renewal Terms", "Auto-Renewal":
            - Extract EACH sub-point as a SEPARATE renewal clause
            - Notice Period: Extract days (e.g., "30 days notice" → notice_period_days: 30)
            - Term Length: Extract renewal period (e.g., "12-month periods")
            - Pricing Adjustment: Extract pricing change terms
            - Auto Extension: Set auto_renew: true if automatic extension mentioned
            - Create multiple clauses if there are multiple renewal terms
            
            TERMINATION CLAUSES - Look for sections like "Termination", "5. Termination", "Early Termination":
            - Extract EACH sub-point as a SEPARATE termination clause
            - Notice Period: Extract days (e.g., "60 days notice" → termination_notice_period: 60)
            - For Cause: Extract immediate termination conditions
            - Early Fee: Extract percentage or amount (e.g., "25%" → early_termination_fee: 25)
            - Transition Support: Extract transition period and requirements
            - Create multiple clauses if there are multiple termination terms
            
            STRUCTURED EXTRACTION:
            - When you see numbered sections (1., 2., 3.), extract each as a separate term/clause
            - When you see bullet points (•, -, *), extract each as a separate term/clause
            - Each term/clause should be granular - don't combine multiple points
            - Be thorough - extract EVERYTHING, not just summaries

            Contract text:
            {ocr_text}

            IMPORTANT OUTPUT FORMAT:
            - Return ONLY valid JSON
            - NO preamble text or explanations
            - NO markdown code blocks
            - Just the raw JSON object starting with {{ and ending with }}
            - Include AT LEAST 3-5 terms and 3-5 clauses in your output
            
            Extract EVERYTHING and return complete JSON with no explanations, preamble, or markdown formatting.
            """

            # Increase max_tokens to allow for complete extraction of all terms and clauses
            extracted_json = simple_ollama(prompt, max_tokens=5000)

            if not extracted_json:
                return {
                    'success': False,
                    'error': 'AI extraction failed - no response from LLaMA'
                }

            try:
                # Use the helper function to clean JSON response
                cleaned = self._clean_json_response(extracted_json)
                
                # Parse JSON
                data = json.loads(cleaned)
                logger.info("[SUCCESS] JSON parsed successfully")

                # Normalize to dict: if model returned a list, pick first object
                if isinstance(data, list):
                    data = data[0] if data and isinstance(data[0], dict) else {}

                # Normalize terms to {category,title,text,is_standard}
                terms = data.get('terms')
                if isinstance(terms, list):
                    normalized_terms = []
                    for t in terms:
                        if isinstance(t, dict):
                            normalized_terms.append({
                                'category': t.get('category') or t.get('term_category') or '',
                                'title': t.get('title') or t.get('term_title') or '',
                                'text': t.get('text') or t.get('term_text') or '',
                                'is_standard': bool(t.get('is_standard', False))
                            })
                        elif isinstance(t, str):
                            normalized_terms.append({'category': 'General', 'title': 'General', 'text': t, 'is_standard': False})
                    data['terms'] = normalized_terms

                # Normalize clauses to {name,type,text,...}
                clauses = data.get('clauses')
                if isinstance(clauses, list):
                    normalized_clauses = []
                    for c in clauses:
                        if isinstance(c, dict):
                            normalized_clauses.append({
                                'name': c.get('name') or c.get('clause_name') or 'Clause',
                                'type': c.get('type') or c.get('clause_type') or 'standard',
                                'text': c.get('text') or c.get('clause_text') or '',
                                'notice_period_days': c.get('notice_period_days'),
                                'auto_renew': c.get('auto_renew'),
                                'renewal_terms': c.get('renewal_terms'),
                                'termination_notice_period': c.get('termination_notice_period'),
                                'early_termination_fee': c.get('early_termination_fee'),
                                'termination_conditions': c.get('termination_conditions'),
                            })
                        elif isinstance(c, str):
                            normalized_clauses.append({'name': 'Clause', 'type': 'standard', 'text': c})
                    data['clauses'] = normalized_clauses

                return {
                    'success': True,
                    'data': data,
                    'confidence': 85.0
                }
            except json.JSONDecodeError as e:
                logger.error(f"[ERROR] Contract JSON parsing error: {e}")
                logger.error(f"Raw AI response: {extracted_json}")
                # Return error instead of fallback data to avoid data mismapping
                return {
                    'success': False,
                    'error': f'Failed to parse AI response: {str(e)}',
                    'raw_response': extracted_json[:500]  # First 500 chars for debugging
                }
        except Exception as e:
            logger.error(f"[ERROR] Contract AI extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def extract_sla_data_with_payload(self, text: str, payload_headers: Dict) -> Dict:
        """Extract SLA data from text using provided payload headers as template"""
        try:
            if not text.strip():
                return {
                    'success': False,
                    'error': 'No text provided for extraction'
                }
            
            if not payload_headers:
                return {
                    'success': False,
                    'error': 'No payload headers provided'
                }
            
            # Create AI prompt using the payload headers as template
            payload_fields = list(payload_headers.keys())
            payload_example = json.dumps(payload_headers, indent=2)
            
            prompt = f"""
            Extract SLA (Service Level Agreement) information from the following document text.
            Use the provided payload template to structure your response.
            
            Payload Template (extract data to fill these fields):
            {payload_example}
            
            Document text to analyze:
            {text}
            
            Instructions:
            1. Extract information from the document text
            2. Fill the payload fields with extracted data
            3. If a field is not found in the document, use appropriate default values
            4. For dates, use YYYY-MM-DD format
            5. For numeric values, provide actual numbers
            6. For arrays (like metrics), extract all relevant items from the document
            7. Return ONLY the JSON object with filled data
            
            Extracted data in JSON format:
            """
            
            # Call LLaMA for extraction with payload template
            extracted_json = simple_ollama(prompt, max_tokens=1500)
            
            if not extracted_json:
                # Fallback: create mock data based on the payload template
                logger.warning("[WARNING] Ollama timeout - using fallback extraction")
                fallback_data = self._create_fallback_data(text, payload_headers)
                
                return {
                    'success': True,
                    'data': fallback_data,
                    'confidence': 0.3,
                    'warning': 'AI extraction timed out, using fallback extraction'
                }
            
            # Try to parse the JSON response
            try:
                # Use the helper function to clean JSON response
                cleaned_response = self._clean_json_response(extracted_json)
                
                extracted_data = json.loads(cleaned_response)
                
                # Validate that we have the expected fields from payload
                validated_data = self._validate_extracted_data(extracted_data, payload_headers)
                
                # Calculate confidence based on extracted fields
                confidence = self._calculate_extraction_confidence(validated_data)
                
                return {
                    'success': True,
                    'data': validated_data,
                    'confidence': confidence,
                    'payload_fields': payload_fields
                }
                
            except json.JSONDecodeError as e:
                logger.error(f"[ERROR] Failed to parse AI response as JSON: {e}")
                logger.error(f"[ERROR] Raw AI response: {extracted_json}")
                
                # Fallback: use payload template with default values
                fallback_data = self._create_fallback_data(text, payload_headers)
                
                return {
                    'success': True,
                    'data': fallback_data,
                    'confidence': 0.2,
                    'warning': 'AI response could not be parsed, using fallback extraction'
                }
                
        except Exception as e:
            logger.error(f"[ERROR] SLA extraction with payload error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _validate_extracted_data(self, extracted_data: Dict, payload_headers: Dict) -> Dict:
        """Validate and merge extracted data with payload template"""
        validated_data = {}
        
        for field, default_value in payload_headers.items():
            if field in extracted_data:
                validated_data[field] = extracted_data[field]
            else:
                validated_data[field] = default_value
        
        return validated_data
    
    def _calculate_extraction_confidence(self, data: Dict) -> float:
        """Calculate confidence score based on extracted fields"""
        total_fields = len(data)
        filled_fields = sum(1 for value in data.values() 
                          if value is not None and value != '' and value != [] and value != {})
        
        if total_fields == 0:
            return 0.0
        
        confidence = (filled_fields / total_fields) * 100
        return min(confidence, 100.0)
    
    def _create_fallback_data(self, text: str, payload_headers: Dict) -> Dict:
        """Create fallback data when AI extraction fails"""
        fallback_data = {}
        
        # Copy the payload template first
        for field, default_value in payload_headers.items():
            fallback_data[field] = default_value
        
        # Try to extract some basic information using patterns
        import re
        
        # Look for SLA name in the text
        if 'sla_name' in payload_headers:
            sla_pattern = r'(?i)(service level agreement|sla).*?([A-Za-z][^.\n]{10,50})'
            sla_match = re.search(sla_pattern, text)
            if sla_match:
                fallback_data['sla_name'] = sla_match.group(2).strip()
            else:
                # Provide a sample SLA name
                fallback_data['sla_name'] = "Sample Service Level Agreement"
        
        # Look for vendor/contract IDs
        if 'vendor_id' in payload_headers:
            fallback_data['vendor_id'] = "1"
        if 'contract_id' in payload_headers:
            fallback_data['contract_id'] = "1"
        
        # Look for SLA type
        if 'sla_type' in payload_headers:
            if 'availability' in text.lower():
                fallback_data['sla_type'] = "AVAILABILITY"
            elif 'performance' in text.lower():
                fallback_data['sla_type'] = "PERFORMANCE"
            else:
                fallback_data['sla_type'] = "GENERAL"
        
        # Look for dates
        date_pattern = r'\d{4}-\d{2}-\d{2}'
        dates = re.findall(date_pattern, text)
        if dates and 'effective_date' in payload_headers:
            fallback_data['effective_date'] = dates[0]
        elif 'effective_date' in payload_headers:
            fallback_data['effective_date'] = "2024-01-01"
            
        if len(dates) > 1 and 'expiry_date' in payload_headers:
            fallback_data['expiry_date'] = dates[1]
        elif 'expiry_date' in payload_headers:
            fallback_data['expiry_date'] = "2024-12-31"
        
        # Look for business service
        if 'business_service_impacted' in payload_headers:
            if 'database' in text.lower():
                fallback_data['business_service_impacted'] = "Database Services"
            elif 'application' in text.lower():
                fallback_data['business_service_impacted'] = "Application Services"
            else:
                fallback_data['business_service_impacted'] = "IT Services"
        
        # Look for compliance framework
        if 'compliance_framework' in payload_headers:
            if 'iso' in text.lower():
                fallback_data['compliance_framework'] = "ISO 27001"
            elif 'sox' in text.lower():
                fallback_data['compliance_framework'] = "SOX Compliance"
            else:
                fallback_data['compliance_framework'] = "Standard Compliance"
        
        # Look for metrics
        if 'metrics' in payload_headers and isinstance(payload_headers['metrics'], list):
            metrics = []
            if 'availability' in text.lower():
                metrics.append({
                    "metric_name": "Service Availability",
                    "target_value": "99.9",
                    "measurement_unit": "%",
                    "measurement_frequency": "MONTHLY",
                    "penalty_clause": "Service credit for downtime",
                    "measurement_methodology": "Automated monitoring"
                })
            if 'response' in text.lower() or 'time' in text.lower():
                metrics.append({
                    "metric_name": "Response Time",
                    "target_value": "200",
                    "measurement_unit": "ms",
                    "measurement_frequency": "DAILY",
                    "penalty_clause": "Performance penalty",
                    "measurement_methodology": "Real-time monitoring"
                })
            if not metrics:
                metrics.append({
                    "metric_name": "Service Quality",
                    "target_value": "95",
                    "measurement_unit": "%",
                    "measurement_frequency": "MONTHLY",
                    "penalty_clause": "Standard penalty",
                    "measurement_methodology": "Quality assessment"
                })
            fallback_data['metrics'] = metrics
        
        return fallback_data
    
    def _extract_basic_sla_info(self, ocr_text: str) -> Dict:
        """Fallback method to extract basic SLA information"""
        try:
            # Simple keyword-based extraction
            text_lower = ocr_text.lower()
            
            # Extract basic information using regex patterns
            import re
            
            # Find dates
            date_pattern = r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}'
            dates = re.findall(date_pattern, ocr_text)
            
            # Find percentages
            percent_pattern = r'(\d+\.?\d*)\s*%'
            percentages = re.findall(percent_pattern, ocr_text)
            
            # Find SLA type
            sla_type = "AVAILABILITY"  # default
            if "performance" in text_lower:
                sla_type = "PERFORMANCE"
            elif "security" in text_lower:
                sla_type = "SECURITY"
            
            # Find priority
            priority = "MEDIUM"  # default
            if "high" in text_lower:
                priority = "HIGH"
            elif "low" in text_lower:
                priority = "LOW"
            
            basic_data = {
                "sla_name": "Extracted SLA Document",
                "vendor_id": "",
                "contract_id": "",
                "sla_type": sla_type,
                "effective_date": dates[0] if dates else "",
                "expiry_date": dates[1] if len(dates) > 1 else "",
                "status": "PENDING",
                "business_service_impacted": "",
                "reporting_frequency": "monthly",
                "baseline_period": "",
                "improvement_targets": {},
                "penalty_threshold": "",
                "credit_threshold": "",
                "measurement_methodology": "",
                "exclusions": "",
                "force_majeure_clauses": "",
                "compliance_framework": "",
                "audit_requirements": "",
                "document_versioning": "v1.0",
                "compliance_score": "",
                "priority": priority,
                "approval_status": "PENDING",
                "metrics": []
            }
            
            return {
                'success': True,
                'data': basic_data,
                'confidence': 40.0
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Basic extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def process_document(self, file_path: str, user_id: str = "ocr-user") -> Dict:
        """Complete document processing pipeline"""
        try:
            logger.info(f"[EMOJI] Starting document processing: {file_path}")
            
            # Step 1: Upload to S3 (continue even if this fails)
            upload_result = self.upload_to_s3(file_path, user_id)
            upload_success = upload_result.get('success', False)
            
            if not upload_success:
                logger.warning(f"[WARNING] S3 upload failed, continuing with OCR processing: {upload_result.get('error', 'Unknown error')}")
            
            # Step 2: Perform OCR
            ocr_result = self.perform_ocr(file_path)
            if not ocr_result['success']:
                return {
                    'success': False,
                    'error': f"OCR failed: {ocr_result['error']}",
                    'upload_info': upload_result if upload_success else None
                }
            
            # Step 3A: Extract SLA structured data with AI (existing flow)
            extraction_result = self.extract_sla_data_with_ai(ocr_result['text'])

            # Step 3B: Extract Contract structured data for frontend auto-fill
            contract_extraction = self.extract_contract_data_with_ai(ocr_result['text'])
            
            return {
                'success': True,
                'upload_info': upload_result if upload_success else {
                    'success': False,
                    'error': upload_result.get('error', 'S3 upload failed'),
                    'file_info': None
                },
                'ocr_result': ocr_result,
                'extraction_result': extraction_result,
                'contract_extraction': contract_extraction,
                # Expose top-level 'data' for frontend convenience, normalized to dict
                'data': (contract_extraction.get('data') if isinstance(contract_extraction, dict) else None),
                'message': 'Document processed successfully' + (' (S3 upload failed)' if not upload_success else '')
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Document processing error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def process_contract_only(self, file_path: str, user_id: str = "ocr-user") -> Dict:
        """Contract-only document processing pipeline (no SLA extraction)."""
        try:
            logger.info(f"[EMOJI] Starting contract-only document processing: {file_path}")
            
            # Step 1: Upload to S3 (continue even if this fails)
            upload_result = self.upload_to_s3(file_path, user_id)
            upload_success = upload_result.get('success', False)
            
            if not upload_success:
                logger.warning(f"[WARNING] S3 upload failed, continuing with OCR processing: {upload_result.get('error', 'Unknown error')}")
            
            # Step 2: Perform OCR
            ocr_result = self.perform_ocr(file_path)
            if not ocr_result['success']:
                return {
                    'success': False,
                    'error': f"OCR failed: {ocr_result['error']}",
                    'upload_info': upload_result if upload_success else None
                }
            
            # Step 3: Extract ONLY contract data (skip SLA)
            contract_extraction = self.extract_contract_data_with_ai(ocr_result['text'])
            
            # Minimal SLA structure for compatibility (SLA unchanged and skipped)
            minimal_sla_result = {
                'success': False,
                'data': {},
                'confidence': 0.0,
                'message': 'SLA extraction skipped (contract-only mode)'
            }
            
            return {
                'success': True,
                'upload_info': upload_result if upload_success else {
                    'success': False,
                    'error': upload_result.get('error', 'S3 upload failed'),
                    'file_info': None
                },
                'ocr_result': ocr_result,
                'extraction_result': minimal_sla_result,
                'contract_extraction': contract_extraction,
                'data': contract_extraction.get('data') if isinstance(contract_extraction, dict) else {},
                'message': 'Contract document processed successfully (SLA extraction skipped)' + (' (S3 upload failed)' if not upload_success else '')
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Contract-only processing error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def extract_bcp_data_with_ai(self, ocr_text: str) -> Dict:
        """Extract unified BCP/DRP data using AI (LLaMA) - all fields for all plan types"""
        try:
            prompt = f"""
            Extract Business Continuity Plan (BCP) or Disaster Recovery Plan (DRP) information from the following document text.
            Return the data in JSON format with these exact fields (use null for fields not found):
            
            {{
                "plan_id": null,
                "purpose_scope": "string or null",
                "regulatory_references": ["string"] or null,
                "critical_services": ["string"] or null,
                "dependencies_internal": ["string"] or null,
                "dependencies_external": ["string"] or null,
                "risk_assessment_summary": "string or null",
                "bia_summary": "string or null",
                "rto_targets": {{"service": "time"}} or null,
                "rpo_targets": {{"service": "time"}} or null,
                "critical_systems": ["string"] or null,
                "critical_applications": ["string"] or null,
                "databases_list": ["string"] or null,
                "supporting_infrastructure": ["string"] or null,
                "third_party_services": ["string"] or null,
                "incident_types": ["string"] or null,
                "alternate_work_locations": ["string"] or null,
                "communication_plan_internal": "string or null",
                "communication_plan_bank": "string or null",
                "roles_responsibilities": ["string"] or null,
                "training_testing_schedule": "string or null",
                "maintenance_review_cycle": "string or null",
                "disaster_scenarios": ["string"] or null,
                "disaster_declaration_process": "string or null",
                "data_backup_strategy": "string or null",
                "recovery_site_details": "string or null",
                "failover_procedures": "string or null",
                "failback_procedures": "string or null",
                "network_recovery_steps": "string or null",
                "application_restoration_order": ["string"] or null,
                "testing_validation_schedule": "string or null"
            }}
            
            Document text:
            {ocr_text}
            
            Extract the information and return only valid JSON:
            """
            
            # Call LLaMA for extraction
            extracted_json = simple_ollama(prompt, max_tokens=2000)
            
            if not extracted_json:
                return {
                    'success': False,
                    'error': 'AI extraction failed - no response from LLaMA'
                }
            
            # Clean and parse the response
            try:
                # Use the helper function to clean JSON response
                cleaned = self._clean_json_response(extracted_json)
                
                # Parse JSON
                extracted_data = json.loads(cleaned)
                
                return {
                    'success': True,
                    'data': extracted_data,
                    'confidence': 85.0
                }
                
            except json.JSONDecodeError as e:
                logger.error(f"[ERROR] BCP JSON parsing error: {e}")
                logger.error(f"Raw AI response: {extracted_json}")
                
                # Try to extract basic information manually
                return self._extract_basic_bcp_info(ocr_text)
                
        except Exception as e:
            logger.error(f"[ERROR] BCP AI extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def extract_drp_data_with_ai(self, ocr_text: str) -> Dict:
        """Extract unified BCP/DRP data using AI (LLaMA) - all fields for all plan types"""
        try:
            prompt = f"""
            Extract Business Continuity Plan (BCP) or Disaster Recovery Plan (DRP) information from the following document text.
            Return the data in JSON format with these exact fields (use null for fields not found):
            
            {{
                "plan_id": null,
                "purpose_scope": "string or null",
                "regulatory_references": ["string"] or null,
                "critical_services": ["string"] or null,
                "dependencies_internal": ["string"] or null,
                "dependencies_external": ["string"] or null,
                "risk_assessment_summary": "string or null",
                "bia_summary": "string or null",
                "rto_targets": {{"service": "time"}} or null,
                "rpo_targets": {{"service": "time"}} or null,
                "critical_systems": ["string"] or null,
                "critical_applications": ["string"] or null,
                "databases_list": ["string"] or null,
                "supporting_infrastructure": ["string"] or null,
                "third_party_services": ["string"] or null,
                "incident_types": ["string"] or null,
                "alternate_work_locations": ["string"] or null,
                "communication_plan_internal": "string or null",
                "communication_plan_bank": "string or null",
                "roles_responsibilities": ["string"] or null,
                "training_testing_schedule": "string or null",
                "maintenance_review_cycle": "string or null",
                "disaster_scenarios": ["string"] or null,
                "disaster_declaration_process": "string or null",
                "data_backup_strategy": "string or null",
                "recovery_site_details": "string or null",
                "failover_procedures": "string or null",
                "failback_procedures": "string or null",
                "network_recovery_steps": "string or null",
                "application_restoration_order": ["string"] or null,
                "testing_validation_schedule": "string or null"
            }}
            
            Document text:
            {ocr_text}
            
            Extract the information and return only valid JSON:
            """
            
            # Call LLaMA for extraction
            extracted_json = simple_ollama(prompt, max_tokens=2000)
            
            if not extracted_json:
                return {
                    'success': False,
                    'error': 'AI extraction failed - no response from LLaMA'
                }
            
            # Clean and parse the response
            try:
                # Use the helper function to clean JSON response
                cleaned = self._clean_json_response(extracted_json)
                
                # Parse JSON
                extracted_data = json.loads(cleaned)
                
                return {
                    'success': True,
                    'data': extracted_data,
                    'confidence': 85.0
                }
                
            except json.JSONDecodeError as e:
                logger.error(f"[ERROR] DRP JSON parsing error: {e}")
                logger.error(f"Raw AI response: {extracted_json}")
                
                # Try to extract basic information manually
                return self._extract_basic_drp_info(ocr_text)
                
        except Exception as e:
            logger.error(f"[ERROR] DRP AI extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _extract_basic_bcp_info(self, ocr_text: str) -> Dict:
        """Fallback method to extract basic BCP information"""
        try:
            # Simple keyword-based extraction
            text_lower = ocr_text.lower()
            
            # Extract basic information using regex patterns
            import re
            
            # Find dates
            date_pattern = r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}'
            dates = re.findall(date_pattern, ocr_text)
            
            # Find percentages
            percent_pattern = r'(\d+\.?\d*)\s*%'
            percentages = re.findall(percent_pattern, ocr_text)
            
            basic_data = {
                "plan_id": None,
                "purpose_scope": "Business Continuity Plan for maintaining operations during disruptions",
                "regulatory_references": ["SOX", "Basel III", "PCI DSS"] if "sox" in text_lower or "basel" in text_lower else None,
                "critical_services": ["Payment Processing", "Customer Service"] if "payment" in text_lower else None,
                "dependencies_internal": ["IT Systems", "HR Department"] if "it" in text_lower else None,
                "dependencies_external": ["Cloud Provider", "Banking Partners"] if "cloud" in text_lower else None,
                "risk_assessment_summary": "Risk assessment covers cyber threats, natural disasters, and operational failures",
                "bia_summary": "Business Impact Analysis identifies critical business functions and recovery requirements",
                "rto_targets": {"Critical Services": "4h", "Support Services": "8h"},
                "rpo_targets": {"Critical Data": "15m", "Support Data": "1h"},
                "critical_systems": None,
                "critical_applications": None,
                "databases_list": None,
                "supporting_infrastructure": None,
                "third_party_services": None,
                "incident_types": ["Cyber Attack", "Natural Disaster", "System Failure"],
                "alternate_work_locations": ["Remote Work", "Backup Office"],
                "communication_plan_internal": "Internal communication follows hierarchical structure with multiple channels",
                "communication_plan_bank": "Customer notifications via website, mobile app, email, and SMS",
                "roles_responsibilities": ["Incident Commander", "Communication Lead", "Technical Lead"],
                "training_testing_schedule": "Annual training, quarterly exercises, semi-annual testing",
                "maintenance_review_cycle": "Quarterly review with annual comprehensive updates",
                "disaster_scenarios": None,
                "disaster_declaration_process": None,
                "data_backup_strategy": None,
                "recovery_site_details": None,
                "failover_procedures": None,
                "failback_procedures": None,
                "network_recovery_steps": None,
                "application_restoration_order": None,
                "testing_validation_schedule": None
            }
            
            return {
                'success': True,
                'data': basic_data,
                'confidence': 40.0
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Basic BCP extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def _extract_basic_drp_info(self, ocr_text: str) -> Dict:
        """Fallback method to extract basic DRP information"""
        try:
            # Simple keyword-based extraction
            text_lower = ocr_text.lower()
            
            # Extract basic information using regex patterns
            import re
            
            # Find dates
            date_pattern = r'\d{4}-\d{2}-\d{2}|\d{1,2}/\d{1,2}/\d{4}'
            dates = re.findall(date_pattern, ocr_text)
            
            # Find percentages
            percent_pattern = r'(\d+\.?\d*)\s*%'
            percentages = re.findall(percent_pattern, ocr_text)
            
            basic_data = {
                "plan_id": None,
                "purpose_scope": "Disaster Recovery Plan for rapid recovery of critical IT systems and infrastructure",
                "regulatory_references": ["SOX", "Basel III", "PCI DSS", "FFIEC"] if "sox" in text_lower or "basel" in text_lower else None,
                "critical_services": None,
                "dependencies_internal": None,
                "dependencies_external": None,
                "risk_assessment_summary": None,
                "bia_summary": None,
                "rto_targets": {"Critical Systems": "2h", "Applications": "4h", "Databases": "1h"},
                "rpo_targets": {"Critical Systems": "30m", "Applications": "1h", "Databases": "15m"},
                "critical_systems": ["Core Banking System", "Payment Gateway", "Database Servers"],
                "critical_applications": ["Loan Management System", "Trading Platform", "Customer Portal"],
                "databases_list": ["Customer Database", "Transaction Database", "Risk Database"],
                "supporting_infrastructure": ["Network", "Storage", "Servers", "Security"],
                "third_party_services": ["Cloud Provider", "SMS Gateway", "Email Service"],
                "incident_types": None,
                "alternate_work_locations": None,
                "communication_plan_internal": None,
                "communication_plan_bank": None,
                "roles_responsibilities": None,
                "training_testing_schedule": None,
                "maintenance_review_cycle": "Monthly review with quarterly updates",
                "disaster_scenarios": ["Data Center Failure", "Network Outage", "Cyber Attack"],
                "disaster_declaration_process": "Three-tier process: Level 1 (Minor), Level 2 (Major), Level 3 (Critical)",
                "data_backup_strategy": "Multi-tier backup with real-time replication and off-site storage",
                "recovery_site_details": "Primary DR site 50 miles away, Secondary DR site 200 miles away",
                "failover_procedures": "Automated failover with manual verification and stakeholder notification",
                "failback_procedures": "Primary system validation, data sync, service migration, performance testing",
                "network_recovery_steps": "Assess damage, activate backup circuits, configure routing, test connectivity",
                "application_restoration_order": ["Core Banking", "Payment Gateway", "Customer Portal", "Risk Management"],
                "testing_validation_schedule": "Monthly DR testing with annual comprehensive exercise"
            }
            
            return {
                'success': True,
                'data': basic_data,
                'confidence': 40.0
            }
            
        except Exception as e:
            logger.error(f"[ERROR] Basic DRP extraction error: {e}")
            return {
                'success': False,
                'error': str(e)
            }


    def process_bcp_drp_document(self, plan_id: int, plan_type: str, file_uri: str) -> Dict:
        """
        Process BCP/DRP document with OCR and AI extraction
        
        Args:
            plan_id: Plan ID from bcp_drp_plans table
            plan_type: 'BCP' or 'DRP'
            file_uri: URI to the document file
            
        Returns:
            Dict with success status and extracted data
        """
        try:
            logger.info(f"[INFO] Processing {plan_type} document for plan {plan_id}")
            logger.info(f"[INFO] File URI: {file_uri}")
            
            # Step 1: Download and process document
            logger.info("[INFO] Step 1: Downloading and processing document")
            
            # Get document content from file URI
            try:
                if file_uri.startswith('http'):
                    # Handle HTTP URLs
                    response = requests.get(file_uri, timeout=30)
                    if response.status_code == 200:
                        document_content = response.content
                        logger.info(f"[INFO] Downloaded document from URL: {len(document_content)} bytes")
                    else:
                        raise Exception(f"Failed to download document: HTTP {response.status_code}")
                else:
                    # Handle local file paths - construct full path
                    import os
                    from django.conf import settings
                    
                    # If it's a relative path, construct the full path
                    if not os.path.isabs(file_uri):
                        # Check if it starts with media URL
                        if file_uri.startswith('uploads/'):
                            # Construct full path using MEDIA_ROOT
                            full_path = os.path.join(settings.MEDIA_ROOT, file_uri)
                        else:
                            # Assume it's relative to MEDIA_ROOT
                            full_path = os.path.join(settings.MEDIA_ROOT, file_uri)
                    else:
                        full_path = file_uri
                    
                    logger.info(f"[INFO] MEDIA_ROOT: {settings.MEDIA_ROOT}")
                    logger.info(f"[INFO] BASE_DIR: {settings.BASE_DIR}")
                    logger.info(f"[INFO] File URI: {file_uri}")
                    logger.info(f"[INFO] Attempting to read file: {full_path}")
                    
                    # Check if file exists
                    if not os.path.exists(full_path):
                        # Try alternative paths
                        alternative_paths = [
                            os.path.join(settings.BASE_DIR, file_uri),
                            os.path.join(settings.MEDIA_ROOT, file_uri),
                            os.path.join(settings.BASE_DIR, 'media', file_uri),
                            os.path.join(settings.BASE_DIR, 'backend', 'media', file_uri),
                            file_uri
                        ]
                        
                        logger.info(f"[INFO] File not found at {full_path}, trying alternative paths...")
                        for alt_path in alternative_paths:
                            logger.info(f"[INFO] Trying path: {alt_path}")
                            if os.path.exists(alt_path):
                                full_path = alt_path
                                logger.info(f"[INFO] Found file at alternative path: {full_path}")
                                break
                        else:
                            # List directory contents for debugging
                            try:
                                media_dir = settings.MEDIA_ROOT
                                if os.path.exists(media_dir):
                                    logger.info(f"[INFO] MEDIA_ROOT directory contents: {os.listdir(media_dir)}")
                                else:
                                    logger.info(f"[INFO] MEDIA_ROOT directory does not exist: {media_dir}")
                            except Exception as debug_e:
                                logger.info(f"[INFO] Could not list MEDIA_ROOT contents: {debug_e}")
                            
                            raise Exception(f"File not found. Tried paths: {[full_path] + alternative_paths}")
                    
                    with open(full_path, 'rb') as f:
                        document_content = f.read()
                        logger.info(f"[INFO] Read local document: {len(document_content)} bytes from {full_path}")
                        
            except Exception as e:
                logger.error(f"[ERROR] Failed to get document content: {e}")
                return {
                    'success': False,
                    'error': f'Failed to get document content: {str(e)}'
                }
            
            # Step 2: Run OCR on document
            logger.info("[INFO] Step 2: Running OCR on document")
            
            try:
                # Detect file extension from file_uri
                file_ext = os.path.splitext(file_uri)[1].lower()
                if not file_ext:
                    file_ext = '.pdf'  # Default to PDF if extension not found
                
                # Save content to temporary file for OCR processing
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                    tmp_file.write(document_content)
                    tmp_file_path = tmp_file.name
                
                ocr_result = self.perform_ocr(tmp_file_path)
                
                # Clean up temporary file
                import os
                try:
                    os.unlink(tmp_file_path)
                except:
                    pass
                
                if not ocr_result['success']:
                    logger.error(f"[ERROR] OCR failed: {ocr_result.get('error')}")
                    return {
                        'success': False,
                        'error': f'OCR processing failed: {ocr_result.get("error")}'
                    }
                
                ocr_text = ocr_result['text']
                ocr_confidence = ocr_result.get('confidence', 0.0)
                
                logger.info(f"[INFO] OCR completed: {len(ocr_text)} characters, confidence: {ocr_confidence}")
                
            except Exception as e:
                logger.error(f"[ERROR] OCR processing error: {e}")
                return {
                    'success': False,
                    'error': f'OCR processing failed: {str(e)}'
                }
            
            # Step 3: AI extraction based on plan type
            logger.info(f"[INFO] Step 3: Running AI extraction for {plan_type}")
            
            try:
                # Use unified extraction method for all plan types
                # Both extract_bcp_data_with_ai and extract_drp_data_with_ai now extract all unified fields
                if plan_type.upper() == 'BCP':
                    extraction_result = self.extract_bcp_data_with_ai(ocr_text)
                elif plan_type.upper() == 'DRP':
                    extraction_result = self.extract_drp_data_with_ai(ocr_text)
                else:
                    # For any other plan type (CRP, etc.), use the unified extraction method
                    # Both methods now extract all fields, so we can use either one
                    extraction_result = self.extract_bcp_data_with_ai(ocr_text)
                
                if not extraction_result['success']:
                    logger.error(f"[ERROR] AI extraction failed: {extraction_result.get('error')}")
                    return {
                        'success': False,
                        'error': f'AI extraction failed: {extraction_result.get("error")}'
                    }
                
                extracted_data = extraction_result['data']
                extraction_confidence = extraction_result.get('confidence', 0.0)
                
                logger.info(f"[SUCCESS] AI extraction completed for {plan_type} with confidence: {extraction_confidence}")
                
            except Exception as e:
                logger.error(f"[ERROR] AI extraction error: {e}")
                return {
                    'success': False,
                    'error': f'AI extraction failed: {str(e)}'
                }
            
            # Step 4: Return results
            logger.info(f"[SUCCESS] Document processing completed for {plan_type} plan {plan_id}")
            
            return {
                'success': True,
                'plan_id': plan_id,
                'plan_type': plan_type,
                'ocr_text': ocr_text,
                'ocr_confidence': ocr_confidence,
                'extracted_data': extracted_data,
                'extraction_confidence': extraction_confidence,
                'confidence': (ocr_confidence + extraction_confidence) / 2
            }
            
        except Exception as e:
            logger.error(f"[ERROR] BCP/DRP document processing error: {e}")
            return {
                'success': False,
                'error': str(e)
            }


# Global service instance
document_service = DocumentProcessingService()
