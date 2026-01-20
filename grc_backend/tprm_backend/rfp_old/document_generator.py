import io
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from django.http import HttpResponse
from django.conf import settings
import os


class RFPDocumentGenerator:
    """Generate RFP documents in Word and PDF formats"""
    
    def __init__(self, rfp_data):
        self.rfp_data = rfp_data
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    def generate_word_document(self):
        """Generate Word document (.docx)"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading(f'Request for Proposal (RFP)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # RFP Number and Title
        doc.add_heading(f'RFP Number: {self.rfp_data.get("rfp_number", "N/A")}', level=1)
        doc.add_heading(f'Title: {self.rfp_data.get("rfp_title", "N/A")}', level=2)
        
        # Basic Information Section
        doc.add_heading('1. Basic Information', level=1)
        
        basic_info = [
            ['RFP Number', self.rfp_data.get('rfp_number', 'N/A')],
            ['RFP Title', self.rfp_data.get('rfp_title', 'N/A')],
            ['RFP Type', self.rfp_data.get('rfp_type', 'N/A')],
            ['Category', self.rfp_data.get('category', 'N/A')],
            ['Description', self.rfp_data.get('description', 'N/A')],
        ]
        
        self._add_table(doc, basic_info, 'Basic Information')
        
        # Financial Information Section
        doc.add_heading('2. Financial Information', level=1)
        
        financial_info = [
            ['Estimated Value', f"{self.rfp_data.get('estimated_value', 'N/A')} {self.rfp_data.get('currency', 'USD')}"],
            ['Budget Range (Min)', f"{self.rfp_data.get('budget_range_min', 'N/A')} {self.rfp_data.get('currency', 'USD')}"],
            ['Budget Range (Max)', f"{self.rfp_data.get('budget_range_max', 'N/A')} {self.rfp_data.get('currency', 'USD')}"],
        ]
        
        self._add_table(doc, financial_info, 'Financial Information')
        
        # Timeline Information Section
        doc.add_heading('3. Timeline Information', level=1)
        
        timeline_info = [
            ['Issue Date', self._format_date(self.rfp_data.get('issue_date'))],
            ['Submission Deadline', self._format_datetime(self.rfp_data.get('submission_deadline'))],
            ['Evaluation Period End', self._format_date(self.rfp_data.get('evaluation_period_end'))],
        ]
        
        self._add_table(doc, timeline_info, 'Timeline Information')
        
        # Evaluation Criteria Section
        doc.add_heading('4. Evaluation Criteria', level=1)
        
        if self.rfp_data.get('evaluation_criteria'):
            criteria_data = [['Criterion', 'Weight (%)', 'Description', 'Veto']]
            for criterion in self.rfp_data['evaluation_criteria']:
                criteria_data.append([
                    criterion.get('criteria_name', 'N/A'),
                    str(criterion.get('weight_percentage', 0)),
                    criterion.get('criteria_description', 'N/A'),
                    'Yes' if criterion.get('veto_enabled') else 'No'
                ])
            
            self._add_table(doc, criteria_data, 'Evaluation Criteria')
        
        # Process Settings Section
        doc.add_heading('5. Process Settings', level=1)
        
        process_info = [
            ['Evaluation Method', self.rfp_data.get('evaluation_method', 'N/A')],
            ['Criticality Level', self.rfp_data.get('criticality_level', 'N/A')],
            ['Geographical Scope', self.rfp_data.get('geographical_scope', 'N/A')],
            ['Allow Late Submissions', 'Yes' if self.rfp_data.get('allow_late_submissions') else 'No'],
            ['Auto-approve when approved', 'Yes' if self.rfp_data.get('auto_approved') else 'No'],
        ]
        
        self._add_table(doc, process_info, 'Process Settings')
        
        # Compliance Requirements Section
        if self.rfp_data.get('compliance_requirements'):
            doc.add_heading('6. Compliance Requirements', level=1)
            compliance_text = self.rfp_data.get('compliance_requirements', '')
            if isinstance(compliance_text, str):
                try:
                    compliance_list = json.loads(compliance_text)
                    if isinstance(compliance_list, list):
                        for req in compliance_list:
                            doc.add_paragraph(f"• {req}")
                    else:
                        doc.add_paragraph(compliance_text)
                except json.JSONDecodeError:
                    doc.add_paragraph(compliance_text)
            else:
                doc.add_paragraph(str(compliance_text))
        
        # Footer
        doc.add_page_break()
        doc.add_heading('Document Information', level=1)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph(f'RFP Status: {self.rfp_data.get("status", "DRAFT")}')
        
        return doc
    
    def generate_pdf_document(self):
        """Generate PDF document"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch, bottomMargin=1*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkblue
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph('Request for Proposal (RFP)', title_style))
        story.append(Spacer(1, 20))
        
        # RFP Number and Title
        story.append(Paragraph(f'<b>RFP Number:</b> {self.rfp_data.get("rfp_number", "N/A")}', styles['Heading2']))
        story.append(Paragraph(f'<b>Title:</b> {self.rfp_data.get("rfp_title", "N/A")}', styles['Heading3']))
        story.append(Spacer(1, 20))
        
        # Basic Information
        story.append(Paragraph('1. Basic Information', heading_style))
        basic_data = [
            ['RFP Number', self.rfp_data.get('rfp_number', 'N/A')],
            ['RFP Title', self.rfp_data.get('rfp_title', 'N/A')],
            ['RFP Type', self.rfp_data.get('rfp_type', 'N/A')],
            ['Category', self.rfp_data.get('category', 'N/A')],
            ['Description', self.rfp_data.get('description', 'N/A')],
        ]
        story.append(self._create_pdf_table(basic_data))
        story.append(Spacer(1, 20))
        
        # Financial Information
        story.append(Paragraph('2. Financial Information', heading_style))
        financial_data = [
            ['Estimated Value', f"{self.rfp_data.get('estimated_value', 'N/A')} {self.rfp_data.get('currency', 'USD')}"],
            ['Budget Range (Min)', f"{self.rfp_data.get('budget_range_min', 'N/A')} {self.rfp_data.get('currency', 'USD')}"],
            ['Budget Range (Max)', f"{self.rfp_data.get('budget_range_max', 'N/A')} {self.rfp_data.get('currency', 'USD')}"],
        ]
        story.append(self._create_pdf_table(financial_data))
        story.append(Spacer(1, 20))
        
        # Timeline Information
        story.append(Paragraph('3. Timeline Information', heading_style))
        timeline_data = [
            ['Issue Date', self._format_date(self.rfp_data.get('issue_date'))],
            ['Submission Deadline', self._format_datetime(self.rfp_data.get('submission_deadline'))],
            ['Evaluation Period End', self._format_date(self.rfp_data.get('evaluation_period_end'))],
        ]
        story.append(self._create_pdf_table(timeline_data))
        story.append(Spacer(1, 20))
        
        # Evaluation Criteria
        story.append(Paragraph('4. Evaluation Criteria', heading_style))
        if self.rfp_data.get('evaluation_criteria'):
            criteria_data = [['Criterion', 'Weight (%)', 'Description', 'Veto']]
            for criterion in self.rfp_data['evaluation_criteria']:
                criteria_data.append([
                    criterion.get('criteria_name', 'N/A'),
                    str(criterion.get('weight_percentage', 0)),
                    criterion.get('criteria_description', 'N/A'),
                    'Yes' if criterion.get('veto_enabled') else 'No'
                ])
            story.append(self._create_pdf_table(criteria_data))
            story.append(Spacer(1, 20))
        
        # Process Settings
        story.append(Paragraph('5. Process Settings', heading_style))
        process_data = [
            ['Evaluation Method', self.rfp_data.get('evaluation_method', 'N/A')],
            ['Criticality Level', self.rfp_data.get('criticality_level', 'N/A')],
            ['Geographical Scope', self.rfp_data.get('geographical_scope', 'N/A')],
            ['Allow Late Submissions', 'Yes' if self.rfp_data.get('allow_late_submissions') else 'No'],
            ['Auto-approve when approved', 'Yes' if self.rfp_data.get('auto_approved') else 'No'],
        ]
        story.append(self._create_pdf_table(process_data))
        story.append(Spacer(1, 20))
        
        # Compliance Requirements
        if self.rfp_data.get('compliance_requirements'):
            story.append(Paragraph('6. Compliance Requirements', heading_style))
            compliance_text = self.rfp_data.get('compliance_requirements', '')
            if isinstance(compliance_text, str):
                try:
                    compliance_list = json.loads(compliance_text)
                    if isinstance(compliance_list, list):
                        for req in compliance_list:
                            story.append(Paragraph(f"• {req}", styles['Normal']))
                    else:
                        story.append(Paragraph(compliance_text, styles['Normal']))
                except json.JSONDecodeError:
                    story.append(Paragraph(compliance_text, styles['Normal']))
            else:
                story.append(Paragraph(str(compliance_text), styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Document Information
        story.append(PageBreak())
        story.append(Paragraph('Document Information', heading_style))
        story.append(Paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}', styles['Normal']))
        story.append(Paragraph(f'RFP Status: {self.rfp_data.get("status", "DRAFT")}', styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _add_table(self, doc, data, title):
        """Add a table to Word document"""
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add data to table
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                # Make header row bold
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        doc.add_paragraph()  # Add space after table
    
    def _create_pdf_table(self, data):
        """Create a table for PDF"""
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        return table
    
    def _format_date(self, date_str):
        """Format date string for display"""
        if not date_str:
            return 'N/A'
        try:
            if 'T' in date_str:
                date_str = date_str.split('T')[0]
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            return date_obj.strftime('%B %d, %Y')
        except:
            return date_str
    
    def _format_datetime(self, datetime_str):
        """Format datetime string for display"""
        if not datetime_str:
            return 'N/A'
        try:
            if 'T' in datetime_str:
                date_obj = datetime.fromisoformat(datetime_str.replace('Z', '+00:00'))
                return date_obj.strftime('%B %d, %Y at %I:%M %p')
            else:
                return self._format_date(datetime_str)
        except:
            return datetime_str


def generate_rfp_document(rfp_data, format_type='word'):
    """
    Generate RFP document in specified format
    
    Args:
        rfp_data: Dictionary containing RFP information
        format_type: 'word' or 'pdf'
    
    Returns:
        HttpResponse with document
    """
    generator = RFPDocumentGenerator(rfp_data)
    rfp_number = rfp_data.get('rfp_number', 'RFP')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if format_type.lower() == 'word':
        doc = generator.generate_word_document()
        
        # Save to BytesIO
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        response = HttpResponse(
            doc_buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{rfp_number}_{timestamp}.docx"'
        
    elif format_type.lower() == 'pdf':
        pdf_buffer = generator.generate_pdf_document()
        
        response = HttpResponse(
            pdf_buffer.getvalue(),
            content_type='application/pdf'
        )
        response['Content-Disposition'] = f'attachment; filename="{rfp_number}_{timestamp}.pdf"'
    
    else:
        raise ValueError("Format must be 'word' or 'pdf'")
    
    return response
