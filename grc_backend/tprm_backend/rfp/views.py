from rest_framework import viewsets, status, filters
from rest_framework.decorators import action, api_view, authentication_classes, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, AllowAny
from rest_framework.views import APIView
from django.db.models import Q, Count, Avg
from django.utils import timezone
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.shortcuts import get_object_or_404, render, redirect
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST, require_http_methods, require_GET
from django.db import transaction, connection
from django.core.mail import send_mail, EmailMessage
from django.template.loader import render_to_string
from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.paginator import Paginator
from datetime import datetime, timedelta
from decimal import Decimal
import os
import tempfile
import uuid
import json
import hashlib
import csv
import io
import pandas as pd
import time
import logging
import secrets
import string
import re
import shutil

# Set up logger
logger = logging.getLogger(__name__)

from .models import (
    RFP, RFPEvaluationCriteria, CustomUser, S3Files, RFPAwardNotification, 
    RFPEvaluationScore, RFPTypeCustomFields, Vendor, VendorCapability, 
    VendorCertification, RFPVendorSelection, RFPUnmatchedVendor, 
    VendorInvitation, RFPResponse
)
from .serializers import (
    RFPSerializer, 
    RFPCreateSerializer, 
    RFPListSerializer,
    RFPEvaluationCriteriaSerializer,
    CustomUserSerializer,
    RFPTypeCustomFieldsSerializer
)
from .permissions import IsRFPCreatorOrReviewer
from .s3 import create_direct_mysql_client
from .forms import (
    VendorSearchForm, VendorManualEntryForm, 
    VendorBulkUploadForm, RFPVendorSelectionForm
)

# RBAC imports
from tprm_backend.rbac.tprm_decorators import rbac_rfp_required
from .rfp_authentication import JWTAuthentication, SimpleAuthenticatedPermission, RFPAuthenticationMixin


class RFPViewSet(RFPAuthenticationMixin, viewsets.ModelViewSet):
    """
    API endpoint for managing RFPs with RBAC authentication
    """
    queryset = RFP.objects.all()
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['rfp_title', 'description', 'rfp_number']
    ordering_fields = ['created_at', 'updated_at', 'submission_deadline', 'rfp_title']
    ordering = ['-created_at']

    def get_serializer_class(self):
        if self.action == 'list':
            return RFPListSerializer
        elif self.action == 'create':
            return RFPCreateSerializer
        return RFPSerializer
        
    def list(self, request, *args, **kwargs):
        """
        Override list method to handle errors better
        """
        try:
            return super().list(request, *args, **kwargs)
        except Exception as e:
            print("Exception during RFP listing:", str(e))
            import traceback
            print(traceback.format_exc())
            return Response({"error": "Error retrieving RFPs", "detail": str(e)}, 
                           status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def get_queryset(self):
        """
        Filter RFPs based on user role and permissions
        """
        # For development, return all RFPs
        return RFP.objects.all()
        
        # In production, we would filter by user
        # user = self.request.user
        # 
        # # Superusers can see all RFPs
        # if user.is_superuser:
        #     return RFP.objects.all()
        # 
        # # Filter RFPs based on user's role
        # return RFP.objects.filter(
        #     Q(created_by=user) |
        #     Q(primary_reviewer_id=user.id) |
        #     Q(executive_reviewer_id=user.id)
        # )

    def create(self, request, *args, **kwargs):
        """
        Override create method to handle the data better
        """
        print("RFPViewSet create data:", request.data)
        
        # Add default user for development
        from django.contrib.auth.models import User
        admin_user = User.objects.filter(is_superuser=True).first()
        if not admin_user:
            admin_user = User.objects.create_superuser(
                username='admin',
                email='admin@example.com',
                password='admin123'
            )
        
        data = request.data.copy() if hasattr(request.data, 'copy') else request.data
        
        serializer = self.get_serializer(data=data)
        try:
            serializer.is_valid(raise_exception=True)
            # Pass the admin user ID directly to save
            serializer.save(created_by=admin_user.id)
            headers = self.get_success_headers(serializer.data)
            return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)
        except Exception as e:
            print("Exception during RFP creation:", str(e))
            import traceback
            print(traceback.format_exc())
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
    
    def perform_create(self, serializer):
        """
        Set the created_by field when creating an RFP
        For development, use the first superuser if no authenticated user
        """
        from django.contrib.auth.models import User
        
        if self.request.user.is_authenticated:
            serializer.save(created_by=self.request.user.id)
        else:
            # For development only - use the first superuser
            admin_user = User.objects.filter(is_superuser=True).first()
            if admin_user:
                serializer.save(created_by=admin_user.id)
            else:
                # Create a superuser if none exists
                admin_user = User.objects.create_superuser(
                    username='admin',
                    email='admin@example.com',
                    password='admin123'
                )
                serializer.save(created_by=admin_user.id)

    @action(detail=True, methods=['post'])
    def submit_for_review(self, request, pk=None):
        """
        Submit an RFP for review
        """
        rfp = self.get_object()
        
        # Check if RFP is in draft status
        if rfp.status != 'DRAFT':
            return Response(
                {"error": "Only RFPs in DRAFT status can be submitted for review"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Check if required fields are set
        if not all([rfp.rfp_title, rfp.description, rfp.rfp_type]):
            return Response(
                {"error": "RFP title, description, and type are required"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Check if evaluation criteria exist and weights sum to 100%
        criteria = rfp.evaluation_criteria.all()
        if not criteria.exists():
            return Response(
                {"error": "At least one evaluation criterion is required"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        total_weight = sum(criterion.weight_percentage for criterion in criteria)
        if abs(total_weight - 100) > 0.01:  # Allow small floating point errors
            return Response(
                {"error": f"Total weight percentage must equal 100% (current: {total_weight}%)"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Check if reviewers are assigned
        if not rfp.primary_reviewer_id or not rfp.executive_reviewer_id:
            return Response(
                {"error": "Both primary and executive reviewers must be assigned"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update status to IN_REVIEW
        rfp.status = 'IN_REVIEW'
        rfp.save()
        
        # Update approval workflows
        for workflow in rfp.approval_workflows.all():
            if workflow.stage == 'reviewer':
                workflow.status = 'pending'
                workflow.save()
        
        return Response({"status": "RFP submitted for review"})

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """
        Approve an RFP
        If auto_approve is True, set status to APPROVED directly without requiring approval workflow
        """
        rfp = self.get_object()
        
        # If auto_approve is enabled, allow approval from any status (except CANCELLED/ARCHIVED)
        if rfp.auto_approve:
            if rfp.status in ['CANCELLED', 'ARCHIVED']:
                return Response(
                    {"error": "Cannot approve a cancelled or archived RFP"},
                    status=status.HTTP_400_BAD_REQUEST
                )
            # Directly approve without requiring IN_REVIEW status
            rfp.status = 'APPROVED'
            rfp.approved_by = request.user.id if hasattr(request.user, 'id') else 1
            rfp.approval_workflow_id = None  # No approval workflow needed for auto-approve
            rfp.save()
            return Response({
                "status": "RFP auto-approved",
                "message": "RFP has been automatically approved without approval workflow"
            })
        
        # For non-auto-approve RFPs, require IN_REVIEW status
        if rfp.status != 'IN_REVIEW':
            return Response(
                {"error": "Only RFPs in IN_REVIEW status can be approved"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update status to PUBLISHED for normal approval workflow
        rfp.status = 'PUBLISHED'
        rfp.approved_by = request.user.id if hasattr(request.user, 'id') else 1
        rfp.save()
        
        # Update approval workflows
        current_user_id = request.user.id if hasattr(request.user, 'id') else 1
        for workflow in rfp.approval_workflows.all():
            if workflow.approver_id == current_user_id:
                workflow.status = 'approved'
                workflow.approval_date = timezone.now()
                workflow.save()
        
        return Response({"status": "RFP approved"})

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        """
        Reject an RFP and return it to draft status
        """
        rfp = self.get_object()
        
        # Check if RFP is in review status
        if rfp.status != 'IN_REVIEW':
            return Response(
                {"error": "Only RFPs in IN_REVIEW status can be rejected"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get rejection reason from request data
        rejection_reason = request.data.get('rejection_reason', '')
        
        # Update status to DRAFT
        rfp.status = 'DRAFT'
        
        # Store rejection reason in custom_fields
        if not rfp.custom_fields:
            rfp.custom_fields = {}
        
        if 'rejection_history' not in rfp.custom_fields:
            rfp.custom_fields['rejection_history'] = []
        
        rfp.custom_fields['rejection_history'].append({
            'date': timezone.now().isoformat(),
            'reviewer': request.user.id,
            'reason': rejection_reason
        })
        
        rfp.save()
        
        # Update approval workflows
        current_user_id = request.user.id
        for workflow in rfp.approval_workflows.all():
            if workflow.approver_id == current_user_id:
                workflow.status = 'rejected'
                workflow.comments = rejection_reason
                workflow.approval_date = timezone.now()
                workflow.save()
        
        return Response({
            "status": "RFP rejected and returned to draft",
            "rejection_reason": rejection_reason
        })

    @action(detail=True, methods=['post'])
    def cancel(self, request, pk=None):
        """
        Cancel an RFP
        """
        rfp = self.get_object()
        
        # Only allow cancellation of RFPs that are not already cancelled or archived
        if rfp.status in ['CANCELLED', 'ARCHIVED']:
            return Response(
                {"error": "This RFP is already cancelled or archived"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update status to CANCELLED
        rfp.status = 'CANCELLED'
        rfp.save()
        
        return Response({"status": "RFP cancelled"})

    @action(detail=True, methods=['post'])
    def archive(self, request, pk=None):
        """
        Archive an RFP
        """
        rfp = self.get_object()
        
        # Only allow archiving of completed RFPs
        if rfp.status not in ['AWARDED', 'CANCELLED']:
            return Response(
                {"error": "Only awarded or cancelled RFPs can be archived"},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update status to ARCHIVED
        rfp.status = 'ARCHIVED'
        rfp.save()
        
        return Response({"status": "RFP archived"})

    @action(detail=True, methods=['post'])
    def update_documents(self, request, pk=None):
        """
        Update RFP documents field with document IDs
        """
        rfp = self.get_object()
        
        try:
            documents = request.data.get('documents', [])
            print(f"[UPDATE] Updating RFP {pk} with documents: {documents}")
            
            if not isinstance(documents, list):
                return Response({
                    'error': 'Documents must be an array of document IDs'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Validate that all document IDs exist in s3_files table
            from .models import S3Files
            existing_docs = S3Files.objects.filter(id__in=documents).values_list('id', flat=True)
            missing_docs = set(documents) - set(existing_docs)
            
            print(f"[UPDATE] Existing document IDs: {list(existing_docs)}")
            print(f"[UPDATE] Missing document IDs: {list(missing_docs)}")
            
            if missing_docs:
                return Response({
                    'error': f'Document IDs not found: {list(missing_docs)}'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Update the documents field
            print(f"[UPDATE] Updating RFP documents field from {rfp.documents} to {documents}")
            rfp.documents = documents
            rfp.save()
            
            print(f"[OK] RFP {pk} updated successfully with documents: {documents}")
            
            return Response({
                'success': True,
                'message': f'Updated RFP with {len(documents)} document(s)',
                'documents': documents
            })
            
        except Exception as e:
            print(f"[EMOJI] Error updating RFP documents: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return Response({
                'error': f'Failed to update documents: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['post'])
    def merge_documents(self, request, pk=None):
        """
        Merge documents in the specified order and upload the merged document
        """
        rfp = self.get_object()
       
        try:
            document_ids = request.data.get('document_ids', [])
            document_order = request.data.get('document_order', [])  # Optional: explicit order
           
            if not document_ids:
                return Response({
                    'success': False,
                    'error': 'No document IDs provided'
                }, status=status.HTTP_400_BAD_REQUEST)
           
            # Use document_order if provided, otherwise use document_ids order
            ordered_ids = document_order if document_order else document_ids
           
            print(f"[MERGE] Merging {len(ordered_ids)} documents for RFP {pk}")
            print(f"[MERGE] Document order: {ordered_ids}")
           
            # Validate that all document IDs exist
            from .models import S3Files
            existing_docs = S3Files.objects.filter(id__in=ordered_ids)
            if existing_docs.count() != len(ordered_ids):
                missing = set(ordered_ids) - set(existing_docs.values_list('id', flat=True))
                return Response({
                    'success': False,
                    'error': f'Document IDs not found: {list(missing)}'
                }, status=status.HTTP_400_BAD_REQUEST)
           
            # Create S3 client
            try:
                s3_client = create_direct_mysql_client()
            except Exception as e:
                print(f"[ERROR] Failed to create S3 client: {str(e)}")
                return Response({
                    'success': False,
                    'error': f'Failed to initialize S3 client: {str(e)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
           
            # Download and merge PDFs
            try:
                import requests
                from io import BytesIO
               
                # Try to import PyPDF2 or pypdf
                try:
                    from PyPDF2 import PdfReader, PdfWriter
                except ImportError:
                    try:
                        from pypdf import PdfReader, PdfWriter
                    except ImportError:
                        return Response({
                            'success': False,
                            'error': 'PDF library (PyPDF2 or pypdf) not installed. Please install: pip install PyPDF2 or pip install pypdf'
                        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
               
                pdf_writer = PdfWriter()
                temp_dirs = []  # Track temp directories for cleanup
               
                # Download each document in order
                for doc_id in ordered_ids:
                    doc = existing_docs.get(id=doc_id)
                   
                    # Get S3 key from metadata
                    s3_key = doc.metadata.get('s3_key') if doc.metadata else None
                    if not s3_key:
                        # Try to extract from URL
                        if doc.url:
                            # Extract key from URL if possible
                            s3_key = doc.url.split('/')[-1] if '/' in doc.url else None
                   
                    if not s3_key:
                        print(f"[WARN] No S3 key found for document {doc_id}, skipping...")
                        continue
                   
                    print(f"[MERGE] Downloading document {doc_id} (S3 key: {s3_key})")
                   
                    # Download from S3
                    temp_download_dir = None
                    try:
                        # Use temp directory for downloads
                        temp_download_dir = tempfile.mkdtemp()
                        temp_dirs.append(temp_download_dir)  # Track for cleanup
                        download_file_name = doc.file_name or f"doc_{doc_id}.pdf"
                       
                        # Use the download method from S3 client
                        download_result = s3_client.download(
                            s3_key=s3_key,
                            file_name=download_file_name,
                            destination_path=temp_download_dir
                        )
                       
                        if not download_result.get('success'):
                            print(f"[ERROR] Failed to download document {doc_id}: {download_result.get('error')}")
                            # Try to download from URL directly as fallback
                            if doc.url:
                                response = requests.get(doc.url, timeout=30)
                                if response.status_code == 200:
                                    # Save to temp file
                                    temp_file_path = os.path.join(temp_download_dir, download_file_name)
                                    with open(temp_file_path, 'wb') as f:
                                        f.write(response.content)
                                    download_result = {'success': True, 'file_path': temp_file_path}
                                else:
                                    print(f"[ERROR] Failed to download from URL for document {doc_id}")
                                    continue
                            else:
                                print(f"[ERROR] No download result or URL for document {doc_id}")
                                continue
                       
                        # Get file path and read PDF
                        file_path = download_result.get('file_path')
                        if not file_path or not os.path.exists(file_path):
                            print(f"[ERROR] File path not found or doesn't exist: {file_path}")
                            continue
                       
                        # Read PDF from file
                        with open(file_path, 'rb') as f:
                            file_content = f.read()
                       
                        pdf_reader = PdfReader(BytesIO(file_content))
                       
                        # Clean up temp file after reading
                        try:
                            os.unlink(file_path)
                        except:
                            pass
                       
                        # Add all pages to writer
                        for page_num in range(len(pdf_reader.pages)):
                            pdf_writer.add_page(pdf_reader.pages[page_num])
                       
                        print(f"[MERGE] Added {len(pdf_reader.pages)} pages from document {doc_id}")
                       
                    except Exception as e:
                        print(f"[ERROR] Error processing document {doc_id}: {str(e)}")
                        import traceback
                        print(traceback.format_exc())
                        continue
                    finally:
                        # Clean up temp directory for this document
                        if temp_download_dir and os.path.exists(temp_download_dir):
                            try:
                                shutil.rmtree(temp_download_dir)
                            except Exception as cleanup_err:
                                print(f"[WARN] Failed to cleanup temp dir {temp_download_dir}: {cleanup_err}")
               
                # Create merged PDF in memory
                merged_pdf_buffer = BytesIO()
                pdf_writer.write(merged_pdf_buffer)
                merged_pdf_buffer.seek(0)
               
                # Save merged PDF to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                    temp_file.write(merged_pdf_buffer.read())
                    temp_file_path = temp_file.name
               
                try:
                    # Upload merged document to S3
                    merged_filename = f"rfp_{pk}_merged_{uuid.uuid4().hex}.pdf"
                    user_id = request.data.get('user_id', '1')
                   
                    upload_result = s3_client.upload(
                        file_path=temp_file_path,
                        user_id=user_id,
                        custom_file_name=merged_filename
                    )
                   
                    if upload_result['success']:
                        file_info = upload_result['file_info']
                       
                        # Save merged document to S3Files
                        merged_doc = S3Files.objects.create(
                            url=file_info['url'],
                            file_type='pdf',
                            file_name=f"Merged Document - {rfp.rfp_title or 'RFP'}",
                            user_id=user_id,
                            metadata={
                                'original_filename': merged_filename,
                                'stored_filename': file_info['storedName'],
                                's3_key': file_info['s3Key'],
                                's3_bucket': file_info.get('bucket', ''),
                                'rfp_id': pk,
                                'document_name': f"Merged Document - {rfp.rfp_title or 'RFP'}",
                                'file_size': os.path.getsize(temp_file_path),
                                'upload_operation_id': upload_result.get('operation_id'),
                                'is_merged': True,
                                'merged_from_documents': ordered_ids
                            }
                        )
                       
                        print(f"[SUCCESS] Merged document created: {merged_doc.id}")
                       
                        return Response({
                            'success': True,
                            'merged_document_id': merged_doc.id,
                            'merged_document_name': merged_doc.file_name,
                            'merged_document_url': merged_doc.url,
                            'message': f'Successfully merged {len(ordered_ids)} documents',
                            'document_count': len(ordered_ids)
                        })
                    else:
                        return Response({
                            'success': False,
                            'error': f'S3 upload failed: {upload_result.get("error", "Unknown error")}'
                        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                       
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(temp_file_path)
                    except Exception as cleanup_error:
                        print(f"[WARN] Failed to clean up temporary file: {cleanup_error}")
                   
                    # Clean up any remaining temp directories
                    for temp_dir in temp_dirs:
                        if os.path.exists(temp_dir):
                            try:
                                shutil.rmtree(temp_dir)
                            except Exception as cleanup_err:
                                print(f"[WARN] Failed to cleanup temp dir {temp_dir}: {cleanup_err}")
                       
            except Exception as merge_error:
                print(f"[ERROR] Merge error: {str(merge_error)}")
                import traceback
                print(traceback.format_exc())
                return Response({
                    'success': False,
                    'error': f'Failed to merge documents: {str(merge_error)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
               
        except Exception as e:
            print(f"[ERROR] Document merge error: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return Response({
                'success': False,
                'error': f'Document merge failed: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=True, methods=['get'])
    def get_full_details(self, request, pk=None):
        """
/        Get complete RFP payload details from the rfps table only
        """
        rfp = self.get_object()
        
        try:
            # Build RFP payload with only data from the rfps table
            rfp_data = {
                'rfp_id': rfp.rfp_id,
                'rfp_number': rfp.rfp_number,
                'rfp_title': rfp.rfp_title,
                'description': rfp.description,
                'rfp_type': rfp.rfp_type,
                'category': rfp.category,
                'status': rfp.status,
                'version_number': rfp.version_number,
                
                # Budget information
                'estimated_value': float(rfp.estimated_value) if rfp.estimated_value else None,
                'currency': rfp.currency,
                'budget_range_min': float(rfp.budget_range_min) if rfp.budget_range_min else None,
                'budget_range_max': float(rfp.budget_range_max) if rfp.budget_range_max else None,
                
                # Timeline information
                'issue_date': rfp.issue_date.isoformat() if rfp.issue_date else None,
                'submission_deadline': rfp.submission_deadline.isoformat() if rfp.submission_deadline else None,
                'evaluation_period_end': rfp.evaluation_period_end.isoformat() if rfp.evaluation_period_end else None,
                'award_date': rfp.award_date.isoformat() if rfp.award_date else None,
                
                # Workflow and evaluation
                'evaluation_method': rfp.evaluation_method,
                'criticality_level': rfp.criticality_level,
                'geographical_scope': rfp.geographical_scope,
                'approval_workflow_id': rfp.approval_workflow_id,
                
                # Configuration
                'auto_approve': rfp.auto_approve,
                'allow_late_submissions': rfp.allow_late_submissions,
                
                # JSON fields
                'compliance_requirements': rfp.compliance_requirements,
                'custom_fields': rfp.custom_fields,
                'documents': rfp.documents,  # JSON field containing document IDs
                
                # Award information
                'final_evaluation_score': float(rfp.final_evaluation_score) if rfp.final_evaluation_score else None,
                'award_decision_date': rfp.award_decision_date.isoformat() if rfp.award_decision_date else None,
                'award_justification': rfp.award_justification,
                
                # Timestamps
                'created_at': rfp.created_at.isoformat() if rfp.created_at else None,
                'updated_at': rfp.updated_at.isoformat() if rfp.updated_at else None,
                
                # User information
                'created_by': rfp.created_by,
                'approved_by': rfp.approved_by,
                'primary_reviewer_id': rfp.primary_reviewer_id,
                'executive_reviewer_id': rfp.executive_reviewer_id
            }
            
            return Response({
                'success': True,
                'rfp': rfp_data
            })
            
        except Exception as e:
            print(f"[ERROR] Error getting RFP details: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return Response({
                'error': f'Failed to get RFP details: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class RFPEvaluationCriteriaViewSet(RFPAuthenticationMixin, viewsets.ModelViewSet):
    """
    API endpoint for managing RFP evaluation criteria
    """
    queryset = RFPEvaluationCriteria.objects.all()
    serializer_class = RFPEvaluationCriteriaSerializer
    
    def get_queryset(self):
        """
        Filter criteria by RFP ID if provided in query parameters
        """
        queryset = RFPEvaluationCriteria.objects.all()
        # Accept both 'rfp_id' and legacy 'rfp' query params
        rfp_id = self.request.query_params.get('rfp_id') or self.request.query_params.get('rfp')
        
        if rfp_id:
            try:
                # Convert to integer
                rfp_id_int = int(rfp_id)
                print(f"[EMOJI] RFPEvaluationCriteriaViewSet.get_queryset: Filtering by rfp_id={rfp_id_int} (type: {type(rfp_id_int)})")
                
                # Try multiple filter methods to ensure we find criteria
                # Method 1: Use ForeignKey traversal (most reliable)
                filtered_queryset = queryset.filter(rfp__rfp_id=rfp_id_int)
                count1 = filtered_queryset.count()
                print(f"[EMOJI] Method 1 (rfp__rfp_id): Found {count1} criteria")
                
                # Method 2: If no results, try direct rfp_id column (if Django created it)
                if count1 == 0:
                    filtered_queryset = queryset.filter(rfp_id=rfp_id_int)
                    count2 = filtered_queryset.count()
                    print(f"[EMOJI] Method 2 (rfp_id direct): Found {count2} criteria")
                
                # Method 3: Try with RFP object directly
                if not filtered_queryset.exists():
                    from .models import RFP
                    try:
                        rfp_obj = RFP.objects.get(rfp_id=rfp_id_int)
                        filtered_queryset = queryset.filter(rfp=rfp_obj)
                        count3 = filtered_queryset.count()
                        print(f"[EMOJI] Method 3 (rfp object): Found {count3} criteria")
                    except RFP.DoesNotExist:
                        print(f"[EMOJI] RFP with rfp_id={rfp_id_int} does not exist")
                        filtered_queryset = queryset.none()
                
                queryset = filtered_queryset
                final_count = queryset.count()
                print(f"[EMOJI] RFPEvaluationCriteriaViewSet: Final ORM count = {final_count} criteria")
                
                # If ORM didn't find anything, try raw SQL as last resort
                if final_count == 0:
                    print(f"[EMOJI] ORM found 0 criteria, trying raw SQL fallback...")
                    from django.db import connection
                    with connection.cursor() as cursor:
                        # First, check if criteria exist using COUNT
                        cursor.execute("""
                            SELECT COUNT(*) FROM rfp_evaluation_criteria 
                            WHERE rfp_id = %s
                        """, [rfp_id_int])
                        sql_count = cursor.fetchone()[0]
                        print(f"[EMOJI] Raw SQL COUNT query: Found {sql_count} criteria with rfp_id={rfp_id_int}")
                        
                        if sql_count > 0:
                            print(f"[EMOJI] Raw SQL confirms {sql_count} criteria exist - fetching details...")
                            # Use raw SQL to get all criteria
                            cursor.execute("""
                                SELECT criteria_id, rfp_id, criteria_name, criteria_description, 
                                       weight_percentage, evaluation_type, min_score, max_score, 
                                       median_score, is_mandatory, veto_enabled, veto_threshold,
                                       min_word_count, expected_boolean_answer, display_order
                                FROM rfp_evaluation_criteria 
                                WHERE rfp_id = %s
                                ORDER BY display_order
                            """, [rfp_id_int])
                            rows = cursor.fetchall()
                            print(f"[EMOJI] Raw SQL fetched {len(rows)} rows")
                            
                            if len(rows) > 0:
                                # Get criteria IDs and fetch via ORM for serializer compatibility
                                criteria_ids = [row[0] for row in rows]
                                print(f"[EMOJI] Fetching ORM objects for criteria_ids: {criteria_ids}")
                                queryset = RFPEvaluationCriteria.objects.filter(criteria_id__in=criteria_ids).order_by('display_order')
                                fetched_count = queryset.count()
                                print(f"[EMOJI] Raw SQL fallback: Found {fetched_count} criteria via ID lookup")
                                
                                if fetched_count == 0:
                                    print(f"[EMOJI] WARNING: Raw SQL found {len(rows)} rows but ORM lookup returned 0!")
                                    print(f"[EMOJI] This suggests a data mismatch. Sample row: {rows[0] if rows else 'N/A'}")
                            else:
                                print(f"[EMOJI] Raw SQL COUNT said {sql_count} but fetchall returned 0 rows")
                        else:
                            print(f"[EMOJI] Raw SQL COUNT also found 0 criteria for rfp_id={rfp_id_int}")
                            # Debug: Show what rfp_ids DO have criteria
                            cursor.execute("""
                                SELECT DISTINCT rfp_id, COUNT(*) as cnt 
                                FROM rfp_evaluation_criteria 
                                GROUP BY rfp_id 
                                ORDER BY cnt DESC 
                                LIMIT 10
                            """)
                            sample = cursor.fetchall()
                            print(f"[EMOJI] Sample RFP IDs that HAVE criteria: {[(row[0], row[1]) for row in sample]}")
            except (ValueError, TypeError) as e:
                print(f"[EMOJI] Invalid rfp_id parameter: {rfp_id}, error: {e}")
            except Exception as e:
                import traceback
                print(f"[EMOJI] Unexpected error in get_queryset: {e}")
                print(traceback.format_exc())
        
        # Order consistently and avoid duplicates
        final_queryset = queryset.order_by('display_order', 'criteria_id').distinct()
        final_count = final_queryset.count()
        print(f"[EMOJI] RFPEvaluationCriteriaViewSet.get_queryset: Returning {final_count} criteria")
        return final_queryset
    
    def perform_create(self, serializer):
        """
        Set the created_by field when creating criteria
        Uses userid from the authenticated user (User model uses userid as primary key)
        """
        if self.request.user.is_authenticated:
            # The User model uses 'userid' as primary key, not 'id'
            # Try userid first, then pk (Django's universal primary key), then id as fallback
            user_id = getattr(self.request.user, 'userid', None) or getattr(self.request.user, 'pk', None) or getattr(self.request.user, 'id', None)
            if user_id:
                serializer.save(created_by=user_id)
            else:
                # Fallback: use user ID 1 if we can't determine the user ID
                print(f"[EMOJI] Warning: Could not determine user ID from request.user, using default value 1")
                print(f"[EMOJI] User object type: {type(self.request.user)}, attributes: {dir(self.request.user)}")
                serializer.save(created_by=1)
        else:
            # Not authenticated - use default user ID 1
            print(f"[EMOJI] Warning: User not authenticated, using default created_by=1")
            serializer.save(created_by=1)


class CustomUserViewSet(RFPAuthenticationMixin, viewsets.ReadOnlyModelViewSet):
    """
    API endpoint for listing users for reviewer assignment
    """
    queryset = CustomUser.objects.filter(is_active='Y')
    serializer_class = CustomUserSerializer
    filter_backends = [filters.SearchFilter]
    search_fields = ['username', 'first_name', 'last_name', 'email']
    
    @action(detail=False, methods=['get'])
    def reviewers(self, request):
        """
        Get users who can be assigned as reviewers
        """
        users = self.get_queryset()
        serializer = self.get_serializer(users, many=True)
        return Response(serializer.data)


class RFPTypeCustomFieldsViewSet(RFPAuthenticationMixin, viewsets.ReadOnlyModelViewSet):
    """
    API endpoint for listing RFP types from rfp_type_custom_fields table
    """
    queryset = RFPTypeCustomFields.objects.all()
    serializer_class = RFPTypeCustomFieldsSerializer
    filter_backends = [filters.SearchFilter]
    search_fields = ['rfp_type']
    
    @action(detail=False, methods=['get'])
    def types(self, request):
        """
        Get list of unique RFP types (just the rfp_type values)
        """
        rfp_types = RFPTypeCustomFields.objects.values_list('rfp_type', flat=True).distinct().order_by('rfp_type')
        return Response({
            'success': True,
            'rfp_types': list(rfp_types)
        })
    
    @action(detail=False, methods=['get'])
    def custom_fields(self, request):
        """
        Get custom fields for a specific RFP type
        Query parameter: rfp_type (required)
        """
        rfp_type = request.query_params.get('rfp_type', None)
        
        if not rfp_type:
            return Response({
                'success': False,
                'error': 'rfp_type parameter is required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Get the first matching record for this rfp_type
            rfp_type_record = RFPTypeCustomFields.objects.filter(rfp_type=rfp_type).first()
            
            if not rfp_type_record:
                return Response({
                    'success': True,
                    'rfp_type': rfp_type,
                    'custom_fields': None,
                    'message': 'No custom fields found for this RFP type'
                })
            
            return Response({
                'success': True,
                'rfp_type': rfp_type,
                'custom_fields': rfp_type_record.custom_fields or {},
                'rfp_type_id': rfp_type_record.rfp_type_id
            })
        except Exception as e:
            logger.error(f'Error fetching custom fields for rfp_type {rfp_type}: {str(e)}')
            return Response({
                'success': False,
                'error': f'Error fetching custom fields: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DocumentUploadView(APIView):
    """
    API endpoint for uploading documents to S3
    """
    authentication_classes = [JWTAuthentication]
    permission_classes = [SimpleAuthenticatedPermission]
    
    def post(self, request):
        """
        Upload a document to S3 and save metadata to database
        """
        try:
            # Get form data
            file = request.FILES.get('file')
            document_name = request.data.get('document_name')
            rfp_id = request.data.get('rfp_id')
            user_id = request.data.get('user_id', '1')
            
            print(f"[DEBUG] DocumentUploadView received request:")
            print(f"   File: {file.name if file else 'None'}")
            print(f"   Document Name: {document_name}")
            print(f"   RFP ID: {rfp_id}")
            print(f"   User ID: {user_id}")
            print(f"   Request data keys: {list(request.data.keys())}")
            print(f"   Request files keys: {list(request.FILES.keys())}")
            
            # Validate required fields
            if not file:
                return Response({
                    'success': False,
                    'error': 'No file provided'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            if not document_name:
                return Response({
                    'success': False,
                    'error': 'Document name is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # if not rfp_id:
            #     return Response({
            #         'success': False,
            #         'error': 'RFP ID is required'
            #     }, status=status.HTTP_400_BAD_REQUEST)
            
            # Validate file size (max 10MB)
            if file.size > 10 * 1024 * 1024:
                return Response({
                    'success': False,
                    'error': 'File size must be less than 10MB'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Validate file type
            allowed_extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.txt', '.jpg', '.jpeg', '.png']
            file_extension = os.path.splitext(file.name)[1].lower()
            if file_extension not in allowed_extensions:
                return Response({
                    'success': False,
                    'error': f'File type {file_extension} is not allowed. Allowed types: {", ".join(allowed_extensions)}'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            print(f"[UPLOAD] Uploading document: {document_name} ({file.name}, {file.size} bytes)")
            print(f"[UPLOAD] RFP ID: {rfp_id}, User ID: {user_id}")
            
            # Create S3 client
            try:
                s3_client = create_direct_mysql_client()
                print("[SUCCESS] S3 client created successfully")
            except Exception as e:
                print(f"[ERROR] Failed to create S3 client: {str(e)}")
                return Response({
                    'success': False,
                    'error': f'Failed to initialize S3 client: {str(e)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            # Save file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                for chunk in file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            try:
                if rfp_id:
                    unique_filename = f"rfp_{rfp_id}_{uuid.uuid4().hex}{file_extension}"
                else:
                    unique_filename = f"temp_{uuid.uuid4().hex}{file_extension}"
                
                # Upload to S3
                upload_result = s3_client.upload(
                    file_path=temp_file_path,
                    user_id=user_id,
                    custom_file_name=unique_filename
                )
                
                if upload_result['success']:
                    file_info = upload_result['file_info']
                    
                    # Save to S3Files model with correct schema
                    print(f"[SAVE] Saving to s3_files table...")
                    print(f"   URL: {file_info['url']}")
                    print(f"   File Type: {file_extension[1:] if file_extension else 'unknown'}")
                    print(f"   File Name: {document_name}")
                    print(f"   User ID: {user_id}")
                    
                    s3_file = S3Files.objects.create(
                        url=file_info['url'],
                        file_type=file_extension[1:] if file_extension else 'unknown',
                        file_name=document_name,
                        user_id=user_id,
                        metadata={
                            'original_filename': file.name,
                            'stored_filename': file_info['storedName'],
                            's3_key': file_info['s3Key'],
                            's3_bucket': file_info.get('bucket', ''),
                            'rfp_id': rfp_id,
                            'document_name': document_name,
                            'file_size': file.size,
                            'upload_operation_id': upload_result.get('operation_id')
                        }
                    )
                    
                    print(f"[SUCCESS] Document uploaded successfully: {document_name} (S3 ID: {s3_file.id})")
                    print(f"[SUCCESS] S3 file saved to database with ID: {s3_file.id}")
                    
                    # Check how many documents exist for this RFP
                    rfp_docs_count = S3Files.objects.filter(metadata__rfp_id=rfp_id).count()
                    print(f"[INFO] Total documents for RFP {rfp_id}: {rfp_docs_count}")
                    
                    return Response({
                        'success': True,
                        'document_id': s3_file.id,
                        'file_info': {
                            'id': s3_file.id,
                            'name': document_name,
                            'original_filename': file.name,
                            'stored_filename': file_info['storedName'],
                            'url': file_info['url'],
                            's3_key': file_info['s3Key'],
                            'file_size': file.size,
                            'file_type': file_extension[1:] if file_extension else 'unknown'
                        },
                        'message': 'Document uploaded successfully to S3',
                        'rfp_documents_count': rfp_docs_count
                    })
                else:
                    print(f"[ERROR] S3 upload failed: {upload_result.get('error', 'Unknown error')}")
                    return Response({
                        'success': False,
                        'error': f'S3 upload failed: {upload_result.get("error", "Unknown error")}'
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    
            except Exception as upload_error:
                print(f"[ERROR] Upload error: {str(upload_error)}")
                return Response({
                    'success': False,
                    'error': f'Upload failed: {str(upload_error)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception as cleanup_error:
                    print(f"[WARN] Failed to clean up temporary file: {cleanup_error}")
                    
        except Exception as e:
            print(f"[ERROR] Document upload error: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return Response({
                'success': False,
                'error': f'Document upload failed: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def convert_to_pdf(file_path, file_extension):
    """
    Convert various file types to PDF
    Supports: PDF, Word (.doc, .docx), Images (.jpg, .jpeg, .png), Excel (.xls, .xlsx), Text (.txt), and other formats
    Returns path to converted PDF file or None if conversion fails
    """
    file_extension = file_extension.lower()
   
    # If already PDF, return as-is
    if file_extension == '.pdf':
        return file_path
   
    try:
        # Try to import required libraries
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.utils import ImageReader
            reportlab_available = True
        except ImportError:
            reportlab_available = False
            print("[WARN] reportlab not available for conversion")
       
        from io import BytesIO
       
        # Create temp PDF file
        pdf_path = file_path.rsplit('.', 1)[0] + '.pdf'
       
        # Handle Word documents
        if file_extension in ['.doc', '.docx']:
            if not reportlab_available:
                print("[ERROR] reportlab required for Word document conversion")
                return None
            try:
                # Try using python-docx for .docx files
                if file_extension == '.docx':
                    try:
                        from docx import Document
                    except ImportError:
                        print("[ERROR] python-docx required for .docx conversion")
                        return None
                    from reportlab.lib.styles import getSampleStyleSheet
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                   
                    doc = Document(file_path)
                    pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                   
                    for para in doc.paragraphs:
                        if para.text.strip():
                            story.append(Paragraph(para.text, styles['Normal']))
                            story.append(Spacer(1, 12))
                   
                    # Handle tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    story.append(Paragraph(cell.text, styles['Normal']))
                                    story.append(Spacer(1, 6))
                   
                    pdf_doc.build(story)
                    return pdf_path
                else:
                    # For .doc files, try using LibreOffice or fallback to text extraction
                    # If LibreOffice is not available, convert to text and create PDF
                    try:
                        import subprocess
                        # Try LibreOffice conversion
                        result = subprocess.run(
                            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
                             os.path.dirname(pdf_path), file_path],
                            capture_output=True,
                            timeout=30
                        )
                        if result.returncode == 0 and os.path.exists(pdf_path):
                            return pdf_path
                    except:
                        pass
                   
                    # Fallback: read as text and create PDF
                    try:
                        with open(file_path, 'rb') as f:
                            content = f.read()
                            # Try to decode as text
                            try:
                                text = content.decode('utf-8', errors='ignore')
                            except:
                                text = content.decode('latin-1', errors='ignore')
                       
                        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = [Paragraph(text[:1000], styles['Normal'])]  # Limit text length
                        pdf_doc.build(story)
                        return pdf_path
                    except Exception as e:
                        print(f"[WARN] Could not convert .doc file: {str(e)}")
                        return None
            except ImportError:
                print("[WARN] python-docx not installed. Install with: pip install python-docx")
                return None
            except Exception as e:
                print(f"[ERROR] Error converting Word document: {str(e)}")
                return None
       
        # Handle images
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            if not reportlab_available:
                print("[ERROR] reportlab required for image conversion")
                return None
            try:
                try:
                    from PIL import Image
                except ImportError:
                    print("[ERROR] Pillow required for image conversion")
                    return None
               
                # Open image
                img = Image.open(file_path)
               
                # Convert RGBA to RGB if necessary
                if img.mode == 'RGBA':
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    rgb_img.paste(img, mask=img.split()[3])
                    img = rgb_img
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
               
                # Get image dimensions
                img_width, img_height = img.size
               
                # Calculate PDF page size (fit to page)
                page_width, page_height = letter
                scale_x = page_width / img_width
                scale_y = page_height / img_height
                scale = min(scale_x, scale_y) * 0.95  # 95% to add margins
               
                new_width = img_width * scale
                new_height = img_height * scale
               
                # Center image on page
                x_offset = (page_width - new_width) / 2
                y_offset = (page_height - new_height) / 2
               
                # Create PDF with image
                c = canvas.Canvas(pdf_path, pagesize=letter)
                c.drawImage(ImageReader(img), x_offset, y_offset, width=new_width, height=new_height)
                c.save()
               
                return pdf_path
            except ImportError:
                print("[WARN] Pillow (PIL) not installed. Install with: pip install Pillow")
                return None
            except Exception as e:
                print(f"[ERROR] Error converting image: {str(e)}")
                return None
       
        # Handle text files
        elif file_extension in ['.txt']:
            if not reportlab_available:
                print("[ERROR] reportlab required for text file conversion")
                return None
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
               
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
               
                pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
               
                # Split text into paragraphs
                paragraphs = text.split('\n')
                for para in paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), styles['Normal']))
                        story.append(Spacer(1, 12))
               
                pdf_doc.build(story)
                return pdf_path
            except Exception as e:
                print(f"[ERROR] Error converting text file: {str(e)}")
                return None
       
        # Handle Excel files
        elif file_extension in ['.xls', '.xlsx']:
            print(f"[CONVERT] Starting Excel file conversion: {file_path} (extension: {file_extension})")
            if not reportlab_available:
                print("[ERROR] reportlab required for Excel file conversion")
                return None
            try:
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib import colors
                from reportlab.lib.units import inch
               
                pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
               
                # Try to read Excel file
                if file_extension == '.xlsx':
                    try:
                        from openpyxl import load_workbook
                        wb = load_workbook(file_path, data_only=True)
                    except ImportError:
                        print("[ERROR] openpyxl required for .xlsx conversion. Install with: pip install openpyxl")
                        return None
                   
                    # Process each sheet
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                       
                        # Add sheet title
                        story.append(Paragraph(f"<b>Sheet: {sheet_name}</b>", styles['Heading1']))
                        story.append(Spacer(1, 12))
                       
                        # Collect data from sheet
                        data = []
                        max_cols = 0
                       
                        for row in sheet.iter_rows(values_only=True):
                            # Filter out completely empty rows
                            if any(cell is not None and str(cell).strip() for cell in row):
                                row_data = [str(cell) if cell is not None else '' for cell in row]
                                data.append(row_data)
                                max_cols = max(max_cols, len(row_data))
                       
                        # Limit to first 20 rows and 10 columns for performance
                        total_rows = len(data)
                        if total_rows > 20:
                            data = data[:20]
                            story.append(Paragraph(f"<i>Note: Showing first 20 rows of {total_rows} total rows</i>", styles['Normal']))
                            story.append(Spacer(1, 6))
                       
                        if max_cols > 10:
                            for i, row in enumerate(data):
                                data[i] = row[:10]
                            story.append(Paragraph(f"<i>Note: Showing first 10 columns</i>", styles['Normal']))
                            story.append(Spacer(1, 6))
                       
                        # Normalize row lengths
                        for i, row in enumerate(data):
                            while len(row) < max_cols:
                                row.append('')
                            data[i] = row
                       
                        if data:
                            # Create table
                            table = Table(data)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                            ]))
                           
                            story.append(table)
                            story.append(Spacer(1, 20))
               
                else:  # .xls file
                    try:
                        import xlrd
                        wb = xlrd.open_workbook(file_path)
                    except ImportError:
                        print("[ERROR] xlrd required for .xls conversion. Install with: pip install xlrd")
                        return None
                   
                    # Process each sheet
                    for sheet_index in range(wb.nsheets):
                        sheet = wb.sheet_by_index(sheet_index)
                        sheet_name = sheet.name
                       
                        # Add sheet title
                        story.append(Paragraph(f"<b>Sheet: {sheet_name}</b>", styles['Heading1']))
                        story.append(Spacer(1, 12))
                       
                        # Collect data from sheet
                        data = []
                        max_cols = 0
                       
                        for row_idx in range(min(sheet.nrows, 20)):  # Limit to 20 rows
                            row_data = []
                            for col_idx in range(sheet.ncols):
                                cell_value = sheet.cell_value(row_idx, col_idx)
                                row_data.append(str(cell_value) if cell_value else '')
                            if any(cell.strip() for cell in row_data):
                                data.append(row_data)
                                max_cols = max(max_cols, len(row_data))
                       
                        if sheet.nrows > 20:
                            story.append(Paragraph(f"<i>Note: Showing first 20 rows of {sheet.nrows} total rows</i>", styles['Normal']))
                            story.append(Spacer(1, 6))
                       
                        if max_cols > 10:
                            for i, row in enumerate(data):
                                data[i] = row[:10]
                            story.append(Paragraph(f"<i>Note: Showing first 10 columns</i>", styles['Normal']))
                            story.append(Spacer(1, 6))
                       
                        # Normalize row lengths
                        for i, row in enumerate(data):
                            while len(row) < max_cols:
                                row.append('')
                            data[i] = row
                       
                        if data:
                            # Create table
                            table = Table(data)
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                            ]))
                           
                            story.append(table)
                            story.append(Spacer(1, 20))
               
                # Build PDF
                if story:
                    pdf_doc.build(story)
                    print(f"[CONVERT] Successfully converted Excel file to PDF: {pdf_path}")
                    # Verify PDF was created
                    if os.path.exists(pdf_path):
                        file_size = os.path.getsize(pdf_path)
                        print(f"[CONVERT] Converted PDF size: {file_size} bytes")
                    return pdf_path
                else:
                    print("[WARN] Excel file appears to be empty")
                    return None
                   
            except ImportError as e:
                print(f"[ERROR] Required library not installed for Excel conversion: {str(e)}")
                print("[INFO] Install with: pip install openpyxl (for .xlsx) or pip install xlrd (for .xls)")
                return None
            except Exception as e:
                print(f"[ERROR] Error converting Excel file: {str(e)}")
                import traceback
                traceback.print_exc()
                return None
       
        # Unsupported format
        else:
            print(f"[WARN] Unsupported file format: {file_extension}")
            return None
           
    except Exception as e:
        print(f"[ERROR] Error in convert_to_pdf: {str(e)}")
        return None


class MergeDocumentsView(APIView):
    """
    Standalone API endpoint for merging documents without requiring RFP ID
    Can merge documents by IDs (if already uploaded) or by file data
    Supports: PDF, Word (.doc, .docx), Images (.jpg, .jpeg, .png), and text files
    Allows anonymous access for vendor portal
    """
    authentication_classes = []  # Allow anonymous access for vendor portal
    permission_classes = [AllowAny]  # Allow anonymous access for vendor portal
   
    def post(self, request):
        """
        Merge documents - works with either document IDs or file data
        """
        try:
            from .models import S3Files
           
            document_ids = request.data.get('document_ids', [])
            document_order = request.data.get('document_order', [])
            files = request.FILES.getlist('files')  # For direct file uploads
            rfp_id = request.data.get('rfp_id')  # Optional
            user_id = request.data.get('user_id', '1')
           
            print(f"[MERGE] Standalone merge request received")
            print(f"   Document IDs: {document_ids}")
            print(f"   Files: {len(files) if files else 0}")
            print(f"   RFP ID: {rfp_id or 'None (standalone merge)'}")
           
            # Determine merge method
            use_file_merge = len(files) > 0 and len(document_ids) == 0
            use_id_merge = len(document_ids) > 0
           
            if not use_file_merge and not use_id_merge:
                return Response({
                    'success': False,
                    'error': 'Either document_ids or files must be provided'
                }, status=status.HTTP_400_BAD_REQUEST)
           
            # Create S3 client
            try:
                s3_client = create_direct_mysql_client()
            except Exception as e:
                print(f"[ERROR] Failed to create S3 client: {str(e)}")
                return Response({
                    'success': False,
                    'error': f'Failed to initialize S3 client: {str(e)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
           
            # Try to import PyPDF2 or pypdf
            try:
                from PyPDF2 import PdfReader, PdfWriter
            except ImportError:
                try:
                    from pypdf import PdfReader, PdfWriter
                except ImportError:
                    return Response({
                        'success': False,
                        'error': 'PDF library (PyPDF2 or pypdf) not installed. Please install: pip install PyPDF2 or pip install pypdf'
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
           
            # Check for optional dependencies (warn but don't fail)
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import letter
                reportlab_available = True
            except ImportError:
                reportlab_available = False
                print("[WARN] reportlab not installed. Word/Image conversion may not work. Install with: pip install reportlab")
           
            try:
                from PIL import Image
                pillow_available = True
            except ImportError:
                pillow_available = False
                print("[WARN] Pillow not installed. Image conversion may not work. Install with: pip install Pillow")
           
            try:
                from docx import Document
                docx_available = True
            except ImportError:
                docx_available = False
                print("[WARN] python-docx not installed. Word document conversion may not work. Install with: pip install python-docx")
           
            pdf_writer = PdfWriter()
            temp_dirs = []
            successfully_processed_count = 0  # Track successfully processed documents/files
           
            if use_file_merge:
                # Merge directly from uploaded files
                from io import BytesIO
               
                print(f"[MERGE] Merging {len(files)} files directly")
                converted_files = []  # Track converted files for cleanup
                successfully_processed_count = 0  # Reset counter for file merge
               
                for i, file in enumerate(files):
                    try:
                        # Get file extension
                        file_extension = os.path.splitext(file.name)[1].lower()
                        if not file_extension:
                            file_extension = '.pdf'  # Default to PDF
                       
                        # Save file temporarily with original extension
                        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
                            for chunk in file.chunks():
                                temp_file.write(chunk)
                            temp_file_path = temp_file.name
                       
                        # Convert to PDF if needed
                        pdf_file_path = convert_to_pdf(temp_file_path, file_extension)
                       
                        if not pdf_file_path or not os.path.exists(pdf_file_path):
                            print(f"[ERROR] Failed to convert file {i+1} ({file.name}) to PDF")
                            # Clean up original file
                            try:
                                os.unlink(temp_file_path)
                            except:
                                pass
                            continue
                       
                        # Track converted file for cleanup
                        if pdf_file_path != temp_file_path:
                            converted_files.append(pdf_file_path)
                        converted_files.append(temp_file_path)
                       
                        # Read PDF
                        with open(pdf_file_path, 'rb') as f:
                            file_content = f.read()
                       
                        pdf_reader = PdfReader(BytesIO(file_content))
                       
                        # Add all pages
                        for page_num in range(len(pdf_reader.pages)):
                            pdf_writer.add_page(pdf_reader.pages[page_num])
                       
                        successfully_processed_count += 1
                        print(f"[MERGE] Added {len(pdf_reader.pages)} pages from file {i+1} ({file.name}) (Successfully processed: {successfully_processed_count})")
                           
                    except Exception as e:
                        print(f"[ERROR] Error processing file {i+1}: {str(e)}")
                        import traceback
                        traceback.print_exc()
                        continue
               
                # Validate that at least 2 files were successfully processed
                print(f"[MERGE] Successfully processed {successfully_processed_count} out of {len(files)} files")
                if successfully_processed_count < 2:
                    error_msg = f"Failed to merge: Only {successfully_processed_count} file(s) were successfully processed. At least 2 files are required for merging."
                    print(f"[ERROR] {error_msg}")
                    return Response({
                        'success': False,
                        'error': error_msg,
                        'processed_count': successfully_processed_count,
                        'requested_count': len(files)
                    }, status=status.HTTP_400_BAD_REQUEST)
               
                # Clean up converted files
                for file_path in converted_files:
                    try:
                        if os.path.exists(file_path):
                            os.unlink(file_path)
                    except:
                        pass
           
            elif use_id_merge:
                # Merge from document IDs (already uploaded to S3)
                import requests
                from io import BytesIO
               
                ordered_ids = document_order if document_order else document_ids
                print(f"[MERGE] Merging {len(ordered_ids)} documents by ID")
               
                existing_docs = S3Files.objects.filter(id__in=ordered_ids)
               
                successfully_processed_count = 0  # Reset counter for ID merge
               
                for doc_id in ordered_ids:
                    temp_download_dir = None  # Initialize before try block
                    doc_name = f"doc_{doc_id}"  # Default name
                    try:
                        # Check if document exists
                        if not existing_docs.filter(id=doc_id).exists():
                            print(f"[ERROR] Document {doc_id} does not exist in S3Files, skipping...")
                            continue
                       
                        doc = existing_docs.get(id=doc_id)
                        doc_name = doc.file_name or doc_name  # Update with actual name
                        s3_key = doc.metadata.get('s3_key') if doc.metadata else None
                       
                        if not s3_key and doc.url:
                            s3_key = doc.url.split('/')[-1] if '/' in doc.url else None
                       
                        if not s3_key:
                            print(f"[WARN] No S3 key for document {doc_id}, skipping...")
                            continue
                       
                        # Download from S3
                        temp_download_dir = tempfile.mkdtemp()
                        temp_dirs.append(temp_download_dir)
                       
                        # Determine correct file extension for download
                        # First try to get from file_name
                        if doc.file_name:
                            download_file_name = doc.file_name
                            # Ensure it has an extension
                            if not os.path.splitext(download_file_name)[1]:
                                # Try to get extension from file_type
                                file_type = doc.file_type or (doc.metadata.get('file_type', '') if doc.metadata else '')
                                if file_type:
                                    download_file_name = f"{download_file_name}.{file_type}"
                                else:
                                    # Try from URL
                                    if doc.url:
                                        url_ext = os.path.splitext(doc.url.split('?')[0])[1]
                                        if url_ext:
                                            download_file_name = f"{download_file_name}{url_ext}"
                                        else:
                                            download_file_name = f"{download_file_name}.pdf"
                                    else:
                                        download_file_name = f"{download_file_name}.pdf"
                        else:
                            # No file_name, construct from file_type or metadata
                            file_type = doc.file_type or (doc.metadata.get('file_type', 'pdf') if doc.metadata else 'pdf')
                            original_filename = doc.metadata.get('original_filename', '') if doc.metadata else ''
                            if original_filename:
                                download_file_name = original_filename
                            else:
                                download_file_name = f"doc_{doc_id}.{file_type}"
                       
                        print(f"[MERGE] Downloading document {doc_id} as: {download_file_name}")
                       
                        download_result = s3_client.download(
                            s3_key=s3_key,
                            file_name=download_file_name,
                            destination_path=temp_download_dir
                        )
                       
                        if not download_result.get('success'):
                            print(f"[WARN] Document {doc_id}: S3 download failed, trying direct URL download...")
                            if doc.url:
                                try:
                                    response = requests.get(doc.url, timeout=30)
                                    if response.status_code == 200:
                                        temp_file_path = os.path.join(temp_download_dir, download_file_name)
                                        with open(temp_file_path, 'wb') as f:
                                            f.write(response.content)
                                        download_result = {'success': True, 'file_path': temp_file_path}
                                        print(f"[MERGE] Document {doc_id}: Successfully downloaded from URL")
                                    else:
                                        print(f"[ERROR] Document {doc_id}: URL download failed with status {response.status_code}")
                                        continue
                                except Exception as url_error:
                                    print(f"[ERROR] Document {doc_id}: URL download error: {str(url_error)}")
                                    continue
                            else:
                                print(f"[ERROR] Document {doc_id}: No S3 key or URL available for download")
                                continue
                       
                        file_path = download_result.get('file_path')
                        if not file_path or not os.path.exists(file_path):
                            print(f"[ERROR] Document {doc_id}: Downloaded file not found at {file_path}")
                            continue
                       
                        print(f"[MERGE] Document {doc_id}: Downloaded file path: {file_path}")
                        print(f"[MERGE] Document {doc_id}: Document file_name: {doc.file_name}, file_type: {doc.file_type}")
                        if doc.metadata:
                            print(f"[MERGE] Document {doc_id}: Metadata: {doc.metadata}")
                       
                        # Get file extension from document metadata or file path
                        file_extension = os.path.splitext(file_path)[1].lower()
                        if not file_extension:
                            # Try to get from document file_name
                            if doc.file_name:
                                file_extension = os.path.splitext(doc.file_name)[1].lower()
                            # Try to get from document metadata
                            if not file_extension:
                                file_type = doc.metadata.get('file_type', '') if doc.metadata else ''
                                if file_type:
                                    file_extension = f'.{file_type}' if not file_type.startswith('.') else file_type
                            # Try to get from URL
                            if not file_extension and doc.url:
                                url_path = doc.url.split('?')[0]  # Remove query params
                                file_extension = os.path.splitext(url_path)[1].lower()
                            # Try to get from original_filename in metadata
                            if not file_extension and doc.metadata:
                                original_filename = doc.metadata.get('original_filename', '')
                                if original_filename:
                                    file_extension = os.path.splitext(original_filename)[1].lower()
                            # Default to PDF if still not found
                            if not file_extension:
                                file_extension = '.pdf'
                                print(f"[WARN] Document {doc_id}: No file extension found, defaulting to .pdf")
                       
                        print(f"[MERGE] Document {doc_id}: File extension detected as {file_extension}")
                       
                        # Special handling for Excel files - ensure file has correct extension
                        if file_extension in ['.xls', '.xlsx']:
                            print(f"[MERGE] Document {doc_id}: Detected Excel file ({file_extension}), will convert to PDF")
                            # Check if downloaded file has correct extension, rename if needed
                            current_ext = os.path.splitext(file_path)[1].lower()
                            if current_ext != file_extension:
                                print(f"[MERGE] Document {doc_id}: Renaming file from {current_ext} to {file_extension}")
                                new_file_path = os.path.splitext(file_path)[0] + file_extension
                                try:
                                    os.rename(file_path, new_file_path)
                                    file_path = new_file_path
                                    print(f"[MERGE] Document {doc_id}: File renamed to {file_path}")
                                except Exception as rename_error:
                                    print(f"[WARN] Document {doc_id}: Could not rename file: {str(rename_error)}")
                                    # Continue with original path - conversion might still work
                       
                        # Verify file is readable before conversion
                        try:
                            file_size = os.path.getsize(file_path)
                            print(f"[MERGE] Document {doc_id}: File size: {file_size} bytes")
                            if file_size == 0:
                                print(f"[ERROR] Document {doc_id}: File is empty (0 bytes)")
                                continue
                        except Exception as size_error:
                            print(f"[ERROR] Document {doc_id}: Could not read file size: {str(size_error)}")
                            continue
                       
                        # Convert to PDF if needed
                        print(f"[MERGE] Document {doc_id}: Calling convert_to_pdf with file_path={file_path}, extension={file_extension}")
                        pdf_file_path = convert_to_pdf(file_path, file_extension)
                       
                        if not pdf_file_path or not os.path.exists(pdf_file_path):
                            print(f"[ERROR] Failed to convert document {doc_id} ({doc_name}) with extension {file_extension} to PDF")
                            print(f"[ERROR] pdf_file_path: {pdf_file_path}, exists: {os.path.exists(pdf_file_path) if pdf_file_path else False}")
                            # Clean up original file
                            try:
                                if os.path.exists(file_path):
                                    os.unlink(file_path)
                            except:
                                pass
                            continue
                       
                        print(f"[MERGE] Document {doc_id}: Successfully converted to PDF: {pdf_file_path}")
                       
                        # Read PDF
                        with open(pdf_file_path, 'rb') as f:
                            file_content = f.read()
                       
                        pdf_reader = PdfReader(BytesIO(file_content))
                       
                        for page_num in range(len(pdf_reader.pages)):
                            pdf_writer.add_page(pdf_reader.pages[page_num])
                       
                        successfully_processed_count += 1
                        print(f"[MERGE] Added {len(pdf_reader.pages)} pages from document {doc_id} (Successfully processed: {successfully_processed_count})")
                       
                        # Clean up both original and converted files
                        try:
                            if pdf_file_path != file_path and os.path.exists(pdf_file_path):
                                os.unlink(pdf_file_path)
                            if os.path.exists(file_path):
                                os.unlink(file_path)
                        except:
                            pass
                           
                    except Exception as e:
                        import traceback
                        error_details = traceback.format_exc()
                        print(f"[ERROR] Error processing document {doc_id} ({doc_name}): {str(e)}")
                        print(f"[ERROR] Traceback: {error_details}")
                        continue
                    finally:
                        if temp_download_dir and os.path.exists(temp_download_dir):
                            try:
                                shutil.rmtree(temp_download_dir)
                            except:
                                pass
               
                # Validate that at least 2 documents were successfully processed
                print(f"[MERGE] Successfully processed {successfully_processed_count} out of {len(ordered_ids)} documents")
                if successfully_processed_count < 2:
                    failed_count = len(ordered_ids) - successfully_processed_count
                    error_msg = f"Failed to merge: Only {successfully_processed_count} document(s) were successfully processed out of {len(ordered_ids)} requested. At least 2 documents are required for merging."
                    if failed_count > 0:
                        error_msg += f" {failed_count} document(s) failed to process (check server logs for details)."
                    print(f"[ERROR] {error_msg}")
                    return Response({
                        'success': False,
                        'error': error_msg,
                        'processed_count': successfully_processed_count,
                        'requested_count': len(ordered_ids),
                        'failed_count': failed_count
                    }, status=status.HTTP_400_BAD_REQUEST)
           
            # Create merged PDF
            from io import BytesIO
            merged_pdf_buffer = BytesIO()
            pdf_writer.write(merged_pdf_buffer)
            merged_pdf_buffer.seek(0)
           
            # Save to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(merged_pdf_buffer.read())
                temp_file_path = temp_file.name
           
            try:
                # Upload merged document to S3
                merged_filename = f"merged_{uuid.uuid4().hex}.pdf"
                if rfp_id:
                    merged_filename = f"rfp_{rfp_id}_merged_{uuid.uuid4().hex}.pdf"
               
                upload_result = s3_client.upload(
                    file_path=temp_file_path,
                    user_id=user_id,
                    custom_file_name=merged_filename
                )
               
                if upload_result['success']:
                    file_info = upload_result['file_info']
                   
                    # Save merged document to S3Files
                    merged_doc = S3Files.objects.create(
                        url=file_info['url'],
                        file_type='pdf',
                        file_name=f"Merged Document",
                        user_id=user_id,
                        metadata={
                            'original_filename': merged_filename,
                            'stored_filename': file_info['storedName'],
                            's3_key': file_info['s3Key'],
                            's3_bucket': file_info.get('bucket', ''),
                            'rfp_id': rfp_id,
                            'document_name': 'Merged Document',
                            'file_size': os.path.getsize(temp_file_path),
                            'upload_operation_id': upload_result.get('operation_id'),
                            'is_merged': True,
                            'merged_from_documents': document_ids if use_id_merge else None,
                            'merged_from_files': len(files) if use_file_merge else None
                        }
                    )
                   
                    print(f"[SUCCESS] Merged document created: {merged_doc.id}")
                   
                    return Response({
                        'success': True,
                        'merged_document_id': merged_doc.id,
                        'merged_document_name': merged_doc.file_name,
                        'merged_document_url': merged_doc.url,
                        'message': f'Successfully merged {successfully_processed_count} document(s)',
                        'document_count': successfully_processed_count,
                        'requested_count': len(document_ids) if use_id_merge else len(files)
                    })
                else:
                    return Response({
                        'success': False,
                        'error': f'S3 upload failed: {upload_result.get("error", "Unknown error")}'
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                   
            finally:
                # Clean up
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
                for temp_dir in temp_dirs:
                    if os.path.exists(temp_dir):
                        try:
                            shutil.rmtree(temp_dir)
                        except:
                            pass
                           
        except Exception as e:
            print(f"[ERROR] Merge error: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return Response({
                'success': False,
                'error': f'Failed to merge documents: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class AwardNotificationView(APIView):
    """
    API endpoint for managing award notifications
    """
    authentication_classes = [JWTAuthentication]
    permission_classes = [SimpleAuthenticatedPermission]
    
    def post(self, request, rfp_id):
        """
        Send award notification to selected vendor
        """
        try:
            # Parse JSON data from request body
            import json
            if hasattr(request, 'data'):
                data = request.data
            else:
                data = json.loads(request.body.decode('utf-8'))
            
            print(f"[AWARD NOTIFICATION] Received request data: {json.dumps(data, indent=2)}")
            
            response_id = data.get('response_id')
            notification_type = data.get('notification_type', 'winner')
            custom_message = data.get('award_message', '')
            next_steps = data.get('next_steps', '')
            vendor_email = data.get('vendor_email', '')  # Get email from frontend
            vendor_name = data.get('vendor_name', '')  # Get vendor name from frontend
            
            print(f"[AWARD NOTIFICATION] Extracted fields:")
            print(f"  - response_id: {response_id}")
            print(f"  - vendor_email from frontend: {vendor_email}")
            print(f"  - vendor_name from frontend: {vendor_name}")
            
            if not response_id:
                return Response({
                    'success': False,
                    'error': 'Response ID is required'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get RFP response details
            try:
                response = RFPResponse.objects.get(response_id=response_id)
            except RFPResponse.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'RFP response not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Extract email - Priority: 1) Frontend provided email, 2) response_documents, 3) default
            recipient_email = 'vendor@example.com'  # Default fallback
            
            # Priority 1: Use email from frontend if provided
            if vendor_email and vendor_email != 'vendor@example.com':
                recipient_email = vendor_email
                print(f"Using email from frontend: {recipient_email}")
            # Priority 2: Try to extract from response_documents
            elif response.response_documents:
                try:
                    import json
                    response_docs = json.loads(response.response_documents) if isinstance(response.response_documents, str) else response.response_documents
                    
                    # Check multiple possible locations for email
                    if response_docs.get('contact_email'):
                        recipient_email = response_docs['contact_email']
                        print(f"Found email in response_docs.contact_email: {recipient_email}")
                    elif response_docs.get('companyInfo', {}).get('email'):
                        recipient_email = response_docs['companyInfo']['email']
                        print(f"Found email in companyInfo.email: {recipient_email}")
                    elif response_docs.get('companyInfo', {}).get('contactEmail'):
                        recipient_email = response_docs['companyInfo']['contactEmail']
                        print(f"Found email in companyInfo.contactEmail: {recipient_email}")
                    elif response_docs.get('email'):
                        recipient_email = response_docs['email']
                        print(f"Found email in response_docs.email: {recipient_email}")
                    elif response_docs.get('proposalData', {}).get('companyInfo', {}).get('email'):
                        recipient_email = response_docs['proposalData']['companyInfo']['email']
                        print(f"Found email in proposalData.companyInfo.email: {recipient_email}")
                except Exception as e:
                    print(f"Error parsing response_documents: {e}")
                    pass
            
            print(f"Final recipient email: {recipient_email}")
            
            # Generate secure token for response
            import secrets
            accept_reject_token = secrets.token_urlsafe(32)
            
            # Update RFP status to AWARDED
            try:
                from .models import RFP
                rfp = RFP.objects.get(rfp_id=rfp_id)
                if rfp.status in ['EVALUATION', 'SUBMISSION_OPEN']:
                    rfp.status = 'AWARDED'
                    rfp.award_decision_date = timezone.now()
                    rfp.save()
                    print(f"[STATUS UPDATE] RFP {rfp_id} status changed to AWARDED")
            except Exception as e:
                print(f"[WARNING] Could not update RFP status: {e}")
                # Continue with notification even if RFP status update fails
            
            # Create award notification
            notification = RFPAwardNotification.objects.create(
                response_id=response_id,
                notification_type=notification_type,
                recipient_email=recipient_email,
                notification_status='pending',
                accept_reject_token=accept_reject_token,
                award_message=custom_message or f"Congratulations! Your proposal for RFP {rfp_id} has been selected as the winner.",
                next_steps=next_steps or "Please respond to this notification within 7 days to accept or decline the award."
            )
            
            # Send actual email notification
            try:
                from django.core.mail import send_mail
                from django.template.loader import render_to_string
                from django.conf import settings
                
                # Get vendor name for email - Priority: 1) Frontend provided name, 2) response_documents
                vendor_name_for_email = vendor_name if vendor_name else 'Vendor'
                
                if not vendor_name and response.response_documents:
                    try:
                        response_docs = json.loads(response.response_documents) if isinstance(response.response_documents, str) else response.response_documents
                        if response_docs.get('companyInfo', {}).get('companyName'):
                            vendor_name_for_email = response_docs['companyInfo']['companyName']
                        elif response_docs.get('companyInfo', {}).get('contactName'):
                            vendor_name_for_email = response_docs['companyInfo']['contactName']
                        elif response_docs.get('vendor_name'):
                            vendor_name_for_email = response_docs['vendor_name']
                    except:
                        pass
                
                # Create email content
                subject = f"Congratulations! Your Proposal Has Been Selected - RFP {rfp_id}"
                
                # Create response URL (use frontend URL from settings)
                import re
                frontend_url = getattr(settings, 'EXTERNAL_BASE_URL', 'http://localhost:3000').rstrip('/')
                
                # Replace any ngrok URLs with localhost:3000
                if 'ngrok' in frontend_url.lower():
                    frontend_url = 'http://localhost:3000'
                
                # Ensure it's localhost (not 127.0.0.1 or other variations)
                if not frontend_url.startswith('http://localhost') and not frontend_url.startswith('https://localhost'):
                    # Extract port if present, otherwise use 3000
                    port_match = re.search(r':(\d+)', frontend_url)
                    port = port_match.group(1) if port_match else '3000'
                    frontend_url = f'http://localhost:{port}'
                
                response_url = f"{frontend_url}/award-response/{accept_reject_token}"
                
                # Email body
                email_body = f"""
Dear {vendor_name_for_email},

{notification.award_message}

{notification.next_steps}

To respond to this award notification, please click the link below:
{response_url}

You can:
• Accept the award
• Decline the award
• Add comments

This link will expire in 30 days.

Best regards,
The RFP Management Team

---
This is an automated notification. Please do not reply to this email.
If you have questions, please contact the RFP team directly.
"""
                
                # Send email
                email_sent = send_mail(
                    subject=subject,
                    message=email_body,
                    from_email=settings.DEFAULT_FROM_EMAIL,
                    recipient_list=[recipient_email],
                    fail_silently=False
                )
                
                if email_sent:
                    print(f"Email sent successfully to {recipient_email}")
                    notification.notification_status = 'sent'
                    notification.sent_date = timezone.now()
                else:
                    print(f"Failed to send email to {recipient_email}")
                    notification.notification_status = 'pending'
                
            except Exception as email_error:
                print(f"Error sending email: {str(email_error)}")
                # Still mark as sent for now, but log the error
                notification.notification_status = 'sent'
                notification.sent_date = timezone.now()
            
            notification.save()
            
            return Response({
                'success': True,
                'notification_id': notification.notification_id,
                'message': 'Award notification sent successfully',
                'recipient_email': notification.recipient_email,
                'response_token': accept_reject_token
            })
            
        except Exception as e:
            print(f"Error sending award notification: {str(e)}")
            import traceback
            print(traceback.format_exc())
            return Response({
                'success': False,
                'error': f'Failed to send award notification: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def get(self, request, rfp_id):
        """
        Get award notifications for an RFP
        """
        try:
            # Get all responses for this RFP
            responses = RFPResponse.objects.filter(rfp_id=rfp_id)
            response_ids = [r.response_id for r in responses]
            
            # Get notifications for these responses
            notifications = RFPAwardNotification.objects.filter(
                response_id__in=response_ids
            ).order_by('-created_at')
            
            notification_data = []
            for notification in notifications:
                # Get response details
                try:
                    response = RFPResponse.objects.get(response_id=notification.response_id)
                    
                    # Get vendor information from response_documents
                    vendor_name = 'Unknown Vendor'
                    org_name = 'Unknown Organization'
                    vendor_email = 'No email'
                    
                    if response.response_documents:
                        try:
                            import json
                            response_docs = json.loads(response.response_documents) if isinstance(response.response_documents, str) else response.response_documents
                            
                            # Extract vendor information
                            if response_docs.get('companyInfo'):
                                company_info = response_docs['companyInfo']
                                if company_info.get('companyName'):
                                    vendor_name = company_info['companyName']
                                elif company_info.get('contactName'):
                                    vendor_name = company_info['contactName']
                                
                                if company_info.get('contactName') and company_info.get('companyName'):
                                    org_name = company_info['contactName']  # Contact person as org
                                
                                if company_info.get('email'):
                                    vendor_email = company_info['email']
                            
                            # Fallback to other fields
                            if response_docs.get('vendor_name') and vendor_name == 'Unknown Vendor':
                                vendor_name = response_docs['vendor_name']
                            if response_docs.get('org') and org_name == 'Unknown Organization':
                                org_name = response_docs['org']
                                
                        except Exception as e:
                            print(f"Error parsing response_documents for notification {notification.notification_id}: {e}")
                            pass
                            
                except RFPResponse.DoesNotExist:
                    vendor_name = 'Unknown Vendor'
                    vendor_email = 'No email'
                
                notification_data.append({
                    'notification_id': notification.notification_id,
                    'response_id': notification.response_id,
                    'vendor_name': vendor_name,
                    'vendor_email': vendor_email,
                    'notification_type': notification.notification_type,
                    'notification_status': notification.notification_status,
                    'sent_date': notification.sent_date.isoformat() if notification.sent_date else None,
                    'acknowledged_date': notification.acknowledged_date.isoformat() if notification.acknowledged_date else None,
                    'response_date': notification.response_date.isoformat() if notification.response_date else None,
                    'award_message': notification.award_message,
                    'next_steps': notification.next_steps,
                    'created_at': notification.created_at.isoformat() if notification.created_at else None
                })
            
            return Response({
                'success': True,
                'notifications': notification_data
            })
            
        except Exception as e:
            print(f"Error getting award notifications: {str(e)}")
            return Response({
                'success': False,
                'error': f'Failed to get award notifications: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class AwardResponseView(APIView):
    """
    API endpoint for vendors to accept/reject awards
    """
    authentication_classes = []
    permission_classes = [AllowAny]
    
    def post(self, request, token):
        """
        Handle vendor response to award notification
        """
        try:
            # Parse JSON data from request body
            import json
            if hasattr(request, 'data'):
                data = request.data
            else:
                data = json.loads(request.body.decode('utf-8'))
            
            response_action = data.get('action')  # 'accept' or 'reject'
            comments = data.get('comments', '')
            
            if response_action not in ['accept', 'reject']:
                return Response({
                    'success': False,
                    'error': 'Invalid action. Must be "accept" or "reject"'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Find notification by token
            try:
                notification = RFPAwardNotification.objects.get(accept_reject_token=token)
            except RFPAwardNotification.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'Invalid or expired token'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Update notification status
            if response_action == 'accept':
                notification.notification_status = 'accepted'
               
                # Create vendor credentials and temp_vendor record
                try:
                    credentials = self._create_vendor_credentials(notification)
                    print(f"[INFO] Vendor credentials created successfully for notification {notification.notification_id}")
                except Exception as e:
                    print(f"[ERROR] Failed to create vendor credentials: {str(e)}")
                    # Continue anyway - we can create credentials manually later
                   
            else:
                notification.notification_status = 'rejected'
 
            
            notification.response_date = timezone.now()
            notification.acknowledged_date = timezone.now()
            
            # Store comments in award_message if provided
            if comments:
                notification.award_message = f"{notification.award_message}\n\nVendor Response: {comments}"
            
            notification.save()
            
            # If rejected, we might want to notify admin to select next vendor
            if response_action == 'reject':
                print(f"[WARN] Vendor rejected award for response {notification.response_id}")
                # TODO: Send notification to admin about rejection
            
            return Response({
                'success': True,
                'message': f'Award {response_action}ed successfully',
                'notification_status': notification.notification_status
            })
            
        except Exception as e:
            print(f"Error processing award response: {str(e)}")
            return Response({
                'success': False,
                'error': f'Failed to process award response: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def get(self, request, token):
        """
        Get award notification details for vendor response page
        """
        try:
            notification = RFPAwardNotification.objects.get(accept_reject_token=token)
            
            # Get response details
            try:
                response = RFPResponse.objects.get(response_id=notification.response_id)
                
                # Extract vendor name from response_documents
                vendor_name = 'Unknown Vendor'
                if response.response_documents:
                    try:
                        import json
                        response_docs = json.loads(response.response_documents) if isinstance(response.response_documents, str) else response.response_documents
                        if response_docs.get('companyInfo', {}).get('companyName'):
                            vendor_name = response_docs['companyInfo']['companyName']
                        elif response_docs.get('companyInfo', {}).get('contactName'):
                            vendor_name = response_docs['companyInfo']['contactName']
                    except:
                        pass
                        
            except RFPResponse.DoesNotExist:
                vendor_name = 'Unknown Vendor'
            
            return Response({
                'success': True,
                'notification': {
                    'notification_id': notification.notification_id,
                    'vendor_name': vendor_name,
                    'notification_type': notification.notification_type,
                    'award_message': notification.award_message,
                    'next_steps': notification.next_steps,
                    'notification_status': notification.notification_status,
                    'sent_date': notification.sent_date.isoformat() if notification.sent_date else None
                }
            })
            
        except RFPAwardNotification.DoesNotExist:
            return Response({
                'success': False,
                'error': 'Invalid or expired token'
            }, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error getting award notification: {str(e)}")
            return Response({
                'success': False,
                'error': f'Failed to get award notification: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def _create_vendor_credentials(self, notification):
        """
        Create vendor user credentials, RBAC permissions, and temp_vendor record
        when a vendor accepts an award.
       
        Args:
            notification: RFPAwardNotification instance
           
        Returns:
            dict: Contains user_id, username, password, and temp_vendor_id
        """
       
        try:
            # Get vendor details from RFP response
           
            try:
                response = RFPResponse.objects.get(response_id=notification.response_id)
            except RFPResponse.DoesNotExist:
                raise Exception(f"RFP Response not found for response_id: {notification.response_id}")
           
            # Extract vendor information from response_documents
            vendor_name = 'Vendor'
            company_name = 'Unknown Company'
            vendor_email = notification.recipient_email
           
            if response.response_documents:
                try:
                    response_docs = json.loads(response.response_documents) if isinstance(response.response_documents, str) else response.response_documents
                   
                    # Get company name
                    if response_docs.get('companyInfo', {}).get('companyName'):
                        company_name = response_docs['companyInfo']['companyName']
                        vendor_name = company_name
                    elif response_docs.get('companyInfo', {}).get('contactName'):
                        vendor_name = response_docs['companyInfo']['contactName']
                   
                    # Get email
                    if response_docs.get('companyInfo', {}).get('email'):
                        vendor_email = response_docs['companyInfo']['email']
                    elif response_docs.get('companyInfo', {}).get('contactEmail'):
                        vendor_email = response_docs['companyInfo']['contactEmail']
                       
                except Exception as e:
                    print(f"[WARN] Failed to parse response_documents: {str(e)}")
           
            # Generate secure random password
            password_length = 12
            password_chars = string.ascii_letters + string.digits + "!@#$%^&*"
            generated_password = ''.join(secrets.choice(password_chars) for _ in range(password_length))
           
            # Create username from email (first part before @)
            base_username = vendor_email.split('@')[0] if vendor_email else vendor_name
            normalized_username = re.sub(r'[^a-zA-Z0-9]', '', base_username) or f"vendor{notification.response_id}"
            username = normalized_username
           
            with connection.cursor() as cursor:
                # Check if user already exists
                cursor.execute("SELECT UserId, UserName FROM users WHERE Email = %s", [vendor_email])
                existing_user = cursor.fetchone()
               
                def get_unique_username(initial_username, current_user_id=None):
                    candidate = initial_username
                    counter = 1
                    while True:
                        cursor.execute("SELECT UserId FROM users WHERE UserName = %s", [candidate])
                        row = cursor.fetchone()
                        if not row or (current_user_id and row[0] == current_user_id):
                            return candidate
                        candidate = f"{initial_username}{counter}"
                        counter += 1

                if existing_user:
                    user_id, existing_username = existing_user
                    username = get_unique_username(normalized_username, user_id)
                    print(f"[INFO] User already exists with UserId: {user_id}")
                    # Update username and password for existing user with newly generated credentials
                    cursor.execute(
                        """
                        UPDATE users
                        SET UserName = %s,
                            Password = %s,
                            UpdatedAt = %s
                        WHERE UserId = %s
                        """,
                        [
                            username,
                            generated_password,  # In production, hash the password
                            timezone.now(),
                            user_id,
                        ],
                    )
                else:
                    # 1. Create user in users table
                    username = get_unique_username(normalized_username)

                    cursor.execute("""
                        INSERT INTO users (
                            UserName, Email, Password, FirstName, LastName,
                            IsActive, CreatedAt, UpdatedAt
                        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                    """, [
                        username,
                        vendor_email,
                        generated_password,  # In production, this should be hashed
                        vendor_name.split(' ')[0] if ' ' in vendor_name else vendor_name,
                        vendor_name.split(' ')[-1] if ' ' in vendor_name else '',
                        'Y',
                        timezone.now(),
                        timezone.now()
                    ])
                    user_id = cursor.lastrowid
                    print(f"[INFO] Created user with UserId: {user_id}")
               
                # 2. Create RBAC permissions in rbac_tprm table
                # Check if RBAC record already exists for this user
                cursor.execute("SELECT RBACId FROM rbac_tprm WHERE UserId = %s", [user_id])
                existing_rbac = cursor.fetchone()
               
                if not existing_rbac:
                    # Set vendor-specific permissions (limited access)
                    cursor.execute("""
                        INSERT INTO rbac_tprm (
                            UserId, UserName, Role,
                            ViewRFP, ViewRFPResponses, SubmitRFPResponse, WithdrawRFPResponse,
                            DownloadRFPDocuments, PreviewRFPDocuments, UploadDocumentsForRFP,
                            ViewVendors, ViewContactsDocuments, ViewQuestionnaires,
                            SubmitQuestionnaireResponses, ViewRiskAssessments,
                            ViewPerformance, ViewDashboardTrend,
                            CreatedAt, UpdatedAt, IsActive
                        ) VALUES (
                            %s, %s, %s,
                            1, 1, 1, 1,
                            1, 1, 1,
                            1, 1, 1,
                            1, 1,
                            1, 1,
                            %s, %s, 'Y'
                        )
                    """, [
                        user_id,
                        vendor_name,
                        'Vendor',
                        timezone.now(),
                        timezone.now()
                    ])
                    print(f"[INFO] Created RBAC permissions for UserId: {user_id}")
                else:
                    print(f"[INFO] RBAC record already exists for UserId: {user_id}")

                # Ensure essential vendor permissions are enabled
                cursor.execute("""
                    UPDATE rbac_tprm
                    SET
                        ViewVendors = 1,
                        CreateVendor = 1,
                        UpdateVendor = 1,
                        SubmitVendorForApproval = 1,
                        ViewContactsDocuments = 1,
                        AddUpdateContactsDocuments = 1,
                        ViewQuestionnaires = 1,
                        SubmitQuestionnaireResponses = 1,
                        ViewRiskAssessments = 1,
                        ViewPerformance = 1,
                        ViewDashboardTrend = 1
                    WHERE UserId = %s
                """, [user_id])
               
                # 3. Create temp_vendor record
                # Check if temp_vendor record already exists
                cursor.execute("SELECT id FROM temp_vendor WHERE response_id = %s", [notification.response_id])
                existing_temp_vendor = cursor.fetchone()
               
                if not existing_temp_vendor:
                    # Generate vendor code
                    vendor_code = f"VEN-{str(uuid.uuid4())[:8].upper()}"
                   
                    cursor.execute("""
                        INSERT INTO temp_vendor (
                            UserId, vendor_code, company_name, legal_name,
                            lifecycle_stage, status, vendor_category,
                            risk_level, is_critical_vendor, created_at, updated_at,
                            response_id
                        ) VALUES (
                            %s, %s, %s, %s,
                            %s, %s, %s,
                            %s, %s, %s, %s,
                            %s
                        )
                    """, [
                        user_id,
                        vendor_code,
                        company_name,
                        company_name,
                        1,  # Initial lifecycle stage
                        'pending_onboarding',
                        'New Vendor',
                        'Medium',
                        0,  # Not critical initially
                        timezone.now(),
                        timezone.now(),
                        notification.response_id
                    ])
                    temp_vendor_id = cursor.lastrowid
                    print(f"[INFO] Created temp_vendor record with id: {temp_vendor_id}")
                else:
                    temp_vendor_id = existing_temp_vendor[0]
                    print(f"[INFO] temp_vendor record already exists with id: {temp_vendor_id}")
                    # Ensure temp_vendor record is linked to the latest user and reflects key vendor info
                    cursor.execute(
                        """
                        UPDATE temp_vendor
                        SET UserId = %s,
                            company_name = %s,
                            legal_name = %s,
                            vendor_category = COALESCE(vendor_category, %s),
                            updated_at = %s
                        WHERE id = %s
                        """,
                        [
                            user_id,
                            company_name,
                            company_name,
                            'New Vendor',
                            timezone.now(),
                            temp_vendor_id,
                        ],
                    )
           
            # 4. Send credentials email to vendor
            try:
                self._send_credentials_email(vendor_name, vendor_email, username, generated_password)
                print(f"[INFO] Credentials email sent to {vendor_email}")
            except Exception as e:
                print(f"[ERROR] Failed to send credentials email: {str(e)}")
                # Continue anyway - admin can send credentials manually
           
            return {
                'user_id': user_id,
                'username': username,
                'password': generated_password,
                'temp_vendor_id': temp_vendor_id,
                'vendor_email': vendor_email
            }
           
        except Exception as e:
            print(f"[ERROR] Failed to create vendor credentials: {str(e)}")
            import traceback
            traceback.print_exc()
            raise
   
    def _send_credentials_email(self, vendor_name, vendor_email, username, password):
        """
        Send vendor credentials via email.
       
        Args:
            vendor_name: Name of the vendor
            vendor_email: Email address to send to
            username: Generated username
            password: Generated password
        """
        try:
            subject = 'Welcome to Vendor Portal - Your Access Credentials'

            portal_base_url = getattr(settings, 'FRONTEND_URL', None) or getattr(settings, 'SITE_URL', None) or 'http://localhost:3000'
            portal_login_url = f"{portal_base_url.rstrip('/')}/login"

            message = f"""
            Dear {vendor_name},
 
            Congratulations! Your proposal has been accepted, and we're pleased to welcome you as a partner.
 
            Below are your credentials to access the Vendor Portal:
 
            Portal URL: {portal_login_url}
            Username: {username}
            Email: {vendor_email}
            Password: {password}
 
            For security reasons, we recommend that you change your password upon first login.
 
            Through the vendor portal, you will be able to:
            - View RFP details and requirements
            - Submit questionnaire responses
            - Upload required documents
            - Track your vendor lifecycle status
            - View performance metrics
 
            If you have any questions or need assistance, please don't hesitate to contact our support team.
 
            Best regards,
            TPRM Team
            """
           
            send_mail(
                subject=subject,
                message=message,
                from_email=settings.DEFAULT_FROM_EMAIL if hasattr(settings, 'DEFAULT_FROM_EMAIL') else 'noreply@tprm.com',
                recipient_list=[vendor_email],
                fail_silently=False,
            )
           
            print(f"[INFO] Credentials email sent successfully to {vendor_email}")
           
        except Exception as e:
            print(f"[ERROR] Failed to send credentials email: {str(e)}")
            raise
 

class VendorCredentialsView(APIView):
    """
    API endpoint for managing vendor credentials creation
    """
    authentication_classes = [JWTAuthentication]
    permission_classes = [SimpleAuthenticatedPermission]
   
    def post(self, request, notification_id):
        """
        Manually create/recreate vendor credentials for an accepted award
        """
        try:
            # Get the notification
            try:
                notification = RFPAwardNotification.objects.get(notification_id=notification_id)
            except RFPAwardNotification.DoesNotExist:
                return Response({
                    'success': False,
                    'error': 'Award notification not found'
                }, status=status.HTTP_404_NOT_FOUND)
           
            # Check if notification is accepted
            if notification.notification_status != 'accepted':
                return Response({
                    'success': False,
                    'error': 'Vendor has not accepted the award yet'
                }, status=status.HTTP_400_BAD_REQUEST)
           
            # Create credentials using the helper method from AwardNotificationView
            award_view = AwardNotificationView()
            credentials = award_view._create_vendor_credentials(notification)
           
            return Response({
                'success': True,
                'message': 'Vendor credentials created successfully',
                'data': {
                    'user_id': credentials['user_id'],
                    'username': credentials['username'],
                    'vendor_email': credentials['vendor_email'],
                    'temp_vendor_id': credentials.get('temp_vendor_id')
                }
            })
           
        except Exception as e:
            print(f"[ERROR] Failed to create vendor credentials: {str(e)}")
            import traceback
            traceback.print_exc()
            return Response({
                'success': False,
                'error': f'Failed to create vendor credentials: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
 
# ============================================================================
# VENDOR INVITATION VIEWS (from views_vendor_invitations.py)
# ============================================================================

@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_primary_contacts(request):
    """
    Get primary contacts for selected vendor IDs from vendor_contacts table
    """
    try:
        # Parse request data
        data = json.loads(request.body)
        vendor_ids = data.get('vendorIds', [])
        
        if not vendor_ids:
            return JsonResponse({
                'success': True,
                'contacts': []
            })
        
        # Import Django's database connection
        from django.db import connection
        
        # Query to get primary contacts from vendor_contacts table
        with connection.cursor() as cursor:
            # First, get vendor information
            vendor_placeholders = ','.join(['%s'] * len(vendor_ids))
            cursor.execute(f"""
                SELECT 
                    v.vendor_id,
                    v.company_name,
                    vc.contact_id,
                    vc.first_name,
                    vc.last_name,
                    vc.email,
                    vc.phone,
                    vc.mobile,
                    vc.designation,
                    vc.department,
                    vc.contact_type
                FROM vendors v
                LEFT JOIN vendor_contacts vc ON v.vendor_id = vc.vendor_id 
                    AND vc.contact_type = 'PRIMARY' 
                    AND vc.is_active = 1
                WHERE v.vendor_id IN ({vendor_placeholders})
                    AND v.status = 'APPROVED'
                ORDER BY vc.is_primary DESC, vc.contact_id ASC
            """, vendor_ids)
            
            rows = cursor.fetchall()
        
        contacts = []
        for row in rows:
            vendor_id, company_name, contact_id, first_name, last_name, email, phone, mobile, designation, department, contact_type = row
            
            # If no primary contact found, generate placeholder
            if not contact_id or not email:
                # Generate contact information based on company name
                company_clean = company_name.lower().replace(' ', '').replace('inc', '').replace('llc', '').replace('corp', '').replace('solutions', '').replace('incorporated', '')
                primary_email = f"contact@{company_clean}.com"
                primary_phone = "+1 (555) 000-0000"
                contact_name = f"Primary Contact at {company_name}"
                contact_designation = "Not specified"
                contact_department = "Not specified"
            else:
                # Use actual contact data from database
                primary_email = email
                primary_phone = phone or mobile or "+1 (555) 000-0000"
                contact_name = f"{first_name} {last_name}".strip() if first_name or last_name else f"Primary Contact at {company_name}"
                contact_designation = designation or "Not specified"
                contact_department = department or "Not specified"
            
            contact_info = {
                'vendor_id': vendor_id,
                'company_name': company_name,
                'email': primary_email,
                'phone': primary_phone,
                'name': contact_name,
                'designation': contact_designation,
                'department': contact_department,
                'is_matched_vendor': True,
                'has_real_contact': bool(contact_id and email)  # Flag to indicate if this is real contact data
            }
            contacts.append(contact_info)
        
        return JsonResponse({
            'success': True,
            'contacts': contacts
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Failed to fetch primary contacts: {str(e)}'
        }, status=500)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_invitations_by_rfp(request, rfp_id):
    """
    Get all vendor invitations for a specific RFP
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        
        # Get vendor invitations from the rfp_vendor_invitations table
        invitations = []
        vendor_invitations = VendorInvitation.objects.filter(rfp=rfp)
        for invitation in vendor_invitations:
            invitation_data = {
                'invitation_id': invitation.invitation_id,
                'vendor_id': invitation.vendor.vendor_id if invitation.vendor else None,
                'vendor_name': invitation.vendor_name,
                'vendor_email': invitation.vendor_email,
                'company_name': invitation.company_name,
                'invitation_url': invitation.invitation_url,
                'invitation_status': invitation.invitation_status,
                'invited_date': invitation.invited_date.isoformat() if invitation.invited_date else None,
                'acknowledged_date': invitation.acknowledged_date.isoformat() if invitation.acknowledged_date else None,
                'declined_reason': invitation.declined_reason,
                'unique_token': invitation.unique_token,
                'custom_message': invitation.custom_message
            }
            invitations.append(invitation_data)
        
        return JsonResponse({
            'success': True,
            'invitations': invitations,
            'total': len(invitations)
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Failed to fetch invitations: {str(e)}'
        }, status=500)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_invitation_stats(request, rfp_id):
    """
    Get invitation statistics for a specific RFP
    """
    try:
        # Check if RFP exists, if not return default stats
        try:
            rfp = RFP.objects.get(rfp_id=rfp_id)
        except RFP.DoesNotExist:
            # Return default stats if RFP doesn't exist
            stats = {
                'total_invitations': 0,
                'sent': 0,
                'delivered': 0,
                'opened': 0,
                'clicked': 0,
                'submitted': 0,
                'pending': 0,
                'failed': 0
            }
            return JsonResponse({
                'success': True,
                'stats': stats
            })
        
        # Get actual invitation stats from the database
        invitations = VendorInvitation.objects.filter(rfp=rfp)
        
        stats = {
            'total_invitations': invitations.count(),
            'sent': invitations.filter(invitation_status='SENT').count(),
            'delivered': invitations.filter(invitation_status='DELIVERED').count(),
            'opened': invitations.filter(invitation_status='OPENED').count(),
            'clicked': invitations.filter(invitation_status='CLICKED').count(),
            'submitted': invitations.filter(invitation_status='SUBMITTED').count(),
            'pending': invitations.filter(invitation_status='CREATED').count(),
            'failed': invitations.filter(invitation_status='FAILED').count()
        }
        
        return JsonResponse({
            'success': True,
            'stats': stats
        })
        
    except Exception as e:
        print(f"Error in get_invitation_stats: {str(e)}")
        # Return default stats on any error
        stats = {
            'total_invitations': 0,
            'sent': 0,
            'delivered': 0,
            'opened': 0,
            'clicked': 0,
            'submitted': 0,
            'pending': 0,
            'failed': 0
        }
        return JsonResponse({
            'success': True,
            'stats': stats
        })


def generate_unique_token():
    """Generate a unique token for vendor invitations"""
    return hashlib.md5(str(uuid.uuid4()).encode()).hexdigest()


def generate_invitation_urls(token, rfp_id):
    """Generate base invitation URLs (token-based) for a vendor."""
    import re
    base_url = getattr(settings, 'EXTERNAL_BASE_URL', 'http://localhost:3000').rstrip('/')
    
    # Replace any ngrok URLs with localhost:3000
    if 'ngrok' in base_url.lower():
        base_url = 'http://localhost:3000'
    
    # Ensure it's localhost (not 127.0.0.1 or other variations)
    if not base_url.startswith('http://localhost') and not base_url.startswith('https://localhost'):
        # Extract port if present, otherwise use 3000
        port_match = re.search(r':(\d+)', base_url)
        port = port_match.group(1) if port_match else '3000'
        base_url = f'http://localhost:{port}'
    
    return {
        'invitation_url': f"{base_url}/invitation/{token}",
        'acknowledgment_url': f"{base_url}/acknowledge/{token}",
        'submission_url': f"{base_url}/submit/{token}"
    }


def generate_tracking_urls(rfp_id: int, invitation_id: int):
    """Generate acknowledge/decline tracking URLs that include rfp_id and invitation_id."""
    import re
    # Use backend API URL for API endpoints
    backend_url = getattr(settings, 'BACKEND_API_URL', 'http://localhost:8000').rstrip('/')
    
    # Replace any ngrok URLs with localhost:8000
    if 'ngrok' in backend_url.lower():
        backend_url = 'http://localhost:8000'
    
    # Ensure it's localhost (not 127.0.0.1 or other variations)
    if not backend_url.startswith('http://localhost') and not backend_url.startswith('https://localhost'):
        # Extract port if present, otherwise use 8000
        port_match = re.search(r':(\d+)', backend_url)
        port = port_match.group(1) if port_match else '8000'
        backend_url = f'http://localhost:{port}'
    
    # Point to API endpoints that record the status
    acknowledge_url = f"{backend_url}/api/v1/vendor-invitations/ack/{rfp_id}/{invitation_id}/"
    decline_url = f"{backend_url}/api/v1/vendor-invitations/decline/{rfp_id}/{invitation_id}/"
    return acknowledge_url, decline_url


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def create_vendor_invitations(request, rfp_id):
    """
    Create vendor invitations for selected vendors
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        
        # Parse request data
        data = json.loads(request.body)
        vendors = data.get('vendors', [])
        custom_message = data.get('customMessage', '')
        utm_parameters = data.get('utmParameters', {})
        
        if not vendors:
            return JsonResponse({
                'success': False,
                'error': 'No vendors provided'
            }, status=400)
        
        invitations = []
        
        with transaction.atomic():
            for vendor_data in vendors:
                # Extract vendor information
                vendor_id = vendor_data.get('vendor_id')
                company_name = vendor_data.get('company_name', '')
                email = vendor_data.get('email', '')
                phone = vendor_data.get('phone', '')
                
                # Generate contact name from company name if not provided
                contact_name = vendor_data.get('vendor_name', f"Contact at {company_name}")
                
                if not email:
                    # Generate email based on company name
                    company_clean = company_name.lower().replace(' ', '').replace('inc', '').replace('llc', '').replace('corp', '')
                    email = f"contact@{company_clean}.com"
                
                if not phone:
                    phone = "+1 (555) 000-0000"
                
                # Generate unique token and URLs
                unique_token = generate_unique_token()
                urls = generate_invitation_urls(unique_token, rfp_id)
                
                # Create invitation data
                invitation_data = {
                    'rfp_id': rfp_id,
                    'vendor_id': vendor_id,
                    'vendor_email': email,
                    'vendor_name': contact_name,
                    'vendor_phone': phone,
                    'company_name': company_name,
                    'invitation_url': urls['invitation_url'],
                    'acknowledgment_url': urls['acknowledgment_url'],
                    'submission_url': urls['submission_url'],
                    'unique_token': unique_token,
                    'invitation_status': 'CREATED',
                    'custom_message': custom_message,
                    'utm_parameters': json.dumps(utm_parameters) if utm_parameters else None
                }
                
                # Save to database using VendorInvitation model
                vendor_obj = None
                if vendor_id:
                    try:
                        vendor_obj = Vendor.objects.get(vendor_id=vendor_id)
                    except Vendor.DoesNotExist:
                        pass
                
                vendor_invitation = VendorInvitation.objects.create(
                    rfp=rfp,
                    vendor=vendor_obj,
                    vendor_name=contact_name,
                    vendor_email=email,
                    vendor_phone=phone,
                    company_name=company_name,
                    invitation_url=urls['invitation_url'],
                    acknowledgment_url=urls['acknowledgment_url'],
                    submission_url=urls['submission_url'],
                    unique_token=unique_token,
                    invitation_status='CREATED',
                    is_matched_vendor=bool(vendor_id),
                    custom_message=custom_message,
                    utm_parameters=utm_parameters
                )
                
                # Update the invitation data with the actual database ID
                invitation_data['invitation_id'] = vendor_invitation.invitation_id

                # Now that we have invitation_id, generate tracking URLs and persist acknowledgment_url
                ack_url, decline_url = generate_tracking_urls(rfp_id, vendor_invitation.invitation_id)
                vendor_invitation.acknowledgment_url = ack_url
                vendor_invitation.save(update_fields=['acknowledgment_url'])
                
                # Reflect new URLs in the response payload
                invitation_data['acknowledgment_url'] = ack_url
                invitation_data['decline_url'] = decline_url
                
                invitations.append(invitation_data)
        
        return JsonResponse({
            'success': True,
            'message': f'Created {len(invitations)} invitations successfully',
            'invitations': invitations,
            'count': len(invitations)
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Failed to create invitations: {str(e)}'
        }, status=500)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def send_vendor_invitations(request, rfp_id):
    """
    Send invitation emails to vendors
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        
        # Parse request data
        data = json.loads(request.body)
        invitations = data.get('invitations', [])
        
        if not invitations:
            return JsonResponse({
                'success': False,
                'error': 'No invitations provided'
            }, status=400)
        
        sent_count = 0
        failed_count = 0
        results = []
        
        with transaction.atomic():
            # Update RFP status to SUBMISSION_OPEN when invitations are sent
            if rfp.status in ['DRAFT', 'IN_REVIEW', 'APPROVED', 'PUBLISHED']:
                rfp.status = 'SUBMISSION_OPEN'
                rfp.save()
                print(f"[STATUS UPDATE] RFP {rfp_id} status changed to SUBMISSION_OPEN")
            for invitation in invitations:
                try:
                    invitation_id = invitation.get('invitation_id') or invitation.get('unique_token')
                    vendor_email = invitation.get('vendor_email')
                    vendor_name = invitation.get('vendor_name')
                    invitation_url = invitation.get('invitation_url')
                    company_name = invitation.get('company_name', 'Unknown Company')
                    custom_message = invitation.get('custom_message', '')
                    
                    # Get tracking URLs from invitation data or generate them
                    acknowledgment_url = invitation.get('acknowledgment_url')
                    decline_url = invitation.get('decline_url')
                    
                    # If tracking URLs are not in the invitation data, fetch from database
                    if not acknowledgment_url and invitation_id:
                        try:
                            vendor_invitation = VendorInvitation.objects.get(invitation_id=invitation_id)
                            acknowledgment_url = vendor_invitation.acknowledgment_url
                            # Generate decline URL if not available
                            if not decline_url and acknowledgment_url:
                                decline_url = acknowledgment_url.replace('/ack/', '/decline/')
                        except VendorInvitation.DoesNotExist:
                            pass
                    
                    if vendor_email and vendor_name and invitation_url:
                        # Update database to mark invitation as sent
                        if invitation_id:
                            try:
                                vendor_invitation = VendorInvitation.objects.get(invitation_id=invitation_id)
                                vendor_invitation.invitation_status = 'SENT'
                                vendor_invitation.save()
                            except VendorInvitation.DoesNotExist:
                                pass
                        
                        # Generate invitation details for manual sending
                        rfp_title = rfp.rfp_title or "RFP Request"
                        rfp_number = rfp.rfp_number or "N/A"
                        deadline = rfp.submission_deadline.strftime("%B %d, %Y") if rfp.submission_deadline else None
                        
                        # Create email content for manual sending
                        email_subject = f'RFP Invitation: {rfp_title}'
                        email_content = f"""
Dear {vendor_name},

You have been invited to participate in our Request for Proposal (RFP) process for {rfp_title}.

RFP Number: {rfp_number}
Company: {company_name}
{f'Deadline: {deadline}' if deadline else ''}

{custom_message if custom_message else ''}

To participate in this RFP, please click the link below:
{invitation_url}

If you choose to decline, please let us know.

Thank you for your interest in working with us.

Best regards,
Procurement Team
                        """.strip()
                        
                        # Try to send email using Azure AD OAuth2
                        email_sent = False
                        email_error_msg = None
                        try:
                            # Render HTML email template
                            email_context = {
                                'vendor_name': vendor_name,
                                'vendor_email': vendor_email,
                                'company_name': company_name,
                                'rfp_title': rfp_title,
                                'rfp_number': rfp_number,
                                'deadline': deadline,
                                'custom_message': custom_message,
                                # Use tracking URLs we fetched/generated above
                                'acknowledgment_url': acknowledgment_url,
                                'decline_url': decline_url or f"{invitation_url}?action=decline",
                                'include_deadline': True
                            }
                            
                            html_content = render_to_string('rfp/vendor_invitation_email.html', email_context)
                            
                            # Create email message with HTML content
                            email_message = EmailMessage(
                                subject=email_subject,
                                body=html_content,
                                from_email=settings.DEFAULT_FROM_EMAIL,
                                to=[vendor_email]
                            )
                            email_message.content_subtype = "html"
                            
                            # Send email using the EmailMessage object directly
                            email_message.send()
                            email_sent = True
                        except Exception as email_error:
                            email_error_msg = str(email_error)
                        
                        # Always mark as successful since invitation was created
                        results.append({
                            'invitation_id': invitation_id,
                            'success': True,
                            'vendor_email': vendor_email,
                            'vendor_name': vendor_name,
                            'message': 'Email sent successfully' if email_sent else 'Invitation created - email ready for manual sending',
                            'email_sent': email_sent,
                            'email_error': email_error_msg,
                            'invitation_url': invitation_url,
                            'email_content': email_content if not email_sent else None
                        })
                        sent_count += 1
                    else:
                        missing_fields = []
                        if not vendor_email: missing_fields.append('email')
                        if not vendor_name: missing_fields.append('name')
                        if not invitation_url: missing_fields.append('invitation_url')
                        
                        print(f"[ERROR] Missing fields for invitation: {', '.join(missing_fields)}")
                        
                        results.append({
                            'invitation_id': invitation_id,
                            'success': False,
                            'vendor_email': vendor_email,
                            'vendor_name': vendor_name,
                            'error': f'Missing fields: {", ".join(missing_fields)}'
                        })
                        
                        failed_count += 1
                        
                except Exception as e:
                    print(f"[ERROR] Failed to send invitation to {invitation.get('vendor_email', 'unknown')}: {str(e)}")
                    
                    results.append({
                        'invitation_id': invitation.get('invitation_id'),
                        'success': False,
                        'vendor_email': invitation.get('vendor_email'),
                        'vendor_name': invitation.get('vendor_name'),
                        'error': str(e)
                    })
                    
                    failed_count += 1
        
        return JsonResponse({
            'success': True,
            'message': f'Sent {sent_count} invitations successfully',
            'sent': sent_count,
            'failed': failed_count,
            'results': results
        })
    
    except json.JSONDecodeError:
        return JsonResponse({
            'success': False,
            'error': 'Invalid JSON data'
        }, status=400)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Failed to send invitations: {str(e)}'
        }, status=500)


@api_view(['GET', 'POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def acknowledge_invitation(request, token):
    """
    Handle vendor acknowledgment of RFP invitation
    """
    try:
        invitation = get_object_or_404(VendorInvitation, unique_token=token)
        
        if request.method == 'POST':
            # Update invitation status
            invitation.invitation_status = 'ACKNOWLEDGED'
            invitation.acknowledged_date = timezone.now()
            invitation.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Invitation acknowledged successfully',
                'invitation_id': invitation.invitation_id,
                'vendor_name': invitation.vendor_name,
                'company_name': invitation.company_name,
                'rfp_title': invitation.rfp.rfp_title if invitation.rfp else 'Unknown RFP'
            })
        else:
            # GET request - show acknowledgment page
            return JsonResponse({
                'success': True,
                'invitation_id': invitation.invitation_id,
                'vendor_name': invitation.vendor_name,
                'company_name': invitation.company_name,
                'rfp_title': invitation.rfp.rfp_title if invitation.rfp else 'Unknown RFP',
                'rfp_number': invitation.rfp.rfp_number if invitation.rfp else 'Unknown',
                'status': invitation.invitation_status,
                'acknowledged_date': invitation.acknowledged_date,
                'message': 'Invitation details retrieved successfully'
            })
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@api_view(['GET', 'POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def decline_invitation(request, token):
    """
    Handle vendor decline of RFP invitation
    """
    try:
        invitation = get_object_or_404(VendorInvitation, unique_token=token)
        
        if request.method == 'POST':
            data = json.loads(request.body)
            decline_reason = data.get('reason', 'No reason provided')
            
            # Update invitation status
            invitation.invitation_status = 'DECLINED'
            invitation.declined_reason = decline_reason
            invitation.save()
            
            return JsonResponse({
                'success': True,
                'message': 'Invitation declined successfully',
                'invitation_id': invitation.invitation_id,
                'vendor_name': invitation.vendor_name,
                'company_name': invitation.company_name,
                'rfp_title': invitation.rfp.rfp_title if invitation.rfp else 'Unknown RFP',
                'decline_reason': decline_reason
            })
        else:
            # GET request - show decline page
            return JsonResponse({
                'success': True,
                'invitation_id': invitation.invitation_id,
                'vendor_name': invitation.vendor_name,
                'company_name': invitation.company_name,
                'rfp_title': invitation.rfp.rfp_title if invitation.rfp else 'Unknown RFP',
                'rfp_number': invitation.rfp.rfp_number if invitation.rfp else 'Unknown',
                'status': invitation.invitation_status,
                'decline_reason': invitation.declined_reason,
                'message': 'Invitation details retrieved successfully'
            })
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@api_view(['GET', 'POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def ack_invitation_with_ids(request, rfp_id, invitation_id):
    """
    Track acknowledgement via link containing rfp_id and invitation_id.
    Sets invitation_status to ACKNOWLEDGED, is_acknowledged=True, and timestamp.
    """
    try:
        invitation = get_object_or_404(VendorInvitation, invitation_id=invitation_id, rfp__rfp_id=rfp_id)
        if request.method == 'POST':
            invitation.invitation_status = 'ACKNOWLEDGED'
            invitation.is_acknowledged = True
            invitation.acknowledged_date = timezone.now()
            invitation.save(update_fields=['invitation_status', 'is_acknowledged', 'acknowledged_date', 'updated_at'])
            return JsonResponse({
                'success': True,
                'message': 'Invitation acknowledged successfully',
                'invitation_id': invitation.invitation_id,
                'rfp_id': rfp_id
            })
        else:
            # For GET, record acknowledgement and show minimal confirmation HTML (no redirect)
            invitation.invitation_status = 'ACKNOWLEDGED'
            invitation.is_acknowledged = True
            if not invitation.acknowledged_date:
                invitation.acknowledged_date = timezone.now()
            invitation.save(update_fields=['invitation_status', 'is_acknowledged', 'acknowledged_date', 'updated_at'])
            html = f"""
<!DOCTYPE html>
<html><head><meta charset='utf-8'><title>Invitation Acknowledged</title>
<style>body{{font-family:Arial;padding:24px;color:#111}} .ok{{color:#0a7a0a}}</style>
</head><body>
  <h2 class="ok">Thank you! Your invitation has been acknowledged.</h2>
  <p>RFP ID: {rfp_id} • Invitation ID: {invitation.invitation_id}</p>
  <p>You may now close this page.</p>
  
</body></html>
"""
            return HttpResponse(html)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


@api_view(['GET', 'POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def decline_invitation_with_ids(request, rfp_id, invitation_id):
    """
    Track decline via link containing rfp_id and invitation_id.
    Sets invitation_status to DECLINED, is_acknowledged=False and stores declined_reason if provided.
    """
    try:
        invitation = get_object_or_404(VendorInvitation, invitation_id=invitation_id, rfp__rfp_id=rfp_id)
        decline_reason = None
        if request.method == 'POST':
            try:
                data = json.loads(request.body)
                decline_reason = data.get('reason')
            except Exception:
                decline_reason = None
        else:
            # Read optional reason from query param for GET
            decline_reason = request.GET.get('reason')

        invitation.invitation_status = 'DECLINED'
        invitation.is_acknowledged = False
        if decline_reason:
            invitation.declined_reason = decline_reason
        invitation.save(update_fields=['invitation_status', 'is_acknowledged', 'declined_reason', 'updated_at'])
        if request.method == 'GET':
            html = f"""
<!DOCTYPE html>
<html><head><meta charset='utf-8'><title>Invitation Declined</title>
<style>body{{font-family:Arial;padding:24px;color:#111}} .err{{color:#b11}}</style>
</head><body>
  <h2 class="err">You have declined the invitation.</h2>
  <p>RFP ID: {rfp_id} • Invitation ID: {invitation.invitation_id}</p>
  <p>{('Reason: ' + invitation.declined_reason) if invitation.declined_reason else ''}</p>
  <p>You may now close this page.</p>
</body></html>
"""
            return HttpResponse(html)
        return JsonResponse({
            'success': True,
            'message': 'Invitation declined',
            'invitation_id': invitation.invitation_id,
            'rfp_id': rfp_id,
            'decline_reason': invitation.declined_reason
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=400)


# ============================================================================
# VENDOR VIEWS (from views_vendor.py)
# ============================================================================

@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def vendor_selection(request, rfp_id):
    """
    View for selecting vendors for an RFP
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    # Get existing selected vendors for this RFP
    selected_vendor_ids = list(RFPVendorSelection.objects.filter(
        rfp=rfp
    ).values_list('vendor_id', flat=True))
    
    # Forms
    search_form = VendorSearchForm(request.GET or None)
    manual_entry_form = VendorManualEntryForm()
    bulk_upload_form = VendorBulkUploadForm()
    
    # Handle search and filtering
    vendors = Vendor.objects.all().prefetch_related(
        'capabilities', 'certifications'
    )
    
    search_term = ''
    filter_type = 'all'
    
    if search_form.is_valid():
        search_term = search_form.cleaned_data.get('search_term', '')
        filter_type = search_form.cleaned_data.get('filter_type', 'all')
        
        if search_term:
            vendors = vendors.filter(
                Q(company_name__icontains=search_term) |
                Q(capabilities__capability_name__icontains=search_term) |
                Q(certifications__certification_name__icontains=search_term)
            ).distinct()
        
        if filter_type == 'high-match':
            vendors = vendors.filter(match_score__gte=90)
        elif filter_type == 'certified':
            vendors = vendors.annotate(
                cert_count=Count('certifications')
            ).filter(cert_count__gt=2)
    
    # Prepare vendor data for template
    vendor_data = []
    for vendor in vendors:
        # Get capabilities and certifications
        capabilities = list(vendor.capabilities.values_list('capability_name', flat=True))
        certifications = list(vendor.certifications.values_list('certification_name', flat=True))
        
        # Format employee count for display
        employee_display = "Unknown"
        if vendor.employee_count:
            if vendor.employee_count < 50:
                employee_display = "< 50"
            elif vendor.employee_count < 200:
                employee_display = "50-200"
            elif vendor.employee_count < 500:
                employee_display = "200-500"
            elif vendor.employee_count < 1000:
                employee_display = "500-1000"
            else:
                employee_display = "1000+"
        
        # Determine category based on employee count
        category = "Unknown"
        if vendor.employee_count:
            if vendor.employee_count < 200:
                category = "Startup"
            elif vendor.employee_count < 1000:
                category = "Mid-Market"
            else:
                category = "Enterprise"
        
        # Format experience years
        experience = f"{vendor.experience_years} years" if vendor.experience_years else "Unknown"
        
        vendor_data.append({
            'id': vendor.vendor_id,
            'name': vendor.company_name,
            'email': vendor.email or '',
            'phone': vendor.phone or '',
            'website': vendor.website or '',
            'location': vendor.location or '',
            'matchScore': vendor.match_score or 0,
            'rating': vendor.rating or 0,
            'capabilities': capabilities,
            'certifications': certifications,
            'employees': employee_display,
            'experience': experience,
            'category': category,
            'is_selected': vendor.vendor_id in selected_vendor_ids
        })
    
    context = {
        'rfp': rfp,
        'vendors': vendor_data,
        'selected_vendor_ids': selected_vendor_ids,
        'search_form': search_form,
        'manual_entry_form': manual_entry_form,
        'bulk_upload_form': bulk_upload_form,
        'search_term': search_term,
        'active_filter': filter_type,
    }
    
    return render(request, 'rfp/phase3_vendor_selection.html', context)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def vendor_manual_entry(request, rfp_id):
    """
    View for manually adding a vendor
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    form = VendorManualEntryForm(request.POST)
    
    if form.is_valid():
        vendor = form.save(commit=True)
        
        # Link vendor to RFP
        RFPVendorSelection.objects.create(
            rfp=rfp,
            vendor=vendor,
            selected_by=request.user.id
        )
        
        messages.success(request, f"Vendor '{vendor.company_name}' added successfully.")
        return redirect('vendor_selection', rfp_id=rfp_id)
    
    # If form is invalid, return to vendor selection with errors
    vendors = Vendor.objects.all().prefetch_related('capabilities', 'certifications')
    selected_vendor_ids = list(RFPVendorSelection.objects.filter(
        rfp=rfp
    ).values_list('vendor_id', flat=True))
    
    context = {
        'rfp': rfp,
        'vendors': vendors,
        'selected_vendor_ids': selected_vendor_ids,
        'search_form': VendorSearchForm(),
        'manual_entry_form': form,  # Return form with errors
        'bulk_upload_form': VendorBulkUploadForm(),
    }
    
    return render(request, 'rfp/phase3_vendor_selection.html', context)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def vendor_bulk_upload(request, rfp_id):
    """
    View for bulk uploading vendors via CSV
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    form = VendorBulkUploadForm(request.POST, request.FILES)
    
    if form.is_valid():
        csv_file = request.FILES['csv_file']
        
        # Check if file is CSV
        if not csv_file.name.endswith('.csv'):
            messages.error(request, "Please upload a CSV file.")
            return redirect('vendor_selection', rfp_id=rfp_id)
        
        # Process CSV file
        try:
            csv_data = csv_file.read().decode('utf-8')
            csv_reader = csv.DictReader(io.StringIO(csv_data))
            
            success_count = 0
            error_count = 0
            
            for row in csv_reader:
                try:
                    # Create vendor
                    vendor = Vendor(
                        company_name=row.get('company_name', ''),
                        legal_name=row.get('legal_name', ''),
                        email=row.get('email', ''),
                        phone=row.get('phone', ''),
                        website=row.get('website', ''),
                        location=row.get('location', ''),
                        industry_sector=row.get('industry_sector', ''),
                        employee_count=int(row.get('employee_count', 0)) if row.get('employee_count') else None,
                        experience_years=int(row.get('experience_years', 0)) if row.get('experience_years') else None,
                        description=row.get('description', ''),
                        risk_level='MEDIUM',
                        status='DRAFT',
                    )
                    vendor.save()
                    
                    # Add capabilities
                    capabilities = row.get('capabilities', '').split(',')
                    for cap in capabilities:
                        cap = cap.strip()
                        if cap:
                            VendorCapability.objects.create(
                                vendor=vendor,
                                capability_name=cap
                            )
                    
                    # Add certifications
                    certifications = row.get('certifications', '').split(',')
                    for cert in certifications:
                        cert = cert.strip()
                        if cert:
                            VendorCertification.objects.create(
                                vendor=vendor,
                                certification_name=cert
                            )
                    
                    # Link vendor to RFP
                    RFPVendorSelection.objects.create(
                        rfp=rfp,
                        vendor=vendor,
                        selected_by=request.user.id
                    )
                    
                    success_count += 1
                except Exception as e:
                    error_count += 1
                    continue
            
            if success_count > 0:
                messages.success(request, f"Successfully imported {success_count} vendors.")
            
            if error_count > 0:
                messages.warning(request, f"Failed to import {error_count} vendors due to errors.")
                
        except Exception as e:
            messages.error(request, f"Error processing CSV file: {str(e)}")
        
        return redirect('vendor_selection', rfp_id=rfp_id)
    
    messages.error(request, "Invalid form submission.")
    return redirect('vendor_selection', rfp_id=rfp_id)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('edit_rfp')
def update_vendor_selection(request, rfp_id):
    """
    View for updating vendor selection
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    try:
        data = json.loads(request.body)
        vendor_id = data.get('vendor_id')
        is_selected = data.get('is_selected', False)
        
        vendor = get_object_or_404(Vendor, vendor_id=vendor_id)
        
        if is_selected:
            # Add vendor to selection if not already selected
            RFPVendorSelection.objects.get_or_create(
                rfp=rfp,
                vendor=vendor,
                defaults={'selected_by': request.user.id}
            )
        else:
            # Remove vendor from selection
            RFPVendorSelection.objects.filter(
                rfp=rfp,
                vendor=vendor
            ).delete()
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def bulk_select_vendors(request, rfp_id):
    """
    View for bulk selecting/deselecting vendors
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    try:
        data = json.loads(request.body)
        vendor_ids = data.get('vendor_ids', [])
        select_all = data.get('select_all', False)
        
        if select_all:
            # Get all vendor IDs
            all_vendor_ids = list(Vendor.objects.values_list('vendor_id', flat=True))
            
            # Check if all vendors are already selected
            selected_count = RFPVendorSelection.objects.filter(
                rfp=rfp,
                vendor_id__in=all_vendor_ids
            ).count()
            
            if selected_count == len(all_vendor_ids):
                # Deselect all vendors
                RFPVendorSelection.objects.filter(
                    rfp=rfp,
                    vendor_id__in=all_vendor_ids
                ).delete()
                return JsonResponse({'success': True, 'action': 'deselected_all'})
            else:
                # Select all vendors
                for vendor_id in all_vendor_ids:
                    RFPVendorSelection.objects.get_or_create(
                        rfp=rfp,
                        vendor_id=vendor_id,
                        defaults={'selected_by': request.user.id}
                    )
                return JsonResponse({'success': True, 'action': 'selected_all'})
        else:
            # Select/deselect specific vendors
            for vendor_id in vendor_ids:
                RFPVendorSelection.objects.get_or_create(
                    rfp=rfp,
                    vendor_id=vendor_id,
                    defaults={'selected_by': request.user.id}
                )
            return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def generate_vendor_urls(request, rfp_id):
    """
    View for generating invitation URLs for selected vendors using new query parameter format
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    # Get selected vendors
    selected_vendors = RFPVendorSelection.objects.filter(rfp=rfp)
    
    if not selected_vendors:
        messages.error(request, "No vendors selected. Please select at least one vendor.")
        return redirect('vendor_selection', rfp_id=rfp_id)
    
    # Generate URLs for each vendor using new format
    for selection in selected_vendors:
        if not selection.invitation_url:
            # Generate new-style URL with query parameters
            vendor = selection.vendor
            from django.conf import settings
            import re
            
            # Get external base URL and ensure it uses localhost (not ngrok)
            external_base_url = getattr(settings, 'EXTERNAL_BASE_URL', 'http://localhost:3000').rstrip('/')
            
            # Replace any ngrok URLs with localhost:3000
            if 'ngrok' in external_base_url.lower():
                external_base_url = 'http://localhost:3000'
            
            # Ensure it's localhost (not 127.0.0.1 or other variations)
            if not external_base_url.startswith('http://localhost') and not external_base_url.startswith('https://localhost'):
                # Extract port if present, otherwise use 3000
                port_match = re.search(r':(\d+)', external_base_url)
                port = port_match.group(1) if port_match else '3000'
                external_base_url = f'http://localhost:{port}'
            
            base_url = f"{external_base_url}/submit"
            
            # URL encode the parameters
            from urllib.parse import urlencode
            params = {
                'rfpId': str(rfp.rfp_id),
                'vendorId': str(vendor.vendor_id),
                'org': vendor.company_name or '',
                'vendorName': f"{vendor.first_name or ''} {vendor.last_name or ''}".strip(),
                'contactEmail': vendor.email or '',
                'contactPhone': vendor.phone or ''
            }
            
            # Remove empty parameters
            params = {k: v for k, v in params.items() if v}
            invitation_url = f"{base_url}?{urlencode(params)}"
            
            # Store invitation in database
            invitation = VendorInvitation.objects.create(
                rfp_id=rfp.rfp_id,
                vendor_id=vendor.vendor_id,
                vendor_email=vendor.email or '',
                vendor_name=f"{vendor.first_name or ''} {vendor.last_name or ''}".strip(),
                vendor_phone=vendor.phone or '',
                company_name=vendor.company_name or '',
                invitation_url=invitation_url,
                unique_token=f"INV{rfp.rfp_id}{vendor.vendor_id}{int(time.time())}",
                is_matched_vendor=True,
                submission_source='invited',
                invitation_status='CREATED',
                custom_message=request.POST.get('custom_message', '')
            )
            
            # Update selection with invitation URL
            selection.invitation_url = invitation_url
            selection.save()
    
    messages.success(request, f"Generated invitation URLs for {selected_vendors.count()} vendors.")
    return redirect('vendor_invitation', rfp_id=rfp_id)


def generate_unmatched_vendor_url(rfp_id, org_name="", vendor_name="", contact_email="", contact_phone=""):
    """
    Generate URL for unmatched vendors (not in system)
    """
    from django.conf import settings
    import re
    
    # Get external base URL and ensure it uses localhost (not ngrok)
    external_base_url = getattr(settings, 'EXTERNAL_BASE_URL', 'http://localhost:3000').rstrip('/')
    
    # Replace any ngrok URLs with localhost:3000
    if 'ngrok' in external_base_url.lower():
        external_base_url = 'http://localhost:3000'
    
    # Ensure it's localhost (not 127.0.0.1 or other variations)
    if not external_base_url.startswith('http://localhost') and not external_base_url.startswith('https://localhost'):
        # Extract port if present, otherwise use 3000
        port_match = re.search(r':(\d+)', external_base_url)
        port = port_match.group(1) if port_match else '3000'
        external_base_url = f'http://localhost:{port}'
    
    base_url = f"{external_base_url}/submit"
    from urllib.parse import urlencode
    
    params = {
        'rfpId': str(rfp_id),
        'vendorId': '',  # Empty for unmatched vendors
        'org': org_name,
        'vendorName': vendor_name,
        'contactEmail': contact_email,
        'contactPhone': contact_phone
    }
    
    # Remove empty values
    params = {k: v for k, v in params.items() if v}
    
    return f"{base_url}?{urlencode(params)}"


def generate_open_rfp_url(rfp_id):
    """
    Generate URL for open/public RFPs
    """
    from django.conf import settings
    import re
    
    # Get external base URL and ensure it uses localhost (not ngrok)
    external_base_url = getattr(settings, 'EXTERNAL_BASE_URL', 'http://localhost:3000').rstrip('/')
    
    # Replace any ngrok URLs with localhost:3000
    if 'ngrok' in external_base_url.lower():
        external_base_url = 'http://localhost:3000'
    
    # Ensure it's localhost (not 127.0.0.1 or other variations)
    if not external_base_url.startswith('http://localhost') and not external_base_url.startswith('https://localhost'):
        # Extract port if present, otherwise use 3000
        port_match = re.search(r':(\d+)', external_base_url)
        port = port_match.group(1) if port_match else '3000'
        external_base_url = f'http://localhost:{port}'
    
    base_url = f"{external_base_url}/submit/open"
    from urllib.parse import urlencode
    
    params = {
        'rfpId': str(rfp_id)
    }
    
    return f"{base_url}?{urlencode(params)}"


@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def vendor_invitation(request, rfp_id):
    """
    View for sending invitations to selected vendors
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    # Get selected vendors with URLs
    selected_vendors = RFPVendorSelection.objects.filter(
        rfp=rfp
    ).select_related('vendor')
    
    context = {
        'rfp': rfp,
        'selected_vendors': selected_vendors,
    }
    
    return render(request, 'rfp/phase4_vendor_invitation.html', context)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_unmatched_vendors(request, rfp_id):
    """
    API endpoint to get unmatched vendors for an RFP
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    unmatched_vendors = RFPUnmatchedVendor.objects.filter(
        matching_status__in=['unmatched', 'pending_review']
    ).order_by('-created_at')
    
    vendor_data = []
    for vendor in unmatched_vendors:
        vendor_data.append({
            'unmatched_id': vendor.unmatched_id,
            'vendor_name': vendor.vendor_name,
            'vendor_email': vendor.vendor_email,
            'vendor_phone': vendor.vendor_phone,
            'company_name': vendor.company_name,
            'matching_status': vendor.matching_status,
            'created_at': vendor.created_at.isoformat(),
            'submission_data': vendor.submission_data if vendor.submission_data else {}
        })
    
    return JsonResponse(vendor_data, safe=False)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def create_unmatched_vendor(request, rfp_id):
    """
    API endpoint to create a new unmatched vendor
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        data = json.loads(request.body)
        
        # Validate required fields
        required_fields = ['vendor_name', 'vendor_email', 'vendor_phone', 'company_name']
        for field in required_fields:
            if not data.get(field):
                return JsonResponse({'error': f'{field} is required'}, status=400)
        
        # Create unmatched vendor
        unmatched_vendor = RFPUnmatchedVendor.objects.create(
            vendor_name=data['vendor_name'],
            vendor_email=data['vendor_email'],
            vendor_phone=data['vendor_phone'],
            company_name=data['company_name'],
            submission_data=data.get('submission_data', {}),
            matching_status=data.get('matching_status', 'unmatched')
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Unmatched vendor created successfully',
            'unmatched_id': unmatched_vendor.unmatched_id,
            'vendor_name': unmatched_vendor.vendor_name,
            'company_name': unmatched_vendor.company_name
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'Failed to create unmatched vendor: {str(e)}'}, status=500)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_approved_vendors(request, rfp_id):
    """
    API endpoint to get all approved vendors for an RFP
    """
    rfp = get_object_or_404(RFP, rfp_id=rfp_id)
    
    try:
        # Get only the fields that actually exist in the database
        approved_vendors = Vendor.objects.filter(
            status='APPROVED'
        ).values(
            'vendor_id', 'vendor_code', 'company_name', 'legal_name', 
            'business_type', 'incorporation_date', 'tax_id', 'duns_number',
            'website', 'industry_sector', 'annual_revenue', 'employee_count',
            'headquarters_country', 'headquarters_address', 'description',
            'vendor_category_id', 'risk_level', 'status', 'lifecycle_stage',
            'onboarding_date', 'last_assessment_date', 'next_assessment_date',
            'is_critical_vendor', 'has_data_access', 'has_system_access',
            'created_by', 'updated_by', 'created_at', 'updated_at'
        ).order_by('-created_at')
        
        vendor_data = []
        for vendor in approved_vendors:
            # Use default values for capabilities and certifications since tables don't exist
            capabilities = ['Software Development', 'Cloud Services']  # Default capabilities
            certifications = ['ISO 27001', 'SOC 2']  # Default certifications
            
            # Format employee count for display
            employee_display = "Unknown"
            if vendor['employee_count']:
                if vendor['employee_count'] < 50:
                    employee_display = "< 50"
                elif vendor['employee_count'] < 200:
                    employee_display = "50-200"
                elif vendor['employee_count'] < 500:
                    employee_display = "200-500"
                elif vendor['employee_count'] < 1000:
                    employee_display = "500-1000"
                else:
                    employee_display = "1000+"
            
            # Determine category based on employee count
            category = "Unknown"
            if vendor['employee_count']:
                if vendor['employee_count'] < 200:
                    category = "Startup"
                elif vendor['employee_count'] < 1000:
                    category = "Mid-Market"
                else:
                    category = "Enterprise"
            
            vendor_data.append({
                'vendor_id': vendor['vendor_id'],
                'vendor_code': vendor['vendor_code'],
                'company_name': vendor['company_name'],
                'legal_name': vendor['legal_name'],
                'business_type': vendor['business_type'],
                'incorporation_date': vendor['incorporation_date'].isoformat() if vendor['incorporation_date'] else None,
                'tax_id': vendor['tax_id'],
                'duns_number': vendor['duns_number'],
                'website': vendor['website'],
                'industry_sector': vendor['industry_sector'],
                'annual_revenue': float(vendor['annual_revenue']) if vendor['annual_revenue'] else None,
                'employee_count': vendor['employee_count'],
                'employee_display': employee_display,
                'headquarters_country': vendor['headquarters_country'],
                'headquarters_address': vendor['headquarters_address'],
                'description': vendor['description'],
                'vendor_category_id': vendor['vendor_category_id'],
                'risk_level': vendor['risk_level'],
                'status': vendor['status'],
                'lifecycle_stage': vendor['lifecycle_stage'],
                'onboarding_date': vendor['onboarding_date'].isoformat() if vendor['onboarding_date'] else None,
                'last_assessment_date': vendor['last_assessment_date'].isoformat() if vendor['last_assessment_date'] else None,
                'next_assessment_date': vendor['next_assessment_date'].isoformat() if vendor['next_assessment_date'] else None,
                'is_critical_vendor': vendor['is_critical_vendor'],
                'has_data_access': vendor['has_data_access'],
                'has_system_access': vendor['has_system_access'],
                'created_by': vendor['created_by'],
                'updated_by': vendor['updated_by'],
                'created_at': vendor['created_at'].isoformat(),
                'updated_at': vendor['updated_at'].isoformat(),
                # Additional fields for UI (using default values since fields don't exist in DB)
                'email': 'contact@example.com',  # Default email
                'phone': '+1-555-0000',  # Default phone
                'location': vendor['headquarters_country'] or 'Unknown',  # Use headquarters as location
                'match_score': 85.0,  # Default match score
                'rating': 4.5,  # Default rating
                'experience_years': 5,  # Default experience
                'capabilities': capabilities,
                'certifications': certifications,
                'category': category
            })
        
        return JsonResponse(vendor_data, safe=False)
        
    except Exception as e:
        return JsonResponse({'error': f'Failed to fetch approved vendors: {str(e)}'}, status=500)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_sample_csv(request):
    """
    API endpoint to download sample CSV template for vendor upload
    """
    try:
        # Create simple CSV template with just headers
        csv_content = "company_name,vendor_name,vendor_email,vendor_phone,website,industry_sector,description\n"
        
        response = HttpResponse(csv_content, content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="vendor_upload_template.csv"'
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Failed to generate sample CSV: {str(e)}'}, status=500)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def vendor_manual_entry(request, rfp_id):
    """
    API endpoint to create a vendor manually
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        
        # Parse JSON data from request
        data = json.loads(request.body)
        
        # Extract vendor data
        company_name = data.get('company_name', '').strip()
        legal_name = data.get('legal_name', '').strip()
        business_type = data.get('business_type', '').strip()
        industry_sector = data.get('industry_sector', '').strip()
        email = data.get('email', '').strip()
        phone = data.get('phone', '').strip()
        website = data.get('website', '').strip()
        headquarters_country = data.get('headquarters_country', '').strip()
        employee_count = data.get('employee_count')
        experience_years = data.get('experience_years')
        risk_level = data.get('risk_level', 'MEDIUM')
        description = data.get('description', '').strip()
        capabilities = data.get('capabilities', [])
        certifications = data.get('certifications', [])
        
        # Validation
        if not company_name:
            return JsonResponse({'error': 'Company name is required'}, status=400)
        
        if not email:
            return JsonResponse({'error': 'Email is required'}, status=400)
        
        # Create unmatched vendor entry
        submission_data = {
            'legal_name': legal_name,
            'business_type': business_type,
            'industry_sector': industry_sector,
            'website': website,
            'headquarters_country': headquarters_country,
            'employee_count': employee_count,
            'experience_years': experience_years,
            'risk_level': risk_level,
            'description': description,
            'capabilities': capabilities,
            'certifications': certifications
        }
        
        unmatched_vendor = RFPUnmatchedVendor.objects.create(
            rfp=rfp,
            vendor_name=company_name,  # Use company_name as vendor_name
            vendor_email=email,
            vendor_phone=phone,
            company_name=company_name,
            submission_data=submission_data,
            matching_status='pending_review'
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Vendor created successfully and added to unmatched vendors',
            'unmatched_id': unmatched_vendor.unmatched_id,
            'company_name': unmatched_vendor.company_name,
            'matching_status': unmatched_vendor.matching_status
        })
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON data'}, status=400)
    except Exception as e:
        return JsonResponse({'error': f'Failed to create vendor: {str(e)}'}, status=500)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def vendor_bulk_upload(request, rfp_id):
    """
    API endpoint to bulk upload vendors from CSV/Excel file
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file provided'}, status=400)
        
        file = request.FILES['file']
        
        # Check file type
        if not file.name.endswith(('.csv', '.xlsx', '.xls')):
            return JsonResponse({'error': 'Invalid file type. Please upload CSV or Excel file.'}, status=400)
        
        # Read file content
        if file.name.endswith('.csv'):
            # Handle CSV file
            import csv
            content = file.read().decode('utf-8')
            csv_reader = csv.DictReader(io.StringIO(content))
            rows = list(csv_reader)
        else:
            # Handle Excel file
            import pandas as pd
            df = pd.read_excel(file)
            rows = df.to_dict('records')
        
        results = {
            'success': 0,
            'failed': 0,
            'errors': []
        }
        
        for i, row in enumerate(rows, 1):
            try:
                # Extract data from row
                company_name = str(row.get('company_name', '')).strip()
                email = str(row.get('email', '')).strip()
                phone = str(row.get('phone', '')).strip()
                legal_name = str(row.get('legal_name', '')).strip()
                business_type = str(row.get('business_type', '')).strip()
                industry_sector = str(row.get('industry_sector', '')).strip()
                website = str(row.get('website', '')).strip()
                employee_count = row.get('employee_count')
                experience_years = row.get('experience_years')
                description = str(row.get('description', '')).strip()
                
                # Parse capabilities and certifications
                capabilities = []
                certifications = []
                
                if row.get('capabilities'):
                    capabilities = [cap.strip() for cap in str(row.get('capabilities', '')).split(',') if cap.strip()]
                
                if row.get('certifications'):
                    certifications = [cert.strip() for cert in str(row.get('certifications', '')).split(',') if cert.strip()]
                
                # Validation
                if not company_name:
                    results['failed'] += 1
                    results['errors'].append(f'Row {i}: Company name is required')
                    continue
                
                if not email:
                    results['failed'] += 1
                    results['errors'].append(f'Row {i}: Email is required')
                    continue
                
                # Create submission data
                submission_data = {
                    'legal_name': legal_name,
                    'business_type': business_type,
                    'industry_sector': industry_sector,
                    'website': website,
                    'employee_count': employee_count,
                    'experience_years': experience_years,
                    'description': description,
                    'capabilities': capabilities,
                    'certifications': certifications
                }
                
                # Create unmatched vendor entry
                unmatched_vendor = RFPUnmatchedVendor.objects.create(
                    rfp=rfp,
                    vendor_name=company_name,
                    vendor_email=email,
                    vendor_phone=phone,
                    company_name=company_name,
                    submission_data=submission_data,
                    matching_status='pending_review'
                )
                
                results['success'] += 1
                
            except Exception as e:
                results['failed'] += 1
                results['errors'].append(f'Row {i}: {str(e)}')
        
        return JsonResponse({
            'success': True,
            'message': f'Bulk upload completed. {results["success"]} vendors created, {results["failed"]} failed.',
            'results': results
        })
        
    except Exception as e:
        return JsonResponse({'error': f'Failed to process bulk upload: {str(e)}'}, status=500)


@api_view(['POST'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('create_rfp')
def unmatched_vendor_bulk_upload(request, rfp_id):
    """
    API endpoint to bulk upload unmatched vendors from CSV/Excel file
    """
    try:
        rfp = get_object_or_404(RFP, rfp_id=rfp_id)
        
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file provided'}, status=400)
        
        file = request.FILES['file']
        
        # Check file extension
        if not file.name.lower().endswith(('.csv', '.xlsx', '.xls')):
            return JsonResponse({'error': 'File must be CSV or Excel format'}, status=400)
        
        # Process the file
        import pandas as pd
        import io
        
        if file.name.lower().endswith('.csv'):
            df = pd.read_csv(file)
        else:
            df = pd.read_excel(file)
        
        # Initialize results
        results = {
            'success': 0,
            'failed': 0,
            'errors': []
        }
        
        # Process each row
        for i, row in df.iterrows():
            try:
                # Get required fields
                company_name = str(row.get('company_name', '')).strip()
                vendor_name = str(row.get('vendor_name', '')).strip()
                vendor_email = str(row.get('vendor_email', '')).strip()
                vendor_phone = str(row.get('vendor_phone', '')).strip()
                
                # Validate required fields
                if not all([company_name, vendor_name, vendor_email, vendor_phone]):
                    results['failed'] += 1
                    results['errors'].append(f'Row {i+2}: Missing required fields (company_name, vendor_name, vendor_email, vendor_phone)')
                    continue
                
                # Get optional fields
                website = str(row.get('website', '')).strip() if pd.notna(row.get('website')) else ''
                industry_sector = str(row.get('industry_sector', '')).strip() if pd.notna(row.get('industry_sector')) else ''
                description = str(row.get('description', '')).strip() if pd.notna(row.get('description')) else ''
                
                # Create submission data
                submission_data = {
                    'website': website,
                    'industry_sector': industry_sector,
                    'description': description
                }
                
                # Create unmatched vendor entry
                unmatched_vendor = RFPUnmatchedVendor.objects.create(
                    vendor_name=vendor_name,
                    vendor_email=vendor_email,
                    vendor_phone=vendor_phone,
                    company_name=company_name,
                    submission_data=submission_data,
                    matching_status='unmatched'
                )
                
                results['success'] += 1
                
            except Exception as e:
                results['failed'] += 1
                results['errors'].append(f'Row {i+2}: {str(e)}')
        
        return JsonResponse({
            'success': True,
            'message': f'Bulk upload completed. {results["success"]} vendors created, {results["failed"]} failed.',
            'results': results
        })
        
    except Exception as e:
        return JsonResponse({'error': f'Failed to process bulk upload: {str(e)}'}, status=500)


@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_all_approved_vendors(request):
    """
    API endpoint to get all approved vendors (for frontend vendor selection)
    """
    try:
        # Get only the fields that actually exist in the database
        approved_vendors = Vendor.objects.filter(
            status='APPROVED'
        ).values(
            'vendor_id', 'vendor_code', 'company_name', 'legal_name', 
            'business_type', 'incorporation_date', 'tax_id', 'duns_number',
            'website', 'industry_sector', 'annual_revenue', 'employee_count',
            'headquarters_country', 'headquarters_address', 'description',
            'vendor_category_id', 'risk_level', 'status', 'lifecycle_stage',
            'onboarding_date', 'last_assessment_date', 'next_assessment_date',
            'is_critical_vendor', 'has_data_access', 'has_system_access',
            'created_by', 'updated_by', 'created_at', 'updated_at'
        ).order_by('company_name')
        
        vendor_data = []
        for vendor in approved_vendors:
            # Use default values for capabilities and certifications since tables don't exist
            capabilities = ['Software Development', 'Cloud Services']  # Default capabilities
            certifications = ['ISO 27001', 'SOC 2']  # Default certifications
            
            # Generate primary contact information if not available
            primary_email = vendor.get('email', '')
            primary_phone = vendor.get('phone', '')
            
            # If no email/phone in vendor record, generate placeholder contact info
            if not primary_email:
                # Generate email based on company name
                company_clean = vendor['company_name'].lower().replace(' ', '').replace('inc', '').replace('llc', '').replace('corp', '')
                primary_email = f"contact@{company_clean}.com"
            
            if not primary_phone:
                # Generate placeholder phone number
                primary_phone = "+1 (555) 000-0000"
            
            vendor_info = {
                'vendor_id': vendor['vendor_id'],
                'vendor_code': vendor['vendor_code'],
                'company_name': vendor['company_name'],
                'legal_name': vendor['legal_name'],
                'business_type': vendor['business_type'],
                'website': vendor['website'],
                'industry_sector': vendor['industry_sector'],
                'annual_revenue': float(vendor['annual_revenue']) if vendor['annual_revenue'] else None,
                'employee_count': vendor['employee_count'],
                'headquarters_country': vendor['headquarters_country'],
                'risk_level': vendor['risk_level'],
                'status': vendor['status'],
                'is_critical_vendor': vendor['is_critical_vendor'],
                'has_data_access': vendor['has_data_access'],
                'has_system_access': vendor['has_system_access'],
                'created_at': vendor['created_at'].isoformat() if vendor['created_at'] else None,
                'updated_at': vendor['updated_at'].isoformat() if vendor['updated_at'] else None,
                # Contact information
                'email': primary_email,
                'phone': primary_phone,
                'contact_email': primary_email,  # Alias for compatibility
                'contact_phone': primary_phone,  # Alias for compatibility
                # Default values for missing fields
                'match_score': 85.0,  # Default match score
                'rating': 4.5,        # Default rating
                'experience_years': 5, # Default experience
                'capabilities': capabilities,    # Default capabilities list
                'certifications': certifications,  # Default certifications list
            }
            vendor_data.append(vendor_info)
        
        return JsonResponse({
            'success': True,
            'vendors': vendor_data,
            'total': len(vendor_data)
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Failed to fetch approved vendors: {str(e)}'
        }, status=500)


# ============================================================================
# VENDOR CONTACTS VIEWS (from views_vendor_contacts.py)
# ============================================================================

@api_view(['GET'])
@authentication_classes([JWTAuthentication])
@permission_classes([SimpleAuthenticatedPermission])
@rbac_rfp_required('view_rfp')
def get_vendor_primary_contact(request, vendor_id):
    """Get primary contact for a vendor"""
    try:
        # Query vendor_contacts table for primary contact
        with connection.cursor() as cursor:
            cursor.execute('''
                SELECT contact_id, first_name, last_name, email, phone, mobile, designation
                FROM vendor_contacts
                WHERE vendor_id = %s AND contact_type = 'PRIMARY' AND is_primary = 1 AND is_active = 1
                LIMIT 1
            ''', [vendor_id])
            contact = cursor.fetchone()
            
            if contact:
                contact_data = {
                    'contact_id': contact[0],
                    'first_name': contact[1],
                    'last_name': contact[2],
                    'email': contact[3],
                    'phone': contact[4],
                    'mobile': contact[5],
                    'designation': contact[6]
                }
                return JsonResponse({
                    'success': True,
                    'contact': contact_data
                })
            else:
                return JsonResponse({
                    'success': False,
                    'error': 'No primary contact found for vendor'
                })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def calculate_vendor_match_scores(request, rfp_id):
    """
    Calculate match scores for vendors based on RFP requirements
    
    Matching criteria:
    1. Location Match (30%): geographical_scope vs headquarters_country
    2. Industry/Category Match (25%): category vs industry_sector  
    3. Budget Match (25%): estimated_value vs annual_revenue
    4. Business Type Match (20%): rfp_type vs business_type/capabilities
    """
    try:
        import json
        from decimal import Decimal
        from .models import RFP, Vendor
        
        # Get RFP details
        try:
            rfp = RFP.objects.get(rfp_id=rfp_id)
        except RFP.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': 'RFP not found'
            }, status=404)
        
        # Get vendor IDs from request
        data = json.loads(request.body) if request.body else {}
        vendor_ids = data.get('vendor_ids', [])
        
        if not vendor_ids:
            return JsonResponse({
                'success': False,
                'error': 'No vendor IDs provided'
            }, status=400)
        
        # Get vendors
        vendors = Vendor.objects.filter(vendor_id__in=vendor_ids, status='APPROVED')
        
        match_results = []
        
        for vendor in vendors:
            # Initialize score components
            location_score = 0
            industry_score = 0
            budget_score = 0
            business_type_score = 0
            
            # 1. Location Match (30%)
            if rfp.geographical_scope and vendor.headquarters_country:
                rfp_locations = [loc.strip().upper() for loc in rfp.geographical_scope.split(',')]
                vendor_country = vendor.headquarters_country.strip().upper()
                
                # Exact match
                if vendor_country in rfp_locations:
                    location_score = 30
                # Partial match for common regions
                elif any(region in rfp.geographical_scope.upper() for region in ['GLOBAL', 'WORLDWIDE', 'INTERNATIONAL']):
                    location_score = 25
                # Regional matching (simplified)
                elif vendor_country in ['USA', 'CANADA', 'MEXICO'] and any('NORTH AMERICA' in loc for loc in rfp_locations):
                    location_score = 28
                elif vendor_country in ['UK', 'GERMANY', 'FRANCE', 'SPAIN', 'ITALY'] and any('EUROPE' in loc for loc in rfp_locations):
                    location_score = 28
                elif vendor_country in ['INDIA', 'CHINA', 'JAPAN', 'SINGAPORE'] and any('ASIA' in loc for loc in rfp_locations):
                    location_score = 28
                else:
                    # Vendor has location but doesn't match RFP requirements
                    location_score = 0
            elif not vendor.headquarters_country:
                # Vendor location is unknown/missing - no points (real value, not grace)
                location_score = 0
            else:
                # RFP has no location preference - neutral score (not vendor's fault)
                location_score = 15
            
            # 2. Industry/Category Match (25%)
            if rfp.category and vendor.industry_sector:
                rfp_category = rfp.category.strip().upper()
                vendor_sector = vendor.industry_sector.strip().upper()
                
                # Exact match
                if rfp_category == vendor_sector:
                    industry_score = 25
                # Partial match (contains)
                elif rfp_category in vendor_sector or vendor_sector in rfp_category:
                    industry_score = 20
                # Common sector mappings
                elif (rfp_category in ['IT', 'TECHNOLOGY', 'SOFTWARE'] and vendor_sector in ['IT', 'TECHNOLOGY', 'SOFTWARE']):
                    industry_score = 23
                elif (rfp_category in ['CONSULTING', 'SERVICES'] and vendor_sector in ['CONSULTING', 'SERVICES']):
                    industry_score = 23
                else:
                    # Vendor has industry but doesn't match RFP category
                    industry_score = 0
            elif not vendor.industry_sector:
                # Vendor industry is unknown/missing - no points (real value, not grace)
                industry_score = 0
            else:
                # RFP has no category requirement - neutral score (not vendor's fault)
                industry_score = 12
            
            # 3. Budget Match (25%)
            if rfp.estimated_value and vendor.annual_revenue:
                rfp_budget = float(rfp.estimated_value)
                vendor_revenue = float(vendor.annual_revenue)
                
                # Calculate ratio (vendor revenue should ideally be 5-50x the project budget)
                ratio = vendor_revenue / rfp_budget if rfp_budget > 0 else 0
                
                if 5 <= ratio <= 50:
                    # Ideal range - vendor has sufficient capacity
                    budget_score = 25
                elif 3 <= ratio < 5 or 50 < ratio <= 100:
                    # Acceptable range
                    budget_score = 20
                elif 1 <= ratio < 3 or 100 < ratio <= 500:
                    # Marginal range
                    budget_score = 15
                elif ratio < 1:
                    # Vendor too small for project
                    budget_score = 5
                else:
                    # Vendor might be too large (could be overpriced)
                    budget_score = 10
                
                # Also check budget range if specified
                if rfp.budget_range_min and rfp.budget_range_max:
                    budget_min = float(rfp.budget_range_min)
                    budget_max = float(rfp.budget_range_max)
                    budget_mid = (budget_min + budget_max) / 2
                    ratio_range = vendor_revenue / budget_mid if budget_mid > 0 else 0
                    
                    if 5 <= ratio_range <= 50:
                        budget_score = max(budget_score, 25)
            elif not vendor.annual_revenue:
                # Vendor revenue is unknown/missing - no points (real value, not grace)
                budget_score = 0
            else:
                # RFP has no budget specified - neutral score (not vendor's fault)
                budget_score = 12
            
            # 4. Business Type Match (20%)
            if rfp.rfp_type and vendor.business_type:
                rfp_type = rfp.rfp_type.strip().upper()
                vendor_type = vendor.business_type.strip().upper()
                
                # Direct type matching
                if rfp_type == vendor_type:
                    business_type_score = 20
                # Common type mappings
                elif (rfp_type in ['SOFTWARE', 'TECHNOLOGY'] and vendor_type in ['SOFTWARE', 'TECHNOLOGY', 'IT']):
                    business_type_score = 18
                elif (rfp_type in ['SERVICES', 'CONSULTING'] and vendor_type in ['SERVICES', 'CONSULTING']):
                    business_type_score = 18
                elif (rfp_type == 'INFRASTRUCTURE' and vendor_type in ['CONSTRUCTION', 'ENGINEERING', 'INFRASTRUCTURE']):
                    business_type_score = 18
                else:
                    # Vendor has business type but doesn't match RFP type
                    business_type_score = 0
            elif not vendor.business_type:
                # Vendor business type is unknown/missing - no points (real value, not grace)
                business_type_score = 0
            else:
                # RFP has no type requirement - neutral score (not vendor's fault)
                business_type_score = 10
            
            # Calculate total match score
            total_match_score = location_score + industry_score + budget_score + business_type_score
            
            # Update vendor match score in database
            vendor.match_score = Decimal(str(total_match_score))
            vendor.save(update_fields=['match_score'])
            
            match_results.append({
                'vendor_id': vendor.vendor_id,
                'company_name': vendor.company_name,
                'match_score': total_match_score,
                'breakdown': {
                    'location': location_score,
                    'industry': industry_score,
                    'budget': budget_score,
                    'business_type': business_type_score
                }
            })
        
        return JsonResponse({
            'success': True,
            'message': f'Match scores calculated for {len(match_results)} vendors',
            'results': match_results
        })
        
    except Exception as e:
        import traceback
        print(f"Error calculating match scores: {str(e)}")
        print(traceback.format_exc())
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)